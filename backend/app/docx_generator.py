from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from io import BytesIO
from datetime import datetime
from typing import Any, List, Dict, Optional, Union
import requests
from . import models
from .utils import append_label_if_missing, clean_spelling_errors, format_no_field
from .letterhead_templates import get_template
import logging

from .docx_generation import (
    format_currency,
    format_currency_words,
    format_currency_aligned,
    format_room_count,
    round_for_say,
    format_material_list,
    )

from .docx_generation import (
    add_section_heading,
    add_market_value_line,
    add_value_rounded_line,
    add_inline_field,
    add_subsection_paragraph,
    format_building_valuation_2line,
    format_addon_compact,
    )

from .docx_generation import (
    add_border_to_paragraph,
    MAP_IMAGE_WIDTH , MAP_IMAGE_MAX_HEIGHT,PROPERTY_PHOTO_WIDTH, PROPERTY_PHOTO_HEIGHT, IMAGE_SPACING_BEFORE,
    IMAGE_SPACING_AFTER , MAJOR_SECTION_SPACE_BEFORE , MAJOR_SECTION_SPACE_AFTER ,

    SUBSECTION_SPACE_BEFORE ,    SUBSECTION_SPACE_AFTER ,

    BODY_PARA_SPACE_BEFORE ,    BODY_PARA_SPACE_AFTER ,

    INLINE_FIELD_SPACE_BEFORE ,    INLINE_FIELD_SPACE_AFTER ,

    SUBHEADING_SPACE_BEFORE ,    SUBHEADING_SPACE_AFTER ,

    INDENTED_CONTENT_SPACE_BEFORE ,    INDENTED_CONTENT_SPACE_AFTER ,    INDENTED_CONTENT_LEFT_INDENT ,

    BOUNDARY_LIST_SPACE_AFTER ,    ACCOMMODATION_ROOM_SPACE_AFTER,

    FONT_SIZE_DOCUMENT_TITLE ,    FONT_SIZE_SECTION_HEADING ,    FONT_SIZE_SUBSECTION_HEADING ,

    FONT_SIZE_BODY ,    FONT_SIZE_INLINE_LABEL,    FONT_SIZE_VALUATION ,

    FONT_SIZE_TABLE_HEADER,    FONT_SIZE_TABLE_CELL ,    FONT_SIZE_INVOICE_TOTAL ,

    FONT_SIZE_CAPTION ,    FONT_SIZE_BANK_HEADER,    FONT_SIZE_BANK_DETAILS,
    FONT_SIZE_SIGNATURE,    FONT_SIZE_CERTIFICATION,
    )

from .docx_generation import (
    calculate_image_dimensions,
    apply_letterbox_to_image,
    )

from .docx_generation import (
    safe_get_json_field,
    safe_get_array_item,
    to_float,
    safe_parse_json_string,
    safe_get_nested,
    )

logger = logging.getLogger(__name__)

Deed_Type = Union[Dict[str, Any], Any]

# ===== LEGAL ASPECTS PARAGRAPH GENERATORS =====

def generate_ownership_paragraph(report: Any) -> str:
    """
    Generate professional ownership paragraph with graceful handling of missing data.

    Template adapts based on available data:
    - Full deed info: Complete ownership statement with deed details
    - Survey plan only: Ownership based on plan identification
    - Minimal data: Fallback to valuation basis statement
    """
    parts: list[str] = []

    # Determine property identification method - USE SAFE HELPERS
    deeds = safe_get_json_field(report, 'deeds', [])
    has_deed = bool(deeds)
    has_plan = bool(safe_get_json_field(report, 'plan_number', None))

    # Owner name (use applicant if available) - USE SAFE HELPER
    owner_name = safe_get_json_field(report, 'applicant_full_name', None)

    # Fallback to placeholder if no owner name
    if not owner_name:
        owner_name = "[Applicant Name Not Provided]"

    # Title/pedigree search statement
    if getattr(report, 'title_search_conducted', None) == 'No' or \
       getattr(report, 'pedigree_search_conducted', None) == 'No':
        parts.append("I did not search regarding the history and pedigree of the property.")

    # Main ownership statement
    if has_deed and owner_name:
        # Deed-based ownership - USE SAFE ARRAY ACCESS
        deed : Deed_Type = safe_get_array_item(deeds, 0, {})
        # Handle both dict and object formats
        deed_type: str = deed.get('deed_type', 'transfer deed') if isinstance(deed, dict) else getattr(deed, 'deed_type', 'transfer deed')
        deed_number = deed.get('deed_number', '') if isinstance(deed, dict) else getattr(deed, 'deed_number', '')
        deed_date = deed.get('deed_date', '') if isinstance(deed, dict) else getattr(deed, 'deed_date', '')
        notary_name = deed.get('notary_name', '') if isinstance(deed, dict) else getattr(deed, 'notary_name', '')
        notary_location = deed.get('notary_location', '') if isinstance(deed, dict) else getattr(deed, 'notary_location', '')

        ownership_stmt = f"{owner_name} claims ownership to the property"

        if deed_number and deed_date:
            ownership_stmt += f" by {deed_type} No:{deed_number} dated {deed_date}"
        else:
            ownership_stmt += f" by {deed_type}"

        if notary_name and notary_location:
            ownership_stmt += f" attested by {notary_name} notary public in {notary_location} district"
        elif notary_name:
            ownership_stmt += f" attested by {notary_name} notary public"

        ownership_stmt += "."
        parts.append(ownership_stmt)

    elif has_plan and owner_name:
        # Plan-based ownership
        plan_number = report.plan_number
        plan_date = getattr(report, 'plan_date', '')
        surveyor = getattr(report, 'licensed_surveyor_name', '')

        ownership_stmt = f"{owner_name} claims ownership to the property by transfer deed for the property identified as per the Survey Plan No: {plan_number}"

        if plan_date:
            ownership_stmt += f" dated {plan_date}"
        if surveyor:
            ownership_stmt += f" prepared by {surveyor}"

        ownership_stmt += "."
        parts.append(ownership_stmt)

    # Encumbrance statement
    property_encumbered = getattr(report, 'property_encumbered', None)
    if property_encumbered == 'Yes':
        encumbrance_type = getattr(report, 'encumbrance_type', 'encumbrance')
        encumbrance_details = getattr(report, 'encumbrance_details', '')

        if encumbrance_type == 'Mortgage' and encumbrance_details:
            parts.append(f"This property is already mortgaged to {encumbrance_details}.")
        elif encumbrance_type == 'Life Interest' and encumbrance_details:
            parts.append(f"This property has a life interest held by {encumbrance_details}.")
        else:
            parts.append(f"This property has {encumbrance_type.lower()} encumbrance.")

    # Valuation basis statement
    ownership_type = getattr(report, 'ownership_type', 'freehold')
    valuation_basis = getattr(report, 'valuation_basis_note', None)

    if not valuation_basis:
        if property_encumbered == 'Yes':
            valuation_basis = "free from all legal encumbrance"
        else:
            valuation_basis = "free from all encumbrances"

    valuation_stmt = f"I valued {ownership_type.lower()} interest of the property {valuation_basis}."
    parts.append(valuation_stmt)

    # Fallback for minimal data
    if not parts or len(parts) == 1:
        return f"I did not search regarding the history and pedigree of the property. Valuation is based on {ownership_type.lower()} title {valuation_basis}."

    return " ".join(parts)


def generate_street_lines_paragraph(report) -> str:
    """
    Generate street lines paragraph with contextual information.
    """
    status = getattr(report, 'street_lines_status', 'not affected')
    status_lower = status.lower()

    parts = []

    # Main status statement
    parts.append(f"Street lines are {status_lower} to the property.")

    # Impact description if provided
    impact_desc = getattr(report, 'street_lines_impact_description', None)
    if impact_desc:
        parts.append(impact_desc)

    # Municipal context
    is_municipal = getattr(report, 'is_municipal_limit', False)
    municipal_context = "within Municipal Council Limits" if is_municipal else "outside Municipal Council Limits"

    # Gazette reference if available
    gazette_ref = getattr(report, 'street_lines_gazette_ref', None)
    gazette_date = getattr(report, 'street_lines_gazette_date', None)

    if gazette_ref and gazette_date:
        parts.append(f"Properties located {municipal_context} are subject to street line regulations as per Gazette No: {gazette_ref} dated {gazette_date}.")
    else:
        # Generic explanatory statement
        if status_lower == 'not affected':
            parts.append("Street lines are affected to the properties which are located along roads within Municipal Council Limits and imposed by a gazette.")
        else:
            parts.append(f"Properties located {municipal_context} are subject to street line regulations.")

    return " ".join(parts)


def generate_building_limits_paragraph(report) -> str:
    """
    Generate building limits paragraph with approval and distance information.
    """
    status = getattr(report, 'building_limits_status', 'not affected')
    status_lower = status.lower()

    parts = []

    # Main status statement
    parts.append(f"Building limits are {status_lower} to the property.")

    # Building plan approval
    plan_approved = getattr(report, 'building_plan_approved', None)
    approval_authority = getattr(report, 'building_approval_authority', None)

    if plan_approved == 'Yes':
        if approval_authority:
            parts.append(f"Building Plan approved by {approval_authority}.")
        else:
            parts.append("Building Plan approved by the local authority.")

    # Distance from road
    distance = getattr(report, 'building_distance_from_road', None)
    if distance:
        parts.append(f"Building limit is {distance} from the access road.")

    # Building within limits
    within_limits = getattr(report, 'building_within_limits', None)
    if within_limits == 'Yes':
        parts.append("Existing building is located inside of the building limits in accordance with the approved building plan.")
    elif within_limits == 'No':
        parts.append("Existing building is located outside of the building limits.")

    # Plan reference
    plan_ref = getattr(report, 'building_plan_reference', None)
    if plan_ref and plan_approved == 'Yes':
        parts.append(f"Building plan reference: {plan_ref}.")

    return " ".join(parts)


def generate_local_authority_paragraph(report) -> str:
    """
    Generate local authority paragraph with rating and administrative information.
    """
    # Check if user provided custom free-text (backward compatibility)
    custom_text = getattr(report, 'local_authority_data', None)
    if custom_text and len(custom_text) > 50:  # User wrote a full paragraph
        return custom_text

    parts = []

    # Administrative location
    pradeshiya_sabha = getattr(report, 'pradeshiya_sabha', None)
    district = getattr(report, 'property_district', None)
    province = getattr(report, 'property_province', None)

    if pradeshiya_sabha or district or province:
        location_parts = []
        if pradeshiya_sabha:
            location_parts.append(pradeshiya_sabha)
        if district:
            location_parts.append(f"{district} District")
        if province:
            location_parts.append(f"{province} Province")

        parts.append(", ".join(location_parts) + ".")

    # Rating status
    is_rated = getattr(report, 'local_authority_rated', None)

    if is_rated == 'Yes':
        parts.append("This area has been rated for the local levy of taxes.")

        # Tax levy details
        tax_levy = getattr(report, 'local_authority_tax_levy', None)
        if tax_levy:
            parts.append(tax_levy)

        # Assessment number
        assessment_num = getattr(report, 'assessment_number', None)
        if assessment_num:
            parts.append(f"Assessment No: {assessment_num}.")
    elif is_rated == 'No':
        parts.append("This area has not been rated for the local levy of taxes.")

    # Fallback to custom text if no structured data
    if not parts and custom_text:
        return custom_text

    return " ".join(parts)


def generate_rent_act_paragraph(report) -> str:
    """
    Generate rent act effectiveness paragraph.
    """
    status = getattr(report, 'rent_act_effectiveness', 'Not affected')

    # Direct mapping from dropdown options
    status_map = {
        'Not affected': "This property is not affected by the rent act.",
        'Subject to Rent Act No. 7 of 1972': "This property is subject to the rent act of No. 7 of 1972.",
        'Subject to Rent Act Amendment No. 26 of 2002': "This property is subject to the rent act of No. 7 of 1972 (Amendment No. 26 of 2002).",
        'Partially affected': "This property is partially affected by the rent control regulations."
    }

    return status_map.get(status, "This property is not affected by the rent act.")


# ===== NARRATIVE GENERATION HELPER FUNCTIONS =====

def _synthesize_location_context(locations: List[str]) -> str:
    """
    Synthesize location context from comparable location descriptions.

    Extracts common patterns and generates appropriate location phrase
    for land values narrative.

    Args:
        locations: List of location descriptions from comparable properties

    Returns:
        Synthesized location context string (e.g., "located in this area",
        "located fronting main highway", "located in Rambukkana area")

    Examples:
        >>> _synthesize_location_context([])
        'in this area'
        >>> _synthesize_location_context(['Rambukkana town', 'Rambukkana center'])
        'located in Rambukkana area'
        >>> _synthesize_location_context(['Fronting Kurunegala Highway'])
        'located fronting Kurunegala Highway'
    """
    if not locations:
        return "in this area"

    # Filter out empty locations
    valid_locations = [loc.strip() for loc in locations if loc and loc.strip()]

    if not valid_locations:
        return "in this area"

    # If single location
    if len(valid_locations) == 1:
        return f"located in {valid_locations[0]}"

    # Find common words (case-insensitive)
    from collections import Counter

    # Tokenize all locations
    all_words = []
    for loc in valid_locations:
        # Extract significant words (ignore common words)
        words = loc.lower().replace(',', '').split()
        all_words.extend(words)

    # Find most common significant words
    word_counts = Counter(all_words)

    # Common stop words to ignore
    stop_words = {'in', 'the', 'a', 'an', 'of', 'to', 'and', 'or', 'near', 'at', 'from', 'this', 'that'}
    significant_words = {word: count for word, count in word_counts.items()
                         if count >= len(valid_locations) * 0.5 and word not in stop_words}

    if significant_words:
        # Use most common word as area reference
        common_area = max(significant_words, key=significant_words.get)
        return f"located in {common_area.title()} area"

    # Check for highway/road frontage mentions
    combined = " ".join(valid_locations).lower()
    if 'highway' in combined:
        # Try to extract highway name if possible
        for loc in valid_locations:
            if 'highway' in loc.lower():
                # Try to extract highway name
                words = loc.split()
                for i, word in enumerate(words):
                    if 'highway' in word.lower() and i > 0:
                        return f"located fronting {' '.join(words[max(0,i-2):i+1])}"
                return "located fronting main highway"

    if 'road' in combined:
        # Check if specific road name is mentioned
        for loc in valid_locations:
            if 'road' in loc.lower() and '-' in loc:
                # Might be like "Kurunegala-Puttalam road"
                return f"located fronting {loc}"
        return "located with road access in this area"

    # Generic fallback
    return "located in this village"


def generate_land_values_paragraph(comparables: List[dict]) -> str:
    """
    Generate professional land values paragraph from comparable properties data.

    Groups comparables by property type and generates narrative description
    following Sri Lankan valuation report conventions.

    Args:
        comparables: List of comparable property dictionaries with keys:
            - property_type: 'Commercial' | 'Residential' | 'Agricultural'
            - location_description: str
            - extent: float (perches)
            - rate_per_perch: float (LKR)

    Returns:
        Formatted paragraph(s) describing land values in natural language.
        Returns empty string if no valid comparables.

    Examples:
        >>> comps = [
        ...     {'property_type': 'Residential', 'extent': 15, 'rate_per_perch': 50000,
        ...      'location_description': 'Rambukkana town'},
        ...     {'property_type': 'Residential', 'extent': 20, 'rate_per_perch': 60000,
        ...      'location_description': 'Rambukkana center'}
        ... ]
        >>> generate_land_values_paragraph(comps)
        'Residential building blocks in extent about 15 – 20 Perches located in Rambukkana area
         are being sold at rates ranging from Rs.50,000/= to Rs.60,000/= Per Perch...'
    """
    # Edge case: Empty or None
    if not comparables:
        return ""

    # Filter out invalid comparables (zero values)
    valid_comparables = [
        c for c in comparables
        if c.get('extent', 0) > 0 and c.get('rate_per_perch', 0) > 0
    ]

    if not valid_comparables:
        return ""  # Return empty string - let market analysis stand alone

    # Handle single comparable specially
    if len(valid_comparables) == 1:
        comp = valid_comparables[0]
        prop_type = comp.get('property_type', 'Residential')

        # Property type descriptors
        type_descriptors = {
            'Residential': 'Residential building blocks',
            'Commercial': 'Commercial building blocks',
            'Agricultural': 'Agricultural land parcels'
        }
        property_descriptor = type_descriptors.get(prop_type, 'Building blocks')

        location = comp.get('location_description', 'this area')
        extent = comp.get('extent', 0)
        rate = comp.get('rate_per_perch', 0)

        return (
            f"{property_descriptor} of about {extent:.0f} Perches "
            f"in {location} is valued at approximately Rs.{rate:,.0f}/= Per Perch."
        )

    # Group by property type
    property_groups = {}
    for comp in valid_comparables:
        prop_type = comp.get('property_type', 'Residential')
        if prop_type not in property_groups:
            property_groups[prop_type] = []
        property_groups[prop_type].append(comp)

    # Property type ordering
    type_order = ['Residential', 'Commercial', 'Agricultural']

    paragraphs = []

    for prop_type in type_order:
        if prop_type not in property_groups:
            continue

        group = property_groups[prop_type]

        # Calculate ranges
        extents = [c['extent'] for c in group]
        rates = [c['rate_per_perch'] for c in group]

        min_extent = min(extents)
        max_extent = max(extents)
        min_rate = min(rates)
        max_rate = max(rates)

        # Format extent range
        if min_extent == max_extent:
            extent_str = f"about {min_extent:.0f} Perches"
        else:
            extent_str = f"about {min_extent:.0f} – {max_extent:.0f} Perches"

        # Format rate range with proper currency formatting
        def format_rate(rate):
            # Format with thousands separator and /= suffix
            return f"Rs.{rate:,.0f}/="

        if min_rate == max_rate:
            rate_str = f"{format_rate(min_rate)} Per Perch"
        else:
            rate_str = f"{format_rate(min_rate)} to {format_rate(max_rate)} Per Perch"

        # Extract location context
        locations = [c.get('location_description', '') for c in group]
        location_context = _synthesize_location_context(locations)

        # Determine property type descriptor
        type_descriptors = {
            'Residential': 'Residential building blocks',
            'Commercial': 'Commercial building blocks',
            'Agricultural': 'Agricultural land parcels'
        }
        property_descriptor = type_descriptors.get(prop_type, 'Building blocks')

        # Calculate rate variance for contextual phrase
        rate_variance = ((max_rate - min_rate) / min_rate * 100) if min_rate > 0 else 0

        if rate_variance < 20:
            context_phrase = "showing stable market conditions"
        elif rate_variance > 50:
            context_phrase = "depending upon the location, topography and nature of the property and convenience etc."
        else:
            context_phrase = "to be influenced by factors affecting the property market"

        # Construct paragraph
        paragraph = (
            f"{property_descriptor} in extent {extent_str} {location_context} "
            f"are being sold at rates ranging from {rate_str}, {context_phrase}."
        )

        paragraphs.append(paragraph)

    return " ".join(paragraphs)


def generate_simplified_certification_text(
    valuer_name: str,
    valuer_designation: str,
    lot_number: Optional[str],
    plan_number: Optional[str],
    plan_date: Optional[str],
    licensed_surveyor_name: Optional[str],
    deeds: Optional[List[dict]],
    property_identification_type: Optional[str]
) -> str:
    """
    Generate simplified single-paragraph certification text.

    New simplified format (replaces old multi-paragraph format):
    "I, [Valuer Name], [Designation], do hereby certify that the property
    inspected by me and valued above is identical to the property depicted as
    [Lot X in] Plan No [number] dated [date] made by [Surveyor], Licensed Surveyor."

    Key behaviors:
    - If lot_number provided and plan exists: "...property depicted as Lot {lot_number} in Plan No..."
    - If no lot_number: "...property depicted as Plan No..."
    - Always use "property depicted as" (not "land depicted as")
    - Only mention plan (not deed) in certification
    - If only deed (no plan): "...property described in {deed_type} No..."
    - If lot_number exists but no plan_number: skip lot, use deed/generic

    Args:
        valuer_name: Name of the valuer
        valuer_designation: Professional designation (e.g., "Chartered Valuer")
        lot_number: Lot number (e.g., "Lot 15", "Lots 1 & 2")
        plan_number: Survey plan number
        plan_date: Survey plan date
        licensed_surveyor_name: Name of surveyor
        deeds: List of deed dictionaries
        property_identification_type: Type of identification

    Returns:
        Formatted single-paragraph certification text
    """
    # Start with valuer identification
    cert_text = f"I, {valuer_name}, {valuer_designation}, do hereby certify that the property inspected by me and valued above is identical to "

    # Determine property identification string
    property_id = None

    # Infer identification type if not provided (backward compatibility)
    if not property_identification_type:
        if plan_number and deeds and len(deeds) > 0:
            property_identification_type = "plan_and_deed"
        elif plan_number:
            property_identification_type = "plan"
        elif deeds and len(deeds) > 0:
            property_identification_type = "deed"

    # Build property identification based on type
    if property_identification_type in ["plan", "plan_and_deed"]:
        # Plan-based identification (ignore deed as per requirements)
        if plan_number:
            # Check if we should include lot number
            # Only include lot if both lot_number and plan_number exist
            if lot_number and lot_number.strip():
                # Format with lot number
                plan_ref = format_no_field("Plan", plan_number)
                property_id = f"the property depicted as Lot {lot_number} in {plan_ref}"
            else:
                # Format without lot number
                plan_ref = format_no_field("Plan", plan_number)
                property_id = f"the property depicted as {plan_ref}"

            # Add plan date
            if plan_date:
                property_id += f" dated {plan_date}"

            # Add surveyor
            if licensed_surveyor_name:
                property_id += f" made by {licensed_surveyor_name}, Licensed Surveyor"

    # If no plan, fall back to deed
    if not property_id and property_identification_type in ["deed", "plan_and_deed", "certificate_of_sale"]:
        # Deed-based identification (fallback)
        has_deed = isinstance(deeds, list) and len(deeds) > 0
        if has_deed:
            deed = deeds[0]
            deed_number = deed.get('deed_number', '') if isinstance(deed, dict) else getattr(deed, 'deed_number', '')
            deed_type = deed.get('deed_type', 'Deed') if isinstance(deed, dict) else getattr(deed, 'deed_type', 'Deed')
            deed_date = deed.get('deed_date', '') if isinstance(deed, dict) else getattr(deed, 'deed_date', '')

            if deed_number:
                property_id = f"the property described in {deed_type} No. {deed_number}"
                if deed_date:
                    property_id += f" dated {deed_date}"

    # Final fallback if no identification data
    if not property_id:
        property_id = "the property described in this report"

    # Complete the certification text
    cert_text += property_id + "."

    return cert_text


# DEPRECATED: Keep for backward compatibility with old code
def generate_certificate_of_identity_text(
    property_identification_type: Optional[str],
    plan_number: Optional[str],
    plan_date: Optional[str],
    deeds: Optional[List[dict]],
    licensed_surveyor_name: Optional[str] = None
) -> Optional[str]:
    """
    DEPRECATED: Use generate_simplified_certification_text instead.

    This function is kept for backward compatibility only.
    Generates old-style Certificate of Identity text.
    """
    if not property_identification_type:
        # Backward compatibility: infer type from available data
        if plan_number and deeds and len(deeds) > 0:
            property_identification_type = "plan_and_deed"
        elif plan_number:
            property_identification_type = "plan"
        elif deeds and len(deeds) > 0:
            property_identification_type = "deed"
        else:
            return None

    # Base text
    base_text = "I certify that the property inspected by me is identical to the property described in "

    # Extract deed information if available
    has_deed = isinstance(deeds, list) and len(deeds) > 0
    deed_text = ""
    if has_deed:
        deed = deeds[0]
        deed_number = deed.get('deed_number', '') if isinstance(deed, dict) else getattr(deed, 'deed_number', '')
        deed_type = deed.get('deed_type', 'Deed') if isinstance(deed, dict) else getattr(deed, 'deed_type', 'Deed')
        deed_date = deed.get('deed_date', '') if isinstance(deed, dict) else getattr(deed, 'deed_date', '')

        if deed_number:
            deed_text = f"{deed_type} No. {deed_number}"
            if deed_date:
                deed_text += f" dated {deed_date}"

    # Generate text based on identification type
    if property_identification_type == "plan":
        if not plan_number:
            return None
        plan_ref = format_no_field("Plan", plan_number)
        text = base_text + plan_ref
        if plan_date:
            text += f" dated {plan_date}"
        if licensed_surveyor_name:
            text += f" made by {licensed_surveyor_name}, Licensed Surveyor"
        text += "."
        return text

    elif property_identification_type == "deed":
        if not deed_text:
            return None
        return base_text + deed_text + "."

    elif property_identification_type == "plan_and_deed":
        if not deed_text and not plan_number:
            return None

        # If only one is available, use whichever exists
        if not deed_text and plan_number:
            plan_ref = format_no_field("Plan", plan_number)
            text = base_text + plan_ref
            if plan_date:
                text += f" dated {plan_date}"
            if licensed_surveyor_name:
                text += f" made by {licensed_surveyor_name}, Licensed Surveyor"
            text += "."
            return text

        if deed_text and not plan_number:
            return base_text + deed_text + "."

        # Both available - combine them
        plan_ref = format_no_field("Plan", plan_number)
        text = base_text + deed_text + " and identified in " + plan_ref
        if plan_date:
            text += f" dated {plan_date}"
        if licensed_surveyor_name:
            text += f" made by {licensed_surveyor_name}, Licensed Surveyor"
        text += "."
        return text

    elif property_identification_type == "certificate_of_sale":
        # Certificate of sale is similar to deed
        if not deed_text:
            return None
        return base_text + deed_text + "."

    # Unknown type or no data
    return None


def add_signature_block(
    doc,
    user,
    valuer_name: Optional[str],
    valuer_designation: Optional[str],
    certification_date: Optional[str]
):
    """
    Add a standardized signature block to the document.

    Format:
    - Spacing before signature
    - Underline ("_" * 40) for signature
    - Optional signature image (if user has one)
    - Valuer name (bold, 9pt)
    - Valuer designation (9pt)
    - Certification date (9pt)

    Args:
        doc: Document object
        user: User model instance
        valuer_name: Name of valuer
        valuer_designation: Professional designation
        certification_date: Date of certification
    """
    # Add spacing before signature
    doc.add_paragraph("\n")

    # Signature line (underline)
    sig_line = doc.add_paragraph("_" * 40)
    sig_line.paragraph_format.space_before = Pt(24)
    sig_line.paragraph_format.space_after = Pt(6)

    # Add signature image if available
    if hasattr(user, 'signature_image') and user.signature_image:
        try:
            response = requests.get(user.signature_image, timeout=10)
            if response.status_code == 200:
                sig_para = doc.add_paragraph()
                sig_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                sig_para.paragraph_format.space_before = Pt(0)
                sig_para.paragraph_format.space_after = Pt(6)
                image_stream = BytesIO(response.content)
                sig_para.add_run().add_picture(image_stream, width=Inches(2))
        except Exception as e:
            logger.warning(f"Failed to add signature image: {e}")

    # Valuer name
    if valuer_name:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(valuer_name)
        run.font.bold = True
        run.font.size = FONT_SIZE_BODY

    # Valuer designation
    if valuer_designation:
        p = doc.add_paragraph(valuer_designation)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = FONT_SIZE_BODY

    # Certification date
    if certification_date:
        p = doc.add_paragraph(certification_date)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = FONT_SIZE_BODY

    doc.add_paragraph()


def get_pronoun(title: Optional[str]) -> Dict[str, str]:
    """
    Determine pronouns based on title.
    Returns: {'subject': 'he'/'she', 'object': 'him'/'her', 'possessive': 'his'/'her'}
    """
    if not title:
        return {'subject': 'he', 'object': 'him', 'possessive': 'his'}

    title_lower = title.lower().strip()

    # Female titles
    if title_lower in ['mrs.', 'mrs', 'miss.', 'miss', 'ms.', 'ms']:
        return {'subject': 'she', 'object': 'her', 'possessive': 'her'}

    # Default to male
    return {'subject': 'he', 'object': 'him', 'possessive': 'his'}

def generate_title_block(report: models.Report) -> str:
    """Generate the title block with support for plan/deed/certificate identification types"""
    lines = []
    lines.append("VALUATION REPORT")
    lines.append("of")

    # Determine identification type (with backward compatibility for old reports)
    id_type = report.property_identification_type
    if not id_type:
        # Infer type from existing data for old reports
        if report.plan_number:
            id_type = "plan"
        elif report.deeds:
            # Old reports with deed data
            id_type = "deed"
        else:
            # Fallback to plan-style if nothing is detected
            id_type = "plan"

    # Generate title based on identification type
    if id_type == "plan" or id_type == "plan_and_deed":
        # Plan-based title (original format)
        # For plan_and_deed: show plan info in header, deed info appears in certification section
        lot_desc = report.lot_number or '[Lot Number]'

        # Remove common prefixes that shouldn't be in lot description
        lot_desc_stripped = lot_desc.strip()
        prefixes_to_remove = ['plan no', 'plan no:', 'lot plan no', 'lot plan no:']
        lot_desc_lower = lot_desc_stripped.lower()
        for prefix in prefixes_to_remove:
            if lot_desc_lower.startswith(prefix):
                lot_desc = lot_desc_stripped[len(prefix):].strip()
                break

        # Ensure it has "Lot" prefix
        if not lot_desc.lower().startswith('lot'):
            lot_desc = f"Lot {lot_desc}"

        plan_num = report.plan_number or '[Plan Number]'
        plan_formatted = format_no_field("Plan", plan_num)

        # Combine date with first line
        plan_date = report.plan_date or '[Date]'
        prop_desc = f"The Property Depicted as {lot_desc} in {plan_formatted} Dated {plan_date}"
        lines.append(prop_desc)

        # Second line starts with "made by"
        surveyor = report.licensed_surveyor_name or '[Surveyor Name]'
        plan_info = f"made by {surveyor} Licensed Surveyor."
        lines.append(plan_info)

    elif id_type == "deed":
        # Deed-based title: show deed information
        deed_type = "Deed"
        deed_number = "[Deed Number]"
        deed_date = "[Date]"

        # Extract deed info from report.deeds JSON array (first deed)
        if report.deeds and isinstance(report.deeds, list) and len(report.deeds) > 0:
            first_deed = report.deeds[0]
            deed_type = first_deed.get('deed_type', 'Deed')
            deed_number = first_deed.get('deed_number', '[Deed Number]')
            deed_date = first_deed.get('deed_date', '[Date]')
        else:
            logger.warning(f"[DOCX] Report {report.id} has deed identification type but no deed data")

        # Format deed number with "No."
        deed_number_formatted = format_no_field("", deed_number, include_label=False)
        prop_desc = f"The Property Depicted as described in {deed_type} No. {deed_number_formatted} dated {deed_date}"
        lines.append(prop_desc)

    elif id_type == "certificate_of_sale":
        # Certificate of sale title: show certificate information (stored as deed internally)
        deed_number = "[Certificate Number]"
        deed_date = "[Date]"

        # Extract certificate info from report.deeds JSON array (certificate stored as deed internally)
        if report.deeds and isinstance(report.deeds, list) and len(report.deeds) > 0:
            first_deed = report.deeds[0]
            deed_number = first_deed.get('deed_number', '[Certificate Number]')
            deed_date = first_deed.get('deed_date', '[Date]')
        else:
            logger.warning(f"[DOCX] Report {report.id} has certificate_of_sale identification type but no deed data")

        # Format certificate number with "No."
        cert_number_formatted = format_no_field("", deed_number, include_label=False)
        prop_desc = f"The Property Depicted as described in Certificate of Sale No. {cert_number_formatted} dated {deed_date}"
        lines.append(prop_desc)

    else:
        # Fallback: Address-based title for old reports without identification type
        address = generate_smart_address(report) or '[Property Address]'
        prop_desc = f"The Property Depicted as {address}"
        lines.append(prop_desc)

    return lines

def generate_applicant_statement(report: models.Report) -> str:
    """Generate the applicant statement paragraph with smart grammar"""

    # Build applicant full description
    applicant_desc = f"{report.applicant_title or ''} {report.applicant_full_name or '[Applicant Name]'}".strip()

    # ID information - handle optional fields gracefully
    if report.applicant_id_number and report.applicant_id_type:
        id_no_formatted = format_no_field("", report.applicant_id_number, include_label=False)
        id_info = f"holder {report.applicant_id_type} {id_no_formatted}"
    else:
        # No ID provided - skip this part of the introduction
        id_info = ""

    # Address parts
    address_parts = []
    if report.applicant_address_line1:
        address_parts.append(report.applicant_address_line1)
    if report.applicant_address_line2:
        address_parts.append(report.applicant_address_line2)
    if report.applicant_district:
        cleaned = clean_spelling_errors(report.applicant_district)
        district_text = append_label_if_missing(cleaned, "District", variants=["district"])
        address_parts.append(district_text)
    if report.applicant_province:
        address_parts.append(report.applicant_province)
    if report.applicant_country:
        address_parts.append(report.applicant_country)

    address_str = ", ".join(address_parts) if address_parts else "[Address]"

    # Get pronouns
    pronouns = get_pronoun(report.applicant_title)

    # Build ownership text (auto-generated based on applicant title)
    ownership_text = f"owned by {pronouns['object']}"

    # Property type
    property_type = report.property_type_valued or "immovable property"

    # Valuation type
    valuation_type = report.valuation_type or "Market Value"

    # Build paragraph - conditionally include ID info
    if id_info:
        paragraph1 = f"This Valuation Report is furnished at the request of {applicant_desc} {id_info} of {address_str}."
    else:
        paragraph1 = f"This Valuation Report is furnished at the request of {applicant_desc} of {address_str}."

    # Second part - wishes to know
    wish_text = f"{applicant_desc} wishes to know the {valuation_type} of {property_type} {ownership_text} in the Democratic Socialist Republic of Sri Lanka."

    # Handle additional owners if any
    if report.has_additional_owner == "yes" and report.additional_owner_names:
        wish_text = wish_text.replace(f"{ownership_text}", f"{ownership_text} & {pronouns['possessive']} family {report.additional_owner_names}")

    return [paragraph1, wish_text]

def generate_organization_side_introduction(report: models.Report) -> List[str]:
    """
    Generate organization-side introduction format.

    Format:
    - Paragraph 1: "At the request of [position] [organization] [address],
                    I am furnishing a Valuation Report of the above property
                    for the [purpose] purpose."
    - Paragraph 2-4: Formatted applicant details
    - Paragraph 5: Inspection date (handled separately in main function)

    Returns:
        List of paragraph strings
    """
    paragraphs = []

    # Paragraph 1: Request statement
    request_parts = []

    if report.submission_recipient_position:
        request_parts.append(report.submission_recipient_position)

    if report.submission_organization:
        request_parts.append(report.submission_organization)

    if report.submission_address:
        request_parts.append(report.submission_address)

    requester_text = ", ".join(request_parts) if request_parts else "[Requesting Organization]"

    # Get purpose text
    purpose = report.valuation_purpose or "[purpose]"

    para1 = f"At the request of {requester_text}, I am furnishing a Valuation Report of the above property for the {purpose} purpose."
    paragraphs.append(para1)

    # Paragraph 2: Applicant info
    applicant_desc = f"{report.applicant_title or ''} {report.applicant_full_name or '[Applicant Name]'}".strip()

    # ID information
    id_text = ""
    if report.applicant_id_number and report.applicant_id_type:
        id_no_formatted = format_no_field("", report.applicant_id_number, include_label=False)
        id_text = f" holder {report.applicant_id_type} {id_no_formatted}"

    para2 = f"Applicant        :-{applicant_desc}{id_text}"
    paragraphs.append(para2)

    # Paragraph 3: Address
    address_parts = []
    if report.applicant_address_line1:
        address_parts.append(report.applicant_address_line1)
    if report.applicant_address_line2:
        address_parts.append(report.applicant_address_line2)
    if report.applicant_district:
        cleaned = clean_spelling_errors(report.applicant_district)
        district_text = append_label_if_missing(cleaned, "District", variants=["district"])
        address_parts.append(district_text)
    if report.applicant_province:
        address_parts.append(report.applicant_province)
    if report.applicant_country:
        address_parts.append(report.applicant_country)

    address_str = ", ".join(address_parts) if address_parts else "[Address]"
    para3 = f"Address          :-{address_str}"
    paragraphs.append(para3)

    # Paragraph 4: Contact Number
    contact = report.applicant_contact_number or "[Contact Number]"
    para4 = f"Contact No       :-{contact}"
    paragraphs.append(para4)

    return paragraphs

def generate_multi_property_concluding_statement(
    report: models.Report,
    user: models.User,
    grand_total: float
) -> List[str]:
    """
    Generate concluding statement for multi-property CLIENT REQUEST reports.

    This function creates the closing statement that appears after the property
    summary table and inspection date, including the total market value in words
    and the valuer's signature.

    Args:
        report: Report model instance
        user: User model instance (for valuer name and credentials)
        grand_total: Total valuation amount from property table

    Returns:
        List of paragraph text strings:
        - [0]: Market value statement with amount in words
        - [1]: Valuer name line (Vlr.[Honorific] [Full Name])
        - [2]: Professional designation line

        Returns empty list if report is organization request.

    Example output:
        [
            "Present Market Value of the Properties claimed by Mr. John Doe & his family...",
            "Vlr.K D A Nimalsiri",
            "Chartered Valuer"
        ]
    """
    # Only generate for client requests
    if report.request_type != 'client_request':
        return []

    # Get pronoun based on applicant title
    pronouns = get_pronoun(report.applicant_title)
    gender_pronoun = pronouns['possessive']  # "his" or "her"

    # Build applicant name
    title = report.applicant_title or "Mr."
    full_name = report.applicant_full_name or "[Applicant Name]"

    # Convert amount to words
    amount_words = format_currency_words(grand_total)

    # Construct statement
    statement = (
        f"Present Market Value of the Properties claimed by {title} {full_name} "
        f"& {gender_pronoun} family in the Democratic Socialist Republic of Sri Lanka "
        f"are in a sum of Lanka Rupees {amount_words} only."
    )

    # Build valuer signature lines
    honorific = user.honorific or ""
    valuer_name = user.full_name or "[Valuer Name]"
    designation = user.professional_designation or "[Professional Designation]"

    # Format valuer name line: "Vlr.Honorific FullName" (remove extra spaces)
    valuer_line = f"Vlr.{honorific} {valuer_name}".replace("  ", " ").strip()

    return [statement, valuer_line, designation]

def generate_deed_description(deeds: Optional[List[Dict]]) -> Optional[str]:
    """Generate deed description text if deeds are provided"""
    if not deeds or len(deeds) == 0:
        return None

    deed_parts = []
    for deed in deeds:
        deed_type = deed.get('deed_type', 'Deed')
        deed_num = deed.get('deed_number', '[Number]')
        deed_date = deed.get('deed_date', '[Date]')

        deed_no_formatted = format_no_field("", deed_num, include_label=False)
        deed_text = f"{deed_type} {deed_no_formatted} dated {deed_date}"

        # Add notary info if available
        if deed.get('notary_name') or deed.get('notary_location'):
            notary_parts = []
            if deed.get('notary_name'):
                notary_parts.append(f"attested by {deed['notary_name']} Notary Public")
            if deed.get('notary_location'):
                cleaned = clean_spelling_errors(deed['notary_location'])
                location_text = append_label_if_missing(cleaned, "District", variants=["district"])
                notary_parts.append(f"in {location_text}")
            if notary_parts:
                deed_text += " " + " ".join(notary_parts)

        deed_parts.append(deed_text)

    if len(deed_parts) == 1:
        return f"requesting a Valuation Report of the immovable properties described in schedules of {deed_parts[0]}."
    else:
        # Multiple deeds - join with &
        deeds_text = " & schedule of ".join(deed_parts)
        return f"requesting a Valuation Report of the immovable properties described in schedules of {deeds_text}."

def generate_submission_statement(report: models.Report) -> Optional[str]:
    """Generate submission destination statement with optional purpose and recipient position"""
    if not report.submission_organization:
        return None

    org_text = ""

    # Add recipient position if available
    if report.submission_recipient_position:
        org_text = f"{report.submission_recipient_position}, "

    org_text += report.submission_organization

    # Add address if available
    if report.submission_address:
        org_text += f", {report.submission_address}"

    statement = f"This Valuation Report is to be submitted to {org_text}"

    # Add purpose if available
    if report.valuation_purpose:
        statement += f" for the purpose of {report.valuation_purpose}"

    statement += "."
    return statement

def generate_situation_text(report: models.Report) -> Optional[str]:
    """
    Generate SITUATION section text from property location data
    Following professional Sri Lankan valuation report format
    """
    print(f"[SITUATION] Report ID: {report.id}")
    print(f"[SITUATION] Village: {report.property_village}")
    print(f"[SITUATION] District: {report.property_district}")
    print(f"[SITUATION] GN Division: {report.grama_niladari_division}")
    print(f"[SITUATION] DS Division: {report.property_divisional_secretariat}")
    print(f"[SITUATION] Korale: {report.korale}")
    print(f"[SITUATION] Pradeshiya Sabha: {report.pradeshiya_sabha}")

    # Need at least some basic location info
    if not (report.property_village or report.property_district):
        print("[SITUATION] No location data found, returning None")
        return None

    # Build the comprehensive situation text parts
    parts = []

    # Start with "The property is situated"
    # Handle Assessment Number if municipal/urban property
    prefix = "The property"
    if report.assessment_number:
        assessment_formatted = format_no_field("Assessment", report.assessment_number, include_label=False)
        prefix += f" to be valued is situated bearing {assessment_formatted}"
        if report.land_traditional_name:
            prefix += f" {report.land_traditional_name}"
    elif report.land_traditional_name:
        prefix += f" is situated at {report.land_traditional_name}"
    else:
        prefix += " is situated"

    # Add village information
    if report.property_village:
        if " is situated" in prefix and prefix.endswith("situated"):
            parts.append(f"in {report.property_village} village")
        else:
            parts.append(f" in {report.property_village} village")

    # Add property number (ward/lot number)
    if report.property_number:
        property_no_formatted = format_no_field("", report.property_number, include_label=False)
        parts.append(f"in {property_no_formatted}")

    # Add location direction if available (e.g., "north-east")
    direction_suffix = ""
    if report.location_direction:
        direction_suffix = f"-{report.location_direction}"

    # Add Grama Niladari Division (with intelligent label handling)
    if report.grama_niladari_division:
        cleaned = clean_spelling_errors(report.grama_niladari_division)
        gn_text = append_label_if_missing(
            cleaned,
            "Grama Niladari division",
            variants=["grama niladari division", "Grama Niladari Division", "GN division", "GN Division"]
        )
        # Inject direction suffix before "Grama Niladari division" if it exists
        if direction_suffix and "Grama Niladari division" in gn_text:
            gn_text = gn_text.replace(" Grama Niladari division", f"{direction_suffix} Grama Niladari division")
        elif direction_suffix:
            # If variant was used, just prepend the direction
            gn_text = f"{gn_text}{direction_suffix}"
        parts.append(gn_text)

    # Add Divisional Secretariat (with intelligent label handling)
    if report.property_divisional_secretariat:
        cleaned = clean_spelling_errors(report.property_divisional_secretariat)
        ds_text = append_label_if_missing(
            cleaned,
            "Divisional Secretariat",
            variants=[
                "divisional secretariat division",
                "Divisional Secretariat Division",
                "divisional secretariat",
                "DS division",
                "DS Division"
            ]
        )
        # Add "division of" wrapper, but check if "division" is already in the text
        if "division" in ds_text.lower():
            parts.append(f"within the {ds_text} of")
        else:
            parts.append(f"within the {ds_text} division of")

    # Add Hathpaththuwa (with intelligent label handling)
    if report.hathpaththuwa:
        cleaned = clean_spelling_errors(report.hathpaththuwa)
        hathpaththuwa_text = append_label_if_missing(cleaned, "Hathpaththuwa", variants=["hathpaththuwa", "hatpattu"])
        parts.append(f"in {hathpaththuwa_text}")

    # Add Korale (with intelligent label handling)
    if report.korale:
        cleaned = clean_spelling_errors(report.korale)
        korale_text = append_label_if_missing(cleaned, "Korale", variants=["korale"])
        parts.append(f"in {korale_text}")

    # Add Pradeshiya Sabha or Municipal Council (with intelligent label handling)
    if report.is_municipal_limit and report.property_district:
        district_cleaned = clean_spelling_errors(report.property_district)
        district_text = append_label_if_missing(district_cleaned, "District", variants=["district"])
        parts.append(f"within the Municipal Council limit of {district_text}")
    elif report.pradeshiya_sabha:
        cleaned = clean_spelling_errors(report.pradeshiya_sabha)
        ps_text = append_label_if_missing(cleaned, "Pradeshiya Sabha", variants=["pradeshiya sabha", "PS"])
        parts.append(f"of {ps_text}")

    # Add District (only if not in municipal limit, with intelligent label handling)
    if report.property_district and not report.is_municipal_limit:
        cleaned = clean_spelling_errors(report.property_district)
        district_text = append_label_if_missing(cleaned, "District", variants=["district"])
        parts.append(f"in {district_text}")

    # Add Province (with intelligent label handling)
    if report.property_province:
        cleaned = clean_spelling_errors(report.property_province)
        province_text = append_label_if_missing(cleaned, "Province", variants=["province"])
        parts.append(province_text)

    # Add "of Sri Lanka" at the end
    if parts:
        parts.append("of Sri Lanka")

    # Assemble the complete sentence
    if parts:
        situation_text = prefix + " " + " ".join(parts) + "."
    else:
        # Fallback to smart address if no detailed location data
        smart_address = generate_smart_address(report)
        if smart_address:
            situation_text = f"The property is situated at {smart_address}."
        else:
            print("[SITUATION] No situation text could be generated")
            return None

    print(f"[SITUATION] Generated text: {situation_text}")
    return situation_text

def generate_smart_address(report: models.Report) -> Optional[str]:
    """
    Generate formatted property address from address components.

    Format: [Property Number], [Village] Village, [GN Division] Grama Niladari Division,
            [DS Division] Divisional Secretariat Division, [District] District, [Province]

    Only includes components that are present.
    """
    components = []

    # Property number (e.g., "No: 1202")
    if report.property_number:
        components.append(report.property_number)

    # Village name
    if report.property_village:
        components.append(f"{report.property_village} Village")

    # Grama Niladari Division
    if report.grama_niladari_division:
        components.append(f"{report.grama_niladari_division} Grama Niladari Division")

    # Divisional Secretariat (if available)
    if hasattr(report, 'property_divisional_secretariat') and report.property_divisional_secretariat:
        components.append(f"{report.property_divisional_secretariat} Divisional Secretariat Division")

    # Hathpaththuwa (if available)
    if report.hathpaththuwa:
        components.append(report.hathpaththuwa)

    # Korale (if available)
    if report.korale:
        components.append(report.korale)

    # Pradeshiya Sabha (if available)
    if report.pradeshiya_sabha:
        components.append(report.pradeshiya_sabha)

    # District (usually required)
    if report.property_district:
        components.append(f"{report.property_district} District")

    # Province
    if report.property_province:
        components.append(report.property_province)

    if not components:
        return None

    # Join with commas
    return ", ".join(components)

def generate_access_text(report: models.Report) -> Optional[str]:
    """Generate ACCESS section text from access directions data"""
    print(f"[ACCESS] Report ID: {report.id}")
    print(f"[ACCESS] Directions Text: {report.access_directions_text}")
    print(f"[ACCESS] Starting Point: {report.access_starting_point_name}")
    print(f"[ACCESS] Distance: {report.access_distance_km}")
    print(f"[ACCESS] Map Image URL: {report.location_map_image_data}")

    if not report.access_directions_text:
        print("[ACCESS] No access directions text found, returning None")
        return None

    print(f"[ACCESS] Returning access text: {report.access_directions_text}")
    return report.access_directions_text


def generate_locality_description(report: models.Report) -> Optional[str]:
    """Generate LOCALITY section text from locality data"""
    print(f"[LOCALITY] Report ID: {report.id}")
    print(f"[LOCALITY] Description Text: {report.locality_description_text}")
    print(f"[LOCALITY] Major Town: {report.major_town_name} ({report.distance_to_major_town_km} km)")
    print(f"[LOCALITY] Nearby Facilities: {len(report.nearby_facilities) if report.nearby_facilities else 0}")

    if not report.locality_description_text:
        print("[LOCALITY] No locality description text found, returning None")
        return None

    print(f"[LOCALITY] Returning locality text: {report.locality_description_text}")
    return report.locality_description_text

def generate_boundary_summary_text(report: models.Report) -> Optional[str]:
    """
    Generate professional boundary summary sentence from boundary types.
    Example output: "The boundaries referred above are in accordance with those shown in plan
    and are well defined and demarcated by brick masonry parapet walls constructed on rubble
    masonry foundation on the north, east, south and the west. There is an Auto roller gate
    entrance to the property."
    """
    # Map boundary type IDs to readable labels
    boundary_type_labels = {
        'brick_walls': 'brick masonry parapet walls',
        'barbed_wire': 'barbed wire fence',
        'live_fence': 'live fence',
        'concrete_posts': 'concrete posts',
        'iron_gate': 'iron gate',
        'rubble_foundation': 'rubble masonry foundation',
        'chain_link': 'chain link fence',
        'wooden_fence': 'wooden fence',
        'stone_wall': 'stone wall',
        'hedge': 'hedge'
    }

    # Entrance type labels
    entrance_labels = {
        'auto_roller_gate': 'an Auto roller gate',
        'iron_gate': 'an Iron gate',
        'wooden_gate': 'a Wooden gate',
        'sliding_gate': 'a Sliding gate',
        'swing_gate': 'a Swing gate',
        'no_gate': None
    }

    parts = []

    # Opening statement
    parts.append("The boundaries referred above are in accordance with those shown in plan and are well defined")

    # Determine boundary types per direction or use general physical boundaries
    boundary_types_per_dir = report.boundary_types_per_direction or {}
    physical_types = report.physical_boundaries_types or []

    if boundary_types_per_dir:
        # Group directions by boundary type
        type_to_directions = {}
        directions_order = ['north', 'northeast', 'east', 'southeast', 'south', 'southwest', 'west', 'northwest']

        for direction in directions_order:
            b_type = boundary_types_per_dir.get(direction)
            if b_type:
                if b_type not in type_to_directions:
                    type_to_directions[b_type] = []
                type_to_directions[b_type].append(direction)

        # Build description for each boundary type group
        boundary_descriptions = []
        for b_type, directions in type_to_directions.items():
            type_label = boundary_type_labels.get(b_type, b_type)
            dir_labels = [d for d in directions]

            if len(dir_labels) == 8:
                dir_text = "on all eight sides"
            elif len(dir_labels) == 4:
                # Check if it's the main 4 directions
                main_dirs = {'north', 'south', 'east', 'west'}
                if set(dir_labels) == main_dirs:
                    dir_text = "on all four sides"
                else:
                    dir_text = "on the " + ", ".join(dir_labels[:-1]) + f" and {dir_labels[-1]}"
            elif len(dir_labels) == 1:
                dir_text = f"on the {dir_labels[0]}"
            else:
                # Format as "north, northeast, east and south"
                dir_text = "on the " + ", ".join(dir_labels[:-1]) + f" and {dir_labels[-1]}"

            boundary_descriptions.append(f"demarcated by {type_label} {dir_text}")

        if boundary_descriptions:
            parts.append(" and " + ", ".join(boundary_descriptions))

    elif physical_types:
        # Use general physical boundary types
        type_labels = [boundary_type_labels.get(t, t) for t in physical_types[:3]]

        if 'rubble_foundation' in physical_types:
            # Special case: "brick walls constructed on rubble foundation"
            main_types = [t for t in type_labels if 'rubble' not in t]
            if main_types:
                parts.append(f" and demarcated by {main_types[0]} constructed on rubble masonry foundation on the north, east, south and the west")
            else:
                parts.append(f" and demarcated by rubble masonry foundation on all sides")
        elif type_labels:
            parts.append(f" and demarcated by {', '.join(type_labels)}")

    parts.append(".")

    # Add entrance/gate information
    entrance_type = report.entrance_type
    if entrance_type and entrance_type != 'no_gate':
        entrance_label = entrance_labels.get(entrance_type)
        if entrance_label:
            parts.append(f" There is {entrance_label} entrance to the property.")

    result = "".join(parts)
    return result if result != "The boundaries referred above are in accordance with those shown in plan and are well defined." else None


def format_list_with_grammar(items: List[str]) -> str:
    """
    Format a list of items with proper Oxford comma grammar.

    Examples:
        ['item1'] -> 'item1'
        ['item1', 'item2'] -> 'item1 and item2'
        ['item1', 'item2', 'item3'] -> 'item1, item2, and item3'
    """
    if not items:
        return ""
    if len(items) == 1:
        return items[0]
    elif len(items) == 2:
        return f"{items[0]} and {items[1]}"
    else:
        return ", ".join(items[:-1]) + f", and {items[-1]}"


def aggregate_accommodation_across_building(building: Dict) -> Dict:
    """
    Get accommodation summary from building (new structure) or aggregate from floors (old structure).

    NEW STRUCTURE: building.accommodation_summary (rooms at building level)
    OLD STRUCTURE: building.floors[].accommodation_summary (rooms per floor)

    Args:
        building: Building dict with either new or old structure

    Returns:
        Dict with aggregated room counts by type
    """
    # NEW STRUCTURE: Check if building has accommodation_summary at building level
    if building.get('accommodation_summary'):
        return building['accommodation_summary']

    # OLD STRUCTURE: Fallback to aggregating from floors (backward compatibility)
    floors = building.get('floors', [])
    totals = {
        'bedrooms': 0, 'bathrooms': 0, 'living_rooms': 0,
        'dining_rooms': 0, 'kitchens': 0, 'pantries': 0,
        'verandahs': 0, 'balconies': 0, 'garages': 0,
        'store_rooms': 0, 'other_rooms': 0
    }

    for floor in floors:
        summary = floor.get('accommodation_summary')
        if summary:
            for key in totals.keys():
                totals[key] += summary.get(key, 0)

    return totals


def deduplicate_water_sources(conveniences: List[str]) -> List[str]:
    """
    Merge duplicate water source entries in conveniences list.

    Business rules:
    - If both 'water' (well) and 'pipe_water' exist: prefer pipe-borne, remove well
    - Rationale: Pipe-borne water is superior/preferred utility

    Args:
        conveniences: List of convenience codes

    Returns:
        Cleaned list with duplicates removed
    """
    has_well = 'water' in conveniences
    has_pipe = 'pipe_water' in conveniences

    if has_well and has_pipe:
        # Prefer pipe-borne, remove well
        return [c for c in conveniences if c != 'water']

    return conveniences


def render_construction_details(doc, building: Dict) -> None:
    """
    Render construction materials section for a building.
    Uses professional sentence-based format for new data structure.
    Maintains backward compatibility with old format.
    """
    construction = building.get('construction_materials', {})
    if not construction:
        return

    sentences = []

    # NEW FORMAT: Check for new enhanced structure
    roof_details = construction.get('roof_details', {})
    wall_details = construction.get('wall_details', {})
    floor_details = construction.get('floor_details', {})

    # Roof Construction (NEW FORMAT)
    if roof_details and (roof_details.get('structure_type') or roof_details.get('covering_material')):
        roof_parts = []

        structure_type = roof_details.get('structure_type')
        if structure_type:
            structure_labels = {
                'timber_frame': 'timber frame',
                'steel_trusses': 'steel trusses',
                'concrete_slab': 'concrete slab',
                'rcc_flat': 'RCC flat',
                'prefab_trusses': 'prefabricated trusses',
                'mixed': 'mixed construction'
            }
            roof_parts.append(structure_labels.get(structure_type, structure_type.replace('_', ' ')))

        covering_materials = roof_details.get('covering_material', [])
        if covering_materials:
            covering_labels = {
                'asbestos_sheets': 'asbestos sheets',
                'clay_tiles': 'clay tiles',
                'concrete_tiles': 'concrete tiles',
                'metal_sheets': 'metal sheets',
                'concrete_flat': 'concrete flat',
                'cadjans': 'cadjans',
                'shingles': 'shingles'
            }
            covering_list = [covering_labels.get(m, m.replace('_', ' ')) for m in covering_materials]
            if covering_list:
                roof_parts.append(f"covered with {format_list_with_grammar(covering_list)}")

        additional = roof_details.get('additional_details')
        if additional:
            roof_parts.append(additional.lower() if additional else '')

        if roof_parts:
            sentences.append(f"The roof is constructed with {' '.join(roof_parts)}.")

    # Wall Construction (NEW FORMAT)
    if wall_details and wall_details.get('material'):
        wall_parts = []

        material = wall_details.get('material')
        if material:
            material_labels = {
                'brick_masonry': 'brick masonry',
                'cement_block': 'cement block',
                'stone_masonry': 'stone masonry',
                'rcc_frame': 'RCC frame',
                'rubble_masonry': 'rubble masonry',
                'mud_walls': 'mud walls',
                'timber_frame': 'timber frame',
                'cadjan': 'cadjan',
                'prefab_panels': 'prefabricated panels',
                'mixed': 'mixed materials'
            }
            wall_parts.append(material_labels.get(material, material.replace('_', ' ')))

        finishes = wall_details.get('finish', [])
        if finishes:
            finish_labels = {
                'cement_plaster_painted': 'cement plaster and painted',
                'lime_plaster': 'lime plaster',
                'tiles': 'tiles',
                'exposed_brick': 'exposed brick',
                'color_washed': 'color washed',
                'textured': 'textured finish',
                'unfinished': 'unfinished'
            }
            finish_list = [finish_labels.get(f, f.replace('_', ' ')) for f in finishes]
            if finish_list:
                wall_parts.append(f"finished with {format_list_with_grammar(finish_list)}")

        if wall_parts:
            sentences.append(f"The walls are {' '.join(wall_parts)}.")

    # Floor Construction (NEW FORMAT)
    if floor_details and floor_details.get('material'):
        floor_materials = floor_details.get('material', [])
        if floor_materials:
            material_labels = {
                'cement': 'cement',
                'tiled': 'tiles',
                'terrazzo': 'terrazzo',
                'timber': 'timber',
                'granite': 'granite',
                'marble': 'marble',
                'vinyl': 'vinyl',
                'polished_concrete': 'polished concrete',
                'earth': 'earth'
            }
            material_list = [material_labels.get(m, m.replace('_', ' ')) for m in floor_materials]

            floor_sentence = f"The floor is {format_list_with_grammar(material_list)} rendered"

            quality = floor_details.get('finish_quality')
            if quality:
                quality_labels = {'basic': 'basic', 'standard': 'good', 'premium': 'premium'}
                floor_sentence += f" with {quality_labels.get(quality, quality)} quality finish"

            sentences.append(f"{floor_sentence}.")

    # Ceiling (KEEP EXISTING - works for both old and new format)
    ceiling = construction.get('ceiling_type', '')
    if ceiling and ceiling != 'none':
        ceiling_labels = {
            'gypsum_board': 'gypsum board',
            'plywood': 'plywood',
            'asbestos': 'asbestos',
            'pvc': 'PVC panels',
            'timber': 'timber planks'
        }
        sentences.append(f"The ceiling is finished with {ceiling_labels.get(ceiling, ceiling.replace('_', ' '))}.")

    # Doors & Windows (NEW FORMAT)
    doors_windows = construction.get('doors_windows_details', {})
    if doors_windows:
        dw_parts = []

        # Windows
        window_frame = doors_windows.get('window_frame_material', [])
        window_glass = doors_windows.get('window_glass_type', [])
        window_security = doors_windows.get('window_security', [])

        if window_frame or window_glass or window_security:
            window_parts = []

            if window_frame:
                frame_labels = {
                    'timber': 'timber',
                    'aluminum': 'aluminum',
                    'upvc': 'UPVC',
                    'steel': 'steel'
                }
                frame_list = [frame_labels.get(f, f) for f in window_frame]
                window_parts.append(f"{format_list_with_grammar(frame_list)} frame windows")

            if window_glass:
                glass_labels = {
                    'clear': 'clear glass',
                    'tinted': 'tinted glass',
                    'frosted': 'frosted glass',
                    'double_glazing': 'double glazing'
                }
                glass_list = [glass_labels.get(g, g) for g in window_glass]
                window_parts.append(f"with {format_list_with_grammar(glass_list)}")

            if window_security:
                security_labels = {
                    'burglar_bars': 'burglar bars',
                    'grills': 'security grills',
                    'security_mesh': 'security mesh'
                }
                security_list = [security_labels.get(s, s) for s in window_security]
                window_parts.append(f"fitted with {format_list_with_grammar(security_list)}")

            if window_parts:
                dw_parts.append(f"The building has {' '.join(window_parts)}.")

        # Doors
        main_door = doors_windows.get('main_door_material')
        internal_door = doors_windows.get('internal_door_material')
        door_security = doors_windows.get('door_security', [])

        if main_door or internal_door:
            door_parts = []

            if main_door:
                main_door_labels = {
                    'solid_timber': 'solid timber',
                    'panel_door': 'timber panel',
                    'metal': 'metal',
                    'upvc': 'UPVC',
                    'glass': 'glass'
                }
                door_parts.append(f"main door is {main_door_labels.get(main_door, main_door.replace('_', ' '))}")

            if internal_door:
                internal_door_labels = {
                    'timber': 'solid timber',
                    'flush_doors': 'flush doors',
                    'panel_doors': 'panel doors',
                    'hollow_core': 'hollow core'
                }
                door_parts.append(f"internal doors are {internal_door_labels.get(internal_door, internal_door.replace('_', ' '))}")

            door_sentence = f"The {' and '.join(door_parts)}"

            if door_security:
                security_labels = {
                    'deadbolt': 'deadbolt locks',
                    'security_locks': 'security locks',
                    'door_chain': 'door chains',
                    'multi_point': 'multi-point locking systems'
                }
                security_list = [security_labels.get(s, s) for s in door_security]
                door_sentence += f" with {format_list_with_grammar(security_list)}"

            dw_parts.append(f"{door_sentence}.")

        sentences.extend(dw_parts)

    # BACKWARD COMPATIBILITY: Fall back to old format if new format not available
    if not sentences:
        # Old format handling (existing logic)
        parts = []

        # Foundation (old)
        foundation = construction.get('foundation_type', '')
        if foundation:
            foundation_labels = {
                'rubble_masonry': 'Rubble masonry',
                'concrete_strip': 'Concrete strip',
                'pile': 'Pile',
                'raft': 'Raft',
                'pad': 'Pad foundation'
            }
            parts.append(f"Foundation: {foundation_labels.get(foundation, foundation.replace('_', ' ').title())}")

        # Walls (old)
        wall = construction.get('wall_construction', {})
        if wall:
            wall_parts = []
            material = wall.get('material', '')
            if material:
                wall_material_labels = {
                    'brick_masonry': 'brick masonry',
                    'cement_block': 'cement block',
                    'stone': 'stone masonry',
                    'rcc_columns': 'RCC columns',
                    'cadjan': 'cadjan',
                    'timber': 'timber frame'
                }
                wall_parts.append(wall_material_labels.get(material, material.replace('_', ' ')))

            thickness = wall.get('thickness', '')
            if thickness:
                thickness_labels = {
                    '4_5_inch': '4.5 inch',
                    '9_inch': '9 inch',
                    '13_5_inch': '13.5 inch'
                }
                wall_parts.append(thickness_labels.get(thickness, thickness.replace('_', ' ')))

            finish = wall.get('finish', '')
            if finish:
                finish_labels = {
                    'plastered_painted': 'plastered and painted',
                    'fair_faced': 'fair-faced',
                    'textured': 'textured finish'
                }
                wall_parts.append(finish_labels.get(finish, finish.replace('_', ' ')))

            if wall_parts:
                parts.append(f"Walls: {', '.join(wall_parts)}")

        # Roof (old)
        roof = construction.get('roof_structure', '')
        if roof:
            roof_labels = {
                'timber_frame': 'Timber frame',
                'steel_trusses': 'Steel trusses',
                'rcc_slab': 'RCC slab',
                'cadjan': 'Cadjan'
            }
            parts.append(f"Roof: {roof_labels.get(roof, roof.replace('_', ' ').title())}")

        # Ceiling (old - already handled above)

        if parts:
            construction_text = "; ".join(parts) + "."
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
            para.paragraph_format.line_spacing = 0.9
            run = para.add_run(construction_text)
            run.font.size = FONT_SIZE_BODY
            run.font.color.rgb = RGBColor(0, 0, 0)
            return

    # Render new format as standalone paragraph (no label)
    if sentences:
        construction_text = " ".join(sentences)
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
        para.paragraph_format.line_spacing = 0.9
        run = para.add_run(construction_text)
        run.font.size = FONT_SIZE_BODY
        run.font.color.rgb = RGBColor(0, 0, 0)


def render_utilities_and_conveniences(doc, building: Dict) -> None:
    """
    Render unified utilities and conveniences section for a building.
    Includes water, electricity, sewage, parking, communication, gas, security, and amenities.
    """
    utilities = building.get('utilities_services', {})
    if not utilities:
        return

    parts = []

    # Water supply
    water = utilities.get('water_supply', '')
    if water:
        water_labels = {
            'Pipe-borne (NWSDB)': 'pipe-borne water (NWSDB)',
            'Well': 'well water',
            'Bore/Tube Well': 'bore/tube well water',
            'Rainwater Harvesting': 'rainwater harvesting',
            # Backward compatibility
            'pipe_borne': 'pipe-borne water (NWSDB)',
            'well': 'well water',
            'tube_well': 'tube well'
        }

        # Handle both string (legacy) and array (new) formats
        if isinstance(water, str):
            # Legacy single value
            parts.append(water_labels.get(water, water.replace('_', ' ')))
        elif isinstance(water, list):
            # New array format - use existing format_list_with_grammar
            mapped = [water_labels.get(w, w.lower()) for w in water]
            parts.append(format_list_with_grammar(mapped))

    # Sewage system
    sewage = utilities.get('sewage_system', '')
    if sewage:
        sewage_labels = {
            'municipal_sewer': 'municipal sewage system',
            'septic_tank': 'septic tank',
            'none': 'no sewage system'
        }
        if sewage != 'none':
            parts.append(sewage_labels.get(sewage, sewage.replace('_', ' ')))

    # Electricity
    elec = utilities.get('electricity', {})
    if elec:
        source = elec.get('source', '')
        if source:
            source_text = 'CEB electricity' if source == 'ceb' else f'{source.title()} electricity'
            if elec.get('three_phase'):
                source_text += ' (three-phase connection)'
            parts.append(source_text)

    # Parking
    parking = utilities.get('parking') or {}
    covered = int(to_float(parking.get('covered_spaces')))
    uncovered = int(to_float(parking.get('uncovered_spaces')))
    total_parking = covered + uncovered
    if total_parking > 0:
        parking_text = f"parking for {total_parking} vehicle{'s' if total_parking != 1 else ''}"
        if covered > 0 and uncovered > 0:
            parking_text += f" ({covered} covered, {uncovered} open)"
        elif covered > 0:
            parking_text += " (covered)"
        parts.append(parking_text)

    # Communication services
    comm_parts = []
    if utilities.get('telephone'):
        comm_parts.append('telephone connection')
    if utilities.get('internet'):
        comm_parts.append('internet connection')
    if comm_parts:
        parts.append(", ".join(comm_parts))

    # Gas connection
    if utilities.get('gas_connection'):
        parts.append('gas connection')

    # Security features
    security = utilities.get('security_features', [])
    if security:
        security_labels = {
            'boundary_wall': 'boundary wall',
            'main_gate': 'main gate',
            'cctv': 'CCTV surveillance',
            'security_lights': 'security lighting'
        }
        security_list = [security_labels.get(s, s.replace('_', ' ')) for s in security]
        if security_list:
            parts.append(", ".join(security_list))

    # Amenities
    amenities = utilities.get('amenities') or {}
    amenity_list = []
    if amenities.get('air_conditioning'):
        amenity_list.append('air conditioning')
    if amenities.get('built_in_wardrobes'):
        amenity_list.append('built-in wardrobes')
    if amenities.get('modern_kitchen'):
        amenity_list.append('modern kitchen fittings')
    if amenity_list:
        parts.append(", ".join(amenity_list))

    # Hot water system
    hot_water = utilities.get('hot_water_system', '')
    if hot_water and hot_water != 'none':
        hot_water_labels = {
            'electric_geyser': 'electric geyser',
            'solar_heater': 'solar water heater',
            'gas_heater': 'gas water heater'
        }
        parts.append(hot_water_labels.get(hot_water, hot_water.replace('_', ' ')))

    if parts:
        # Capitalize first letter
        utilities_text = "; ".join(parts)
        utilities_text = utilities_text[0].upper() + utilities_text[1:] + "."
        add_inline_field(doc, "Utilities & Conveniences", utilities_text)


def generate_multi_property_report_docx(report: models.Report, user: models.User) -> BytesIO:
    """
    Generate a multi-property valuation report DOCX file.

    Structure:
    1. Summary Page - Property listing with totals (COMPLETED properties only)
    2. Individual Property Sections - Full details for each COMPLETED property
       - Each property includes its own certification section
    3. Invoice Section - Professional fees

    Features:
    - Filters properties by status='completed' (draft properties excluded)
    - Per-property certification (allows different valuers per property)
    - Properties appear in user-defined order (from property_order field)

    Args:
        report: Report model instance with is_multi_property=True
        user: User model instance

    Returns:
        BytesIO object containing the DOCX document

    Raises:
        ValueError: If report has no properties or all properties are drafts
    """
    try:
        logger.info(f"[MULTI-PROPERTY DOCX] Starting generation for report ID {report.id}")
        logger.info(f"[MULTI-PROPERTY DOCX] Property count: {report.property_count}")

        # Get all properties ordered by property_order
        all_properties = report.properties  # Uses the @property helper method
        if not all_properties:
            raise ValueError("Multi-property report has no properties")

        # Filter to only completed properties (exclude drafts)
        properties = [prop for prop in all_properties if prop.status == 'completed']

        logger.info(f"[MULTI-PROPERTY DOCX] Total properties: {len(all_properties)}, Completed: {len(properties)}, Draft: {len(all_properties) - len(properties)}")

        # Validate that at least one completed property exists
        if not properties:
            raise ValueError("Cannot generate report: All properties are in draft status. Please complete at least one property before generating the report.")

        # Create document
        doc = Document()

        # ===== LETTERHEAD =====
        template_id = user.preferred_letterhead_template or 'classic'
        template = get_template(template_id)
        template.render_letterhead(doc, user, report)

        # ===== SUMMARY PAGE =====
        logger.info("[MULTI-PROPERTY DOCX] Generating summary page")

        # Title: "VALUATION REPORT" - Centered and bold (matching residential/bare land format)
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.paragraph_format.space_before = Pt(0)
        title_para.paragraph_format.space_after = Pt(12)
        title_run = title_para.add_run("VALUATION REPORT")
        title_run.bold = True
        title_run.font.size = Pt(14)  # Updated from 11 to 14
        title_run.font.underline = True  # Add underline
        title_run.font.color.rgb = RGBColor(0, 0, 0)

        # Add spacing before applicant statements (matching residential/bare land)
        spacing_para1 = doc.add_paragraph()
        spacing_para1.paragraph_format.space_before = Pt(8)
        spacing_para1.paragraph_format.space_after = Pt(0)

        # Determine which introduction format to use based on request_type
        if report.request_type == 'organization_request':
            # Organization-side format
            org_intro_paragraphs = generate_organization_side_introduction(report)
            for statement in org_intro_paragraphs:
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
                para.paragraph_format.line_spacing = 0.9

                # Parse label and value for proper formatting
                if ":-" in statement:
                    # This is a label-value pair (Applicant/Address/Contact No)
                    label, value = statement.split(":-", 1)

                    # Add bold label
                    run_label = para.add_run(label.strip())
                    run_label.font.size = FONT_SIZE_BODY
                    run_label.font.bold = True
                    run_label.font.color.rgb = RGBColor(0, 0, 0)

                    # Add separator with proper spacing
                    run_sep = para.add_run(" : ")
                    run_sep.font.size = FONT_SIZE_BODY
                    run_sep.font.color.rgb = RGBColor(0, 0, 0)

                    # Add value
                    run_value = para.add_run(value)
                    run_value.font.size = FONT_SIZE_BODY
                    run_value.font.color.rgb = RGBColor(0, 0, 0)
                else:
                    # This is the introductory paragraph (paragraph 1)
                    run = para.add_run(statement)
                    run.font.size = FONT_SIZE_BODY
                    run.font.color.rgb = RGBColor(0, 0, 0)
        else:
            # Client-side format (default for backward compatibility)
            applicant_statements = generate_applicant_statement(report)
            for statement in applicant_statements:
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
                para.paragraph_format.line_spacing = 0.9
                run = para.add_run(statement)
                run.font.size = FONT_SIZE_BODY
                run.font.color.rgb = RGBColor(0, 0, 0)

            # Add deed/certificate description if applicable (matching residential/bare land logic)
            should_show_deed_sentence = (
                report.property_identification_type in ['deed', 'certificate_of_sale']
                and report.has_deed_info == "yes"
                and report.deeds
            )

            # Backward compatibility: Old reports (NULL type) with deed data should show sentence
            if not report.property_identification_type and report.has_deed_info == "yes" and report.deeds:
                should_show_deed_sentence = True

            if should_show_deed_sentence:
                deed_text = generate_deed_description(report.deeds)
                if deed_text:
                    para = doc.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
                    para.paragraph_format.line_spacing = 0.9
                    run = para.add_run(deed_text)
                    run.font.size = FONT_SIZE_BODY
                    run.font.color.rgb = RGBColor(0, 0, 0)

            # Add submission statement (matching residential/bare land format)
            submission_text = generate_submission_statement(report)
            if submission_text:
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
                para.paragraph_format.line_spacing = 0.9
                run = para.add_run(submission_text)
                run.font.size = FONT_SIZE_BODY
                run.font.color.rgb = RGBColor(0, 0, 0)

        # Add multi-property specific statement about the summary table
        summary_intro_para = doc.add_paragraph()
        summary_intro_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        summary_intro_para.paragraph_format.space_before = Pt(6)
        summary_intro_para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
        summary_intro_para.paragraph_format.line_spacing = 0.9
        summary_intro_run = summary_intro_para.add_run(
            f"This report comprises the valuation of {len(properties)} properties. "
            "The summary of properties and their respective valuations are presented in the following table:"
        )
        summary_intro_run.font.size = FONT_SIZE_BODY
        summary_intro_run.font.color.rgb = RGBColor(0, 0, 0)

        # Property list table
        logger.info("[MULTI-PROPERTY DOCX] Creating property list table")

        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Set custom column widths: #(8%), Location(25%), Description(35%), Extent(12%), Value(20%)
        column_widths = [Inches(0.52), Inches(1.625), Inches(2.275), Inches(0.78), Inches(1.3)]
        for idx_col, width in enumerate(column_widths):
            table.columns[idx_col].width = width

        # Header row
        header_cells = table.rows[0].cells
        headers = ['#', 'Location', 'Description', 'Extent', 'Value (LKR)']
        for i, header_text in enumerate(headers):
            cell = header_cells[i]
            cell.text = header_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = FONT_SIZE_TABLE_HEADER
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Property rows
        grand_total = 0
        for idx, prop in enumerate(properties, start=1):
            row_cells = table.add_row().cells

            # Property number
            row_cells[0].text = str(idx)
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Location
            location = prop.property_village or "N/A"
            if prop.property_district and prop.property_district != location:
                location += f", {prop.property_district}"
            row_cells[1].text = location

            # Description (Lot/Plan)
            desc = ""
            if prop.lot_number:
                lot = prop.lot_number
                desc = f"Lot {lot}"
            if prop.plan_number:
                if desc:
                    desc += f", Plan {prop.plan_number}"
                else:
                    desc = f"Plan {prop.plan_number}"
            if not desc:
                desc = "N/A"
            row_cells[2].text = desc

            # Extent
            extent = prop.land_extent_formatted or "N/A"
            row_cells[3].text = extent

            # Value
            value = prop.valuation_market_value or 0
            grand_total += float(value)
            row_cells[4].text = format_currency(value)
            row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Format all cells
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = FONT_SIZE_TABLE_CELL

        # Total row
        total_row_cells = table.add_row().cells
        total_row_cells[0].merge(total_row_cells[3])
        total_row_cells[0].text = "GRAND TOTAL:"
        total_row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for paragraph in total_row_cells[0].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = FONT_SIZE_BODY

        total_row_cells[4].text = format_currency(grand_total)
        total_row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for paragraph in total_row_cells[4].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = FONT_SIZE_BODY

        # Add spacing after table
        doc.add_paragraph()

        # Inspection date (matching residential/bare land format)
        if report.inspection_date:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
            para.paragraph_format.line_spacing = 0.9

            # Bold label
            run_label = para.add_run("Date of Inspection")
            run_label.font.size = FONT_SIZE_BODY
            run_label.font.bold = True
            run_label.font.color.rgb = RGBColor(0, 0, 0)

            # Separator with proper spacing
            run_sep = para.add_run(" : ")
            run_sep.font.size = FONT_SIZE_BODY
            run_sep.font.color.rgb = RGBColor(0, 0, 0)

            # Regular date value
            run_date = para.add_run(report.inspection_date)
            run_date.font.size = FONT_SIZE_BODY
            run_date.font.color.rgb = RGBColor(0, 0, 0)

        # Add concluding statement for CLIENT REQUESTS only (after Date of Inspection)
        concluding_parts = generate_multi_property_concluding_statement(report, user, grand_total)
        if concluding_parts:
            # Statement paragraph (reduced spacing after Date of Inspection)
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
            para.paragraph_format.line_spacing = 0.9
            run = para.add_run(concluding_parts[0])
            run.font.size = FONT_SIZE_BODY
            run.font.color.rgb = RGBColor(0, 0, 0)

            # Valuer name (bold) - increased space for signature
            name_para = doc.add_paragraph()
            name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            name_para.paragraph_format.space_before = Pt(30)
            name_para.paragraph_format.space_after = Pt(2)
            name_run = name_para.add_run(concluding_parts[1])
            name_run.font.size = FONT_SIZE_BODY
            name_run.font.bold = True
            name_run.font.color.rgb = RGBColor(0, 0, 0)

            # Professional designation
            desig_para = doc.add_paragraph()
            desig_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            desig_para.paragraph_format.space_before = Pt(0)
            desig_para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
            desig_run = desig_para.add_run(concluding_parts[2])
            desig_run.font.size = FONT_SIZE_BODY
            desig_run.font.color.rgb = RGBColor(0, 0, 0)

        # Add special note if applicable (matching residential/bare land format)
        if report.has_special_note == "yes" and report.special_note_text:
            # Note label
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = SUBHEADING_SPACE_AFTER
            para.paragraph_format.line_spacing = 0.9
            run = para.add_run("Note:")
            run.bold = True
            run.font.size = FONT_SIZE_BODY
            run.font.color.rgb = RGBColor(0, 0, 0)

            # Note text
            note_para = doc.add_paragraph()
            note_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            note_para.paragraph_format.space_before = Pt(0)
            note_para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
            note_para.paragraph_format.line_spacing = 0.9
            note_run = note_para.add_run(report.special_note_text)
            note_run.font.size = FONT_SIZE_BODY
            note_run.font.color.rgb = RGBColor(0, 0, 0)

        # REMOVED: Forced page break - let natural pagination handle page breaks
        # This allows better space utilization, especially for organization requests
        # where the table ends earlier and there's often significant space remaining on page 1
        # doc.add_page_break()

        # ===== INDIVIDUAL PROPERTY SECTIONS =====
        logger.info("[MULTI-PROPERTY DOCX] Generating individual property sections")

        for prop_idx, prop in enumerate(properties, start=1):
            logger.info(f"[MULTI-PROPERTY DOCX] Generating section for property {prop_idx}/{len(properties)}")

            # Professional property description (NO "PROPERTY 1" header - removed per user request)
            # Format: "The Property Depicted as Lot [X] in Plan No: [Y]"
            if prop.lot_number and prop.plan_number:
                # Extract just the lot number/identifier (remove "Plan No" prefix if present)
                lot_desc = prop.lot_number.strip() if prop.lot_number else ''

                # Remove common prefixes that shouldn't be in lot description
                prefixes_to_remove = ['plan no', 'plan no:', 'lot plan no', 'lot plan no:']
                lot_desc_lower = lot_desc.lower()
                for prefix in prefixes_to_remove:
                    if lot_desc_lower.startswith(prefix):
                        lot_desc = lot_desc[len(prefix):].strip()
                        break

                # Ensure lot description has "Lot" prefix (if it doesn't already)
                if not lot_desc.lower().startswith('lot'):
                    lot_desc = f"Lot {lot_desc}"

                plan_formatted = format_no_field("Plan", prop.plan_number)

                # Combine date with first line
                plan_date = prop.plan_date or ''
                if plan_date:
                    prop_desc = f"The Property Depicted as {lot_desc} in {plan_formatted} Dated {plan_date}"
                else:
                    prop_desc = f"The Property Depicted as {lot_desc} in {plan_formatted}"

                prop_desc_para = doc.add_paragraph()
                prop_desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                prop_desc_para.paragraph_format.space_before = Pt(0)
                prop_desc_para.paragraph_format.space_after = Pt(0)  # No spacing
                prop_desc_para.paragraph_format.line_spacing = 1.0  # Single line spacing
                prop_desc_run = prop_desc_para.add_run(prop_desc)
                prop_desc_run.bold = True
                prop_desc_run.font.size = Pt(12)  # Updated from 10 to 12
                prop_desc_run.font.underline = True  # Add underline
                prop_desc_run.font.color.rgb = RGBColor(0, 0, 0)

                # Second line starts with "made by" - no spacing between lines
                if prop.licensed_surveyor_name:
                    plan_info = f"made by {prop.licensed_surveyor_name} Licensed Surveyor."
                    plan_info_para = doc.add_paragraph()
                    plan_info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    plan_info_para.paragraph_format.space_before = Pt(0)  # No spacing
                    plan_info_para.paragraph_format.space_after = Pt(12)
                    plan_info_para.paragraph_format.line_spacing = 1.0  # Single line spacing
                    plan_info_run = plan_info_para.add_run(plan_info)
                    plan_info_run.bold = True
                    plan_info_run.font.size = Pt(12)  # Updated from 10 to 12
                    plan_info_run.font.underline = True  # Add underline
                    plan_info_run.font.color.rgb = RGBColor(0, 0, 0)
                else:
                    # Just add spacing if no plan info
                    spacing_para = doc.add_paragraph()
                    spacing_para.paragraph_format.space_after = Pt(12)
            else:
                # Fallback to location-based description if no plan info
                location = prop.property_village or "N/A"
                if prop.property_district:
                    location += f", {prop.property_district}"
                prop_loc_para = doc.add_paragraph()
                prop_loc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                prop_loc_para.paragraph_format.space_after = Pt(12)
                prop_loc_run = prop_loc_para.add_run(location)
                prop_loc_run.font.size = FONT_SIZE_BODY
                prop_loc_run.font.color.rgb = RGBColor(0, 0, 0)

            # Generate property sections (comprehensive full sections)
            _generate_property_sections(doc, prop, report, user)

            # NO page breaks between properties - continuous flow
            # User will manually adjust page breaks as needed

        # ===== INVOICE SECTION =====
        if report.invoice_data:
            logger.info("[MULTI-PROPERTY DOCX] Generating invoice section")
            # Page break handled inside _generate_invoice_section()
            _generate_invoice_section(doc, report.invoice_data, user, report)

        # Note: Per-property certification now included in individual property sections
        # This allows different valuers to certify different properties in the same report

        # ===== SAVE DOCUMENT =====
        logger.info("[MULTI-PROPERTY DOCX] Saving document")
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        if file_stream.getvalue() == b'':
            raise ValueError("Generated multi-property document is empty")

        logger.info(f"[MULTI-PROPERTY DOCX] Document generated successfully, size: {len(file_stream.getvalue())} bytes")
        return file_stream

    except Exception as e:
        logger.error(f"[MULTI-PROPERTY DOCX] Error: {e}", exc_info=True)
        raise


def _generate_property_sections(doc, prop, report, user):
    """
    Generate comprehensive sections for a single property within a multi-property report.

    Generates full professional valuation report structure:
    1. SITUATION
    2. ACCESS (if available)
    3. IDENTIFICATION & BOUNDARIES
    4. DESCRIPTION (land + buildings + photos)
    5. LOCALITY (if available)
    6. LEGAL ASPECTS (if available)
    7. LAND VALUES / COMPARABLES (if available)
    8. VALUATION
    9. CERTIFICATION (per-property)

    Properties flow continuously without page breaks between them.
    """
    section_num = 1

    # ===== 1.0 SITUATION SECTION =====
    situation_text = generate_situation_text(prop)
    if situation_text:
        add_section_heading(doc, f"{section_num}.0", "SITUATION")
        section_num += 1

        situation_para = doc.add_paragraph()
        situation_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        situation_para.paragraph_format.space_before = Pt(0)
        situation_para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
        situation_para.paragraph_format.line_spacing = 0.9
        situation_run = situation_para.add_run(situation_text)
        situation_run.font.size = FONT_SIZE_BODY
        situation_run.font.color.rgb = RGBColor(0, 0, 0)

    # ===== 2.0 ACCESS SECTION (CONDITIONAL) =====
    access_text = generate_access_text(prop)
    if access_text:
        add_section_heading(doc, f"{section_num}.0", "ACCESS")
        section_num += 1

        access_para = doc.add_paragraph()
        access_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        access_para.paragraph_format.space_before = Pt(0)
        access_para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
        access_para.paragraph_format.line_spacing = 0.9
        access_run = access_para.add_run(access_text)
        access_run.font.size = FONT_SIZE_BODY
        access_run.font.color.rgb = RGBColor(0, 0, 0)

        # Add coordinates if available
        if prop.property_latitude and prop.property_longitude:
            coord_para = doc.add_paragraph()
            coord_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            coord_para.paragraph_format.space_before = Pt(6)
            coord_para.paragraph_format.space_after = Pt(6)
            coord_para.paragraph_format.line_spacing = 0.9

            coord_label = coord_para.add_run("Property Location Coordinates: ")
            coord_label.bold = True
            coord_label.font.size = FONT_SIZE_BODY
            coord_label.font.color.rgb = RGBColor(0, 0, 0)

            lat_value = float(prop.property_latitude)
            lng_value = float(prop.property_longitude)
            coord_text = coord_para.add_run(f"{lat_value:.6f}, {lng_value:.6f}")
            coord_text.font.size = FONT_SIZE_BODY
            coord_text.font.color.rgb = RGBColor(0, 0, 0)

        # Add map image if available
        if prop.location_map_image_data:
            try:
                map_url = prop.location_map_image_data
                logger.info(f"[DOCX-MULTI] Fetching map image for property {prop.id} (URL length={len(map_url)})")
                response = requests.get(map_url, timeout=30)
                if response.status_code == 200:
                    map_spacing_para = doc.add_paragraph()
                    map_spacing_para.paragraph_format.space_before = IMAGE_SPACING_BEFORE
                    map_spacing_para.paragraph_format.space_after = Pt(0)

                    map_para = doc.add_paragraph()
                    map_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    map_para.paragraph_format.space_before = Pt(0)
                    map_para.paragraph_format.space_after = IMAGE_SPACING_AFTER

                    image_stream = BytesIO(response.content)
                    dimensions = calculate_image_dimensions(
                        image_stream,
                        MAP_IMAGE_WIDTH,
                        MAP_IMAGE_MAX_HEIGHT
                    )
                    map_para.add_run().add_picture(image_stream, **dimensions)
                    logger.info(f"[DOCX-MULTI] Successfully added map image for property {prop.id}")
                else:
                    logger.warning(f"[DOCX-MULTI] Failed to fetch map image for property {prop.id}: HTTP {response.status_code}")
            except Exception as e:
                logger.warning(f"[DOCX-MULTI] Error adding map image for property {prop.id}: {str(e)}")
        else:
            logger.info(f"[DOCX-MULTI] No location_map_image_data for property {prop.id}")

    # ===== 3.0 IDENTIFICATION OF PROPERTY SECTION =====
    has_property_header_data = (
        prop.land_traditional_name or
        prop.land_extent_formatted or
        prop.boundaries or
        prop.physical_boundaries_types or
        prop.physical_boundaries_description or
        prop.boundaries_summary_text
    )

    if has_property_header_data:
        add_section_heading(doc, f"{section_num}.0", "IDENTIFICATION OF PROPERTY")
        section_num += 1

        # Traditional Land Name
        if prop.land_traditional_name:
            name_para = doc.add_paragraph()
            name_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            name_para.paragraph_format.space_before = Pt(0)
            name_para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
            name_para.paragraph_format.line_spacing = 0.9
            name_label = name_para.add_run("Name of Land: ")
            name_label.bold = True
            name_label.font.size = FONT_SIZE_BODY
            name_label.font.color.rgb = RGBColor(0, 0, 0)
            name_value = name_para.add_run(f'"{prop.land_traditional_name}"')
            name_value.font.size = FONT_SIZE_BODY
            name_value.font.color.rgb = RGBColor(0, 0, 0)

        # Survey Plan Information
        if prop.lot_number and prop.plan_number:
            plan_para = doc.add_paragraph()
            plan_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            plan_para.paragraph_format.space_before = Pt(0)
            plan_para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
            plan_para.paragraph_format.line_spacing = 0.9
            plan_label = plan_para.add_run("Survey Plan: ")
            plan_label.bold = True
            plan_label.font.size = FONT_SIZE_BODY
            plan_label.font.color.rgb = RGBColor(0, 0, 0)

            # Extract just the lot number/identifier (remove "Plan No" prefix if present)
            lot_desc = prop.lot_number.strip() if prop.lot_number else ''

            # Remove common prefixes that shouldn't be in lot description
            prefixes_to_remove = ['plan no', 'plan no:', 'lot plan no', 'lot plan no:']
            lot_desc_lower = lot_desc.lower()
            for prefix in prefixes_to_remove:
                if lot_desc_lower.startswith(prefix):
                    lot_desc = lot_desc[len(prefix):].strip()
                    break

            # Ensure lot description has "Lot" prefix
            if not lot_desc.lower().startswith('lot'):
                lot_desc = f"Lot {lot_desc}"

            # Format as "Lot X in Plan No: Y dated [date] made by [surveyor], Licensed Surveyor"
            plan_formatted = format_no_field("Plan", prop.plan_number)
            plan_text = f"{lot_desc} in {plan_formatted}"
            if prop.plan_date:
                plan_text += f" dated {prop.plan_date}"
            if prop.licensed_surveyor_name:
                plan_text += f" made by {prop.licensed_surveyor_name}, Licensed Surveyor"
            plan_text += "."

            plan_value = plan_para.add_run(plan_text)
            plan_value.font.size = FONT_SIZE_BODY
            plan_value.font.color.rgb = RGBColor(0, 0, 0)

        # Land Extent
        if prop.land_extent_formatted:
            extent_para = doc.add_paragraph()
            extent_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            extent_para.paragraph_format.space_before = Pt(0)
            extent_para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
            extent_para.paragraph_format.line_spacing = 0.9
            extent_label = extent_para.add_run("Extent: ")
            extent_label.bold = True
            extent_label.font.size = FONT_SIZE_BODY
            extent_label.font.color.rgb = RGBColor(0, 0, 0)
            extent_value = extent_para.add_run(prop.land_extent_formatted)
            extent_value.font.size = FONT_SIZE_BODY
            extent_value.font.color.rgb = RGBColor(0, 0, 0)

        # Boundaries - Add detailed boundary descriptions
        boundaries = safe_get_json_field(prop, 'boundaries', None)
        if boundaries:
            boundary_para = doc.add_paragraph()
            boundary_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            boundary_para.paragraph_format.space_before = Pt(6)
            boundary_para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
            boundary_para.paragraph_format.line_spacing = 0.9

            boundary_label = boundary_para.add_run("Boundaries:\n")
            boundary_label.bold = True
            boundary_label.font.size = FONT_SIZE_BODY
            boundary_label.font.color.rgb = RGBColor(0, 0, 0)

            # Add each direction boundary (4 main + 4 optional diagonal)
            for direction in ['North', 'North-East', 'East', 'South-East', 'South', 'South-West', 'West', 'North-West']:
                direction_key = direction.lower().replace('-', '')  # 'North-East' -> 'northeast'
                if isinstance(boundaries, dict) and direction_key in boundaries:
                    boundary_data = boundaries[direction_key]
                    description = boundary_data.get('description', '') if isinstance(boundary_data, dict) else str(boundary_data)

                    # Skip diagonal boundaries if empty
                    if not description and direction_key in ['northeast', 'southeast', 'southwest', 'northwest']:
                        continue

                    if description:
                        boundary_text = f"{direction}: {description}\n"
                        run = boundary_para.add_run(boundary_text)
                        run.font.size = FONT_SIZE_BODY
                        run.font.color.rgb = RGBColor(0, 0, 0)

    # ===== 4.0 DESCRIPTION OF PROPERTY SECTION =====
    add_section_heading(doc, f"{section_num}.0", "DESCRIPTION OF PROPERTY")
    section_num += 1

    # Land Description
    land_desc_text = prop.land_description_text
    if land_desc_text:
        desc_para = doc.add_paragraph()
        desc_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        desc_para.paragraph_format.space_before = Pt(0)
        desc_para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
        desc_para.paragraph_format.line_spacing = 0.9
        desc_run = desc_para.add_run(land_desc_text)
        desc_run.font.size = FONT_SIZE_BODY
        desc_run.font.color.rgb = RGBColor(0, 0, 0)

    # Buildings (if applicable - skip for bare land) - USE SAME FORMAT AS STANDALONE REPORTS
    buildings = safe_get_json_field(prop, 'buildings', [])
    if buildings and prop.property_type != 'bare_land':
        for building_idx, building in enumerate(buildings, 1):
            # Use numbered subsection heading (4.1, 4.2, etc.) - SAME AS STANDALONE
            building_number = f"{section_num - 1}.{building_idx}"
            building_name = building.get('building_name', f'Building {building_idx}')

            # Add building subsection heading using the same helper
            add_section_heading(doc, building_number, building_name)

            # === CONSTRUCTION DETAILS - USE SAME PROFESSIONAL FORMAT AS STANDALONE ===
            render_construction_details(doc, building)

            # === PROFESSIONAL BUILDING DESCRIPTION - SAME AS STANDALONE ===
            # Opening paragraph: Construction materials and general description
            if building.get('building_description_text'):
                opening_para = doc.add_paragraph()
                opening_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                opening_para.paragraph_format.space_before = Pt(0)
                opening_para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
                opening_para.paragraph_format.line_spacing = 0.9
                opening_run = opening_para.add_run(building.get('building_description_text'))
                opening_run.font.size = FONT_SIZE_BODY
                opening_run.font.color.rgb = RGBColor(0, 0, 0)
            else:
                # Build opening description from materials (same logic as standalone)
                roof_types = building.get('roof_types', [])
                wall_types = building.get('wall_types', [])
                floor_types = building.get('floor_types', [])

                if roof_types or wall_types or floor_types:
                    desc_text = "This is constructed with "

                    # Roof
                    if roof_types:
                        roof_labels = {
                            'asbestos': 'asbestos sheets', 'tile': 'tiles',
                            'metal': 'metal sheets', 'concrete': 'concrete flat',
                            'cadjan': 'cadjans', 'zinc': 'zinc sheets'
                        }
                        roof_text = format_material_list(roof_types, roof_labels)
                        desc_text += f"{roof_text} on "

                    # Roof structure
                    stories = building.get('stories', 1)
                    if stories == 1:
                        desc_text += "timber frame roof "
                    else:
                        desc_text += f"{stories} storey structure "

                    # Walls
                    if wall_types:
                        wall_labels = {
                            'brick': 'brick masonry', 'cement_block': 'cement block',
                            'stone': 'stone', 'mud': 'mud', 'timber': 'timber',
                            'cadjan': 'cadjan', 'rcc_columns': 'RCC columns',
                            'rubble_masonry': 'rubble masonry'
                        }
                        wall_text = format_material_list(wall_types, wall_labels)
                        desc_text += f"supported by {wall_text} walls "

                    # Foundation/floor
                    if floor_types:
                        floor_labels = {
                            'cement': 'cement', 'tile': 'tiles', 'terrazzo': 'terrazzo',
                            'timber': 'timber', 'granite': 'granite', 'earth': 'earth',
                            'concrete': 'concrete'
                        }
                        floor_text = format_material_list(floor_types, floor_labels)
                        desc_text += f"and the floor is {floor_text} rendered"

                    desc_text += "."

                    opening_para = doc.add_paragraph()
                    opening_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    opening_para.paragraph_format.space_before = Pt(0)
                    opening_para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
                    opening_para.paragraph_format.line_spacing = 0.9
                    opening_run = opening_para.add_run(desc_text)
                    opening_run.font.size = FONT_SIZE_BODY
                    opening_run.font.color.rgb = RGBColor(0, 0, 0)

            # === ACCOMMODATION - USE SAME AGGREGATED FORMAT AS STANDALONE ===
            # Check if building has accommodation data (new structure at building level or old structure at floor level)
            has_new_building_format = building.get('accommodation_summary') is not None
            floors = building.get('floors', [])
            has_old_floor_format = any(floor.get('accommodation_summary') for floor in floors)

            if has_new_building_format or has_old_floor_format:
                # Get accommodation summary (supports both new and old structures)
                aggregated_rooms = aggregate_accommodation_across_building(building)

                room_parts = []

                if aggregated_rooms.get('bedrooms', 0) > 0:
                    count = aggregated_rooms['bedrooms']
                    room_parts.append(format_room_count(count, 'bedroom'))

                if aggregated_rooms.get('bathrooms', 0) > 0:
                    count = aggregated_rooms['bathrooms']
                    room_parts.append(format_room_count(count, 'bathroom'))

                if aggregated_rooms.get('living_rooms', 0) > 0:
                    count = aggregated_rooms['living_rooms']
                    room_parts.append(format_room_count(count, 'living room'))

                if aggregated_rooms.get('dining_rooms', 0) > 0:
                    count = aggregated_rooms['dining_rooms']
                    room_parts.append(format_room_count(count, 'dining room'))

                if aggregated_rooms.get('kitchens', 0) > 0:
                    count = aggregated_rooms['kitchens']
                    room_parts.append(format_room_count(count, 'kitchen'))

                if aggregated_rooms.get('pantries', 0) > 0:
                    count = aggregated_rooms['pantries']
                    room_parts.append(format_room_count(count, 'pantry', 'pantries'))

                if aggregated_rooms.get('verandahs', 0) > 0:
                    count = aggregated_rooms['verandahs']
                    room_parts.append(format_room_count(count, 'verandah'))

                if aggregated_rooms.get('balconies', 0) > 0:
                    count = aggregated_rooms['balconies']
                    room_parts.append(format_room_count(count, 'balcony', 'balconies'))

                if aggregated_rooms.get('garages', 0) > 0:
                    count = aggregated_rooms['garages']
                    room_parts.append(format_room_count(count, 'garage'))

                if aggregated_rooms.get('store_rooms', 0) > 0:
                    count = aggregated_rooms['store_rooms']
                    room_parts.append(format_room_count(count, 'store room'))

                if aggregated_rooms.get('other_rooms', 0) > 0:
                    count = aggregated_rooms['other_rooms']
                    room_parts.append(format_room_count(count, 'other room'))

                if room_parts:
                    accommodation_text = f"The property comprises {format_list_with_grammar(room_parts)}."
                    add_inline_field(doc, "Accommodation", accommodation_text)

            # === FLOOR AREA BREAKDOWN - CONDITIONAL FORMAT ===
            total_area = building.get('total_floor_area', 0)
            if floors and (total_area or any(f.get('floor_area', 0) > 0 for f in floors)):
                # Filter floors with area > 0
                floors_with_area = [f for f in floors if f.get('floor_area', 0) > 0]

                if len(floors_with_area) == 1:
                    # SINGLE FLOOR: Use inline format
                    floor = floors_with_area[0]
                    floor_area = floor.get('floor_area', 0)
                    area_text = f"{floor_area:,.0f} square feet"
                    add_inline_field(doc, "Floor area", area_text)
                else:
                    # MULTIPLE FLOORS: Use table format
                    floor_area_para = doc.add_paragraph()
                    floor_area_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    floor_area_para.paragraph_format.space_before = INLINE_FIELD_SPACE_BEFORE
                    floor_area_para.paragraph_format.space_after = Pt(2)
                    floor_area_para.paragraph_format.line_spacing = 0.9

                    label_run = floor_area_para.add_run("Floor area:")
                    label_run.bold = True
                    label_run.font.size = FONT_SIZE_BODY
                    label_run.font.color.rgb = RGBColor(0, 0, 0)

                    for floor in floors:
                        floor_name = floor.get('floor_name', 'Floor')
                        floor_area = floor.get('floor_area', 0)

                        if floor_area and floor_area > 0:
                            floor_line_para = doc.add_paragraph()
                            floor_line_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            floor_line_para.paragraph_format.space_before = Pt(0)
                            floor_line_para.paragraph_format.space_after = Pt(1)
                            floor_line_para.paragraph_format.line_spacing = 0.9
                            floor_line_para.paragraph_format.left_indent = Inches(0.5)

                            floor_name_padded = f"{floor_name:<25}"
                            floor_line_run = floor_line_para.add_run(f"{floor_name_padded}{floor_area:>10,.0f} square feet")
                            floor_line_run.font.size = FONT_SIZE_BODY
                            floor_line_run.font.color.rgb = RGBColor(0, 0, 0)

                    if total_area and total_area > 0:
                        total_para = doc.add_paragraph()
                        total_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        total_para.paragraph_format.space_before = Pt(0)
                        total_para.paragraph_format.space_after = INLINE_FIELD_SPACE_AFTER
                        total_para.paragraph_format.line_spacing = 0.9
                        total_para.paragraph_format.left_indent = Inches(0.5)

                        total_padded = f"{'Total':<25}"
                        total_run = total_para.add_run(f"{total_padded}{total_area:>10,.0f} square feet")
                        total_run.bold = True
                        total_run.font.size = FONT_SIZE_BODY
                        total_run.font.color.rgb = RGBColor(0, 0, 0)

            # === AGE AND CONDITION - USE SAME FORMAT AS STANDALONE ===
            building_age = building.get('building_age')
            condition = building.get('condition', '')
            if building_age or condition:
                condition_labels = {
                    'excellent': 'excellent', 'good': 'good', 'fair': 'fair',
                    'poor': 'poor', 'dilapidated': 'dilapidated'
                }

                parts = []
                if building_age is not None and building_age > 0:
                    year_word = "year" if building_age == 1 else "years"
                    parts.append(f"{building_age} {year_word} old")
                if condition:
                    parts.append(f"condition is {condition_labels.get(condition, condition)}")

                age_text = "; ".join(parts) + "." if parts else "Information not provided."
                add_inline_field(doc, "Age and condition", age_text)

            # === UTILITIES AND CONVENIENCES - UNIFIED SECTION ===
            render_utilities_and_conveniences(doc, building)

            # === OCCUPATION (BUILDING-LEVEL) ===
            # Primary: Use building-level occupier data
            if building.get('occupier_name'):
                relationship = building.get('occupier_relationship')

                # Special case for vacant buildings
                if relationship == 'vacant':
                    occupier_text = "The building is currently vacant."
                else:
                    occupier_text = f"The building is occupied by {building.get('occupier_name')}"

                    if relationship:
                        rel_labels = {
                            'owner': 'the owner',
                            'tenant': 'a tenant',
                            'family_member': 'a family member',
                            'caretaker': 'caretaker'
                        }
                        occupier_text += f" who is {rel_labels.get(relationship, relationship)}."
                    else:
                        occupier_text += "."

                add_inline_field(doc, "Occupation", occupier_text, space_after=Pt(6))

            # Fallback: Support old reports with property-level occupier
            elif prop.occupier_name:
                occupier_text = f"The property is occupied by {prop.occupier_name}"
                if prop.occupier_relationship:
                    rel_labels = {
                        'owner': 'the owner',
                        'tenant': 'a tenant',
                        'family_member': 'a family member',
                        'caretaker': 'caretaker'
                    }
                    occupier_text += f" who is {rel_labels.get(prop.occupier_relationship, prop.occupier_relationship)}."
                else:
                    occupier_text += "."

                add_inline_field(doc, "Occupation", occupier_text, space_after=Pt(6))

            # === BUILDING PHOTOS - USE SAME GRID TABLE FORMAT AS STANDALONE ===
            building_photos = building.get('building_photos', [])
            if building_photos and len(building_photos) > 0:
                sorted_photos = sorted(building_photos, key=lambda x: x.get('order', 0))

                import base64
                import re

                num_photos = len(sorted_photos)
                photos_per_row = 3

                idx = 0
                while idx < num_photos:
                    remaining = num_photos - idx
                    if remaining >= photos_per_row:
                        photos_in_row = photos_per_row
                    elif remaining == 1 and idx > 0:
                        photos_in_row = 1
                    else:
                        photos_in_row = remaining

                    table = doc.add_table(rows=2, cols=photos_in_row)
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER

                    for row in table.rows:
                        for cell in row.cells:
                            cell.width = Inches(6.5 / photos_in_row)
                            tc = cell._element
                            tcPr = tc.get_or_add_tcPr()
                            tcBorders = parse_xml(
                                r'<w:tcBorders %s>'
                                r'<w:top w:val="none"/>'
                                r'<w:left w:val="none"/>'
                                r'<w:bottom w:val="none"/>'
                                r'<w:right w:val="none"/>'
                                r'</w:tcBorders>' % nsdecls('w')
                            )
                            tcPr.append(tcBorders)

                    for i in range(photos_in_row):
                        if idx >= num_photos:
                            break

                        photo = sorted_photos[idx]
                        try:
                            image_data = photo.get('image_data', '')
                            caption = photo.get('caption', '')

                            if image_data:
                                # Handle both data URI and raw base64
                                if image_data.startswith('data:image'):
                                    base64_match = re.search(r'base64,(.+)', image_data)
                                    if not base64_match:
                                        idx += 1
                                        continue
                                    base64_data = base64_match.group(1)
                                else:
                                    base64_data = image_data

                                image_bytes = base64.b64decode(base64_data)
                                image_stream = BytesIO(image_bytes)

                                img_width = Inches(2.0)
                                dimensions = calculate_image_dimensions(
                                    image_stream,
                                    img_width,
                                    PROPERTY_PHOTO_HEIGHT
                                )

                                image_stream.seek(0)
                                cell = table.rows[0].cells[i]
                                cell_para = cell.paragraphs[0]
                                cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                cell_para.add_run().add_picture(image_stream, **dimensions)

                                caption_cell = table.rows[1].cells[i]
                                caption_para = caption_cell.paragraphs[0]
                                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                caption_para.paragraph_format.space_before = Pt(2)
                                caption_para.paragraph_format.space_after = Pt(2)

                                caption_text = f"Fig. {idx + 1}"
                                if caption:
                                    caption_text += f": {caption}"

                                caption_run = caption_para.add_run(caption_text)
                                caption_run.font.size = FONT_SIZE_CAPTION
                                caption_run.font.italic = True
                                caption_run.font.color.rgb = RGBColor(60, 60, 60)

                                logger.info(f"[MULTI-PROPERTY] Added building photo {idx + 1}")

                        except Exception as e:
                            logger.error(f"[MULTI-PROPERTY] Error adding photo {idx + 1}: {str(e)}")

                        idx += 1

                    spacing_para = doc.add_paragraph()
                    spacing_para.paragraph_format.space_after = Pt(8)

                final_spacing = doc.add_paragraph()
                final_spacing.paragraph_format.space_after = IMAGE_SPACING_AFTER

    # === DEVELOPMENT FEASIBILITY / ONGOING CONSTRUCTION (Bare Land only) ===
    if prop.property_type == 'bare_land' and prop.ongoing_construction_notes:
        construction_para = doc.add_paragraph()
        construction_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        construction_para.paragraph_format.space_before = Pt(4)
        construction_para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
        construction_para.paragraph_format.line_spacing = 0.9

        label_run = construction_para.add_run("Development Feasibility: ")
        label_run.font.bold = True
        label_run.font.size = FONT_SIZE_BODY
        label_run.font.color.rgb = RGBColor(0, 0, 0)

        value_run = construction_para.add_run(prop.ongoing_construction_notes)
        value_run.font.size = FONT_SIZE_BODY
        value_run.font.color.rgb = RGBColor(0, 0, 0)

    # === TOPOGRAPHICAL FEATURES (Bare Land only) ===
    if prop.property_type == 'bare_land':
        topo_parts = []

        if prop.elevation_changes:
            topo_parts.append(f"elevation changes are {prop.elevation_changes.replace('_', ' ')}")
        if prop.drainage_pattern:
            topo_parts.append(f"drainage pattern is {prop.drainage_pattern.replace('_', ' ')}")
        if prop.vegetation_type:
            topo_parts.append(f"vegetation type is {prop.vegetation_type.replace('_', ' ')}")
        if prop.natural_features:
            topo_parts.append(f"natural features include {prop.natural_features}")

        if topo_parts:
            topo_text = "The land has " + ", ".join(topo_parts) + "."
            topo_para = doc.add_paragraph()
            topo_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            topo_para.paragraph_format.space_before = Pt(4)
            topo_para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
            topo_para.paragraph_format.line_spacing = 0.9

            label_run = topo_para.add_run("Topographical Features: ")
            label_run.font.bold = True
            label_run.font.size = FONT_SIZE_BODY
            label_run.font.color.rgb = RGBColor(0, 0, 0)

            value_run = topo_para.add_run(topo_text)
            value_run.font.size = FONT_SIZE_BODY
            value_run.font.color.rgb = RGBColor(0, 0, 0)

    # Property Photos (embedded in DESCRIPTION section) - USE SAME GRID TABLE FORMAT AS STANDALONE
    if prop.property_type == 'bare_land':
        property_photos = safe_get_json_field(prop, 'property_photos', [])
        if property_photos and len(property_photos) > 0:
            sorted_photos = sorted(property_photos, key=lambda x: x.get('order', 0))

            import base64
            import re

            num_photos = len(sorted_photos)
            photos_per_row = 3

            idx = 0
            while idx < num_photos:
                remaining = num_photos - idx
                if remaining >= photos_per_row:
                    photos_in_row = photos_per_row
                elif remaining == 1 and idx > 0:
                    photos_in_row = 1
                else:
                    photos_in_row = remaining

                table = doc.add_table(rows=2, cols=photos_in_row)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                for row in table.rows:
                    for cell in row.cells:
                        cell.width = Inches(6.5 / photos_in_row)
                        tc = cell._element
                        tcPr = tc.get_or_add_tcPr()
                        tcBorders = parse_xml(
                            r'<w:tcBorders %s>'
                            r'<w:top w:val="none"/>'
                            r'<w:left w:val="none"/>'
                            r'<w:bottom w:val="none"/>'
                            r'<w:right w:val="none"/>'
                            r'</w:tcBorders>' % nsdecls('w')
                        )
                        tcPr.append(tcBorders)

                for i in range(photos_in_row):
                    if idx >= num_photos:
                        break

                    photo = sorted_photos[idx]
                    try:
                        image_data = photo.get('image_data', '')
                        caption = photo.get('caption', '')

                        if image_data:
                            # Handle both data URI and raw base64
                            if image_data.startswith('data:image'):
                                base64_match = re.search(r'base64,(.+)', image_data)
                                if not base64_match:
                                    idx += 1
                                    continue
                                base64_data = base64_match.group(1)
                            else:
                                base64_data = image_data

                            image_bytes = base64.b64decode(base64_data)
                            image_stream = BytesIO(image_bytes)

                            img_width = Inches(2.0)
                            dimensions = calculate_image_dimensions(
                                image_stream,
                                img_width,
                                PROPERTY_PHOTO_HEIGHT
                            )

                            image_stream.seek(0)
                            cell = table.rows[0].cells[i]
                            cell_para = cell.paragraphs[0]
                            cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            cell_para.add_run().add_picture(image_stream, **dimensions)

                            caption_cell = table.rows[1].cells[i]
                            caption_para = caption_cell.paragraphs[0]
                            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            caption_para.paragraph_format.space_before = Pt(2)
                            caption_para.paragraph_format.space_after = Pt(2)

                            caption_text = f"Fig. {idx + 1}"
                            if caption:
                                caption_text += f": {caption}"

                            caption_run = caption_para.add_run(caption_text)
                            caption_run.font.size = FONT_SIZE_CAPTION
                            caption_run.font.italic = True
                            caption_run.font.color.rgb = RGBColor(60, 60, 60)

                            logger.info(f"[MULTI-PROPERTY BARE LAND] Added property photo {idx + 1}")

                    except Exception as e:
                        logger.error(f"[MULTI-PROPERTY BARE LAND] Error adding photo {idx + 1}: {str(e)}")

                    idx += 1

                spacing_para = doc.add_paragraph()
                spacing_para.paragraph_format.space_after = Pt(8)

            final_spacing = doc.add_paragraph()
            final_spacing.paragraph_format.space_after = IMAGE_SPACING_AFTER

    # ===== 5.0 LOCALITY SECTION (CONDITIONAL) =====
    locality_text = generate_locality_description(prop)
    if locality_text or prop.locality_description_text:
        add_section_heading(doc, f"{section_num}.0", "LOCALITY")
        section_num += 1

        final_locality_text = prop.locality_description_text or locality_text

        locality_para = doc.add_paragraph()
        locality_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        locality_para.paragraph_format.space_before = Pt(0)
        locality_para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
        locality_para.paragraph_format.line_spacing = 0.9
        locality_run = locality_para.add_run(final_locality_text)
        locality_run.font.size = FONT_SIZE_BODY
        locality_run.font.color.rgb = RGBColor(0, 0, 0)

    # ===== 6.0 LEGAL ASPECTS SECTION (CONDITIONAL) =====
    has_legal_data = (
        prop.ownership_type or prop.street_lines_status or prop.building_limits_status or
        prop.local_authority_data or prop.rent_act_effectiveness or
        prop.title_search_conducted or prop.property_encumbered or
        prop.street_lines_gazette_ref or prop.building_plan_approved or
        prop.local_authority_rated
    )

    if has_legal_data:
        add_section_heading(doc, f"{section_num}.0", "LEGAL ASPECTS")
        section_num += 1

        # (a) Ownership
        if prop.ownership_type or hasattr(prop, 'deeds') or hasattr(prop, 'plan_number'):
            ownership_para = generate_ownership_paragraph(prop)
            add_subsection_paragraph(doc, "(a)", "Ownership", ownership_para)

        # (b) Street lines (skip for bare_land)
        if prop.street_lines_status and prop.property_type != 'bare_land':
            street_para = generate_street_lines_paragraph(prop)
            add_subsection_paragraph(doc, "(b)", "Street lines", street_para)

        # (c) Building limits (skip for bare_land)
        if prop.building_limits_status and prop.property_type != 'bare_land':
            building_para = generate_building_limits_paragraph(prop)
            add_subsection_paragraph(doc, "(c)", "Building limits", building_para)

        # (d) Local authority data
        if prop.local_authority_data or prop.pradeshiya_sabha or prop.local_authority_rated:
            authority_para = generate_local_authority_paragraph(prop)
            add_subsection_paragraph(doc, "(d)", "Local authority data", authority_para)

        # (e) Rent act effectiveness
        if prop.rent_act_effectiveness:
            rent_para = generate_rent_act_paragraph(prop)
            add_subsection_paragraph(doc, "(e)", "Rent act effectiveness", rent_para)

        doc.add_paragraph()  # Spacing

    # ===== 7.0 LAND VALUES IN THE AREA (CONDITIONAL) =====
    comparables = safe_get_json_field(prop, 'comparable_properties', [])
    # Show section if either comparables OR market analysis exists
    if comparables or prop.land_market_analysis:
        add_section_heading(doc, f"{section_num}.0", "LAND VALUES IN THE AREA")
        section_num += 1

        land_values_text = generate_land_values_paragraph(comparables)
        if land_values_text:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(land_values_text)
            run.font.size = FONT_SIZE_BODY

            # Average rate
            avg_rate = sum(c.get('rate_per_perch', 0) for c in comparables) / len(comparables) if comparables else 0
            if avg_rate > 0:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run(f"Average Rate: LKR {avg_rate:,.2f} per perch")
                run.font.bold = True
                run.font.size = FONT_SIZE_BODY

        if prop.land_market_analysis:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run(prop.land_market_analysis)
            run.font.size = FONT_SIZE_BODY

        doc.add_paragraph()

    # ===== 8.0 VALUATION SECTION =====
    if prop.valuation_total_land_value or prop.valuation_buildings_data:
        add_section_heading(doc, f"{section_num}.0", "VALUATION OF THE PROPERTY")
        section_num += 1

        # Land valuation
        if prop.valuation_total_land_value:
            extent = prop.valuation_land_extent or prop.land_extent_perches or 0
            rate = prop.valuation_rate_per_perch or 0
            land_value = prop.valuation_total_land_value

            p = doc.add_paragraph()

            # Add tab stop for right alignment
            tab_stops = p.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(6.0), WD_TAB_ALIGNMENT.RIGHT)

            text = f"Land – {extent:,.2f} perches @ {format_currency(rate)} per perch\t= {format_currency_aligned(land_value)}"
            run = p.add_run(text)
            run.font.size = FONT_SIZE_VALUATION
            p.paragraph_format.space_after = Pt(3)

        # Buildings valuation (skip for bare_land)
        total_buildings_value = 0
        buildings_insurance_values = []

        if prop.property_type != 'bare_land' and prop.valuation_buildings_data:
            buildings_data = safe_get_json_field(prop, 'valuation_buildings_data', [])

            for idx, bldg in enumerate(buildings_data, 1):
                building_name = bldg.get('building_name', f'Building {idx}')
                subtotal = bldg.get('subtotal', 0)

                # Calculate total floor area from components
                components = bldg.get('components', [])
                total_floor_area = sum(comp.get('floor_area', 0) for comp in components)

                # Calculate average rate per sq.ft
                avg_rate = subtotal / total_floor_area if total_floor_area > 0 else 0

                # Check if depreciation data exists
                has_depreciation = bldg.get('depreciation_amount') is not None and to_float(bldg.get('depreciation_amount', 0)) > 0

                if has_depreciation:
                    depreciation_rate = to_float(bldg.get('depreciation_rate_percent', 0))
                    depreciated_value = to_float(bldg.get('depreciated_value', subtotal))

                    # NEW FORMAT: 2-line building with inline depreciation
                    format_building_valuation_2line(
                        doc,
                        building_name,
                        total_floor_area,
                        avg_rate,
                        depreciation_rate,
                        depreciated_value
                    )

                    building_value = depreciated_value
                else:
                    # No depreciation: Use single-line format
                    p = doc.add_paragraph()

                    # Add tab stop for right alignment
                    tab_stops = p.paragraph_format.tab_stops
                    tab_stops.add_tab_stop(Inches(6.0), WD_TAB_ALIGNMENT.RIGHT)

                    text = f"{building_name} – {total_floor_area:,.0f} sq.ft @ {format_currency(avg_rate)} per square foot\t= {format_currency_aligned(subtotal)}"
                    run = p.add_run(text)
                    run.font.size = FONT_SIZE_VALUATION
                    p.paragraph_format.space_after = Pt(3)
                    building_value = to_float(subtotal)

                total_buildings_value += building_value
                # Insurance always uses replacement cost (undepreciated)
                buildings_insurance_values.append({
                    'name': building_name,
                    'value': to_float(subtotal)
                })

        # Addons/Improvements (support both field names)
        total_addons_value = 0
        if prop.valuation_addons:
            addons = safe_get_json_field(prop, 'valuation_addons', [])
            for addon in addons:
                # Support both 'description' and 'item_name' fields
                addon_description = addon.get('description') or addon.get('item_name', 'Add-on')
                addon_value = to_float(addon.get('value', 0))

                format_addon_compact(doc, addon_description, addon_value)
                total_addons_value += addon_value

        # Calculate market values
        land_val = to_float(prop.valuation_total_land_value)
        market_value_calculated = land_val + total_buildings_value + total_addons_value
        market_value_rounded = round_for_say(market_value_calculated)

        # Determine if we should show "Market Value of the property" section
        has_buildings_or_addons = (total_buildings_value > 0) or (total_addons_value > 0)

        if has_buildings_or_addons:
            # Show "Market Value of the property" line (with double underline)
            add_market_value_line(doc, market_value_calculated, has_blank_before=True)

        # Always show "Value rounded off" line
        add_value_rounded_line(doc, market_value_rounded)

        # === NEW: SUMMARY OF THE VALUATION (previously missing) ===
        # Check if valuation type is "Forced Sale Value" to show forced sale fields
        show_forced_sale = report.valuation_type == "Forced Sale Value"

        if show_forced_sale:
            forced_sale_percentage = prop.valuation_forced_sale_percentage or 90
            forced_sale_value = market_value_rounded * (forced_sale_percentage / 100)

        p = doc.add_paragraph()
        run = p.add_run("SUMMARY OF THE VALUATION")
        run.font.bold = True
        run.font.underline = True
        run.font.size = FONT_SIZE_BODY
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)

        # Open Market Value
        p = doc.add_paragraph()
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(3.5), WD_TAB_ALIGNMENT.LEFT)
        tab_stops.add_tab_stop(Inches(3.7), WD_TAB_ALIGNMENT.LEFT)
        text = f"Open Market Value of the property\t:\t{format_currency(market_value_rounded)}"
        run = p.add_run(text)
        run.font.size = FONT_SIZE_VALUATION
        p.paragraph_format.space_after = Pt(3)

        # Forced Sale Value - only show when valuation type is "Forced Sale Value"
        if show_forced_sale:
            p = doc.add_paragraph()
            tab_stops = p.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(3.5), WD_TAB_ALIGNMENT.LEFT)
            tab_stops.add_tab_stop(Inches(3.7), WD_TAB_ALIGNMENT.LEFT)
            text = f"Forced Sale Value of the property\t:\t{format_currency(forced_sale_value)}"
            run = p.add_run(text)
            run.font.size = FONT_SIZE_VALUATION
            p.paragraph_format.space_after = Pt(3)

        # Insurance Value (NEW INLINE FORMAT)
        if buildings_insurance_values:
            for building_ins in buildings_insurance_values:
                p = doc.add_paragraph()
                tab_stops = p.paragraph_format.tab_stops
                tab_stops.add_tab_stop(Inches(3.5), WD_TAB_ALIGNMENT.LEFT)
                tab_stops.add_tab_stop(Inches(3.7), WD_TAB_ALIGNMENT.LEFT)
                text = f"Insurance Value of the {building_ins['name']}\t:\t{format_currency(building_ins['value'])}"
                run = p.add_run(text)
                run.font.size = FONT_SIZE_VALUATION
                p.paragraph_format.space_after = Pt(2)

            doc.add_paragraph()  # Final spacing

    # ===== 9.0 CERTIFICATION SECTION (SIMPLIFIED FORMAT - REPORT-LEVEL) =====
    add_section_heading(doc, f"{section_num}.0", "CERTIFICATION")
    section_num += 1

    # Use report-level certification - simplified format
    if report.certification_text:
        # Use custom certification text if provided (override mode)
        cert_text = report.certification_text.strip()
    elif report.certification_valuer_name and report.certification_valuer_designation:
        # Auto-generate simplified certification using report-level data
        # Multi-property uses single unified certification
        cert_text = generate_simplified_certification_text(
            valuer_name=report.certification_valuer_name,
            valuer_designation=report.certification_valuer_designation,
            lot_number=report.lot_number,
            plan_number=report.plan_number,
            plan_date=report.plan_date,
            licensed_surveyor_name=report.licensed_surveyor_name,
            deeds=safe_get_json_field(report, 'deeds', []),
            property_identification_type=report.property_identification_type
        )
    else:
        # Fallback if no valuer info
        cert_text = "I hereby certify that I have personally inspected this property and the valuation stated herein represents my professional opinion of the market value as of the date of inspection."

    # Render certification paragraph
    cert_para = doc.add_paragraph()
    cert_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    cert_para.paragraph_format.line_spacing = 0.9
    cert_para.paragraph_format.space_after = Pt(12)
    cert_run = cert_para.add_run(cert_text)
    cert_run.font.size = FONT_SIZE_BODY
    cert_run.font.color.rgb = RGBColor(0, 0, 0)

    # Signature block
    add_signature_block(
        doc=doc,
        user=user,
        valuer_name=report.certification_valuer_name,
        valuer_designation=report.certification_valuer_designation,
        certification_date=report.certification_date
    )


def migrate_invoice_data(invoice_dict):
    """Migrate old invoice structure (qty, unit_price) to new structure (direct total)"""
    if not invoice_dict:
        return invoice_dict

    # Check if migration needed
    items = invoice_dict.get('items', [])
    needs_migration = any('quantity' in item or 'unit_price' in item for item in items)

    if not needs_migration:
        return invoice_dict  # Already new format

    # Migrate items - keep total, remove quantity/unit_price
    migrated_items = [
        {'description': item.get('description', ''), 'total': item.get('total', 0)}
        for item in items
    ]

    # Convert bank_details (string) → manual_bank_details
    return {
        'items': migrated_items,
        'subtotal': invoice_dict.get('subtotal', 0),
        'traveling_charges': invoice_dict.get('traveling_charges'),
        'discount': invoice_dict.get('discount'),
        'total': invoice_dict.get('total', 0),
        'bank_account_ids': [],
        'manual_bank_details': invoice_dict.get('bank_details')  # Old string field
    }


def _generate_invoice_section(doc, invoice_data, user, report=None):
    """Generate professional fee invoice section with letterhead and recipient address"""
    # Add page break for new invoice page
    doc.add_page_break()

    # ===== LETTERHEAD =====
    # Render letterhead matching the rest of the report
    template_id = user.preferred_letterhead_template or 'classic'
    template = get_template(template_id)
    template.render_letterhead(doc, user, report)

    # Add spacing after letterhead
    doc.add_paragraph()

    # ===== RECIPIENT ADDRESS =====
    if report:
        # Extract recipient information from report
        recipient_para = doc.add_paragraph()
        recipient_run = recipient_para.add_run("To:\n")
        recipient_run.bold = True
        recipient_run.font.size = FONT_SIZE_VALUATION

        # Add applicant name if available
        if hasattr(report, 'applicant_full_name') and report.applicant_full_name:
            name_run = recipient_para.add_run(f"{report.applicant_full_name}\n")
            name_run.font.size = FONT_SIZE_VALUATION

        # Add applicant address if available
        address_parts = []
        if hasattr(report, 'applicant_address_line1') and report.applicant_address_line1:
            address_parts.append(report.applicant_address_line1)
        if hasattr(report, 'applicant_address_line2') and report.applicant_address_line2:
            address_parts.append(report.applicant_address_line2)

        if address_parts:
            address_run = recipient_para.add_run(f"{', '.join(address_parts)}\n")
            address_run.font.size = FONT_SIZE_VALUATION

        # Add spacing after recipient address
        doc.add_paragraph()

    # Invoice heading
    invoice_heading = doc.add_paragraph()
    invoice_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    invoice_heading.paragraph_format.space_before = Pt(0)
    invoice_heading.paragraph_format.space_after = Pt(12)
    invoice_run = invoice_heading.add_run("PROFESSIONAL FEES")
    invoice_run.bold = True
    invoice_run.font.size = FONT_SIZE_BODY
    invoice_run.font.color.rgb = RGBColor(0, 0, 0)

    # Parse and migrate invoice_data
    if isinstance(invoice_data, str):
        import json
        invoice_dict = json.loads(invoice_data)
    else:
        invoice_dict = invoice_data

    # Migrate old format to new format
    invoice_dict = migrate_invoice_data(invoice_dict)

    # Invoice table - SIMPLIFIED (2 columns)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set column widths
    table.columns[0].width = Inches(4.5)  # Description - wider
    table.columns[1].width = Inches(2.0)  # Total

    # Header row
    header_cells = table.rows[0].cells
    headers = ['Description', 'Total (LKR)']
    for i, header_text in enumerate(headers):
        cell = header_cells[i]
        cell.text = header_text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = FONT_SIZE_TABLE_HEADER
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 1 else WD_ALIGN_PARAGRAPH.LEFT
        # Shading for header
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        cell_xml = cell._element
        cell_properties = cell_xml.get_or_add_tcPr()
        shading_element = OxmlElement('w:shd')
        shading_element.set(qn('w:fill'), 'D9D9D9')  # Light gray
        cell_properties.append(shading_element)

    # Item rows
    for item in invoice_dict.get('items', []):
        row_cells = table.add_row().cells
        row_cells[0].text = item.get('description', '')
        row_cells[1].text = format_currency(item.get('total', 0))
        row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = FONT_SIZE_TABLE_CELL

    # Subtotal row
    subtotal_row = table.add_row().cells
    subtotal_row[0].text = "Subtotal:"
    subtotal_row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    subtotal_row[1].text = format_currency(invoice_dict.get('subtotal', 0))
    subtotal_row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Traveling charges if present
    if invoice_dict.get('traveling_charges'):
        travel_row = table.add_row().cells
        travel_row[0].text = "Traveling Charges:"
        travel_row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        travel_row[1].text = format_currency(invoice_dict.get('traveling_charges', 0))
        travel_row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Discount if present
    if invoice_dict.get('discount'):
        discount_row = table.add_row().cells
        discount_row[0].text = "Discount:"
        discount_row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        discount_row[1].text = f"-{format_currency(invoice_dict.get('discount', 0))}"
        discount_row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Grand Total row
    total_row = table.add_row().cells
    total_row[0].text = "GRAND TOTAL:"
    total_row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for paragraph in total_row[0].paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = FONT_SIZE_INVOICE_TOTAL

    total_row[1].text = format_currency(invoice_dict.get('total', 0))
    total_row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for paragraph in total_row[1].paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = FONT_SIZE_INVOICE_TOTAL

    # Shade total row
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    for cell in [total_row[0], total_row[1]]:
        cell_xml = cell._element
        cell_properties = cell_xml.get_or_add_tcPr()
        shading_element = OxmlElement('w:shd')
        shading_element.set(qn('w:fill'), 'E6E6E6')
        cell_properties.append(shading_element)

    doc.add_paragraph()  # Spacing

    # Bank Details - NEW: Show all selected accounts
    bank_account_ids = invoice_dict.get('bank_account_ids', [])
    manual_bank_details = invoice_dict.get('manual_bank_details')

    if bank_account_ids and len(bank_account_ids) > 0 and user:
        # Fetch user's bank accounts
        user_accounts = user.bank_accounts or []
        selected_accounts = [acc for acc in user_accounts if acc['id'] in bank_account_ids]

        if selected_accounts:
            bank_heading = doc.add_paragraph()
            bank_heading_run = bank_heading.add_run("Bank Account Details:")
            bank_heading_run.bold = True
            bank_heading_run.font.size = FONT_SIZE_BANK_HEADER
            bank_heading.paragraph_format.space_before = Pt(6)
            bank_heading.paragraph_format.space_after = Pt(3)

            # Render each account
            for account in selected_accounts:
                acc_para = doc.add_paragraph(style='List Bullet')
                acc_para.paragraph_format.left_indent = Inches(0.25)
                acc_para.paragraph_format.line_spacing = 1.0
                acc_run = acc_para.add_run(
                    f"{account['bank_name']}\n"
                    f"Account Number: {account['account_number']}\n"
                    f"Branch: {account['branch_name']}"
                )
                acc_run.font.size = FONT_SIZE_BANK_DETAILS

    elif manual_bank_details:
        # Fallback to manual entry
        bank_heading = doc.add_paragraph()
        bank_heading_run = bank_heading.add_run("Bank Account Details:")
        bank_heading_run.bold = True
        bank_heading_run.font.size = FONT_SIZE_BANK_HEADER
        bank_heading.paragraph_format.space_before = Pt(6)
        bank_heading.paragraph_format.space_after = Pt(3)

        bank_para = doc.add_paragraph()
        bank_para.paragraph_format.line_spacing = 1.0
        bank_run = bank_para.add_run(manual_bank_details)
        bank_run.font.size = FONT_SIZE_BANK_DETAILS

    # Signature space
    doc.add_paragraph()
    doc.add_paragraph()
    signature_para = doc.add_paragraph()
    signature_para.paragraph_format.space_before = Pt(24)
    signature_run = signature_para.add_run("_" * 40)
    signature_run.font.size = FONT_SIZE_SIGNATURE

    sig_label = doc.add_paragraph()
    sig_label_run = sig_label.add_run("Signature")
    sig_label_run.font.size = FONT_SIZE_SIGNATURE
    sig_label_run.font.italic = True


# ===== VEHICLE REPORT GENERATION =====

def generate_vehicle_report_docx(report: models.Report, user: models.User) -> BytesIO:
    """
    Generate a formatted DOCX file for a vehicle valuation report.

    Args:
        report: Report model instance with vehicle data
        user: User model instance

    Returns:
        BytesIO object containing the DOCX document
    """
    try:
        # Create a new Document
        doc = Document()

        # ===== LETTERHEAD =====
        template_id = user.preferred_letterhead_template or 'classic'
        template = get_template(template_id)
        template.render_letterhead(doc, user, report)

        # Get primary vehicle
        vehicle = report.primary_vehicle
        if not vehicle:
            # Try to get from associations
            if report.vehicle_associations and len(report.vehicle_associations) > 0:
                vehicle = report.vehicle_associations[0].vehicle
            else:
                raise ValueError("No vehicle found for this report")

        # ===== TITLE BLOCK =====
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run("VEHICLE VALUATION REPORT")
        title_run.bold = True
        title_run.font.size = Pt(14)
        title_run.font.underline = True

        # Vehicle identification line
        vehicle_desc = f"{vehicle.make or 'Unknown'} {vehicle.model or ''} ({vehicle.registration_number or 'Unregistered'})"
        desc_para = doc.add_paragraph()
        desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        desc_run = desc_para.add_run(vehicle_desc)
        desc_run.bold = True
        desc_run.font.size = FONT_SIZE_BODY

        doc.add_paragraph()  # Spacing

        # ===== HEADER INFORMATION =====
        header_table = doc.add_table(rows=6, cols=2)
        header_table.style = 'Table Grid'

        header_data = [
            ("Purpose of Valuation", report.valuation_purpose or "Market Value Assessment"),
            ("Requested By", report.applicant_full_name or ""),
            ("Report Date", report.report_date or ""),
            ("Inspection Date", report.inspection_date or ""),
            ("Folio Number", report.folio_number or ""),
            ("Inspection Place", report.inspection_place or ""),
        ]

        for i, (label, value) in enumerate(header_data):
            row = header_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = str(value) if value else ""
            # Bold the label
            for paragraph in row.cells[0].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = FONT_SIZE_BODY
            for paragraph in row.cells[1].paragraphs:
                for run in paragraph.runs:
                    run.font.size = FONT_SIZE_BODY

        doc.add_paragraph()  # Spacing

        # ===== VEHICLE IDENTIFICATION SECTION =====
        id_heading = doc.add_paragraph()
        id_heading_run = id_heading.add_run("1. VEHICLE IDENTIFICATION")
        id_heading_run.bold = True
        id_heading_run.font.size = Pt(12)

        id_table = doc.add_table(rows=14, cols=2)
        id_table.style = 'Table Grid'

        id_data = [
            ("Registration Number", vehicle.registration_number),
            ("Provincial Council", vehicle.provincial_council),
            ("Class of Vehicle", vehicle.class_of_vehicle),
            ("Make", vehicle.make),
            ("Model", vehicle.model),
            ("Year of Manufacture", str(vehicle.year_of_manufacture) if vehicle.year_of_manufacture else ""),
            ("Date of First Registration", vehicle.date_of_first_registration),
            ("Body Colour", vehicle.body_colour),
            ("Chassis Number", vehicle.chassis_number),
            ("Engine Number", vehicle.engine_number),
            ("Cylinder Capacity", f"{vehicle.cylinder_capacity} cc" if vehicle.cylinder_capacity else ""),
            ("Fuel Type", vehicle.fuel_type),
            ("Mileage", f"{vehicle.mileage:,} {vehicle.mileage_unit or 'km'}" if vehicle.mileage else ""),
            ("Country of Origin", vehicle.country_of_origin),
        ]

        for i, (label, value) in enumerate(id_data):
            row = id_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = str(value) if value else ""
            for paragraph in row.cells[0].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = FONT_SIZE_BODY
            for paragraph in row.cells[1].paragraphs:
                for run in paragraph.runs:
                    run.font.size = FONT_SIZE_BODY

        doc.add_paragraph()  # Spacing

        # ===== ENGINE & TRANSMISSION SECTION =====
        engine_heading = doc.add_paragraph()
        engine_heading_run = engine_heading.add_run("2. ENGINE & TRANSMISSION")
        engine_heading_run.bold = True
        engine_heading_run.font.size = Pt(12)

        engine_table = doc.add_table(rows=6, cols=2)
        engine_table.style = 'Table Grid'

        engine_data = [
            ("Engine Type", vehicle.engine_type),
            ("Transmission", vehicle.transmission),
            ("Wheel Drive", vehicle.wheel_drive),
            ("Running Condition", vehicle.running_condition),
            ("Engine Condition", vehicle.engine_condition),
            ("Gear Box Condition", vehicle.gear_box_condition),
        ]

        for i, (label, value) in enumerate(engine_data):
            row = engine_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = str(value) if value else ""
            for paragraph in row.cells[0].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = FONT_SIZE_BODY
            for paragraph in row.cells[1].paragraphs:
                for run in paragraph.runs:
                    run.font.size = FONT_SIZE_BODY

        doc.add_paragraph()  # Spacing

        # ===== BODY & CONDITION SECTION =====
        body_heading = doc.add_paragraph()
        body_heading_run = body_heading.add_run("3. BODY & CONDITION")
        body_heading_run.bold = True
        body_heading_run.font.size = Pt(12)

        body_table = doc.add_table(rows=5, cols=2)
        body_table.style = 'Table Grid'

        body_data = [
            ("Body Condition", vehicle.body_condition),
            ("Chassis Condition", vehicle.chassis_condition),
            ("Upholstery Condition", vehicle.upholstery_condition),
            ("Underside Condition", vehicle.underside_condition),
            ("Clutch Status", vehicle.clutch_status),
        ]

        for i, (label, value) in enumerate(body_data):
            row = body_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = str(value) if value else ""
            for paragraph in row.cells[0].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = FONT_SIZE_BODY
            for paragraph in row.cells[1].paragraphs:
                for run in paragraph.runs:
                    run.font.size = FONT_SIZE_BODY

        doc.add_paragraph()  # Spacing

        # ===== PARTS AVAILABILITY SECTION =====
        parts_heading = doc.add_paragraph()
        parts_heading_run = parts_heading.add_run("4. PARTS AVAILABILITY")
        parts_heading_run.bold = True
        parts_heading_run.font.size = Pt(12)

        parts_table = doc.add_table(rows=3, cols=2)
        parts_table.style = 'Table Grid'

        parts_data = [
            ("Body Parts", vehicle.body_parts_status),
            ("Engine Parts", vehicle.engine_parts_status),
            ("Accessories", vehicle.accessories_status),
        ]

        for i, (label, value) in enumerate(parts_data):
            row = parts_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = str(value) if value else ""
            for paragraph in row.cells[0].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = FONT_SIZE_BODY
            for paragraph in row.cells[1].paragraphs:
                for run in paragraph.runs:
                    run.font.size = FONT_SIZE_BODY

        doc.add_paragraph()  # Spacing

        # ===== BRAKES & SAFETY SECTION =====
        brakes_heading = doc.add_paragraph()
        brakes_heading_run = brakes_heading.add_run("5. BRAKES & SAFETY")
        brakes_heading_run.bold = True
        brakes_heading_run.font.size = Pt(12)

        brakes_table = doc.add_table(rows=4, cols=2)
        brakes_table.style = 'Table Grid'

        brakes_data = [
            ("Foot Brake Condition", vehicle.foot_brake_condition),
            ("Disc Brake", "Available" if vehicle.disc_brake_available else "Not Available"),
            ("Parking Brake Condition", vehicle.parking_brake_condition),
            ("ABS", "Available" if vehicle.abs_available else "Not Available"),
        ]

        for i, (label, value) in enumerate(brakes_data):
            row = brakes_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = str(value) if value else ""
            for paragraph in row.cells[0].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = FONT_SIZE_BODY
            for paragraph in row.cells[1].paragraphs:
                for run in paragraph.runs:
                    run.font.size = FONT_SIZE_BODY

        doc.add_paragraph()  # Spacing

        # ===== FEATURES SECTION =====
        features = safe_get_json_field(vehicle, 'features', {})
        if features:
            features_heading = doc.add_paragraph()
            features_heading_run = features_heading.add_run("6. FEATURES & AMENITIES")
            features_heading_run.bold = True
            features_heading_run.font.size = Pt(12)

            features_list = []
            if features.get('air_condition'):
                features_list.append("Air Condition")
            if features.get('dual_air_condition'):
                features_list.append("Dual Air Condition")
            if features.get('power_mirror'):
                features_list.append("Power Mirror")
            if features.get('power_window'):
                features_list.append("Power Window")
            if features.get('power_steering'):
                features_list.append("Power Steering")
            if features.get('airbag'):
                features_list.append(f"Airbag (x{features.get('num_airbags', 1)})")

            if features_list:
                features_para = doc.add_paragraph()
                features_para.add_run("Available Features: ").bold = True
                features_para.add_run(", ".join(features_list))

            if features.get('seats'):
                seats_para = doc.add_paragraph()
                seats_para.add_run("Seating Capacity: ").bold = True
                seats_para.add_run(str(features.get('seats')))

            if features.get('doors'):
                doors_para = doc.add_paragraph()
                doors_para.add_run("Number of Doors: ").bold = True
                doors_para.add_run(str(features.get('doors')))

        doc.add_paragraph()  # Spacing

        # ===== TYRES SECTION =====
        tyres = safe_get_json_field(vehicle, 'tyres', {})
        if tyres:
            tyres_heading = doc.add_paragraph()
            tyres_heading_run = tyres_heading.add_run("7. TYRES")
            tyres_heading_run.bold = True
            tyres_heading_run.font.size = Pt(12)

            tyres_table = doc.add_table(rows=5, cols=3)
            tyres_table.style = 'Table Grid'

            # Header row
            tyres_table.rows[0].cells[0].text = ""
            tyres_table.rows[0].cells[1].text = "Front"
            tyres_table.rows[0].cells[2].text = "Rear"
            for cell in tyres_table.rows[0].cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.bold = True

            front = tyres.get('front', {})
            rear = tyres.get('rear', {})

            tyre_rows = [
                ("Brand", front.get('brand', ''), rear.get('brand', '')),
                ("Size", front.get('size', ''), rear.get('size', '')),
                ("Tread %", f"{front.get('tread_percent', '')}%" if front.get('tread_percent') else "", f"{rear.get('tread_percent', '')}%" if rear.get('tread_percent') else ""),
                ("Condition", front.get('condition', ''), rear.get('condition', '')),
            ]

            for i, (label, front_val, rear_val) in enumerate(tyre_rows, start=1):
                row = tyres_table.rows[i]
                row.cells[0].text = label
                row.cells[1].text = str(front_val) if front_val else ""
                row.cells[2].text = str(rear_val) if rear_val else ""
                for para in row.cells[0].paragraphs:
                    for run in para.runs:
                        run.font.bold = True

            # Spare tyre and replacement
            spare_para = doc.add_paragraph()
            spare_para.add_run("Spare Tyre: ").bold = True
            spare_para.add_run("Available" if tyres.get('spare_available') else "Not Available")

            if tyres.get('need_replacement'):
                replace_para = doc.add_paragraph()
                replace_para.add_run("Note: ").bold = True
                replace_para.add_run("Tyres need replacement")

        doc.add_paragraph()  # Spacing

        # ===== ELECTRICAL & LIGHTS SECTION =====
        electrical = safe_get_json_field(vehicle, 'electrical', {})
        lights = safe_get_json_field(vehicle, 'lights', {})

        if electrical or lights:
            elec_heading = doc.add_paragraph()
            elec_heading_run = elec_heading.add_run("8. ELECTRICAL & LIGHTS")
            elec_heading_run.bold = True
            elec_heading_run.font.size = Pt(12)

            if electrical:
                elec_para = doc.add_paragraph()
                elec_para.add_run("Electrical System: ").bold = True
                elec_items = []
                if electrical.get('starter'):
                    elec_items.append("Starter ✓")
                if electrical.get('horn'):
                    elec_items.append("Horn ✓")
                if electrical.get('wiper'):
                    elec_items.append("Wiper ✓")
                if electrical.get('battery_condition'):
                    elec_items.append(f"Battery ({electrical.get('battery_condition')})")
                elec_para.add_run(", ".join(elec_items) if elec_items else "N/A")

            if lights:
                lights_para = doc.add_paragraph()
                lights_para.add_run("Lights: ").bold = True
                light_items = []
                if lights.get('head'):
                    light_items.append("Head ✓")
                if lights.get('dim'):
                    light_items.append("Dim ✓")
                if lights.get('signal'):
                    light_items.append("Signal ✓")
                if lights.get('parking'):
                    light_items.append("Parking ✓")
                if lights.get('reverse'):
                    light_items.append("Reverse ✓")
                if lights.get('meter'):
                    light_items.append("Meter ✓")
                lights_para.add_run(", ".join(light_items) if light_items else "N/A")

        doc.add_paragraph()  # Spacing

        # ===== HISTORY SECTION =====
        history_heading = doc.add_paragraph()
        history_heading_run = history_heading.add_run("9. HISTORY & REPAIRS")
        history_heading_run.bold = True
        history_heading_run.font.size = Pt(12)

        history_table = doc.add_table(rows=4, cols=2)
        history_table.style = 'Table Grid'

        history_data = [
            ("Accident History", "Yes" if vehicle.has_accidents else "No"),
            ("Repair History", "Yes" if vehicle.has_repairs else "No"),
            ("Repairs Needed Within Year", "Yes" if vehicle.needs_repairs_within_year else "No"),
            ("Body Parts Replaced", "Yes" if vehicle.body_parts_replaced else "No"),
        ]

        for i, (label, value) in enumerate(history_data):
            row = history_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = str(value) if value else ""
            for paragraph in row.cells[0].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = FONT_SIZE_BODY
            for paragraph in row.cells[1].paragraphs:
                for run in paragraph.runs:
                    run.font.size = FONT_SIZE_BODY

        doc.add_paragraph()  # Spacing

        # ===== VALUATION SECTION =====
        val_heading = doc.add_paragraph()
        val_heading_run = val_heading.add_run("10. VALUATION")
        val_heading_run.bold = True
        val_heading_run.font.size = Pt(12)

        val_table = doc.add_table(rows=4, cols=2)
        val_table.style = 'Table Grid'

        val_data = [
            ("Purchase Price", f"Rs. {to_float(vehicle.purchase_price):,.2f}" if vehicle.purchase_price else ""),
            ("Brand New Price", f"Rs. {to_float(vehicle.brand_new_price):,.2f}" if vehicle.brand_new_price else ""),
            ("Market Value", f"Rs. {to_float(vehicle.market_value):,.2f}" if vehicle.market_value else ""),
            ("Forced Sale Value", f"Rs. {to_float(vehicle.forced_sale_value):,.2f}" if vehicle.forced_sale_value else ""),
        ]

        for i, (label, value) in enumerate(val_data):
            row = val_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = str(value) if value else ""
            for paragraph in row.cells[0].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = FONT_SIZE_BODY
            for paragraph in row.cells[1].paragraphs:
                for run in paragraph.runs:
                    run.font.size = FONT_SIZE_BODY

        doc.add_paragraph()  # Spacing

        # Valuation Summary
        if vehicle.valuation_summary:
            summary_para = doc.add_paragraph()
            summary_para.add_run("Valuation Summary: ").bold = True
            summary_para.add_run(vehicle.valuation_summary)

        doc.add_paragraph()  # Spacing

        # ===== OFFICE USE SECTION (Conditional) =====
        if report.is_office_use:
            office_data = safe_get_json_field(vehicle, 'office_data', {})
            if office_data:
                office_heading = doc.add_paragraph()
                office_heading_run = office_heading.add_run("FOR OFFICE USE ONLY")
                office_heading_run.bold = True
                office_heading_run.font.size = Pt(12)

                office_table = doc.add_table(rows=3, cols=2)
                office_table.style = 'Table Grid'

                office_rows = [
                    ("Civil No", office_data.get('civil_no', '')),
                    ("Military No", office_data.get('military_no', '')),
                    ("Approval Position", office_data.get('approval_position', '')),
                ]

                for i, (label, value) in enumerate(office_rows):
                    row = office_table.rows[i]
                    row.cells[0].text = label
                    row.cells[1].text = str(value) if value else ""

            # Past Valuations Table
            past_valuations = safe_get_json_field(vehicle, 'past_valuations', [])
            if past_valuations:
                past_heading = doc.add_paragraph()
                past_heading.paragraph_format.space_before = Pt(12)
                past_heading_run = past_heading.add_run("Previous Assessment")
                past_heading_run.bold = True
                past_heading_run.font.size = Pt(11)

                past_table = doc.add_table(rows=len(past_valuations) + 1, cols=5)
                past_table.style = 'Table Grid'

                # Header row
                headers = ["S/N", "Vehicle No - Civil", "Vehicle No - Military", "Year", "Value (Rs)"]
                for j, header in enumerate(headers):
                    past_table.rows[0].cells[j].text = header
                    for para in past_table.rows[0].cells[j].paragraphs:
                        for run in para.runs:
                            run.font.bold = True

                # Data rows
                for i, pv in enumerate(past_valuations, start=1):
                    past_table.rows[i].cells[0].text = str(pv.get('serial', i))
                    past_table.rows[i].cells[1].text = str(pv.get('civil_no', ''))
                    past_table.rows[i].cells[2].text = str(pv.get('military_no', ''))
                    past_table.rows[i].cells[3].text = str(pv.get('year', ''))
                    past_table.rows[i].cells[4].text = f"Rs. {to_float(pv.get('value')):,.2f}" if pv.get('value') else ""

        doc.add_paragraph()  # Spacing

        # ===== VEHICLE PHOTOS =====
        vehicle_photos = safe_get_json_field(vehicle, 'vehicle_photos', [])
        if vehicle_photos:
            photos_heading = doc.add_paragraph()
            photos_heading_run = photos_heading.add_run("VEHICLE PHOTOGRAPHS")
            photos_heading_run.bold = True
            photos_heading_run.font.size = Pt(12)

            for photo in vehicle_photos:
                image_data = photo.get('image_data', '')
                caption = photo.get('caption', '')

                if image_data:
                    try:
                        # Handle base64 data URL
                        if ',' in image_data:
                            image_data = image_data.split(',')[1]

                        import base64
                        image_bytes = base64.b64decode(image_data)
                        image_stream = BytesIO(image_bytes)

                        # Add image
                        doc.add_picture(image_stream, width=Inches(4))

                        # Add caption if exists
                        if caption:
                            cap_para = doc.add_paragraph()
                            cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            cap_run = cap_para.add_run(caption)
                            cap_run.font.italic = True
                            cap_run.font.size = Pt(10)

                        doc.add_paragraph()  # Spacing
                    except Exception as img_error:
                        logger.warning(f"Could not add vehicle photo: {img_error}")

        # ===== SIGNATURE SECTION =====
        doc.add_paragraph()
        doc.add_paragraph()
        sig_para = doc.add_paragraph()
        sig_para.paragraph_format.space_before = Pt(24)
        sig_para.add_run("_" * 40)

        date_para = doc.add_paragraph()
        date_para.add_run("Date: _________________")

        sig_label = doc.add_paragraph()
        sig_label_run = sig_label.add_run("Signature")
        sig_label_run.font.italic = True

        # Save to BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        logger.info("[DOCX] Vehicle report generated successfully")
        return buffer

    except Exception as e:
        logger.error(f"[DOCX] Error generating vehicle report: {str(e)}")
        import traceback
        logger.error(f"[DOCX] Traceback: {traceback.format_exc()}")
        raise


def generate_user_data_docx(report: models.Report, user: Optional[models.User] = None) -> BytesIO:
    """
    Generate a formatted DOCX file from report and user data with comprehensive error handling.
    Routes to multi-property generator if report has multiple properties.

    Args:
        report: Report model instance
        user: Optional user model (uses report.user if not provided)

    Returns:
        BytesIO object containing the DOCX document

    Raises:
        ValueError: If report data is incomplete or malformed
        AttributeError: If required fields are missing
        TypeError: If data types are incorrect
    """
    # Use the user from the report relationship if not provided
    if user is None:
        user = report.user

    # Route to appropriate generator based on report type
    if hasattr(report, 'report_type') and report.report_type == 'vehicle':
        logger.info("[DOCX] Generating vehicle report")
        return generate_vehicle_report_docx(report, user)

    # Route to multi-property generator if applicable
    if hasattr(report, 'is_multi_property') and report.is_multi_property:
        logger.info(f"[DOCX] Generating multi-property report with {report.property_count} properties")
        return generate_multi_property_report_docx(report, user)

    # Continue with single-property generation
    logger.info("[DOCX] Generating single-property report")

    # Normalize report_type for legacy data
    if not report.report_type:
        report.report_type = 'residential_property'
        logger.warning(f"[DOCX] Report {report.id} missing report_type, defaulting to residential_property")

    # Validate report_type
    if report.report_type not in ['residential_property', 'bare_land', 'multi_property', 'vehicle']:
        logger.warning(f"[DOCX] Unknown report_type: {report.report_type}, treating as residential_property")
        report.report_type = 'residential_property'

    try:

        # Create a new Document
        doc = Document()

        # ===== LETTERHEAD =====
        # Use template system for letterhead rendering
        template_id = user.preferred_letterhead_template or 'classic'
        template = get_template(template_id)
        template.render_letterhead(doc, user, report)

        # ===== OPENING SECTION =====

        # Generate and add title block (centered)
        title_lines = generate_title_block(report)
        for i, line in enumerate(title_lines):
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)

            # Remove spacing between property description lines (lines 2 and 3)
            if i == 2 or i == 3:  # Property description lines
                para.paragraph_format.line_spacing = 1.0  # Single line spacing, no gap
            else:
                para.paragraph_format.line_spacing = 0.9

            run = para.add_run(line)
            if i == 0:  # "VALUATION REPORT"
                run.bold = True
                run.font.size = Pt(14)  # Updated from 11 to 14
                run.font.underline = True  # Add underline
            elif i == 2 or i == 3:  # Both property description lines
                run.bold = True
                run.font.size = FONT_SIZE_BODY
                run.font.underline = True  # Add underline for continuous block effect
            else:
                run.font.size = FONT_SIZE_BODY
                if i == 1:  # "of" line
                    run.font.underline = True  # Add underline for continuous block effect
            run.font.color.rgb = RGBColor(0, 0, 0)

        # Add spacing before applicant statements
        spacing_para1 = doc.add_paragraph()
        spacing_para1.paragraph_format.space_before = Pt(8)
        spacing_para1.paragraph_format.space_after = Pt(0)

        # Determine which introduction format to use based on request_type
        if report.request_type == 'organization_request':
            # Organization-side format
            org_intro_paragraphs = generate_organization_side_introduction(report)
            for statement in org_intro_paragraphs:
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
                para.paragraph_format.line_spacing = 0.9

                # Parse label and value for proper formatting
                if ":-" in statement:
                    # This is a label-value pair (Applicant/Address/Contact No)
                    label, value = statement.split(":-", 1)

                    # Add bold label
                    run_label = para.add_run(label.strip())
                    run_label.font.size = FONT_SIZE_BODY
                    run_label.font.bold = True
                    run_label.font.color.rgb = RGBColor(0, 0, 0)

                    # Add separator with proper spacing
                    run_sep = para.add_run(" : ")
                    run_sep.font.size = FONT_SIZE_BODY
                    run_sep.font.color.rgb = RGBColor(0, 0, 0)

                    # Add value
                    run_value = para.add_run(value)
                    run_value.font.size = FONT_SIZE_BODY
                    run_value.font.color.rgb = RGBColor(0, 0, 0)
                else:
                    # This is the introductory paragraph (paragraph 1)
                    run = para.add_run(statement)
                    run.font.size = FONT_SIZE_BODY
                    run.font.color.rgb = RGBColor(0, 0, 0)
        else:
            # Client-side format (default for backward compatibility)
            applicant_statements = generate_applicant_statement(report)
            for statement in applicant_statements:
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
                para.paragraph_format.line_spacing = 0.9
                run = para.add_run(statement)
                run.font.size = FONT_SIZE_BODY
                run.font.color.rgb = RGBColor(0, 0, 0)

            # Add deed/certificate description ONLY if not plan-based identification
            # (Plan-based reports don't show the deed sentence for cleaner appearance)
            should_show_deed_sentence = (
                report.property_identification_type in ['deed', 'certificate_of_sale']
                and report.has_deed_info == "yes"
                and report.deeds
            )

            # Backward compatibility: Old reports (NULL type) with deed data should show sentence
            if not report.property_identification_type and report.has_deed_info == "yes" and report.deeds:
                should_show_deed_sentence = True

            if should_show_deed_sentence:
                deed_text = generate_deed_description(report.deeds)
                if deed_text:
                    para = doc.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
                    para.paragraph_format.line_spacing = 0.9
                    run = para.add_run(deed_text)
                    run.font.size = FONT_SIZE_BODY
                    run.font.color.rgb = RGBColor(0, 0, 0)

            # Add submission statement
            submission_text = generate_submission_statement(report)
            if submission_text:
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
                para.paragraph_format.line_spacing = 0.9
                run = para.add_run(submission_text)
                run.font.size = FONT_SIZE_BODY
                run.font.color.rgb = RGBColor(0, 0, 0)

        # Add inspection date
        if report.inspection_date:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
            para.paragraph_format.line_spacing = 0.9

            # Bold label
            run_label = para.add_run("Date of Inspection:")
            run_label.font.size = FONT_SIZE_BODY
            run_label.font.bold = True
            run_label.font.color.rgb = RGBColor(0, 0, 0)

            # Regular date value
            run_date = para.add_run(f" {report.inspection_date}")
            run_date.font.size = FONT_SIZE_BODY
            run_date.font.color.rgb = RGBColor(0, 0, 0)

        # Add special note if applicable
        if report.has_special_note == "yes" and report.special_note_text:
            # Note label
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = SUBHEADING_SPACE_AFTER
            para.paragraph_format.line_spacing = 0.9
            run = para.add_run("Note:")
            run.bold = True
            run.font.size = FONT_SIZE_BODY
            run.font.color.rgb = RGBColor(0, 0, 0)

            # Note text
            note_para = doc.add_paragraph()
            note_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            note_para.paragraph_format.space_before = Pt(0)
            note_para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
            note_para.paragraph_format.line_spacing = 0.9
            note_run = note_para.add_run(report.special_note_text)
            note_run.font.size = FONT_SIZE_BODY
            note_run.font.color.rgb = RGBColor(0, 0, 0)

        # ===== 1.0 SITUATION SECTION (First after opening) =====
        situation_text = generate_situation_text(report)
        if situation_text:
            # Add numbered section heading
            add_section_heading(doc, "1.0", "SITUATION")

            # SITUATION text
            situation_para = doc.add_paragraph()
            situation_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            situation_para.paragraph_format.space_before = Pt(0)
            situation_para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
            situation_para.paragraph_format.line_spacing = 0.9
            situation_run = situation_para.add_run(situation_text)
            situation_run.font.size = FONT_SIZE_BODY
            situation_run.font.color.rgb = RGBColor(0, 0, 0)

        # ===== 2.0 ACCESS SECTION (After SITUATION) =====
        access_text = generate_access_text(report)
        locality_text = generate_locality_description(report)

        if access_text:
            # Add numbered section heading
            add_section_heading(doc, "2.0", "ACCESS")

            # ACCESS text
            access_para = doc.add_paragraph()
            access_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            access_para.paragraph_format.space_before = Pt(0)
            access_para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
            access_para.paragraph_format.line_spacing = 0.9
            access_run = access_para.add_run(access_text)
            access_run.font.size = FONT_SIZE_BODY
            access_run.font.color.rgb = RGBColor(0, 0, 0)

            # Add coordinates paragraph if available
            if report.property_latitude and report.property_longitude:
                # Add small spacing
                coord_para = doc.add_paragraph()
                coord_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                coord_para.paragraph_format.space_before = Pt(6)
                coord_para.paragraph_format.space_after = Pt(6)
                coord_para.paragraph_format.line_spacing = 0.9

                # Coordinate label (bold)
                coord_label = coord_para.add_run("Property Location Coordinates: ")
                coord_label.bold = True
                coord_label.font.size = FONT_SIZE_BODY
                coord_label.font.color.rgb = RGBColor(0, 0, 0)

                # Coordinate values (format to 6 decimal places)
                lat_value = float(report.property_latitude)
                lng_value = float(report.property_longitude)
                coord_text = coord_para.add_run(f"{lat_value:.6f}, {lng_value:.6f}")
                coord_text.font.size = FONT_SIZE_BODY
                coord_text.font.color.rgb = RGBColor(0, 0, 0)

            # Add map image if available (embedded within ACCESS section)
            if report.location_map_image_data:
                try:
                    # Fetch the image from URL
                    map_url = report.location_map_image_data
                    logger.info(f"[DOCX] Fetching map image from URL (length={len(map_url)})")
                    logger.debug(f"[DOCX] Map URL: {map_url[:200]}...")  # Log first 200 chars for debugging

                    response = requests.get(map_url, timeout=30)
                    if response.status_code == 200:
                        # Add spacing before map
                        map_spacing_para = doc.add_paragraph()
                        map_spacing_para.paragraph_format.space_before = IMAGE_SPACING_BEFORE
                        map_spacing_para.paragraph_format.space_after = Pt(0)

                        # Add map image
                        map_para = doc.add_paragraph()
                        map_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        map_para.paragraph_format.space_before = Pt(0)
                        map_para.paragraph_format.space_after = IMAGE_SPACING_AFTER

                        # Add image with calculated dimensions (maintains aspect ratio)
                        image_stream = BytesIO(response.content)
                        dimensions = calculate_image_dimensions(
                            image_stream,
                            MAP_IMAGE_WIDTH,
                            MAP_IMAGE_MAX_HEIGHT
                        )
                        map_para.add_run().add_picture(image_stream, **dimensions)

                        logger.info(f"[DOCX] Successfully added map image to document (size={len(response.content)} bytes)")
                    else:
                        logger.warning(f"[DOCX] Failed to fetch map image: HTTP {response.status_code}")
                        logger.debug(f"[DOCX] Response body: {response.text[:500] if response.text else 'empty'}")
                except Exception as e:
                    logger.error(f"[DOCX] Error adding map image: {str(e)}")
                    # Continue without map if error occurs
            else:
                logger.info(f"[DOCX] No location_map_image_data available for report {report.id}")

        # NOTE: LOCALITY section moved to 6.0 (after PHOTOGRAPHS)

        # ===== 3.0 IDENTIFICATION OF PROPERTY SECTION =====
        # This section appears if any property header data is available
        has_property_header_data = (
            report.land_traditional_name or
            report.land_extent_formatted or
            report.boundaries or
            report.physical_boundaries_types or
            report.physical_boundaries_description or
            report.boundaries_summary_text
        )

        if has_property_header_data:
            # Add numbered section heading
            add_section_heading(doc, "3.0", "IDENTIFICATION OF PROPERTY")

            # Traditional Land Name
            if report.land_traditional_name:
                name_para = doc.add_paragraph()
                name_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                name_para.paragraph_format.space_before = Pt(0)
                name_para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
                name_para.paragraph_format.line_spacing = 0.9
                name_label = name_para.add_run("Name of Land: ")
                name_label.bold = True
                name_label.font.size = FONT_SIZE_BODY
                name_label.font.color.rgb = RGBColor(0, 0, 0)
                name_value = name_para.add_run(f'"{report.land_traditional_name}"')
                name_value.font.size = FONT_SIZE_BODY
                name_value.font.color.rgb = RGBColor(0, 0, 0)

            # Survey Plan Information
            if report.lot_number and report.plan_number:
                plan_para = doc.add_paragraph()
                plan_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                plan_para.paragraph_format.space_before = Pt(0)
                plan_para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
                plan_para.paragraph_format.line_spacing = 0.9
                plan_label = plan_para.add_run("Survey Plan: ")
                plan_label.bold = True
                plan_label.font.size = FONT_SIZE_BODY
                plan_label.font.color.rgb = RGBColor(0, 0, 0)

                # Extract just the lot number/identifier (remove "Plan No" prefix if present)
                lot_desc = report.lot_number.strip() if report.lot_number else ''

                # Remove common prefixes that shouldn't be in lot description
                prefixes_to_remove = ['plan no', 'plan no:', 'lot plan no', 'lot plan no:']
                lot_desc_lower = lot_desc.lower()
                for prefix in prefixes_to_remove:
                    if lot_desc_lower.startswith(prefix):
                        lot_desc = lot_desc[len(prefix):].strip()
                        break

                # Ensure lot description has "Lot" prefix
                if not lot_desc.lower().startswith('lot'):
                    lot_desc = f"Lot {lot_desc}"

                # Format as "Lot X in Plan No: Y dated [date] made by [surveyor], Licensed Surveyor"
                plan_formatted = format_no_field("Plan", report.plan_number)
                plan_text = f"{lot_desc} in {plan_formatted}"
                if report.plan_date:
                    plan_text += f" dated {report.plan_date}"
                if report.licensed_surveyor_name:
                    plan_text += f" made by {report.licensed_surveyor_name}, Licensed Surveyor"
                plan_text += "."

                plan_value = plan_para.add_run(plan_text)
                plan_value.font.size = FONT_SIZE_BODY
                plan_value.font.color.rgb = RGBColor(0, 0, 0)

            # Land Extent
            if report.land_extent_formatted:
                extent_para = doc.add_paragraph()
                extent_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                extent_para.paragraph_format.space_before = Pt(0)
                extent_para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
                extent_para.paragraph_format.line_spacing = 0.9
                extent_label = extent_para.add_run("Extent: ")
                extent_label.bold = True
                extent_label.font.size = FONT_SIZE_BODY
                extent_label.font.color.rgb = RGBColor(0, 0, 0)

                extent_text = report.land_extent_formatted
                if report.land_extent_hectares:
                    extent_text += f" [{report.land_extent_hectares:.4f} Hectares]"
                if report.land_extent_square_meters:
                    extent_text += f" [{report.land_extent_square_meters:.2f} m²]"

                extent_value = extent_para.add_run(extent_text)
                extent_value.font.size = FONT_SIZE_BODY
                extent_value.font.color.rgb = RGBColor(0, 0, 0)

            # Boundaries - Enhanced Format with all details
            if report.boundaries:
                # Add spacing before boundaries
                boundaries_spacing = doc.add_paragraph()
                boundaries_spacing.paragraph_format.space_before = SUBHEADING_SPACE_BEFORE
                boundaries_spacing.paragraph_format.space_after = SUBHEADING_SPACE_AFTER

                # Boundaries subheading
                boundaries_heading = doc.add_paragraph()
                boundaries_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                boundaries_heading.paragraph_format.space_before = Pt(0)
                boundaries_heading.paragraph_format.space_after = SUBHEADING_SPACE_AFTER
                boundaries_heading.paragraph_format.line_spacing = 0.9
                boundaries_heading_run = boundaries_heading.add_run("Boundaries:")
                boundaries_heading_run.bold = True
                boundaries_heading_run.font.size = FONT_SIZE_BODY
                boundaries_heading_run.font.color.rgb = RGBColor(0, 0, 0)

                # Physical boundary type labels mapping
                boundary_type_labels = {
                    'brick_walls': 'Brick Walls',
                    'barbed_wire': 'Barbed Wire Fence',
                    'live_fence': 'Live Fence',
                    'concrete_posts': 'Concrete Posts',
                    'iron_gate': 'Iron Gate',
                    'rubble_foundation': 'Rubble Foundation',
                    'chain_link': 'Chain Link Fence',
                    'wooden_fence': 'Wooden Fence',
                    'stone_wall': 'Stone Wall',
                    'hedge': 'Hedge',
                    'none': 'None'
                }

                # Enhanced format for each boundary direction (4 main + 4 optional diagonal)
                directions = ['north', 'northeast', 'east', 'southeast', 'south', 'southwest', 'west', 'northwest']
                direction_labels = ['North', 'North-East', 'East', 'South-East', 'South', 'South-West', 'West', 'North-West']

                # SAFE: Get boundaries with None check
                boundaries = safe_get_json_field(report, 'boundaries', {})
                if not boundaries:
                    boundaries = {}

                for direction, label in zip(directions, direction_labels):
                    boundary_data = boundaries.get(direction, {}) if isinstance(boundaries, dict) else {}

                    # Skip diagonal boundaries if they have no description
                    if direction in ['northeast', 'southeast', 'southwest', 'northwest']:
                        if not boundary_data or not boundary_data.get('description'):
                            continue

                    # Main description line: "North  : Lot 7" or "North-East: Lot 7"
                    boundary_line = f"{label:<11} : "
                    boundary_text = boundary_data.get('description') or 'Not specified'

                    # Create paragraph for main boundary description
                    boundary_para = doc.add_paragraph()
                    boundary_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    boundary_para.paragraph_format.space_before = Pt(0)
                    boundary_para.paragraph_format.space_after = Pt(1)
                    boundary_para.paragraph_format.line_spacing = 0.9
                    boundary_para.paragraph_format.left_indent = Inches(0.5)

                    boundary_run = boundary_para.add_run(boundary_line + boundary_text)
                    boundary_run.font.size = FONT_SIZE_BODY
                    boundary_run.font.color.rgb = RGBColor(0, 0, 0)

                    # Additional boundary details (if any)
                    additional_details = []

                    # Physical Boundary Type (from boundary_types_per_direction)
                    boundary_types_per_dir = report.boundary_types_per_direction or {}
                    physical_type = boundary_types_per_dir.get(direction)
                    if physical_type:
                        type_label = boundary_type_labels.get(physical_type, physical_type.replace('_', ' ').title())
                        additional_details.append(f"Type: {type_label}")

                    # Length (optional)
                    length = boundary_data.get('length')
                    if length:
                        additional_details.append(f"Length: {length}")

                    # Adjoins (optional)
                    adjoins = boundary_data.get('adjoins')
                    if adjoins:
                        additional_details.append(f"Adjoins: {adjoins}")

                    # Additional Notes (optional)
                    notes = boundary_data.get('notes')
                    if notes:
                        additional_details.append(f"Notes: {notes}")

                    # Render additional details as sub-items if any exist
                    if additional_details:
                        for detail in additional_details:
                            detail_para = doc.add_paragraph()
                            detail_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            detail_para.paragraph_format.space_before = Pt(0)
                            detail_para.paragraph_format.space_after = Pt(1)
                            detail_para.paragraph_format.line_spacing = 0.9
                            detail_para.paragraph_format.left_indent = Inches(0.75)  # Extra indent for sub-details

                            detail_run = detail_para.add_run(f"  • {detail}")
                            detail_run.font.size = Pt(8)
                            detail_run.font.color.rgb = RGBColor(80, 80, 80)  # Slightly gray for secondary info

                    # Add small spacing after each boundary direction
                    if additional_details:
                        spacing_para = doc.add_paragraph()
                        spacing_para.paragraph_format.space_before = Pt(0)
                        spacing_para.paragraph_format.space_after = Pt(2)

            # Boundary Summary Sentence (Professional description)
            # Use stored text or auto-generate from boundary types per direction
            boundary_summary = report.boundaries_summary_text
            if not boundary_summary and (report.boundary_types_per_direction or report.physical_boundaries_types):
                boundary_summary = generate_boundary_summary_text(report)

            if boundary_summary:
                summary_para = doc.add_paragraph()
                summary_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                summary_para.paragraph_format.space_before = Pt(6)
                summary_para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
                summary_para.paragraph_format.line_spacing = 0.9
                summary_run = summary_para.add_run(boundary_summary)
                summary_run.font.size = FONT_SIZE_BODY
                summary_run.font.color.rgb = RGBColor(0, 0, 0)


        # ===== 4.0 DESCRIPTION OF PROPERTY SECTION =====
        # Check if there's any property description data
        # Split into land data vs building data for better bare land support
        has_land_data = (
            report.land_description_text or
            report.land_shape or
            report.soil_type or
            report.land_type or
            report.water_table_depth or
            report.flood_risk or
            report.ongoing_construction_notes or
            # Topographical features for bare land
            report.elevation_changes or
            report.drainage_pattern or
            report.vegetation_type or
            report.natural_features
        )

        has_building_data = (
            report.buildings and
            isinstance(report.buildings, list) and
            len(report.buildings) > 0 and
            report.report_type != 'bare_land'
        )

        has_occupier_data = report.occupier_name

        # Show section 4.0 if there's ANY property description data
        if has_land_data or has_building_data or has_occupier_data:
            # Add numbered section heading
            add_section_heading(doc, "4.0", "DESCRIPTION OF PROPERTY")

            # === LAND DESCRIPTION ===
            # Use auto-generated text if available, otherwise build from individual fields
            if report.land_description_text:
                land_para = doc.add_paragraph()
                land_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                land_para.paragraph_format.space_before = Pt(0)
                land_para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
                land_para.paragraph_format.line_spacing = 0.9
                land_run = land_para.add_run(report.land_description_text)
                land_run.font.size = FONT_SIZE_BODY
                land_run.font.color.rgb = RGBColor(0, 0, 0)
            else:
                # Build land description from individual fields
                land_parts = []

                # Shape and type
                if report.land_shape:
                    shape_labels = {
                        'rectangular': 'rectangular', 'square': 'square', 'triangular': 'triangular',
                        'irregular': 'irregular', 'L_shaped': 'L-shaped', 'trapezoidal': 'trapezoidal'
                    }
                    land_parts.append(f"The land is {shape_labels.get(report.land_shape, report.land_shape)} in shape")

                if report.land_type:
                    type_labels = {
                        'flat': 'flat terrain', 'sloped': 'sloped terrain', 'hilly': 'hilly terrain',
                        'waterfront': 'waterfront', 'paddy': 'paddy land', 'garden': 'garden land'
                    }
                    if land_parts:
                        land_parts.append(f"with {type_labels.get(report.land_type, report.land_type)}")
                    else:
                        land_parts.append(f"The land has {type_labels.get(report.land_type, report.land_type)}")

                # Land level
                if report.land_level:
                    level_labels = {
                        'above_road': 'above road level', 'at_road': 'at road level',
                        'below_road': 'below road level', 'varied': 'varied levels'
                    }
                    land_parts.append(f"and is {level_labels.get(report.land_level, report.land_level)}")
                    if report.land_level_difference:
                        land_parts.append(f"by approximately {report.land_level_difference} feet")

                # Soil type
                if report.soil_type:
                    soil_labels = {
                        'sandy': 'sandy', 'clay': 'clay', 'loam': 'loam', 'laterite': 'laterite',
                        'rocky': 'rocky', 'gravel': 'gravel', 'alluvial': 'alluvial'
                    }
                    land_parts.append(f". The soil is {soil_labels.get(report.soil_type, report.soil_type)}")

                # Water table
                if report.water_table_depth:
                    land_parts.append(f"with water table at approximately {report.water_table_depth} feet depth")

                # Risks
                risk_parts = []
                if report.flood_risk and report.flood_risk != 'none':
                    risk_labels = {'low': 'low', 'moderate': 'moderate', 'high': 'high', 'very_high': 'very high'}
                    risk_parts.append(f"flood risk is {risk_labels.get(report.flood_risk, report.flood_risk)}")
                if report.inundation_risk and report.inundation_risk != 'none':
                    risk_labels = {'low': 'low', 'moderate': 'moderate', 'high': 'high', 'very_high': 'very high'}
                    risk_parts.append(f"inundation risk is {risk_labels.get(report.inundation_risk, report.inundation_risk)}")
                if report.earth_slip_risk and report.earth_slip_risk != 'none':
                    risk_labels = {'low': 'low', 'moderate': 'moderate', 'high': 'high', 'very_high': 'very high'}
                    risk_parts.append(f"earth slip risk is {risk_labels.get(report.earth_slip_risk, report.earth_slip_risk)}")

                if risk_parts:
                    land_parts.append(". The " + ", ".join(risk_parts))

                # Land condition
                if report.land_condition:
                    condition_labels = {
                        'well_maintained': 'well maintained', 'developed': 'developed',
                        'undeveloped': 'undeveloped', 'overgrown': 'overgrown', 'cleared': 'cleared'
                    }
                    land_parts.append(f". The land is {condition_labels.get(report.land_condition, report.land_condition)}")

                if land_parts:
                    land_text = " ".join(land_parts)
                    if not land_text.endswith("."):
                        land_text += "."

                    land_para = doc.add_paragraph()
                    land_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    land_para.paragraph_format.space_before = Pt(0)
                    land_para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
                    land_para.paragraph_format.line_spacing = 0.9
                    land_run = land_para.add_run(land_text)
                    land_run.font.size = FONT_SIZE_BODY
                    land_run.font.color.rgb = RGBColor(0, 0, 0)

            # === DEVELOPMENT FEASIBILITY / ONGOING CONSTRUCTION (Bare Land Reports) ===
            if report.ongoing_construction_notes and report.report_type == 'bare_land':
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
                p.paragraph_format.line_spacing = 0.9

                run = p.add_run("Development Feasibility: ")
                run.font.bold = True
                run.font.size = FONT_SIZE_BODY
                run.font.color.rgb = RGBColor(0, 0, 0)

                run = p.add_run(report.ongoing_construction_notes)
                run.font.size = FONT_SIZE_BODY
                run.font.color.rgb = RGBColor(0, 0, 0)

            # === TOPOGRAPHICAL FEATURES (Bare Land specific) ===
            if report.report_type == 'bare_land':
                topo_parts = []

                if report.elevation_changes:
                    topo_parts.append(f"elevation changes are {report.elevation_changes.replace('_', ' ')}")
                if report.drainage_pattern:
                    topo_parts.append(f"drainage pattern is {report.drainage_pattern.replace('_', ' ')}")
                if report.vegetation_type:
                    topo_parts.append(f"vegetation type is {report.vegetation_type.replace('_', ' ')}")
                if report.natural_features:
                    topo_parts.append(f"natural features include {report.natural_features}")

                if topo_parts:
                    topo_text = "The land has " + ", ".join(topo_parts) + "."
                    topo_para = doc.add_paragraph()
                    topo_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    topo_para.paragraph_format.space_before = Pt(4)
                    topo_para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
                    topo_para.paragraph_format.line_spacing = 0.9

                    topo_label = topo_para.add_run("Topographical Features: ")
                    topo_label.font.bold = True
                    topo_label.font.size = FONT_SIZE_BODY

                    topo_value = topo_para.add_run(topo_text)
                    topo_value.font.size = FONT_SIZE_BODY

            # === PROPERTY PHOTOS (Bare Land - embedded in Section 4.0) ===
            if report.report_type == 'bare_land' and report.property_photos:
                property_photos = safe_get_json_field(report, 'property_photos', [])

                if property_photos and len(property_photos) > 0:
                    # Sort photos by order
                    sorted_photos = sorted(property_photos, key=lambda x: x.get('order', 0))

                    # Import modules needed for photo processing
                    import base64
                    import re

                    num_photos = len(sorted_photos)
                    photos_per_row = 3

                    # Process photos in rows with same logic as building photos
                    idx = 0
                    while idx < num_photos:
                        # Determine how many photos in this row
                        remaining = num_photos - idx
                        if remaining >= photos_per_row:
                            photos_in_row = photos_per_row
                        elif remaining == 1 and idx > 0:
                            # Last single photo - center it
                            photos_in_row = 1
                        else:
                            # 2 photos or first row - use all remaining
                            photos_in_row = remaining

                        # Create a table for this row of photos with captions
                        table = doc.add_table(rows=2, cols=photos_in_row)
                        table.alignment = WD_TABLE_ALIGNMENT.CENTER

                        # Remove table borders for clean look
                        for row in table.rows:
                            for cell in row.cells:
                                cell.width = Inches(6.5 / photos_in_row)
                                # Remove all borders
                                tc = cell._element
                                tcPr = tc.get_or_add_tcPr()
                                tcBorders = parse_xml(
                                    r'<w:tcBorders %s>'
                                    r'<w:top w:val="none"/>'
                                    r'<w:left w:val="none"/>'
                                    r'<w:bottom w:val="none"/>'
                                    r'<w:right w:val="none"/>'
                                    r'</w:tcBorders>' % nsdecls('w')
                                )
                                tcPr.append(tcBorders)

                        # Add photos to first row of table
                        for i in range(photos_in_row):
                            if idx >= num_photos:
                                break

                            photo = sorted_photos[idx]
                            try:
                                image_data = photo.get('image_data', '')
                                caption = photo.get('caption', '')

                                # Handle both data URI and raw base64 formats
                                if image_data:
                                    # Check if it's a data URI
                                    if image_data.startswith('data:image'):
                                        # Extract base64 from data URI
                                        base64_match = re.search(r'base64,(.+)', image_data)
                                        if not base64_match:
                                            idx += 1
                                            continue
                                        base64_data = base64_match.group(1)
                                    else:
                                        # Assume it's raw base64
                                        base64_data = image_data

                                    # Decode image
                                    image_bytes = base64.b64decode(base64_data)
                                    image_stream = BytesIO(image_bytes)

                                    # Use uniform 2-inch width
                                    img_width = Inches(2.0)
                                    dimensions = calculate_image_dimensions(
                                        image_stream,
                                        img_width,
                                        PROPERTY_PHOTO_HEIGHT
                                    )

                                    # Add image to table cell
                                    image_stream.seek(0)
                                    cell = table.rows[0].cells[i]
                                    cell_para = cell.paragraphs[0]
                                    cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    cell_para.add_run().add_picture(image_stream, **dimensions)

                                    # Add caption to second row with proper styling
                                    caption_cell = table.rows[1].cells[i]
                                    caption_para = caption_cell.paragraphs[0]
                                    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    caption_para.paragraph_format.space_before = Pt(2)
                                    caption_para.paragraph_format.space_after = Pt(2)

                                    caption_text = f"Fig. {idx + 1}"
                                    if caption:
                                        caption_text += f": {caption}"

                                    caption_run = caption_para.add_run(caption_text)
                                    caption_run.font.size = FONT_SIZE_CAPTION
                                    caption_run.font.italic = True
                                    caption_run.font.color.rgb = RGBColor(60, 60, 60)

                                    logger.info(f"[DOCX] Added property photo {idx + 1}")

                            except Exception as e:
                                logger.error(f"[DOCX] Error adding property photo {idx + 1}: {str(e)}")

                            idx += 1

                        # Add spacing after photo table
                        spacing_para = doc.add_paragraph()
                        spacing_para.paragraph_format.space_after = Pt(8)

                    # Add final spacing after all photos
                    final_spacing = doc.add_paragraph()
                    final_spacing.paragraph_format.space_after = IMAGE_SPACING_AFTER

            # === BUILDING DETAILS (Direct numbering: 4.1, 4.2, 4.3) ===
            # Use has_building_data to avoid redundant checks and respect report_type
            if has_building_data:
                buildings = safe_get_json_field(report, 'buildings', [])
                for idx, building in enumerate(buildings):
                    building_number = f"4.{idx + 1}"
                    building_name = building.get('building_name', f'Building {idx + 1}')

                    # Add building subsection heading
                    add_section_heading(doc, building_number, building_name)

                    # === CONSTRUCTION DETAILS (STANDALONE PARAGRAPH - NO LABEL) ===
                    # This comes FIRST, directly under the building heading
                    render_construction_details(doc, building)

                    # === PROFESSIONAL STRUCTURED FORMAT (All buildings) ===

                    # Use custom description text if provided
                    if building.get('building_description_text'):
                        opening_para = doc.add_paragraph()
                        opening_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        opening_para.paragraph_format.space_before = Pt(0)
                        opening_para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
                        opening_para.paragraph_format.line_spacing = 0.9
                        opening_run = opening_para.add_run(building.get('building_description_text'))
                        opening_run.font.size = FONT_SIZE_BODY
                        opening_run.font.color.rgb = RGBColor(0, 0, 0)
                    else:
                        # Build opening description from materials
                        roof_types = building.get('roof_types', [])
                        wall_types = building.get('wall_types', [])
                        floor_types = building.get('floor_types', [])

                        if roof_types or wall_types or floor_types:
                            desc_text = "This is constructed with "

                            # Roof
                            if roof_types:
                                roof_labels = {
                                    'asbestos': 'asbestos sheets', 'tile': 'tiles',
                                    'metal': 'metal sheets', 'concrete': 'concrete flat',
                                    'cadjan': 'cadjans', 'zinc': 'zinc sheets'
                                }
                                roof_text = format_material_list(roof_types, roof_labels)
                                desc_text += f"{roof_text} on "

                            # Roof structure
                            stories = building.get('stories', 1)
                            if stories == 1:
                                desc_text += "timber frame roof "
                            else:
                                desc_text += f"{stories} storey structure "

                            # Walls
                            if wall_types:
                                wall_labels = {
                                    'brick': 'brick masonry', 'cement_block': 'cement block',
                                    'stone': 'stone', 'mud': 'mud', 'timber': 'timber',
                                    'cadjan': 'cadjan', 'rcc_columns': 'RCC columns',
                                    'rubble_masonry': 'rubble masonry'
                                }
                                wall_text = format_material_list(wall_types, wall_labels)
                                desc_text += f"supported by {wall_text} walls "

                            # Foundation/floor
                            if floor_types:
                                floor_labels = {
                                    'cement': 'cement', 'tile': 'tiles', 'terrazzo': 'terrazzo',
                                    'timber': 'timber', 'granite': 'granite', 'earth': 'earth',
                                    'concrete': 'concrete'
                                }
                                floor_text = format_material_list(floor_types, floor_labels)
                                desc_text += f"and the floor is {floor_text} rendered"

                            desc_text += "."

                            opening_para = doc.add_paragraph()
                            opening_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            opening_para.paragraph_format.space_before = Pt(0)
                            opening_para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
                            opening_para.paragraph_format.line_spacing = 0.9
                            opening_run = opening_para.add_run(desc_text)
                            opening_run.font.size = FONT_SIZE_BODY
                            opening_run.font.color.rgb = RGBColor(0, 0, 0)

                    # === ACCOMMODATION ===
                    # Check if building has accommodation data (new structure at building level or old structure at floor level)
                    has_new_building_format = building.get('accommodation_summary') is not None
                    floors = building.get('floors', [])
                    has_old_floor_format = any(floor.get('accommodation_summary') for floor in floors)

                    if has_new_building_format or has_old_floor_format:
                        # Get accommodation summary (supports both new and old structures)
                        aggregated_rooms = aggregate_accommodation_across_building(building)

                        room_parts = []

                        # Build room counts list from aggregated data using format_room_count()
                        if aggregated_rooms.get('bedrooms', 0) > 0:
                            count = aggregated_rooms['bedrooms']
                            room_parts.append(format_room_count(count, 'bedroom'))

                        if aggregated_rooms.get('bathrooms', 0) > 0:
                            count = aggregated_rooms['bathrooms']
                            room_parts.append(format_room_count(count, 'bathroom'))

                        if aggregated_rooms.get('living_rooms', 0) > 0:
                            count = aggregated_rooms['living_rooms']
                            room_parts.append(format_room_count(count, 'living room'))

                        if aggregated_rooms.get('dining_rooms', 0) > 0:
                            count = aggregated_rooms['dining_rooms']
                            room_parts.append(format_room_count(count, 'dining room'))

                        if aggregated_rooms.get('kitchens', 0) > 0:
                            count = aggregated_rooms['kitchens']
                            room_parts.append(format_room_count(count, 'kitchen'))

                        if aggregated_rooms.get('pantries', 0) > 0:
                            count = aggregated_rooms['pantries']
                            room_parts.append(format_room_count(count, 'pantry', 'pantries'))

                        if aggregated_rooms.get('verandahs', 0) > 0:
                            count = aggregated_rooms['verandahs']
                            room_parts.append(format_room_count(count, 'verandah'))

                        if aggregated_rooms.get('balconies', 0) > 0:
                            count = aggregated_rooms['balconies']
                            room_parts.append(format_room_count(count, 'balcony', 'balconies'))

                        if aggregated_rooms.get('garages', 0) > 0:
                            count = aggregated_rooms['garages']
                            room_parts.append(format_room_count(count, 'garage'))

                        if aggregated_rooms.get('store_rooms', 0) > 0:
                            count = aggregated_rooms['store_rooms']
                            room_parts.append(format_room_count(count, 'store room'))

                        if aggregated_rooms.get('other_rooms', 0) > 0:
                            count = aggregated_rooms['other_rooms']
                            room_parts.append(format_room_count(count, 'other room'))

                        if room_parts:
                            # Create single-line accommodation text
                            accommodation_text = f"The property comprises {format_list_with_grammar(room_parts)}."
                            add_inline_field(doc, "Accommodation", accommodation_text)


                    # === FLOOR AREA (CONDITIONAL FORMAT) ===
                    floors = building.get('floors', [])
                    total_area = building.get('total_floor_area', 0)
                    if floors and (total_area or any(f.get('floor_area', 0) > 0 for f in floors)):
                        # Filter floors with area > 0
                        floors_with_area = [f for f in floors if f.get('floor_area', 0) > 0]

                        if len(floors_with_area) == 1:
                            # SINGLE FLOOR: Use inline format
                            floor = floors_with_area[0]
                            floor_area = floor.get('floor_area', 0)
                            area_text = f"{floor_area:,.0f} square feet"
                            add_inline_field(doc, "Floor area", area_text)
                        else:
                            # MULTIPLE FLOORS: Use table format
                            # Add "Floor Area" label (bold, inline)
                            floor_area_para = doc.add_paragraph()
                            floor_area_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            floor_area_para.paragraph_format.space_before = INLINE_FIELD_SPACE_BEFORE
                            floor_area_para.paragraph_format.space_after = Pt(2)
                            floor_area_para.paragraph_format.line_spacing = 0.9

                            label_run = floor_area_para.add_run("Floor area:")
                            label_run.bold = True
                            label_run.font.size = FONT_SIZE_BODY
                            label_run.font.color.rgb = RGBColor(0, 0, 0)

                            # Add each floor's area (indented with aligned columns)
                            for floor in floors:
                                floor_name = floor.get('floor_name', 'Floor')
                                floor_area = floor.get('floor_area', 0)

                                if floor_area and floor_area > 0:
                                    floor_line_para = doc.add_paragraph()
                                    floor_line_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                    floor_line_para.paragraph_format.space_before = Pt(0)
                                    floor_line_para.paragraph_format.space_after = Pt(1)
                                    floor_line_para.paragraph_format.line_spacing = 0.9
                                    floor_line_para.paragraph_format.left_indent = Inches(0.5)

                                    # Use fixed-width padding for column alignment
                                    floor_name_padded = f"{floor_name:<25}"
                                    floor_line_run = floor_line_para.add_run(f"{floor_name_padded}{floor_area:>10,.0f} square feet")
                                    floor_line_run.font.size = FONT_SIZE_BODY
                                    floor_line_run.font.color.rgb = RGBColor(0, 0, 0)

                            # Add total (indented with aligned columns)
                            if total_area and total_area > 0:
                                total_para = doc.add_paragraph()
                                total_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                total_para.paragraph_format.space_before = Pt(0)
                                total_para.paragraph_format.space_after = INLINE_FIELD_SPACE_AFTER
                                total_para.paragraph_format.line_spacing = 0.9
                                total_para.paragraph_format.left_indent = Inches(0.5)

                                # Use fixed-width padding for column alignment
                                total_padded = f"{'Total':<25}"
                                total_run = total_para.add_run(f"{total_padded}{total_area:>10,.0f} square feet")
                                total_run.bold = True
                                total_run.font.size = FONT_SIZE_BODY
                                total_run.font.color.rgb = RGBColor(0, 0, 0)

                    # === AGE AND CONDITION (INLINE FORMAT) ===
                    building_age = building.get('building_age')
                    condition = building.get('condition', '')
                    if building_age or condition:
                        condition_labels = {
                            'excellent': 'excellent',
                            'good': 'good',
                            'fair': 'fair',
                            'poor': 'poor',
                            'dilapidated': 'dilapidated'
                        }

                        parts = []
                        if building_age is not None and building_age > 0:
                            year_word = "year" if building_age == 1 else "years"
                            parts.append(f"{building_age} {year_word} old")
                        if condition:
                            parts.append(f"condition is {condition_labels.get(condition, condition)}")

                        age_text = "; ".join(parts) + "." if parts else "Information not provided."

                        add_inline_field(doc, "Age and condition", age_text)

                    # === UTILITIES AND CONVENIENCES ===
                    render_utilities_and_conveniences(doc, building)

                    # === OCCUPATION (BUILDING-LEVEL) ===
                    # Primary: Use building-level occupier data
                    if building.get('occupier_name'):
                        relationship = building.get('occupier_relationship')

                        # Special case for vacant buildings
                        if relationship == 'vacant':
                            occupier_text = "The building is currently vacant."
                        else:
                            occupier_text = f"The building is occupied by {building.get('occupier_name')}"

                            if relationship:
                                rel_labels = {
                                    'owner': 'the owner',
                                    'tenant': 'a tenant',
                                    'family_member': 'a family member',
                                    'caretaker': 'caretaker'
                                }
                                occupier_text += f" who is {rel_labels.get(relationship, relationship)}."
                            else:
                                occupier_text += "."

                        add_inline_field(doc, "Occupation", occupier_text, space_after=Pt(6))

                    # Fallback: Support old reports with property-level occupier
                    elif report.occupier_name:
                        occupier_text = f"The property is occupied by {report.occupier_name}"
                        if report.occupier_relationship:
                            rel_labels = {
                                'owner': 'the owner',
                                'tenant': 'a tenant',
                                'family_member': 'a family member',
                                'caretaker': 'caretaker'
                            }
                            occupier_text += f" who is {rel_labels.get(report.occupier_relationship, report.occupier_relationship)}."
                        else:
                            occupier_text += "."

                        add_inline_field(doc, "Occupation", occupier_text, space_after=Pt(6))

                    # === BUILDING PHOTOS (3-column grid layout - NO SUBHEADING) ===
                    building_photos = building.get('building_photos', [])
                    if building_photos and len(building_photos) > 0:
                        # Sort photos by order
                        sorted_photos = sorted(building_photos, key=lambda x: x.get('order', 0))

                        # Modern flexible photo grid layout using tables for proper caption alignment
                        import base64
                        import re


                        num_photos = len(sorted_photos)
                        photos_per_row = 3

                        # Process photos in rows
                        idx = 0
                        while idx < num_photos:
                            # Determine how many photos in this row
                            remaining = num_photos - idx
                            if remaining >= photos_per_row:
                                photos_in_row = photos_per_row
                            elif remaining == 1 and idx > 0:
                                # Last single photo - center it
                                photos_in_row = 1
                            else:
                                # 2 photos or first row - use all remaining
                                photos_in_row = remaining

                            # Create a table for this row of photos with captions
                            table = doc.add_table(rows=2, cols=photos_in_row)
                            table.alignment = WD_TABLE_ALIGNMENT.CENTER

                            # Remove table borders
                            for row in table.rows:
                                for cell in row.cells:
                                    cell.width = Inches(6.5 / photos_in_row)
                                    # Remove all borders
                                    tc = cell._element
                                    tcPr = tc.get_or_add_tcPr()
                                    tcBorders = parse_xml(
                                        r'<w:tcBorders %s>'
                                        r'<w:top w:val="none"/>'
                                        r'<w:left w:val="none"/>'
                                        r'<w:bottom w:val="none"/>'
                                        r'<w:right w:val="none"/>'
                                        r'</w:tcBorders>' % nsdecls('w')
                                    )
                                    tcPr.append(tcBorders)

                            # Add photos to first row of table
                            for i in range(photos_in_row):
                                if idx >= num_photos:
                                    break

                                photo = sorted_photos[idx]
                                try:
                                    image_data = photo.get('image_data', '')
                                    caption = photo.get('caption', '')

                                    if not image_data or not image_data.startswith('data:image'):
                                        idx += 1
                                        continue

                                    # Decode base64 image
                                    base64_match = re.search(r'base64,(.+)', image_data)
                                    if not base64_match:
                                        idx += 1
                                        continue

                                    base64_data = base64_match.group(1)
                                    image_bytes = base64.b64decode(base64_data)
                                    image_stream = BytesIO(image_bytes)

                                    # Use uniform 2-inch width for all photos
                                    img_width = Inches(2.0)

                                    dimensions = calculate_image_dimensions(
                                        image_stream,
                                        img_width,
                                        PROPERTY_PHOTO_HEIGHT
                                    )

                                    # Add image to table cell
                                    image_stream.seek(0)
                                    cell = table.rows[0].cells[i]
                                    cell_para = cell.paragraphs[0]
                                    cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    cell_para.add_run().add_picture(image_stream, **dimensions)

                                    # Add caption to second row
                                    caption_cell = table.rows[1].cells[i]
                                    caption_para = caption_cell.paragraphs[0]
                                    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    caption_para.paragraph_format.space_before = Pt(2)
                                    caption_para.paragraph_format.space_after = Pt(2)

                                    caption_text = f"Fig. {idx + 1}"
                                    if caption:
                                        caption_text += f": {caption}"

                                    caption_run = caption_para.add_run(caption_text)
                                    caption_run.font.size = FONT_SIZE_CAPTION
                                    caption_run.font.italic = True
                                    caption_run.font.color.rgb = RGBColor(60, 60, 60)

                                    print(f"[DOCX] Added building photo {idx + 1} for {building_name}")

                                except Exception as e:
                                    print(f"[DOCX] Error adding photo {idx + 1}: {str(e)}")

                                idx += 1

                            # Add spacing after photo table
                            spacing_para = doc.add_paragraph()
                            spacing_para.paragraph_format.space_after = Pt(8)

                        # Add final spacing after all photos
                        final_spacing = doc.add_paragraph()
                        final_spacing.paragraph_format.space_after = IMAGE_SPACING_AFTER

                    # === ADDITIONAL STRUCTURES ===
                    additional_structures = building.get('additional_structures_description', '')
                    if additional_structures and additional_structures.strip():
                        # Add subheading
                        add_structures_heading = doc.add_paragraph()
                        add_structures_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        add_structures_heading.paragraph_format.space_before = SUBHEADING_SPACE_BEFORE
                        add_structures_heading.paragraph_format.space_after = SUBHEADING_SPACE_AFTER
                        add_structures_heading.paragraph_format.line_spacing = 0.9
                        add_structures_run = add_structures_heading.add_run("Additional Structures")
                        add_structures_run.bold = True
                        add_structures_run.font.size = FONT_SIZE_BODY
                        add_structures_run.font.color.rgb = RGBColor(0, 0, 0)

                        # Add description paragraph
                        structures_para = doc.add_paragraph()
                        structures_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        structures_para.paragraph_format.space_before = Pt(0)
                        structures_para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
                        structures_para.paragraph_format.line_spacing = 0.9
                        structures_para.paragraph_format.left_indent = INDENTED_CONTENT_LEFT_INDENT

                        structures_run = structures_para.add_run(additional_structures.strip())
                        structures_run.font.size = FONT_SIZE_BODY
                        structures_run.font.color.rgb = RGBColor(0, 0, 0)

        # ===== 5.0 LOCALITY SECTION (Moved to last position) =====
        if locality_text:
            # Add numbered section heading
            add_section_heading(doc, "5.0", "LOCALITY")

            # LOCALITY text (now a single concise paragraph from updated AI prompt)
            locality_para = doc.add_paragraph()
            locality_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            locality_para.paragraph_format.space_before = Pt(0)
            locality_para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
            locality_para.paragraph_format.line_spacing = 0.9
            locality_run = locality_para.add_run(locality_text)
            locality_run.font.size = FONT_SIZE_BODY
            locality_run.font.color.rgb = RGBColor(0, 0, 0)

        # ===== 6.0 LEGAL ASPECTS =====
        # Show section if any legal aspect data exists (including new extended fields)
        has_legal_data = (
            report.ownership_type or report.street_lines_status or report.building_limits_status or
            report.local_authority_data or report.rent_act_effectiveness or
            # Extended fields
            report.title_search_conducted or report.property_encumbered or
            report.street_lines_gazette_ref or report.building_plan_approved or
            report.local_authority_rated
        )

        if has_legal_data:
            add_section_heading(doc, "6.0", "LEGAL ASPECTS")

            # (a) Ownership - Generate professional paragraph
            if report.ownership_type or hasattr(report, 'deeds') or hasattr(report, 'plan_number'):
                ownership_para = generate_ownership_paragraph(report)
                add_subsection_paragraph(doc, "(a)", "Ownership", ownership_para)

            # (b) Street lines - Generate contextual paragraph (SKIP for bare_land)
            if report.street_lines_status and report.report_type != 'bare_land':
                street_para = generate_street_lines_paragraph(report)
                add_subsection_paragraph(doc, "(b)", "Street lines", street_para)

            # (c) Building limits - Generate detailed paragraph (SKIP for bare_land)
            if report.building_limits_status and report.report_type != 'bare_land':
                building_para = generate_building_limits_paragraph(report)
                add_subsection_paragraph(doc, "(c)", "Building limits", building_para)

            # (d) Local authority data - Generate administrative paragraph
            if report.local_authority_data or report.pradeshiya_sabha or report.local_authority_rated:
                authority_para = generate_local_authority_paragraph(report)
                add_subsection_paragraph(doc, "(d)", "Local authority data", authority_para)

            # (e) Rent act effectiveness - Generate regulation paragraph
            if report.rent_act_effectiveness:
                rent_para = generate_rent_act_paragraph(report)
                add_subsection_paragraph(doc, "(e)", "Rent act effectiveness", rent_para)

            doc.add_paragraph()  # Spacing

        # ===== 7.0 LAND VALUES IN THE AREA =====
        # Show section if either comparables OR market analysis exists
        if report.comparable_properties or report.land_market_analysis:
            add_section_heading(doc, "7.0", "LAND VALUES IN THE AREA")

            # Parse JSON if needed
            # SAFE: Parse JSON with error handling
            comparables = safe_parse_json_string(report.comparable_properties, [])

            if comparables:
                # Generate paragraph from comparables
                land_values_text = generate_land_values_paragraph(comparables)

                if land_values_text:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.space_after = Pt(6)
                    run = p.add_run(land_values_text)
                    run.font.size = FONT_SIZE_BODY

                    # Add average rate note
                    avg_rate = sum(c.get('rate_per_perch', 0) for c in comparables) / len(comparables) if comparables else 0
                    if avg_rate > 0:
                        p = doc.add_paragraph()
                        p.paragraph_format.space_before = Pt(6)
                        p.paragraph_format.space_after = Pt(6)
                        run = p.add_run(f"Average Rate: LKR {avg_rate:,.2f} per perch")
                        run.font.bold = True
                        run.font.size = FONT_SIZE_BODY

            if report.land_market_analysis:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.space_after = Pt(12)
                run = p.add_run(report.land_market_analysis)
                run.font.size = FONT_SIZE_BODY

            doc.add_paragraph()

        # ===== 8.0 VALUATION =====
        if report.valuation_total_land_value or report.valuation_buildings_data:
            add_section_heading(doc, "8.0", "VALUATION OF THE PROPERTY")

            # === VALUATION BREAKDOWN ===

            # Land valuation line
            if report.valuation_total_land_value:
                extent = report.valuation_land_extent or report.land_extent_perches or 0
                rate = report.valuation_rate_per_perch or 0
                land_value = report.valuation_total_land_value

                p = doc.add_paragraph()

                # Add tab stop for right alignment
                tab_stops = p.paragraph_format.tab_stops
                tab_stops.add_tab_stop(Inches(6.0), WD_TAB_ALIGNMENT.RIGHT)

                text = f"Land – {extent:,.2f} perches @ {format_currency(rate)} per perch\t= {format_currency_aligned(land_value)}"
                run = p.add_run(text)
                run.font.size = FONT_SIZE_VALUATION
                p.paragraph_format.space_after = Pt(3)

            # Buildings valuation (SKIP for bare_land)
            total_depreciated_buildings_value = 0
            buildings_insurance_values = []  # Store per-building insurance values

            if report.valuation_buildings_data and report.report_type != 'bare_land':
                # SAFE: Parse JSON with error handling
                buildings_data = safe_parse_json_string(report.valuation_buildings_data, [])

                for idx, bldg in enumerate(buildings_data, 1):
                    building_name = bldg.get('building_name', f'Building {idx}')
                    subtotal = bldg.get('subtotal', 0)

                    # Calculate total floor area from components
                    components = bldg.get('components', [])
                    total_floor_area = sum(comp.get('floor_area', 0) for comp in components)

                    # Calculate average rate per sq.ft
                    avg_rate = subtotal / total_floor_area if total_floor_area > 0 else 0

                    # Check if depreciation data exists
                    has_depreciation = bldg.get('depreciation_amount') is not None and to_float(bldg.get('depreciation_amount', 0)) > 0

                    if has_depreciation:
                        depreciation_rate = to_float(bldg.get('depreciation_rate_percent', 0))
                        depreciated_value = to_float(bldg.get('depreciated_value', subtotal))

                        # NEW FORMAT: 2-line building with inline depreciation
                        format_building_valuation_2line(
                            doc,
                            building_name,
                            total_floor_area,
                            avg_rate,
                            depreciation_rate,
                            depreciated_value
                        )

                        building_value = depreciated_value
                    else:
                        # No depreciation: Use single-line format (backward compatible)
                        p = doc.add_paragraph()

                        # Add tab stop for right alignment
                        tab_stops = p.paragraph_format.tab_stops
                        tab_stops.add_tab_stop(Inches(6.0), WD_TAB_ALIGNMENT.RIGHT)

                        text = f"{building_name} – {total_floor_area:,.0f} sq.ft @ {format_currency(avg_rate)} per square foot\t= {format_currency_aligned(subtotal)}"
                        run = p.add_run(text)
                        run.font.size = FONT_SIZE_VALUATION
                        p.paragraph_format.space_after = Pt(3)
                        building_value = to_float(subtotal)

                    total_depreciated_buildings_value += building_value
                    # Insurance always uses replacement cost (undepreciated)
                    buildings_insurance_values.append({
                        'name': building_name,
                        'value': to_float(subtotal)  # Replacement cost for insurance
                    })

            # Add-ons (NEW compact format)
            total_addons_value = 0
            if report.valuation_addons:
                # SAFE: Parse JSON with error handling
                addons = safe_parse_json_string(report.valuation_addons, [])

                if addons:
                    # No header, just show add-ons in compact format
                    for addon in addons:
                        format_addon_compact(
                            doc,
                            addon.get('description', 'Add-on'),
                            to_float(addon.get('value', 0))
                        )
                        total_addons_value += to_float(addon.get('value', 0))

            # Calculate market values
            land_value = to_float(report.valuation_total_land_value)
            market_value_calculated = land_value + to_float(total_depreciated_buildings_value) + to_float(total_addons_value)
            market_value_rounded = round_for_say(market_value_calculated)

            # Determine if we should show "Market Value of the property" section
            # SKIP for bare land with ONLY land value (no buildings, no add-ons)
            has_buildings_or_addons = (total_depreciated_buildings_value > 0) or (total_addons_value > 0)

            if has_buildings_or_addons:
                # Show "Market Value of the property" line (with double underline)
                add_market_value_line(doc, market_value_calculated, has_blank_before=True)

            # Always show "Value rounded off" line
            add_value_rounded_line(doc, market_value_rounded)

            # Check if valuation type is "Forced Sale Value" to show forced sale fields
            show_forced_sale = report.valuation_type == "Forced Sale Value"

            if show_forced_sale:
                # Calculate forced sale value for summary section
                forced_sale_percentage = report.valuation_forced_sale_percentage or 90
                forced_sale_value = market_value_rounded * (forced_sale_percentage / 100)

            # === SUMMARY OF THE VALUATION ===
            p = doc.add_paragraph()
            run = p.add_run("SUMMARY OF THE VALUATION")
            run.font.bold = True
            run.font.underline = True
            run.font.size = FONT_SIZE_BODY
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)

            # Open Market Value
            p = doc.add_paragraph()
            tab_stops = p.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(3.5), WD_TAB_ALIGNMENT.LEFT)
            tab_stops.add_tab_stop(Inches(3.7), WD_TAB_ALIGNMENT.LEFT)
            text = f"Open Market Value of the property\t:\t{format_currency(market_value_rounded)}"
            run = p.add_run(text)
            run.font.size = FONT_SIZE_VALUATION
            p.paragraph_format.space_after = Pt(3)

            # Forced Sale Value - only show when valuation type is "Forced Sale Value"
            if show_forced_sale:
                p = doc.add_paragraph()
                tab_stops = p.paragraph_format.tab_stops
                tab_stops.add_tab_stop(Inches(3.5), WD_TAB_ALIGNMENT.LEFT)
                tab_stops.add_tab_stop(Inches(3.7), WD_TAB_ALIGNMENT.LEFT)
                text = f"Forced Sale Value of the property\t:\t{format_currency(forced_sale_value)}"
                run = p.add_run(text)
                run.font.size = FONT_SIZE_VALUATION
                p.paragraph_format.space_after = Pt(3)

            # Insurance Value (NEW INLINE FORMAT) - Only show if there are buildings
            if buildings_insurance_values:
                for building_ins in buildings_insurance_values:
                    p = doc.add_paragraph()
                    tab_stops = p.paragraph_format.tab_stops
                    tab_stops.add_tab_stop(Inches(3.5), WD_TAB_ALIGNMENT.LEFT)
                    tab_stops.add_tab_stop(Inches(3.7), WD_TAB_ALIGNMENT.LEFT)
                    text = f"Insurance Value of the {building_ins['name']}\t:\t{format_currency(building_ins['value'])}"
                    run = p.add_run(text)
                    run.font.size = FONT_SIZE_VALUATION
                    p.paragraph_format.space_after = Pt(2)

                doc.add_paragraph()  # Final spacing

        # ===== 9.0 CERTIFICATION (SIMPLIFIED FORMAT) =====
        if report.certification_text or report.certification_valuer_name:
            add_section_heading(doc, "9.0", "CERTIFICATION")

            # Certification text - single paragraph
            if report.certification_text:
                # Use custom certification text if provided (override mode)
                cert_text = report.certification_text.strip()
            elif report.certification_valuer_name and report.certification_valuer_designation:
                # Auto-generate simplified certification
                cert_text = generate_simplified_certification_text(
                    valuer_name=report.certification_valuer_name,
                    valuer_designation=report.certification_valuer_designation,
                    lot_number=report.lot_number,
                    plan_number=report.plan_number,
                    plan_date=report.plan_date,
                    licensed_surveyor_name=report.licensed_surveyor_name,
                    deeds=safe_get_json_field(report, 'deeds', []),
                    property_identification_type=report.property_identification_type
                )
            else:
                # Fallback if no valuer info
                cert_text = "I hereby certify that I have personally inspected this property and the valuation stated herein represents my professional opinion of the market value as of the date of inspection."

            # Render certification paragraph
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run(cert_text)
            run.font.size = FONT_SIZE_BODY

            # Signature block
            add_signature_block(
                doc=doc,
                user=user,
                valuer_name=report.certification_valuer_name,
                valuer_designation=report.certification_valuer_designation,
                certification_date=report.certification_date
            )

        # ===== INVOICE SECTION =====
        if report.invoice_data:
            logger.info("[DOCX] Generating invoice section")
            _generate_invoice_section(doc, report.invoice_data, user, report)

        # ===== SAVE DOCUMENT =====
        # This MUST be outside the certification block!
        logger.info("[DOCX] About to save document to BytesIO")
        # Save to BytesIO
        file_stream = BytesIO()
        logger.info("[DOCX] Created BytesIO stream")
        doc.save(file_stream)
        logger.info("[DOCX] Document saved to stream")
        file_stream.seek(0)
        logger.info("[DOCX] Stream position reset to 0")

        # Safety check
        if file_stream is None or file_stream.getvalue() == b'':
            raise ValueError("Generated document is empty or None")

        logger.info(f"Document generated successfully, size: {len(file_stream.getvalue())} bytes")
        return file_stream

    except AttributeError as e:
        logger.error(f"Missing required field in report generation: {e}", exc_info=True)
        raise ValueError(f"Report data incomplete: Missing field - {str(e)}")
    except IndexError as e:
        logger.error(f"Array access error in report generation: {e}", exc_info=True)
        raise ValueError(f"Report data malformed: Array indexing error - {str(e)}")
    except TypeError as e:
        logger.error(f"Type error in report data: {e}", exc_info=True)
        raise ValueError(f"Report data has incorrect types: {str(e)}")
    except Exception as e:
        logger.error(f"Unexpected error in DOCX generation: {e}", exc_info=True)
        raise

def get_filename_for_user(report: models.Report) -> str:
    """Generate a safe filename for the report document"""
    # Use applicant name if available, otherwise use valuer name
    name = report.applicant_full_name if report.applicant_full_name else report.user.full_name
    # Clean the name for use in filename
    safe_name = "".join(c for c in name if c.isalnum() or c in (' ', '-', '_')).strip()
    safe_name = safe_name.replace(' ', '_')
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    report_type = report.report_type.replace('_', '-')
    return f"{report_type}_{safe_name}_{timestamp}.docx"
