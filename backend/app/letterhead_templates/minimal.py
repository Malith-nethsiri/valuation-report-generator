"""
Minimal Elegant letterhead template.

This template features ultra-clean design with no borders,
centered name, and compact footer contact information.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

from .base import BaseLetterheadTemplate, TemplateMetadata


class MinimalTemplate(BaseLetterheadTemplate):
    """
    Minimal Elegant letterhead template.

    Features:
    - Centered name only at top
    - No borders (whitespace separation only)
    - Compact footer with pipe separators
    - Professional fonts: Name 11pt bold, credentials 9pt, contact 7pt
    - Clean, spacious, elegant design
    """

    def get_metadata(self) -> TemplateMetadata:
        """Return metadata for the Minimal template."""
        return TemplateMetadata(
            template_id="minimal",
            name="Minimal Elegant",
            description="Ultra-clean centered design with no borders and compact contact footer",
            category="minimal"
        )

    def render_letterhead(self, doc: Document, user, report) -> None:
        """
        Render the Minimal letterhead design.

        Args:
            doc: python-docx Document object
            user: User model instance
            report: Report model instance
        """
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        # Professional Name (Centered, Bold, Larger)
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_para.paragraph_format.space_before = Pt(0)
        name_para.paragraph_format.space_after = Pt(2)
        name_run = name_para.add_run()
        if user.honorific:
            name_run.add_text(f"{user.honorific} ")
        name_run.add_text(user.full_name)
        name_run.bold = True
        name_run.font.size = Pt(14)
        name_run.font.color.rgb = RGBColor(0, 0, 0)

        # Academic Qualifications (Centered)
        if user.academic_qualifications:
            qual_para = doc.add_paragraph()
            qual_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            qual_para.paragraph_format.space_before = Pt(0)
            qual_para.paragraph_format.space_after = Pt(1)
            qual_run = qual_para.add_run(user.academic_qualifications)
            qual_run.font.size = Pt(11)
            qual_run.font.color.rgb = RGBColor(0, 0, 0)

        # Professional Designation (Centered)
        if user.professional_designation:
            desig_para = doc.add_paragraph()
            desig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            desig_para.paragraph_format.space_before = Pt(0)
            desig_para.paragraph_format.space_after = Pt(1)
            desig_run = desig_para.add_run(user.professional_designation)
            desig_run.font.size = Pt(11)
            desig_run.font.color.rgb = RGBColor(0, 0, 0)

        # Extra spacing for clean separation
        spacing_para = doc.add_paragraph()
        spacing_para.paragraph_format.space_before = Pt(10)
        spacing_para.paragraph_format.space_after = Pt(10)

        # Compact Footer Contact Line (Centered, pipe-separated)
        footer_parts = []

        # Email
        if user.email:
            footer_parts.append(user.email)

        # Primary phone
        if user.phone_primary:
            footer_parts.append(user.phone_primary)

        # Address (abbreviated)
        address_parts = []
        if user.locality:
            address_parts.append(user.locality)
        elif user.village:
            address_parts.append(user.village)
        if address_parts:
            footer_parts.append(", ".join(address_parts))

        # Membership
        if user.membership_level or user.membership_number:
            member_text = []
            if user.membership_level:
                member_text.append(user.membership_level)
            if user.membership_number:
                member_text.append(user.membership_number)
            footer_parts.append(" ".join(member_text))

        # Panel Valuer
        if user.panel_valuer_banks and len(user.panel_valuer_banks) > 0:
            footer_parts.append(f"Panel Valuer: {', '.join(user.panel_valuer_banks)}")

        # Add footer line
        if footer_parts:
            footer_para = doc.add_paragraph()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_para.paragraph_format.space_before = Pt(2)
            footer_para.paragraph_format.space_after = Pt(6)
            footer_run = footer_para.add_run(" | ".join(footer_parts))
            footer_run.font.size = Pt(11)
            footer_run.font.color.rgb = RGBColor(64, 64, 64)  # Slightly gray

        # Reference and Date (Centered, on separate lines)
        ref_para = doc.add_paragraph()
        ref_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ref_para.paragraph_format.space_before = Pt(0)
        ref_para.paragraph_format.space_after = Pt(1)
        ref_label = ref_para.add_run("Ref: ")
        ref_label.font.size = Pt(11)
        ref_label.font.color.rgb = RGBColor(0, 0, 0)
        if report.report_reference:
            ref_value = ref_para.add_run(report.report_reference)
            ref_value.font.size = Pt(11)
            ref_value.font.color.rgb = RGBColor(0, 0, 0)

        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.paragraph_format.space_before = Pt(0)
        date_para.paragraph_format.space_after = Pt(0)
        date_label = date_para.add_run("Date: ")
        date_label.font.size = Pt(11)
        date_label.font.color.rgb = RGBColor(0, 0, 0)
        if report.report_date:
            date_value = date_para.add_run(report.report_date)
        else:
            date_value = date_para.add_run(datetime.now().strftime('%Y-%m-%d'))
        date_value.font.size = Pt(11)
        date_value.font.color.rgb = RGBColor(0, 0, 0)

        # Add spacing after letterhead
        spacing_para2 = doc.add_paragraph()
        spacing_para2.paragraph_format.space_before = Pt(8)
        spacing_para2.paragraph_format.space_after = Pt(0)
