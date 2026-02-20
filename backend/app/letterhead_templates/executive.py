"""
Executive Bold letterhead template.

This template features a bold, authoritative design with
thick double-line borders and larger fonts.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

from .base import BaseLetterheadTemplate, TemplateMetadata
from ..docx_generation.styling import add_border_to_paragraph



class ExecutiveTemplate(BaseLetterheadTemplate):
    """
    Executive Bold letterhead template.

    Features:
    - Centered layout with emphasis
    - Thick double-line borders (top and bottom)
    - Condensed single-line contact format
    - Professional fonts: Name 12pt bold, credentials 9pt bold, contact 7pt
    - Authoritative, prestigious design
    """

    def get_metadata(self) -> TemplateMetadata:
        """Return metadata for the Executive template."""
        return TemplateMetadata(
            template_id="executive",
            name="Executive Bold",
            description="Authoritative design with thick double-line borders and bold typography",
            category="executive"
        )

    def render_letterhead(self, doc: Document, user, report) -> None:
        """
        Render the Executive letterhead design.

        Args:
            doc: python-docx Document object
            user: User model instance
            report: Report model instance
        """
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        # Double-line top border
        border_para1 = doc.add_paragraph()
        border_para1.paragraph_format.space_before = Pt(0)
        border_para1.paragraph_format.space_after = Pt(0)
        add_border_to_paragraph(border_para1, "top", 8, "000000")

        border_para2 = doc.add_paragraph()
        border_para2.paragraph_format.space_before = Pt(3)
        border_para2.paragraph_format.space_after = Pt(4)
        add_border_to_paragraph(border_para2, "top", 2, "000000")

        # Professional Name (Centered, Bold, Large)
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_para.paragraph_format.space_before = Pt(0)
        name_para.paragraph_format.space_after = Pt(1)
        name_para.paragraph_format.line_spacing = 0.8
        name_run = name_para.add_run()
        if user.honorific:
            name_run.add_text(f"{user.honorific} ")
        name_run.add_text(user.full_name)
        name_run.bold = True
        name_run.font.size = Pt(14)
        name_run.font.color.rgb = RGBColor(0, 0, 0)

        # Academic Qualifications (Centered, Bold)
        if user.academic_qualifications:
            qual_para = doc.add_paragraph()
            qual_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            qual_para.paragraph_format.space_before = Pt(0)
            qual_para.paragraph_format.space_after = Pt(1)
            qual_para.paragraph_format.line_spacing = 0.8
            qual_run = qual_para.add_run(user.academic_qualifications)
            qual_run.bold = True
            qual_run.font.size = Pt(11)
            qual_run.font.color.rgb = RGBColor(0, 0, 0)

        # Professional Designation (Centered, Bold)
        if user.professional_designation:
            desig_para = doc.add_paragraph()
            desig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            desig_para.paragraph_format.space_before = Pt(0)
            desig_para.paragraph_format.space_after = Pt(2)
            desig_para.paragraph_format.line_spacing = 0.8
            desig_run = desig_para.add_run(user.professional_designation)
            desig_run.bold = True
            desig_run.font.size = Pt(11)
            desig_run.font.color.rgb = RGBColor(0, 0, 0)

        # Condensed contact block (Centered, single line)
        contact_parts = []

        # Residence address (abbreviated)
        if user.locality:
            contact_parts.append(user.locality)

        # Phones
        if user.phone_primary:
            contact_parts.append(f"Tel: {user.phone_primary}")

        # Email
        if user.email:
            contact_parts.append(user.email)

        # Membership
        if user.membership_level or user.membership_number:
            member_text = []
            if user.membership_level:
                member_text.append(user.membership_level)
            if user.membership_number:
                member_text.append(user.membership_number)
            contact_parts.append(" | ".join(member_text))

        # Panel Valuer
        if user.panel_valuer_banks and len(user.panel_valuer_banks) > 0:
            contact_parts.append(f"Panel Valuer: {', '.join(user.panel_valuer_banks)}")

        # Add contact line
        if contact_parts:
            contact_para = doc.add_paragraph()
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_para.paragraph_format.space_before = Pt(2)
            contact_para.paragraph_format.space_after = Pt(3)
            contact_run = contact_para.add_run(" | ".join(contact_parts))
            contact_run.font.size = Pt(11)
            contact_run.font.color.rgb = RGBColor(0, 0, 0)

        # Double-line bottom border
        border_para3 = doc.add_paragraph()
        border_para3.paragraph_format.space_before = Pt(2)
        border_para3.paragraph_format.space_after = Pt(0)
        add_border_to_paragraph(border_para3, "bottom", 2, "000000")

        border_para4 = doc.add_paragraph()
        border_para4.paragraph_format.space_before = Pt(3)
        border_para4.paragraph_format.space_after = Pt(3)
        add_border_to_paragraph(border_para4, "bottom", 8, "000000")

        # Reference and Date Line (Bold labels, larger font)
        ref_table = doc.add_table(rows=1, cols=2)
        ref_table.autofit = False

        ref_cell = ref_table.rows[0].cells[0]
        ref_cell_para = ref_cell.paragraphs[0]
        ref_cell_para.paragraph_format.space_before = Pt(0)
        ref_cell_para.paragraph_format.space_after = Pt(0)
        ref_label = ref_cell_para.add_run("Ref: ")
        ref_label.bold = True
        ref_label.font.size = Pt(11)
        ref_label.font.color.rgb = RGBColor(0, 0, 0)

        if report.report_reference:
            ref_value = ref_cell_para.add_run(report.report_reference)
            ref_value.font.size = Pt(11)
            ref_value.font.color.rgb = RGBColor(0, 0, 0)

        date_cell = ref_table.rows[0].cells[1]
        date_cell_para = date_cell.paragraphs[0]
        date_cell_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_cell_para.paragraph_format.space_before = Pt(0)
        date_cell_para.paragraph_format.space_after = Pt(0)
        date_label = date_cell_para.add_run("Date: ")
        date_label.bold = True
        date_label.font.size = Pt(11)
        date_label.font.color.rgb = RGBColor(0, 0, 0)

        if report.report_date:
            date_value = date_cell_para.add_run(report.report_date)
        else:
            date_value = date_cell_para.add_run(datetime.now().strftime('%Y-%m-%d'))
        date_value.font.size = Pt(11)
        date_value.font.color.rgb = RGBColor(0, 0, 0)

        # Add spacing after letterhead
        spacing_para = doc.add_paragraph()
        spacing_para.paragraph_format.space_before = Pt(8)
        spacing_para.paragraph_format.space_after = Pt(0)
