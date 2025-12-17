"""
Compact Professional letterhead template.

This template features space-efficient design with smaller fonts
and tighter spacing to maximize content area.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

from .base import BaseLetterheadTemplate, TemplateMetadata


def add_border_to_paragraph(paragraph, border_position="bottom", size=12, color="000000"):
    """Add a border to a paragraph"""
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')

    border = OxmlElement(f'w:{border_position}')
    border.set(qn('w:val'), 'single')
    border.set(qn('w:sz'), str(size))
    border.set(qn('w:space'), '1')
    border.set(qn('w:color'), color)

    pBdr.append(border)
    pPr.append(pBdr)


class CompactTemplate(BaseLetterheadTemplate):
    """
    Compact Professional letterhead template.

    Features:
    - Centered layout, space-efficient
    - Thin borders (4pt weight)
    - Two-column contact with smaller fonts
    - Professional fonts: Name 8pt bold, credentials 6-7pt, contact 6pt
    - Very tight spacing (0.7 line height)
    - Space-efficient, fits more content
    """

    def get_metadata(self) -> TemplateMetadata:
        """Return metadata for the Compact template."""
        return TemplateMetadata(
            template_id="compact",
            name="Compact Professional",
            description="Space-efficient design with smaller fonts and tight spacing",
            category="compact"
        )

    def render_letterhead(self, doc: Document, user, report) -> None:
        """
        Render the Compact letterhead design.

        Args:
            doc: python-docx Document object
            user: User model instance
            report: Report model instance
        """
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        # Professional Name (Centered, Bold, Compact)
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_para.paragraph_format.space_before = Pt(0)
        name_para.paragraph_format.space_after = Pt(0)
        name_para.paragraph_format.line_spacing = 0.7
        name_run = name_para.add_run()
        if user.honorific:
            name_run.add_text(f"{user.honorific} ")
        name_run.add_text(user.full_name)
        name_run.bold = True
        name_run.font.size = Pt(8)
        name_run.font.color.rgb = RGBColor(0, 0, 0)

        # Academic Qualifications (Centered, Compact)
        if user.academic_qualifications:
            qual_para = doc.add_paragraph()
            qual_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            qual_para.paragraph_format.space_before = Pt(0)
            qual_para.paragraph_format.space_after = Pt(0)
            qual_para.paragraph_format.line_spacing = 0.7
            qual_run = qual_para.add_run(user.academic_qualifications)
            qual_run.font.size = Pt(7)
            qual_run.font.color.rgb = RGBColor(0, 0, 0)

        # Professional Designation (Centered, Bold, Compact)
        if user.professional_designation:
            desig_para = doc.add_paragraph()
            desig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            desig_para.paragraph_format.space_before = Pt(0)
            desig_para.paragraph_format.space_after = Pt(0)
            desig_para.paragraph_format.line_spacing = 0.7
            desig_run = desig_para.add_run(user.professional_designation)
            desig_run.bold = True
            desig_run.font.size = Pt(7)
            desig_run.font.color.rgb = RGBColor(0, 0, 0)

        # Membership Information (Centered, Compact)
        if user.membership_level or user.membership_number:
            member_para = doc.add_paragraph()
            member_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            member_para.paragraph_format.space_before = Pt(0)
            member_para.paragraph_format.space_after = Pt(0)
            member_para.paragraph_format.line_spacing = 0.7
            member_text = []
            if user.membership_level:
                member_text.append(user.membership_level)
            if user.membership_number:
                member_text.append(user.membership_number)
            member_run = member_para.add_run(" | ".join(member_text))
            member_run.font.size = Pt(6)
            member_run.font.color.rgb = RGBColor(0, 0, 0)

        # Panel Valuer Status (Centered, Compact)
        if user.panel_valuer_banks and len(user.panel_valuer_banks) > 0:
            panel_para = doc.add_paragraph()
            panel_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            panel_para.paragraph_format.space_before = Pt(0)
            panel_para.paragraph_format.space_after = Pt(0)
            panel_para.paragraph_format.line_spacing = 0.7
            banks_text = ", ".join(user.panel_valuer_banks)
            panel_run = panel_para.add_run(f"Panel Valuer: {banks_text}")
            panel_run.font.size = Pt(6)
            panel_run.font.color.rgb = RGBColor(0, 0, 0)

        # Thin border separator
        separator_para = doc.add_paragraph()
        separator_para.paragraph_format.space_before = Pt(1)
        separator_para.paragraph_format.space_after = Pt(1)
        add_border_to_paragraph(separator_para, "bottom", 4, "000000")

        # Contact Information Section (Two columns, smaller fonts)
        contact_table = doc.add_table(rows=1, cols=2)
        contact_table.autofit = False

        # RESIDENCE Column
        residence_cell = contact_table.rows[0].cells[0]
        residence_para = residence_cell.paragraphs[0]
        residence_para.paragraph_format.space_before = Pt(0)
        residence_para.paragraph_format.space_after = Pt(0)
        residence_para.paragraph_format.line_spacing = 0.7
        residence_heading = residence_para.add_run("RESIDENCE")
        residence_heading.bold = True
        residence_heading.font.size = Pt(6)
        residence_heading.font.color.rgb = RGBColor(0, 0, 0)

        # Build residential address
        res_address_parts = []
        if user.house_number:
            res_address_parts.append(user.house_number)
        if user.locality:
            res_address_parts.append(user.locality)
        if res_address_parts:
            res_run = residence_para.add_run("\n" + ", ".join(res_address_parts))
            res_run.font.size = Pt(6)
            res_run.font.color.rgb = RGBColor(0, 0, 0)

        # Residential phones
        if user.phone_primary:
            phone_run = residence_para.add_run(f"\n{user.phone_primary}")
            phone_run.font.size = Pt(6)
            phone_run.font.color.rgb = RGBColor(0, 0, 0)

        # OFFICE Column
        office_cell = contact_table.rows[0].cells[1]
        office_para = office_cell.paragraphs[0]
        office_para.paragraph_format.space_before = Pt(0)
        office_para.paragraph_format.space_after = Pt(0)
        office_para.paragraph_format.line_spacing = 0.7
        office_heading = office_para.add_run("OFFICE")
        office_heading.bold = True
        office_heading.font.size = Pt(6)
        office_heading.font.color.rgb = RGBColor(0, 0, 0)

        # Build office address
        office_address_parts = []
        if user.office_department:
            office_address_parts.append(user.office_department)
        if user.office_region:
            office_address_parts.append(user.office_region)
        if office_address_parts:
            office_run = office_para.add_run("\n" + ", ".join(office_address_parts))
            office_run.font.size = Pt(6)
            office_run.font.color.rgb = RGBColor(0, 0, 0)

        # Office phone
        if user.office_phone:
            office_phone_run = office_para.add_run(f"\n{user.office_phone}")
            office_phone_run.font.size = Pt(6)
            office_phone_run.font.color.rgb = RGBColor(0, 0, 0)

        # Email (Centered, Compact)
        if user.email:
            email_para = doc.add_paragraph()
            email_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            email_para.paragraph_format.space_before = Pt(1)
            email_para.paragraph_format.space_after = Pt(1)
            email_run = email_para.add_run(user.email)
            email_run.font.size = Pt(6)
            email_run.font.color.rgb = RGBColor(0, 0, 0)

        # Thin border separator
        separator_para2 = doc.add_paragraph()
        separator_para2.paragraph_format.space_before = Pt(1)
        separator_para2.paragraph_format.space_after = Pt(1)
        add_border_to_paragraph(separator_para2, "bottom", 4, "000000")

        # Reference and Date Line (Small font)
        ref_table = doc.add_table(rows=1, cols=2)
        ref_table.autofit = False

        ref_cell = ref_table.rows[0].cells[0]
        ref_cell_para = ref_cell.paragraphs[0]
        ref_cell_para.paragraph_format.space_before = Pt(0)
        ref_cell_para.paragraph_format.space_after = Pt(0)
        ref_label = ref_cell_para.add_run("Ref: ")
        ref_label.bold = True
        ref_label.font.size = Pt(7)
        ref_label.font.color.rgb = RGBColor(0, 0, 0)

        if report.report_reference:
            ref_value = ref_cell_para.add_run(report.report_reference)
            ref_value.font.size = Pt(7)
            ref_value.font.color.rgb = RGBColor(0, 0, 0)

        date_cell = ref_table.rows[0].cells[1]
        date_cell_para = date_cell.paragraphs[0]
        date_cell_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_cell_para.paragraph_format.space_before = Pt(0)
        date_cell_para.paragraph_format.space_after = Pt(0)
        date_label = date_cell_para.add_run("Date: ")
        date_label.bold = True
        date_label.font.size = Pt(7)
        date_label.font.color.rgb = RGBColor(0, 0, 0)

        if report.report_date:
            date_value = date_cell_para.add_run(report.report_date)
        else:
            date_value = date_cell_para.add_run(datetime.now().strftime('%Y-%m-%d'))
        date_value.font.size = Pt(7)
        date_value.font.color.rgb = RGBColor(0, 0, 0)

        # Minimal spacing after letterhead
        spacing_para = doc.add_paragraph()
        spacing_para.paragraph_format.space_before = Pt(6)
        spacing_para.paragraph_format.space_after = Pt(0)
