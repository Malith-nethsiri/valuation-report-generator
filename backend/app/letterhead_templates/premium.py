"""
Premium Signature letterhead template.

This template features sophisticated design with asymmetric borders
and stacked contact format with decorative elements.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

from .base import BaseLetterheadTemplate, TemplateMetadata


def add_border_to_paragraph(paragraph, border_position="bottom", size=12, color="000000"):
    """Add a border to a paragraph"""
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')

    border = OxmlElement(f'w:{border_position}')
    border.set(qn('w:val'), 'single')
    border.set(qn('w:sz'), str(size))
    border.set(qn('w:space'), '1')
    border.set(qn('w:color'), color)

    pBdr.append(border)
    pPr.append(pBdr)


class PremiumTemplate(BaseLetterheadTemplate):
    """
    Premium Signature letterhead template.

    Features:
    - Left-aligned name with decorative elements
    - Asymmetric borders (thick top 8pt, thin bottom 3pt)
    - Stacked contact format with bullet symbols
    - Professional fonts: Name 11pt bold, mixed 7-9pt
    - Strategic spacing for visual hierarchy
    - Sophisticated, unique design
    """

    def get_metadata(self) -> TemplateMetadata:
        """Return metadata for the Premium template."""
        return TemplateMetadata(
            template_id="premium",
            name="Premium Signature",
            description="Sophisticated design with asymmetric borders and stacked contact format",
            category="premium"
        )

    def render_letterhead(self, doc: Document, user, report) -> None:
        """
        Render the Premium letterhead design.

        Args:
            doc: python-docx Document object
            user: User model instance
            report: Report model instance
        """
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        # Thick top border
        border_para = doc.add_paragraph()
        border_para.paragraph_format.space_before = Pt(0)
        border_para.paragraph_format.space_after = Pt(4)
        add_border_to_paragraph(border_para, "top", 8, "000000")

        # Professional Name (Left-aligned, Bold, Large)
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        name_para.paragraph_format.space_before = Pt(0)
        name_para.paragraph_format.space_after = Pt(2)
        name_run = name_para.add_run()
        if user.honorific:
            name_run.add_text(f"{user.honorific} ")
        name_run.add_text(user.full_name)
        name_run.bold = True
        name_run.font.size = Pt(11)
        name_run.font.color.rgb = RGBColor(0, 0, 0)

        # Academic Qualifications (Left-aligned, Medium)
        if user.academic_qualifications:
            qual_para = doc.add_paragraph()
            qual_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            qual_para.paragraph_format.space_before = Pt(0)
            qual_para.paragraph_format.space_after = Pt(1)
            qual_run = qual_para.add_run(user.academic_qualifications)
            qual_run.font.size = Pt(9)
            qual_run.font.color.rgb = RGBColor(0, 0, 0)

        # Professional Designation (Left-aligned, Medium Bold)
        if user.professional_designation:
            desig_para = doc.add_paragraph()
            desig_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            desig_para.paragraph_format.space_before = Pt(0)
            desig_para.paragraph_format.space_after = Pt(2)
            desig_run = desig_para.add_run(user.professional_designation)
            desig_run.bold = True
            desig_run.font.size = Pt(9)
            desig_run.font.color.rgb = RGBColor(0, 0, 0)

        # Stacked contact information with bullet points
        # Membership
        if user.membership_level or user.membership_number:
            member_para = doc.add_paragraph()
            member_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            member_para.paragraph_format.space_before = Pt(1)
            member_para.paragraph_format.space_after = Pt(0)
            member_para.paragraph_format.left_indent = Inches(0.2)
            member_text = []
            if user.membership_level:
                member_text.append(user.membership_level)
            if user.membership_number:
                member_text.append(user.membership_number)
            member_run = member_para.add_run(f"• {' | '.join(member_text)}")
            member_run.font.size = Pt(7)
            member_run.font.color.rgb = RGBColor(0, 0, 0)

        # Panel Valuer
        if user.panel_valuer_banks and len(user.panel_valuer_banks) > 0:
            panel_para = doc.add_paragraph()
            panel_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            panel_para.paragraph_format.space_before = Pt(0)
            panel_para.paragraph_format.space_after = Pt(0)
            panel_para.paragraph_format.left_indent = Inches(0.2)
            banks_text = ", ".join(user.panel_valuer_banks)
            panel_run = panel_para.add_run(f"• Panel Valuer: {banks_text}")
            panel_run.font.size = Pt(7)
            panel_run.font.color.rgb = RGBColor(0, 0, 0)

        # Residence Address
        res_address_parts = []
        if user.house_number:
            res_address_parts.append(user.house_number)
        if user.area_development:
            res_address_parts.append(user.area_development)
        if user.locality:
            res_address_parts.append(user.locality)
        if res_address_parts:
            res_para = doc.add_paragraph()
            res_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            res_para.paragraph_format.space_before = Pt(0)
            res_para.paragraph_format.space_after = Pt(0)
            res_para.paragraph_format.left_indent = Inches(0.2)
            res_run = res_para.add_run(f"• Residence: {', '.join(res_address_parts)}")
            res_run.font.size = Pt(7)
            res_run.font.color.rgb = RGBColor(0, 0, 0)

        # Phone
        if user.phone_primary:
            phone_para = doc.add_paragraph()
            phone_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            phone_para.paragraph_format.space_before = Pt(0)
            phone_para.paragraph_format.space_after = Pt(0)
            phone_para.paragraph_format.left_indent = Inches(0.2)
            phone_run = phone_para.add_run(f"• Tel: {user.phone_primary}")
            phone_run.font.size = Pt(7)
            phone_run.font.color.rgb = RGBColor(0, 0, 0)

        # Email
        if user.email:
            email_para = doc.add_paragraph()
            email_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            email_para.paragraph_format.space_before = Pt(0)
            email_para.paragraph_format.space_after = Pt(2)
            email_para.paragraph_format.left_indent = Inches(0.2)
            email_run = email_para.add_run(f"• Email: {user.email}")
            email_run.font.size = Pt(7)
            email_run.font.color.rgb = RGBColor(0, 0, 0)

        # Office (if provided)
        if user.office_department or user.office_region or user.office_phone:
            office_parts = []
            if user.office_department:
                office_parts.append(user.office_department)
            if user.office_region:
                office_parts.append(user.office_region)
            office_para = doc.add_paragraph()
            office_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            office_para.paragraph_format.space_before = Pt(0)
            office_para.paragraph_format.space_after = Pt(0)
            office_para.paragraph_format.left_indent = Inches(0.2)
            office_text = f"• Office: {', '.join(office_parts)}" if office_parts else ""
            if user.office_phone:
                if office_text:
                    office_text += f" | Tel: {user.office_phone}"
                else:
                    office_text = f"• Office Tel: {user.office_phone}"
            if office_text:
                office_run = office_para.add_run(office_text)
                office_run.font.size = Pt(7)
                office_run.font.color.rgb = RGBColor(0, 0, 0)

        # Thin bottom border (asymmetric design)
        separator_para = doc.add_paragraph()
        separator_para.paragraph_format.space_before = Pt(3)
        separator_para.paragraph_format.space_after = Pt(2)
        add_border_to_paragraph(separator_para, "bottom", 3, "000000")

        # Reference and Date Line (Indented)
        ref_para = doc.add_paragraph()
        ref_para.paragraph_format.space_before = Pt(0)
        ref_para.paragraph_format.space_after = Pt(0)
        ref_para.paragraph_format.left_indent = Inches(0.2)

        ref_label = ref_para.add_run("Ref: ")
        ref_label.bold = True
        ref_label.font.size = Pt(8)
        ref_label.font.color.rgb = RGBColor(0, 0, 0)

        if report.report_reference:
            ref_value = ref_para.add_run(report.report_reference)
            ref_value.font.size = Pt(8)
            ref_value.font.color.rgb = RGBColor(0, 0, 0)

        # Add spacing and date on right
        ref_para.add_run("\t\t\t\t")

        date_label = ref_para.add_run("Date: ")
        date_label.bold = True
        date_label.font.size = Pt(8)
        date_label.font.color.rgb = RGBColor(0, 0, 0)

        if report.report_date:
            date_value = ref_para.add_run(report.report_date)
        else:
            date_value = ref_para.add_run(datetime.now().strftime('%Y-%m-%d'))
        date_value.font.size = Pt(8)
        date_value.font.color.rgb = RGBColor(0, 0, 0)

        # Add spacing after letterhead
        spacing_para = doc.add_paragraph()
        spacing_para.paragraph_format.space_before = Pt(8)
        spacing_para.paragraph_format.space_after = Pt(0)
