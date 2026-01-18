"""
Modern Professional letterhead template.

This template features a left-aligned layout with a single thick border
and horizontal contact information flow.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

from .base import BaseLetterheadTemplate, TemplateMetadata


def add_border_to_paragraph(paragraph, border_position="bottom", size=12, color="000000"):
    """Add a border to a paragraph"""
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')

    border = OxmlElement(f'w:{border_position}')
    border.set(qn('w:val'), 'single')
    border.set(qn('w:sz'), str(size))
    border.set(qn('w:space'), '1')
    border.set(qn('w:color'), color)

    pBdr.append(border)
    pPr.append(pBdr)


class ModernTemplate(BaseLetterheadTemplate):
    """
    Modern Professional letterhead template.

    Features:
    - Left-aligned layout
    - Single thick top border (8pt weight)
    - Horizontal contact information flow
    - Professional fonts: Name 10pt bold, credentials 8pt, contact 7pt
    - Contemporary, asymmetric design
    """

    def get_metadata(self) -> TemplateMetadata:
        """Return metadata for the Modern template."""
        return TemplateMetadata(
            template_id="modern",
            name="Modern Professional",
            description="Contemporary left-aligned layout with single thick border and horizontal contact flow",
            category="modern"
        )

    def render_letterhead(self, doc: Document, user, report) -> None:
        """
        Render the Modern letterhead design.

        Args:
            doc: python-docx Document object
            user: User model instance
            report: Report model instance
        """
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        # Top thick border
        border_para = doc.add_paragraph()
        border_para.paragraph_format.space_before = Pt(0)
        border_para.paragraph_format.space_after = Pt(4)
        add_border_to_paragraph(border_para, "top", 8, "000000")

        # Professional Name (Left-aligned, Bold, Larger)
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        name_para.paragraph_format.space_before = Pt(0)
        name_para.paragraph_format.space_after = Pt(1)
        name_run = name_para.add_run()
        if user.honorific:
            name_run.add_text(f"{user.honorific} ")
        name_run.add_text(user.full_name)
        name_run.bold = True
        name_run.font.size = Pt(14)
        name_run.font.color.rgb = RGBColor(0, 0, 0)

        # Academic Qualifications (Left-aligned)
        if user.academic_qualifications:
            qual_para = doc.add_paragraph()
            qual_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            qual_para.paragraph_format.space_before = Pt(0)
            qual_para.paragraph_format.space_after = Pt(1)
            qual_run = qual_para.add_run(user.academic_qualifications)
            qual_run.font.size = Pt(11)
            qual_run.font.color.rgb = RGBColor(0, 0, 0)

        # Professional Designation (Left-aligned, Bold)
        if user.professional_designation:
            desig_para = doc.add_paragraph()
            desig_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            desig_para.paragraph_format.space_before = Pt(0)
            desig_para.paragraph_format.space_after = Pt(1)
            desig_run = desig_para.add_run(user.professional_designation)
            desig_run.bold = True
            desig_run.font.size = Pt(11)
            desig_run.font.color.rgb = RGBColor(0, 0, 0)

        # Membership Information (Left-aligned)
        if user.membership_level or user.membership_number:
            member_para = doc.add_paragraph()
            member_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            member_para.paragraph_format.space_before = Pt(0)
            member_para.paragraph_format.space_after = Pt(1)
            member_text = []
            if user.membership_level:
                member_text.append(user.membership_level)
            if user.membership_number:
                member_text.append(user.membership_number)
            member_run = member_para.add_run(" | ".join(member_text))
            member_run.font.size = Pt(11)
            member_run.font.color.rgb = RGBColor(0, 0, 0)

        # Panel Valuer Status (Left-aligned)
        if user.panel_valuer_banks and len(user.panel_valuer_banks) > 0:
            panel_para = doc.add_paragraph()
            panel_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            panel_para.paragraph_format.space_before = Pt(0)
            panel_para.paragraph_format.space_after = Pt(2)
            banks_text = ", ".join(user.panel_valuer_banks)
            panel_run = panel_para.add_run(f"Panel Valuer: {banks_text}")
            panel_run.font.size = Pt(11)
            panel_run.font.color.rgb = RGBColor(0, 0, 0)

        # Horizontal Contact Information (Single line with pipes)
        contact_parts = []

        # Residence
        res_address_parts = []
        if user.house_number:
            res_address_parts.append(user.house_number)
        if user.area_development:
            res_address_parts.append(user.area_development)
        if user.locality:
            res_address_parts.append(user.locality)
        if res_address_parts:
            contact_parts.append(f"Residence: {', '.join(res_address_parts)}")

        # Phones
        phones = []
        if user.phone_primary:
            phones.append(user.phone_primary)
        if user.phone_secondary:
            phones.append(user.phone_secondary)
        if phones:
            contact_parts.append(f"Tel: {' / '.join(phones)}")

        # Office
        if user.office_department or user.office_region or user.office_street_city:
            office_parts = []
            if user.office_department:
                office_parts.append(user.office_department)
            if user.office_region:
                office_parts.append(user.office_region)
            if user.office_street_city:
                office_parts.append(user.office_street_city)
            contact_parts.append(f"Office: {', '.join(office_parts)}")

        # Office phone
        if user.office_phone:
            contact_parts.append(f"Office Tel: {user.office_phone}")

        # Email
        if user.email:
            contact_parts.append(f"Email: {user.email}")

        # Add contact line
        if contact_parts:
            contact_para = doc.add_paragraph()
            contact_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            contact_para.paragraph_format.space_before = Pt(2)
            contact_para.paragraph_format.space_after = Pt(2)
            contact_run = contact_para.add_run(" | ".join(contact_parts))
            contact_run.font.size = Pt(11)
            contact_run.font.color.rgb = RGBColor(0, 0, 0)

        # Bottom separator
        separator_para = doc.add_paragraph()
        separator_para.paragraph_format.space_before = Pt(2)
        separator_para.paragraph_format.space_after = Pt(2)
        add_border_to_paragraph(separator_para, "bottom", 6, "000000")

        # Reference and Date Line (same line, left and right edges)
        ref_date_para = doc.add_paragraph()
        ref_date_para.paragraph_format.space_before = Pt(0)
        ref_date_para.paragraph_format.space_after = Pt(0)

        # Ref on the left
        ref_label = ref_date_para.add_run("Ref: ")
        ref_label.bold = True
        ref_label.font.size = Pt(11)
        ref_label.font.color.rgb = RGBColor(0, 0, 0)

        if report.report_reference:
            ref_value = ref_date_para.add_run(report.report_reference)
            ref_value.font.size = Pt(11)
            ref_value.font.color.rgb = RGBColor(0, 0, 0)

        # Add spacing tabs to push date to the right
        ref_date_para.add_run("\t\t\t\t")

        # Date on the right
        date_label = ref_date_para.add_run("Date: ")
        date_label.bold = True
        date_label.font.size = Pt(11)
        date_label.font.color.rgb = RGBColor(0, 0, 0)

        if report.report_date:
            date_value = ref_date_para.add_run(report.report_date)
        else:
            date_value = ref_date_para.add_run(datetime.now().strftime('%Y-%m-%d'))
        date_value.font.size = Pt(11)
        date_value.font.color.rgb = RGBColor(0, 0, 0)

        # Add spacing after letterhead
        spacing_para = doc.add_paragraph()
        spacing_para.paragraph_format.space_before = Pt(8)
        spacing_para.paragraph_format.space_after = Pt(0)
