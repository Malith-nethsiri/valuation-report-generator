"""
Classic Professional letterhead template.

This template represents the traditional, centered layout with dual borders
and structured two-column contact sections. This is the original letterhead design.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

from .base import BaseLetterheadTemplate, TemplateMetadata
from ..docx_generation.styling import add_border_to_paragraph


class ClassicTemplate(BaseLetterheadTemplate):
    """
    Classic Professional letterhead template.

    Features:
    - Centered layout throughout
    - Dual horizontal borders (6pt weight)
    - Two-column contact section (Residence | Office)
    - Professional fonts: Name 9pt bold, credentials 7-8pt, contact 7pt
    - Line spacing 0.9-1.0 for compact layout
    - Two-column Ref/Date table
    """

    def get_metadata(self) -> TemplateMetadata:
        """Return metadata for the Classic template."""
        return TemplateMetadata(
            template_id="classic",
            name="Classic Professional",
            description="Traditional centered layout with dual borders and structured contact sections",
            category="traditional"
        )

    def render_letterhead(self, doc: Document, user, report) -> None:
        """
        Render the Classic letterhead design.

        Args:
            doc: python-docx Document object
            user: User model instance
            report: Report model instance
        """
        # Set document margins to be smaller at top
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        # Professional Name & Title (Centered, Bold)
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_para.paragraph_format.space_before = Pt(0)
        name_para.paragraph_format.space_after = Pt(0)
        name_para.paragraph_format.line_spacing = 1.0
        name_run = name_para.add_run()
        if user.honorific:
            name_run.add_text(f"{user.honorific} ")
        name_run.add_text(user.full_name)
        name_run.bold = True
        name_run.font.size = Pt(14)
        name_run.font.color.rgb = RGBColor(0, 0, 0)

        # Academic Qualifications (Centered)
        if user.academic_qualifications:
            qual_para = doc.add_paragraph()
            qual_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            qual_para.paragraph_format.space_before = Pt(0)
            qual_para.paragraph_format.space_after = Pt(0)
            qual_para.paragraph_format.line_spacing = 0.9
            qual_run = qual_para.add_run(user.academic_qualifications)
            qual_run.font.size = Pt(11)
            qual_run.font.color.rgb = RGBColor(0, 0, 0)

        # Professional Designation (Centered, Bold)
        if user.professional_designation:
            desig_para = doc.add_paragraph()
            desig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            desig_para.paragraph_format.space_before = Pt(0)
            desig_para.paragraph_format.space_after = Pt(0)
            desig_para.paragraph_format.line_spacing = 0.9
            desig_run = desig_para.add_run(user.professional_designation)
            desig_run.bold = True
            desig_run.font.size = Pt(11)
            desig_run.font.color.rgb = RGBColor(0, 0, 0)

        # Membership Information (Centered)
        if user.membership_level or user.membership_number:
            member_para = doc.add_paragraph()
            member_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            member_para.paragraph_format.space_before = Pt(0)
            member_para.paragraph_format.space_after = Pt(0)
            member_para.paragraph_format.line_spacing = 0.9
            member_text = []
            if user.membership_level:
                member_text.append(user.membership_level)
            if user.membership_number:
                member_text.append(user.membership_number)
            member_run = member_para.add_run(" | ".join(member_text))
            member_run.font.size = Pt(11)
            member_run.font.color.rgb = RGBColor(0, 0, 0)

        # Panel Valuer Status (Centered)
        if user.panel_valuer_banks and len(user.panel_valuer_banks) > 0:
            panel_para = doc.add_paragraph()
            panel_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            panel_para.paragraph_format.space_before = Pt(0)
            panel_para.paragraph_format.space_after = Pt(0)
            panel_para.paragraph_format.line_spacing = 0.9
            banks_text = ", ".join(user.panel_valuer_banks)
            panel_run = panel_para.add_run(f"Panel Valuer: {banks_text}")
            panel_run.font.size = Pt(11)
            panel_run.font.color.rgb = RGBColor(0, 0, 0)

        # Border separator
        separator_para = doc.add_paragraph()
        separator_para.paragraph_format.space_before = Pt(1)
        separator_para.paragraph_format.space_after = Pt(1)
        separator_para.paragraph_format.line_spacing = 1.0
        add_border_to_paragraph(separator_para, "bottom", 6, "000000")

        # Contact Information Section (Two columns)
        contact_table = doc.add_table(rows=1, cols=2)
        contact_table.autofit = False
        contact_table.allow_autofit = False

        # RESIDENCE Column
        residence_cell = contact_table.rows[0].cells[0]
        residence_para = residence_cell.paragraphs[0]
        residence_para.paragraph_format.space_before = Pt(0)
        residence_para.paragraph_format.space_after = Pt(0)
        residence_para.paragraph_format.line_spacing = 0.9
        residence_heading = residence_para.add_run("RESIDENCE")
        residence_heading.bold = True
        residence_heading.font.size = Pt(11)
        residence_heading.font.color.rgb = RGBColor(0, 0, 0)

        # Build residential address - combine into one line
        res_address_parts = []
        if user.house_number:
            res_address_parts.append(user.house_number)
        if user.area_development:
            res_address_parts.append(user.area_development)
        if user.village:
            res_address_parts.append(user.village)
        if user.locality:
            res_address_parts.append(user.locality)

        if res_address_parts:
            res_run = residence_para.add_run("\n" + ", ".join(res_address_parts))
            res_run.font.size = Pt(11)
            res_run.font.color.rgb = RGBColor(0, 0, 0)

        # Residential phones - combine with separator
        phones = []
        if user.phone_primary:
            phones.append(user.phone_primary)
        if user.phone_secondary:
            phones.append(user.phone_secondary)
        if phones:
            phone_run = residence_para.add_run("\n" + " / ".join(phones))
            phone_run.font.size = Pt(11)
            phone_run.font.color.rgb = RGBColor(0, 0, 0)

        # OFFICE Column
        office_cell = contact_table.rows[0].cells[1]
        office_para = office_cell.paragraphs[0]
        office_para.paragraph_format.space_before = Pt(0)
        office_para.paragraph_format.space_after = Pt(0)
        office_para.paragraph_format.line_spacing = 0.9
        office_heading = office_para.add_run("OFFICE")
        office_heading.bold = True
        office_heading.font.size = Pt(11)
        office_heading.font.color.rgb = RGBColor(0, 0, 0)

        # Build office address - combine into one line
        office_address_parts = []
        if user.office_department:
            office_address_parts.append(user.office_department)
        if user.office_region:
            office_address_parts.append(user.office_region)
        if user.office_street_city:
            office_address_parts.append(user.office_street_city)

        if office_address_parts:
            office_run = office_para.add_run("\n" + ", ".join(office_address_parts))
            office_run.font.size = Pt(11)
            office_run.font.color.rgb = RGBColor(0, 0, 0)

        # Office phone
        if user.office_phone:
            office_phone_run = office_para.add_run(f"\n{user.office_phone}")
            office_phone_run.font.size = Pt(11)
            office_phone_run.font.color.rgb = RGBColor(0, 0, 0)

        # Email (Centered below contact info)
        if user.email:
            email_para = doc.add_paragraph()
            email_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            email_para.paragraph_format.space_before = Pt(1)
            email_para.paragraph_format.space_after = Pt(1)
            email_para.paragraph_format.line_spacing = 0.9
            email_run = email_para.add_run(user.email)
            email_run.font.size = Pt(11)
            email_run.font.color.rgb = RGBColor(0, 0, 0)
            email_run.underline = True

        # Border separator
        separator_para2 = doc.add_paragraph()
        separator_para2.paragraph_format.space_before = Pt(1)
        separator_para2.paragraph_format.space_after = Pt(1)
        separator_para2.paragraph_format.line_spacing = 1.0
        add_border_to_paragraph(separator_para2, "bottom", 6, "000000")

        # Reference and Date Line
        ref_table = doc.add_table(rows=1, cols=2)
        ref_table.autofit = False

        ref_cell = ref_table.rows[0].cells[0]
        ref_cell_para = ref_cell.paragraphs[0]
        ref_cell_para.paragraph_format.space_before = Pt(0)
        ref_cell_para.paragraph_format.space_after = Pt(0)
        ref_cell_para.paragraph_format.line_spacing = 0.9
        ref_label = ref_cell_para.add_run("Ref: ")
        ref_label.bold = True
        ref_label.font.size = Pt(11)
        ref_label.font.color.rgb = RGBColor(0, 0, 0)

        if report.report_reference:
            ref_value = ref_cell_para.add_run(report.report_reference)
            ref_value.font.size = Pt(11)
            ref_value.font.color.rgb = RGBColor(0, 0, 0)

        date_cell = ref_table.rows[0].cells[1]
        date_cell_para = date_cell.paragraphs[0]
        date_cell_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_cell_para.paragraph_format.space_before = Pt(0)
        date_cell_para.paragraph_format.space_after = Pt(0)
        date_cell_para.paragraph_format.line_spacing = 0.9
        date_label = date_cell_para.add_run("Date: ")
        date_label.bold = True
        date_label.font.size = Pt(11)
        date_label.font.color.rgb = RGBColor(0, 0, 0)

        if report.report_date:
            date_value = date_cell_para.add_run(report.report_date)
            date_value.font.size = Pt(11)
            date_value.font.color.rgb = RGBColor(0, 0, 0)
        else:
            date_value = date_cell_para.add_run(datetime.now().strftime('%Y-%m-%d'))
            date_value.font.size = Pt(11)
            date_value.font.color.rgb = RGBColor(0, 0, 0)

        # Add a small spacing after letterhead before content
        spacing_para = doc.add_paragraph()
        spacing_para.paragraph_format.space_before = Pt(8)
        spacing_para.paragraph_format.space_after = Pt(0)
        spacing_para.paragraph_format.line_spacing = 1.0
