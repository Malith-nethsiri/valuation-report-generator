"""
Building accommodation aggregation and construction/utilities rendering for DOCX generation.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from io import BytesIO
from datetime import datetime
from typing import Any, List, Dict, Optional, Union
import requests
from .. import models
from ..utils import append_label_if_missing, clean_spelling_errors, format_no_field
from ..letterhead_templates import get_template
import logging

from .formatting import (
    format_currency, format_currency_words, format_currency_aligned,
    format_room_count, round_for_say, format_material_list,
)
from .paragraph_builders import (
    add_section_heading, add_market_value_line, add_value_rounded_line,
    add_inline_field, add_subsection_paragraph, format_building_valuation_2line, format_addon_compact,
)
from .styling import (
    add_border_to_paragraph,
    MAP_IMAGE_WIDTH, MAP_IMAGE_MAX_HEIGHT, PROPERTY_PHOTO_WIDTH, PROPERTY_PHOTO_HEIGHT,
    IMAGE_SPACING_BEFORE, IMAGE_SPACING_AFTER, MAJOR_SECTION_SPACE_BEFORE, MAJOR_SECTION_SPACE_AFTER,
    SUBSECTION_SPACE_BEFORE, SUBSECTION_SPACE_AFTER, BODY_PARA_SPACE_BEFORE, BODY_PARA_SPACE_AFTER,
    INLINE_FIELD_SPACE_BEFORE, INLINE_FIELD_SPACE_AFTER, SUBHEADING_SPACE_BEFORE, SUBHEADING_SPACE_AFTER,
    INDENTED_CONTENT_SPACE_BEFORE, INDENTED_CONTENT_SPACE_AFTER, INDENTED_CONTENT_LEFT_INDENT,
    BOUNDARY_LIST_SPACE_AFTER, ACCOMMODATION_ROOM_SPACE_AFTER,
    FONT_SIZE_DOCUMENT_TITLE, FONT_SIZE_SECTION_HEADING, FONT_SIZE_SUBSECTION_HEADING,
    FONT_SIZE_BODY, FONT_SIZE_INLINE_LABEL, FONT_SIZE_VALUATION,
    FONT_SIZE_TABLE_HEADER, FONT_SIZE_TABLE_CELL, FONT_SIZE_INVOICE_TOTAL,
    FONT_SIZE_CAPTION, FONT_SIZE_BANK_HEADER, FONT_SIZE_BANK_DETAILS,
    FONT_SIZE_SIGNATURE, FONT_SIZE_CERTIFICATION,
)
from .images import calculate_image_dimensions, apply_letterbox_to_image
from .helpers import (
    safe_get_json_field, safe_get_array_item, to_float, safe_parse_json_string, safe_get_nested,
)

logger = logging.getLogger(__name__)
Deed_Type = Union[Dict[str, Any], Any]

def aggregate_accommodation_across_building(building: Dict) -> Dict:
    """
    Get accommodation summary from building (new structure) or aggregate from floors (old structure).

    NEW STRUCTURE: building.accommodation_summary (rooms at building level)
    OLD STRUCTURE: building.floors[].accommodation_summary (rooms per floor)

    Args:
        building: Building dict with either new or old structure

    Returns:
        Dict with aggregated room counts by type
    """
    # NEW STRUCTURE: Check if building has accommodation_summary at building level
    if building.get('accommodation_summary'):
        return building['accommodation_summary']

    # OLD STRUCTURE: Fallback to aggregating from floors (backward compatibility)
    floors = building.get('floors', [])
    totals = {
        'bedrooms': 0, 'bathrooms': 0, 'living_rooms': 0,
        'dining_rooms': 0, 'kitchens': 0, 'pantries': 0,
        'verandahs': 0, 'balconies': 0, 'garages': 0,
        'store_rooms': 0, 'other_rooms': 0
    }

    for floor in floors:
        summary = floor.get('accommodation_summary')
        if summary:
            for key in totals.keys():
                totals[key] += summary.get(key, 0)

    return totals


def deduplicate_water_sources(conveniences: List[str]) -> List[str]:
    """
    Merge duplicate water source entries in conveniences list.

    Business rules:
    - If both 'water' (well) and 'pipe_water' exist: prefer pipe-borne, remove well
    - Rationale: Pipe-borne water is superior/preferred utility

    Args:
        conveniences: List of convenience codes

    Returns:
        Cleaned list with duplicates removed
    """
    has_well = 'water' in conveniences
    has_pipe = 'pipe_water' in conveniences

    if has_well and has_pipe:
        # Prefer pipe-borne, remove well
        return [c for c in conveniences if c != 'water']

    return conveniences


def render_construction_details(doc, building: Dict) -> None:
    """
    Render construction materials section for a building.
    Uses professional sentence-based format for new data structure.
    Maintains backward compatibility with old format.
    """
    construction = building.get('construction_materials', {})
    if not construction:
        return

    sentences = []

    # NEW FORMAT: Check for new enhanced structure
    roof_details = construction.get('roof_details', {})
    wall_details = construction.get('wall_details', {})
    floor_details = construction.get('floor_details', {})

    # Roof Construction (NEW FORMAT)
    if roof_details and (roof_details.get('structure_type') or roof_details.get('covering_material')):
        roof_parts = []

        structure_type = roof_details.get('structure_type')
        if structure_type:
            structure_labels = {
                'timber_frame': 'timber frame',
                'steel_trusses': 'steel trusses',
                'concrete_slab': 'concrete slab',
                'rcc_flat': 'RCC flat',
                'prefab_trusses': 'prefabricated trusses',
                'mixed': 'mixed construction'
            }
            roof_parts.append(structure_labels.get(structure_type, structure_type.replace('_', ' ')))

        covering_materials = roof_details.get('covering_material', [])
        if covering_materials:
            covering_labels = {
                'asbestos_sheets': 'asbestos sheets',
                'clay_tiles': 'clay tiles',
                'concrete_tiles': 'concrete tiles',
                'metal_sheets': 'metal sheets',
                'concrete_flat': 'concrete flat',
                'cadjans': 'cadjans',
                'shingles': 'shingles'
            }
            covering_list = [covering_labels.get(m, m.replace('_', ' ')) for m in covering_materials]
            if covering_list:
                roof_parts.append(f"covered with {format_list_with_grammar(covering_list)}")

        additional = roof_details.get('additional_details')
        if additional:
            roof_parts.append(additional.lower() if additional else '')

        if roof_parts:
            sentences.append(f"The roof is constructed with {' '.join(roof_parts)}.")

    # Wall Construction (NEW FORMAT)
    if wall_details and wall_details.get('material'):
        wall_parts = []

        material = wall_details.get('material')
        if material:
            material_labels = {
                'brick_masonry': 'brick masonry',
                'cement_block': 'cement block',
                'stone_masonry': 'stone masonry',
                'rcc_frame': 'RCC frame',
                'rubble_masonry': 'rubble masonry',
                'mud_walls': 'mud walls',
                'timber_frame': 'timber frame',
                'cadjan': 'cadjan',
                'prefab_panels': 'prefabricated panels',
                'mixed': 'mixed materials'
            }
            wall_parts.append(material_labels.get(material, material.replace('_', ' ')))

        finishes = wall_details.get('finish', [])
        if finishes:
            finish_labels = {
                'cement_plaster_painted': 'cement plaster and painted',
                'lime_plaster': 'lime plaster',
                'tiles': 'tiles',
                'exposed_brick': 'exposed brick',
                'color_washed': 'color washed',
                'textured': 'textured finish',
                'unfinished': 'unfinished'
            }
            finish_list = [finish_labels.get(f, f.replace('_', ' ')) for f in finishes]
            if finish_list:
                wall_parts.append(f"finished with {format_list_with_grammar(finish_list)}")

        if wall_parts:
            sentences.append(f"The walls are {' '.join(wall_parts)}.")

    # Floor Construction (NEW FORMAT)
    if floor_details and floor_details.get('material'):
        floor_materials = floor_details.get('material', [])
        if floor_materials:
            material_labels = {
                'cement': 'cement',
                'tiled': 'tiles',
                'terrazzo': 'terrazzo',
                'timber': 'timber',
                'granite': 'granite',
                'marble': 'marble',
                'vinyl': 'vinyl',
                'polished_concrete': 'polished concrete',
                'earth': 'earth'
            }
            material_list = [material_labels.get(m, m.replace('_', ' ')) for m in floor_materials]

            floor_sentence = f"The floor is {format_list_with_grammar(material_list)} rendered"

            quality = floor_details.get('finish_quality')
            if quality:
                quality_labels = {'basic': 'basic', 'standard': 'good', 'premium': 'premium'}
                floor_sentence += f" with {quality_labels.get(quality, quality)} quality finish"

            sentences.append(f"{floor_sentence}.")

    # Ceiling (KEEP EXISTING - works for both old and new format)
    ceiling = construction.get('ceiling_type', '')
    if ceiling and ceiling != 'none':
        ceiling_labels = {
            'gypsum_board': 'gypsum board',
            'plywood': 'plywood',
            'asbestos': 'asbestos',
            'pvc': 'PVC panels',
            'timber': 'timber planks'
        }
        sentences.append(f"The ceiling is finished with {ceiling_labels.get(ceiling, ceiling.replace('_', ' '))}.")

    # Doors & Windows (NEW FORMAT)
    doors_windows = construction.get('doors_windows_details', {})
    if doors_windows:
        dw_parts = []

        # Windows
        window_frame = doors_windows.get('window_frame_material', [])
        window_glass = doors_windows.get('window_glass_type', [])
        window_security = doors_windows.get('window_security', [])

        if window_frame or window_glass or window_security:
            window_parts = []

            if window_frame:
                frame_labels = {
                    'timber': 'timber',
                    'aluminum': 'aluminum',
                    'upvc': 'UPVC',
                    'steel': 'steel'
                }
                frame_list = [frame_labels.get(f, f) for f in window_frame]
                window_parts.append(f"{format_list_with_grammar(frame_list)} frame windows")

            if window_glass:
                glass_labels = {
                    'clear': 'clear glass',
                    'tinted': 'tinted glass',
                    'frosted': 'frosted glass',
                    'double_glazing': 'double glazing'
                }
                glass_list = [glass_labels.get(g, g) for g in window_glass]
                window_parts.append(f"with {format_list_with_grammar(glass_list)}")

            if window_security:
                security_labels = {
                    'burglar_bars': 'burglar bars',
                    'grills': 'security grills',
                    'security_mesh': 'security mesh'
                }
                security_list = [security_labels.get(s, s) for s in window_security]
                window_parts.append(f"fitted with {format_list_with_grammar(security_list)}")

            if window_parts:
                dw_parts.append(f"The building has {' '.join(window_parts)}.")

        # Doors
        main_door = doors_windows.get('main_door_material')
        internal_door = doors_windows.get('internal_door_material')
        door_security = doors_windows.get('door_security', [])

        if main_door or internal_door:
            door_parts = []

            if main_door:
                main_door_labels = {
                    'solid_timber': 'solid timber',
                    'panel_door': 'timber panel',
                    'metal': 'metal',
                    'upvc': 'UPVC',
                    'glass': 'glass'
                }
                door_parts.append(f"main door is {main_door_labels.get(main_door, main_door.replace('_', ' '))}")

            if internal_door:
                internal_door_labels = {
                    'timber': 'solid timber',
                    'flush_doors': 'flush doors',
                    'panel_doors': 'panel doors',
                    'hollow_core': 'hollow core'
                }
                door_parts.append(f"internal doors are {internal_door_labels.get(internal_door, internal_door.replace('_', ' '))}")

            door_sentence = f"The {' and '.join(door_parts)}"

            if door_security:
                security_labels = {
                    'deadbolt': 'deadbolt locks',
                    'security_locks': 'security locks',
                    'door_chain': 'door chains',
                    'multi_point': 'multi-point locking systems'
                }
                security_list = [security_labels.get(s, s) for s in door_security]
                door_sentence += f" with {format_list_with_grammar(security_list)}"

            dw_parts.append(f"{door_sentence}.")

        sentences.extend(dw_parts)

    # BACKWARD COMPATIBILITY: Fall back to old format if new format not available
    if not sentences:
        # Old format handling (existing logic)
        parts = []

        # Foundation (old)
        foundation = construction.get('foundation_type', '')
        if foundation:
            foundation_labels = {
                'rubble_masonry': 'Rubble masonry',
                'concrete_strip': 'Concrete strip',
                'pile': 'Pile',
                'raft': 'Raft',
                'pad': 'Pad foundation'
            }
            parts.append(f"Foundation: {foundation_labels.get(foundation, foundation.replace('_', ' ').title())}")

        # Walls (old)
        wall = construction.get('wall_construction', {})
        if wall:
            wall_parts = []
            material = wall.get('material', '')
            if material:
                wall_material_labels = {
                    'brick_masonry': 'brick masonry',
                    'cement_block': 'cement block',
                    'stone': 'stone masonry',
                    'rcc_columns': 'RCC columns',
                    'cadjan': 'cadjan',
                    'timber': 'timber frame'
                }
                wall_parts.append(wall_material_labels.get(material, material.replace('_', ' ')))

            thickness = wall.get('thickness', '')
            if thickness:
                thickness_labels = {
                    '4_5_inch': '4.5 inch',
                    '9_inch': '9 inch',
                    '13_5_inch': '13.5 inch'
                }
                wall_parts.append(thickness_labels.get(thickness, thickness.replace('_', ' ')))

            finish = wall.get('finish', '')
            if finish:
                finish_labels = {
                    'plastered_painted': 'plastered and painted',
                    'fair_faced': 'fair-faced',
                    'textured': 'textured finish'
                }
                wall_parts.append(finish_labels.get(finish, finish.replace('_', ' ')))

            if wall_parts:
                parts.append(f"Walls: {', '.join(wall_parts)}")

        # Roof (old)
        roof = construction.get('roof_structure', '')
        if roof:
            roof_labels = {
                'timber_frame': 'Timber frame',
                'steel_trusses': 'Steel trusses',
                'rcc_slab': 'RCC slab',
                'cadjan': 'Cadjan'
            }
            parts.append(f"Roof: {roof_labels.get(roof, roof.replace('_', ' ').title())}")

        # Ceiling (old - already handled above)

        if parts:
            construction_text = "; ".join(parts) + "."
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
            para.paragraph_format.line_spacing = 0.9
            run = para.add_run(construction_text)
            run.font.size = FONT_SIZE_BODY
            run.font.color.rgb = RGBColor(0, 0, 0)
            return

    # Render new format as standalone paragraph (no label)
    if sentences:
        construction_text = " ".join(sentences)
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
        para.paragraph_format.line_spacing = 0.9
        run = para.add_run(construction_text)
        run.font.size = FONT_SIZE_BODY
        run.font.color.rgb = RGBColor(0, 0, 0)


def render_utilities_and_conveniences(doc, building: Dict) -> None:
    """
    Render unified utilities and conveniences section for a building.
    Includes water, electricity, sewage, parking, communication, gas, security, and amenities.
    """
    utilities = building.get('utilities_services', {})
    if not utilities:
        return

    parts = []

    # Water supply
    water = utilities.get('water_supply', '')
    if water:
        water_labels = {
            'Pipe-borne (NWSDB)': 'pipe-borne water (NWSDB)',
            'Well': 'well water',
            'Bore/Tube Well': 'bore/tube well water',
            'Rainwater Harvesting': 'rainwater harvesting',
            # Backward compatibility
            'pipe_borne': 'pipe-borne water (NWSDB)',
            'well': 'well water',
            'tube_well': 'tube well'
        }

        # Handle both string (legacy) and array (new) formats
        if isinstance(water, str):
            # Legacy single value
            parts.append(water_labels.get(water, water.replace('_', ' ')))
        elif isinstance(water, list):
            # New array format - use existing format_list_with_grammar
            mapped = [water_labels.get(w, w.lower()) for w in water]
            parts.append(format_list_with_grammar(mapped))

    # Sewage system
    sewage = utilities.get('sewage_system', '')
    if sewage:
        sewage_labels = {
            'municipal_sewer': 'municipal sewage system',
            'septic_tank': 'septic tank',
            'none': 'no sewage system'
        }
        if sewage != 'none':
            parts.append(sewage_labels.get(sewage, sewage.replace('_', ' ')))

    # Electricity
    elec = utilities.get('electricity', {})
    if elec:
        source = elec.get('source', '')
        if source:
            source_text = 'CEB electricity' if source == 'ceb' else f'{source.title()} electricity'
            if elec.get('three_phase'):
                source_text += ' (three-phase connection)'
            parts.append(source_text)

    # Parking
    parking = utilities.get('parking') or {}
    covered = int(to_float(parking.get('covered_spaces')))
    uncovered = int(to_float(parking.get('uncovered_spaces')))
    total_parking = covered + uncovered
    if total_parking > 0:
        parking_text = f"parking for {total_parking} vehicle{'s' if total_parking != 1 else ''}"
        if covered > 0 and uncovered > 0:
            parking_text += f" ({covered} covered, {uncovered} open)"
        elif covered > 0:
            parking_text += " (covered)"
        parts.append(parking_text)

    # Communication services
    comm_parts = []
    if utilities.get('telephone'):
        comm_parts.append('telephone connection')
    if utilities.get('internet'):
        comm_parts.append('internet connection')
    if comm_parts:
        parts.append(", ".join(comm_parts))

    # Gas connection
    if utilities.get('gas_connection'):
        parts.append('gas connection')

    # Security features
    security = utilities.get('security_features', [])
    if security:
        security_labels = {
            'boundary_wall': 'boundary wall',
            'main_gate': 'main gate',
            'cctv': 'CCTV surveillance',
            'security_lights': 'security lighting'
        }
        security_list = [security_labels.get(s, s.replace('_', ' ')) for s in security]
        if security_list:
            parts.append(", ".join(security_list))

    # Amenities
    amenities = utilities.get('amenities') or {}
    amenity_list = []
    if amenities.get('air_conditioning'):
        amenity_list.append('air conditioning')
    if amenities.get('built_in_wardrobes'):
        amenity_list.append('built-in wardrobes')
    if amenities.get('modern_kitchen'):
        amenity_list.append('modern kitchen fittings')
    if amenity_list:
        parts.append(", ".join(amenity_list))

    # Hot water system
    hot_water = utilities.get('hot_water_system', '')
    if hot_water and hot_water != 'none':
        hot_water_labels = {
            'electric_geyser': 'electric geyser',
            'solar_heater': 'solar water heater',
            'gas_heater': 'gas water heater'
        }
        parts.append(hot_water_labels.get(hot_water, hot_water.replace('_', ' ')))

    if parts:
        # Capitalize first letter
        utilities_text = "; ".join(parts)
        utilities_text = utilities_text[0].upper() + utilities_text[1:] + "."
        add_inline_field(doc, "Utilities & Conveniences", utilities_text)


