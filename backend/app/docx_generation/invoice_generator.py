"""
Invoice section generation for DOCX valuation reports.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from io import BytesIO
from datetime import datetime
from typing import Any, List, Dict, Optional, Union
import requests
from .. import models
from ..utils import append_label_if_missing, clean_spelling_errors, format_no_field
from ..letterhead_templates import get_template
import logging

from .formatting import (
    format_currency, format_currency_words, format_currency_aligned,
    format_room_count, round_for_say, format_material_list,
)
from .paragraph_builders import (
    add_section_heading, add_market_value_line, add_value_rounded_line,
    add_inline_field, add_subsection_paragraph, format_building_valuation_2line, format_addon_compact,
)
from .styling import (
    add_border_to_paragraph,
    MAP_IMAGE_WIDTH, MAP_IMAGE_MAX_HEIGHT, PROPERTY_PHOTO_WIDTH, PROPERTY_PHOTO_HEIGHT,
    IMAGE_SPACING_BEFORE, IMAGE_SPACING_AFTER, MAJOR_SECTION_SPACE_BEFORE, MAJOR_SECTION_SPACE_AFTER,
    SUBSECTION_SPACE_BEFORE, SUBSECTION_SPACE_AFTER, BODY_PARA_SPACE_BEFORE, BODY_PARA_SPACE_AFTER,
    INLINE_FIELD_SPACE_BEFORE, INLINE_FIELD_SPACE_AFTER, SUBHEADING_SPACE_BEFORE, SUBHEADING_SPACE_AFTER,
    INDENTED_CONTENT_SPACE_BEFORE, INDENTED_CONTENT_SPACE_AFTER, INDENTED_CONTENT_LEFT_INDENT,
    BOUNDARY_LIST_SPACE_AFTER, ACCOMMODATION_ROOM_SPACE_AFTER,
    FONT_SIZE_DOCUMENT_TITLE, FONT_SIZE_SECTION_HEADING, FONT_SIZE_SUBSECTION_HEADING,
    FONT_SIZE_BODY, FONT_SIZE_INLINE_LABEL, FONT_SIZE_VALUATION,
    FONT_SIZE_TABLE_HEADER, FONT_SIZE_TABLE_CELL, FONT_SIZE_INVOICE_TOTAL,
    FONT_SIZE_CAPTION, FONT_SIZE_BANK_HEADER, FONT_SIZE_BANK_DETAILS,
    FONT_SIZE_SIGNATURE, FONT_SIZE_CERTIFICATION,
)
from .images import calculate_image_dimensions, apply_letterbox_to_image
from .helpers import (
    safe_get_json_field, safe_get_array_item, to_float, safe_parse_json_string, safe_get_nested,
)

logger = logging.getLogger(__name__)
Deed_Type = Union[Dict[str, Any], Any]

def migrate_invoice_data(invoice_dict):
    """Migrate old invoice structure (qty, unit_price) to new structure (direct total)"""
    if not invoice_dict:
        return invoice_dict

    # Check if migration needed
    items = invoice_dict.get('items', [])
    needs_migration = any('quantity' in item or 'unit_price' in item for item in items)

    if not needs_migration:
        return invoice_dict  # Already new format

    # Migrate items - keep total, remove quantity/unit_price
    migrated_items = [
        {'description': item.get('description', ''), 'total': item.get('total', 0)}
        for item in items
    ]

    # Convert bank_details (string) → manual_bank_details
    return {
        'items': migrated_items,
        'subtotal': invoice_dict.get('subtotal', 0),
        'traveling_charges': invoice_dict.get('traveling_charges'),
        'discount': invoice_dict.get('discount'),
        'total': invoice_dict.get('total', 0),
        'bank_account_ids': [],
        'manual_bank_details': invoice_dict.get('bank_details')  # Old string field
    }


def generate_invoice_section(doc, invoice_data, user, report=None):
    """Generate professional fee invoice section with letterhead and recipient address"""
    # Add page break for new invoice page
    doc.add_page_break()

    # ===== LETTERHEAD =====
    # Render letterhead matching the rest of the report
    template_id = user.preferred_letterhead_template or 'classic'
    template = get_template(template_id)
    template.render_letterhead(doc, user, report)

    # Add spacing after letterhead
    doc.add_paragraph()

    # ===== RECIPIENT ADDRESS =====
    if report:
        # Extract recipient information from report
        recipient_para = doc.add_paragraph()
        recipient_run = recipient_para.add_run("To:\n")
        recipient_run.bold = True
        recipient_run.font.size = FONT_SIZE_VALUATION

        # Add applicant name if available
        if hasattr(report, 'applicant_full_name') and report.applicant_full_name:
            name_run = recipient_para.add_run(f"{report.applicant_full_name}\n")
            name_run.font.size = FONT_SIZE_VALUATION

        # Add applicant address if available
        address_parts = []
        if hasattr(report, 'applicant_address_line1') and report.applicant_address_line1:
            address_parts.append(report.applicant_address_line1)
        if hasattr(report, 'applicant_address_line2') and report.applicant_address_line2:
            address_parts.append(report.applicant_address_line2)

        if address_parts:
            address_run = recipient_para.add_run(f"{', '.join(address_parts)}\n")
            address_run.font.size = FONT_SIZE_VALUATION

        # Add spacing after recipient address
        doc.add_paragraph()

    # Invoice heading
    invoice_heading = doc.add_paragraph()
    invoice_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    invoice_heading.paragraph_format.space_before = Pt(0)
    invoice_heading.paragraph_format.space_after = Pt(12)
    invoice_run = invoice_heading.add_run("PROFESSIONAL FEES")
    invoice_run.bold = True
    invoice_run.font.size = FONT_SIZE_BODY
    invoice_run.font.color.rgb = RGBColor(0, 0, 0)

    # Parse and migrate invoice_data
    if isinstance(invoice_data, str):
        import json
        invoice_dict = json.loads(invoice_data)
    else:
        invoice_dict = invoice_data

    # Migrate old format to new format
    invoice_dict = migrate_invoice_data(invoice_dict)

    # Invoice table - SIMPLIFIED (2 columns)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set column widths
    table.columns[0].width = Inches(4.5)  # Description - wider
    table.columns[1].width = Inches(2.0)  # Total

    # Header row
    header_cells = table.rows[0].cells
    headers = ['Description', 'Total (LKR)']
    for i, header_text in enumerate(headers):
        cell = header_cells[i]
        cell.text = header_text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = FONT_SIZE_TABLE_HEADER
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 1 else WD_ALIGN_PARAGRAPH.LEFT
        # Shading for header
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        cell_xml = cell._element
        cell_properties = cell_xml.get_or_add_tcPr()
        shading_element = OxmlElement('w:shd')
        shading_element.set(qn('w:fill'), 'D9D9D9')  # Light gray
        cell_properties.append(shading_element)

    # Item rows
    for item in invoice_dict.get('items', []):
        row_cells = table.add_row().cells
        row_cells[0].text = item.get('description', '')
        row_cells[1].text = format_currency(item.get('total', 0))
        row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = FONT_SIZE_TABLE_CELL

    # Subtotal row
    subtotal_row = table.add_row().cells
    subtotal_row[0].text = "Subtotal:"
    subtotal_row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    subtotal_row[1].text = format_currency(invoice_dict.get('subtotal', 0))
    subtotal_row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Traveling charges if present
    if invoice_dict.get('traveling_charges'):
        travel_row = table.add_row().cells
        travel_row[0].text = "Traveling Charges:"
        travel_row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        travel_row[1].text = format_currency(invoice_dict.get('traveling_charges', 0))
        travel_row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Discount if present
    if invoice_dict.get('discount'):
        discount_row = table.add_row().cells
        discount_row[0].text = "Discount:"
        discount_row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        discount_row[1].text = f"-{format_currency(invoice_dict.get('discount', 0))}"
        discount_row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Grand Total row
    total_row = table.add_row().cells
    total_row[0].text = "GRAND TOTAL:"
    total_row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for paragraph in total_row[0].paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = FONT_SIZE_INVOICE_TOTAL

    total_row[1].text = format_currency(invoice_dict.get('total', 0))
    total_row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for paragraph in total_row[1].paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = FONT_SIZE_INVOICE_TOTAL

    # Shade total row
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    for cell in [total_row[0], total_row[1]]:
        cell_xml = cell._element
        cell_properties = cell_xml.get_or_add_tcPr()
        shading_element = OxmlElement('w:shd')
        shading_element.set(qn('w:fill'), 'E6E6E6')
        cell_properties.append(shading_element)

    doc.add_paragraph()  # Spacing

    # Bank Details - NEW: Show all selected accounts
    bank_account_ids = invoice_dict.get('bank_account_ids', [])
    manual_bank_details = invoice_dict.get('manual_bank_details')

    if bank_account_ids and len(bank_account_ids) > 0 and user:
        # Fetch user's bank accounts
        user_accounts = user.bank_accounts or []
        selected_accounts = [acc for acc in user_accounts if acc['id'] in bank_account_ids]

        if selected_accounts:
            bank_heading = doc.add_paragraph()
            bank_heading_run = bank_heading.add_run("Bank Account Details:")
            bank_heading_run.bold = True
            bank_heading_run.font.size = FONT_SIZE_BANK_HEADER
            bank_heading.paragraph_format.space_before = Pt(6)
            bank_heading.paragraph_format.space_after = Pt(3)

            # Render each account
            for account in selected_accounts:
                acc_para = doc.add_paragraph(style='List Bullet')
                acc_para.paragraph_format.left_indent = Inches(0.25)
                acc_para.paragraph_format.line_spacing = 1.0
                acc_run = acc_para.add_run(
                    f"{account['bank_name']}\n"
                    f"Account Number: {account['account_number']}\n"
                    f"Branch: {account['branch_name']}"
                )
                acc_run.font.size = FONT_SIZE_BANK_DETAILS

    elif manual_bank_details:
        # Fallback to manual entry
        bank_heading = doc.add_paragraph()
        bank_heading_run = bank_heading.add_run("Bank Account Details:")
        bank_heading_run.bold = True
        bank_heading_run.font.size = FONT_SIZE_BANK_HEADER
        bank_heading.paragraph_format.space_before = Pt(6)
        bank_heading.paragraph_format.space_after = Pt(3)

        bank_para = doc.add_paragraph()
        bank_para.paragraph_format.line_spacing = 1.0
        bank_run = bank_para.add_run(manual_bank_details)
        bank_run.font.size = FONT_SIZE_BANK_DETAILS

    # Signature space
    doc.add_paragraph()
    doc.add_paragraph()
    signature_para = doc.add_paragraph()
    signature_para.paragraph_format.space_before = Pt(24)
    signature_run = signature_para.add_run("_" * 40)
    signature_run.font.size = FONT_SIZE_SIGNATURE

    sig_label = doc.add_paragraph()
    sig_label_run = sig_label.add_run("Signature")
    sig_label_run.font.size = FONT_SIZE_SIGNATURE
    sig_label_run.font.italic = True


