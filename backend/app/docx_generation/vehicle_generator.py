"""
Vehicle valuation report DOCX generation.
Mirrors the official Sri Lankan vehicle valuation physical form (2 pages).
"""

import base64
import logging
from io import BytesIO
from typing import Any, List, Optional

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from .. import models
from ..letterhead_templates import get_template
from .helpers import safe_get_json_field, to_float
from .images import calculate_image_dimensions

logger = logging.getLogger(__name__)


# ─────────────────────────────────────────────────────────────────────────────
# Private helpers
# ─────────────────────────────────────────────────────────────────────────────

def _val(v: Any) -> str:
    """Return str(v) or '-' when v is None / empty / 'None'."""
    if v is None or str(v).strip() in ('', 'None', 'null'):
        return '-'
    return str(v).strip()


def _chk(v: Any) -> str:
    """Boolean → '✓' when True, blank when False (no dash)."""
    return '✓' if v else ''


def _yn(v: Any) -> str:
    """Boolean → 'Yes' / 'No'."""
    return 'Yes' if v else 'No'


def _fmt_rs(v: Any) -> str:
    """Format monetary value → 'Rs. 1,234,567.00' or '-' when absent."""
    f = to_float(v)
    if f == 0.0 and v in (None, '', 0, 0.0):
        return '-'
    return f'Rs. {f:,.2f}'


def _cell(cell: Any, text: str, bold: bool = False, size: int = 11) -> None:
    """Set cell text, bold, and font size."""
    cell.text = text
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = bold
            run.font.size = Pt(size)


def _set_col_widths(table: Any, widths: List[float]) -> None:
    """Set column widths (inches) for every cell in each column."""
    for row in table.rows:
        for i, w in enumerate(widths):
            if i < len(row.cells):
                row.cells[i].width = Inches(w)


def _set_row_min_height(row: Any, height_cm: float) -> None:
    """Set a minimum row height via XML (atLeast rule)."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(Cm(height_cm).twips)))
    trHeight.set(qn('w:hRule'), 'atLeast')
    trPr.append(trHeight)


# ─────────────────────────────────────────────────────────────────────────────
# Page 1
# ─────────────────────────────────────────────────────────────────────────────

def _add_page_1(doc: Document, vehicle: models.Vehicle, report: models.Report) -> None:

    # Block 1 — Report Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run('Valuation Report Of Vehicle')
    tr.bold = True
    tr.font.size = Pt(14)
    tr.font.underline = True

    # Block 2 — Header Info Table (4-col, 5 rows)
    # widths: 1.57 | 1.57 | 1.57 | 1.56
    hdr = doc.add_table(rows=5, cols=4)
    hdr.style = 'Table Grid'
    _set_col_widths(hdr, [1.57, 1.57, 1.57, 1.56])

    r0 = hdr.rows[0]
    _cell(r0.cells[0], 'Purpose Of Valuation :', bold=True)
    r0.cells[1].merge(r0.cells[3])
    _cell(r0.cells[1], _val(report.valuation_purpose))

    r1 = hdr.rows[1]
    _cell(r1.cells[0], 'Requested By :', bold=True)
    r1.cells[1].merge(r1.cells[3])
    _cell(r1.cells[1], _val(report.applicant_full_name))

    r2 = hdr.rows[2]
    _cell(r2.cells[0], 'Date :', bold=True)
    _cell(r2.cells[1], _val(report.report_date))
    _cell(r2.cells[2], 'Folio No :', bold=True)
    _cell(r2.cells[3], _val(report.folio_number))

    r3 = hdr.rows[3]
    _cell(r3.cells[0], 'Inspection Place :', bold=True)
    r3.cells[1].merge(r3.cells[3])
    _cell(r3.cells[1], _val(report.inspection_place))

    r4 = hdr.rows[4]
    _cell(r4.cells[0], 'Inspection Date :', bold=True)
    r4.cells[1].merge(r4.cells[3])
    _cell(r4.cells[1], _val(report.inspection_date))

    # Block 3 — Vehicle Photos
    ph_heading = doc.add_paragraph()
    ph_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    phr = ph_heading.add_run('Vehicle Photos')
    phr.bold = True
    phr.font.size = Pt(11)

    vehicle_photos = safe_get_json_field(vehicle, 'vehicle_photos', [])
    if vehicle_photos:
        outer = doc.add_table(rows=1, cols=1)
        outer.style = 'Table Grid'
        _set_col_widths(outer, [6.27])
        outer_cell = outer.cell(0, 0)
        inner = outer_cell.add_table(rows=0, cols=2)
        inner.style = 'Table Grid'
        _set_col_widths(inner, [3.135, 3.135])

        for idx in range(0, len(vehicle_photos), 2):
            row = inner.add_row()
            for col, photo in enumerate(vehicle_photos[idx:idx + 2]):
                cell = row.cells[col]
                img_data = photo.get('image_data', '') if isinstance(photo, dict) else getattr(photo, 'image_data', '')
                caption = photo.get('caption', '') if isinstance(photo, dict) else getattr(photo, 'caption', '')
                if img_data:
                    try:
                        if ',' in img_data:
                            img_data = img_data.split(',')[1]
                        image_bytes = base64.b64decode(img_data)
                        image_stream = BytesIO(image_bytes)
                        dims = calculate_image_dimensions(image_stream, max_width=6 / 2.54, max_height=6 / 2.54)
                        image_stream.seek(0)
                        para = cell.paragraphs[0]
                        run = para.add_run()
                        run.add_picture(image_stream, **dims)
                        if caption:
                            cap_para = cell.add_paragraph()
                            cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            cap_run = cap_para.add_run(caption)
                            cap_run.italic = True
                            cap_run.font.size = Pt(9)
                    except Exception as img_err:
                        logger.warning(f'Could not add vehicle photo: {img_err}')
    else:
        ph_table = doc.add_table(rows=1, cols=1)
        ph_table.style = 'Table Grid'
        _set_col_widths(ph_table, [6.27])
        _set_row_min_height(ph_table.rows[0], 5)

    # Block 4 — Description Of Vehicle heading
    desc_h = doc.add_paragraph()
    desc_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc_h.paragraph_format.space_before = Pt(6)
    desc_h.paragraph_format.space_after = Pt(6)
    dhr = desc_h.add_run('Description Of Vehicle')
    dhr.bold = True
    dhr.font.size = Pt(11)

    # Block 5 — Mileage Row (full-width, 2-col)
    mileage_val = f'{vehicle.mileage:,} {vehicle.mileage_unit or "km"}' if vehicle.mileage else '-'
    mil = doc.add_table(rows=1, cols=2)
    mil.style = 'Table Grid'
    _set_col_widths(mil, [3.50, 2.77])
    _cell(mil.rows[0].cells[0], 'Mileage reading during inspection :', bold=True)
    _cell(mil.rows[0].cells[1], mileage_val)

    # Block 6 — Vehicle Identification Grid (4-col, 9 rows)
    # Label cols 1.75" | Value cols 1.39" alternating; total = 1.75+1.39+1.75+1.38 = 6.27
    # Owner name and NIC are sourced from book_data (OCR-extracted CR Book fields)
    book_data = safe_get_json_field(vehicle, 'book_data', None)
    owner_name = _val(book_data.get('owner_name') if isinstance(book_data, dict) else None)
    owner_nic = _val(book_data.get('owner_nic') if isinstance(book_data, dict) else None)

    id_rows = [
        ('Registered Owner',            owner_name,                                'Owner NIC',               owner_nic),
        ('Provincial Council',         _val(vehicle.provincial_council),         'Registration Number',     _val(vehicle.registration_number)),
        ('Class of Vehicle',            _val(vehicle.class_of_vehicle),            'Body Colour',             _val(vehicle.body_colour)),
        ('Chessis No',                  _val(vehicle.chassis_number),              'Engine Number',           _val(vehicle.engine_number)),
        ('Status',                      _val(vehicle.vehicle_status),              'Country of Origin',       _val(vehicle.country_of_origin)),
        ('Make',                        _val(vehicle.make),                        'Model',                   _val(vehicle.model)),
        ('Date of First Registration',  _val(vehicle.date_of_first_registration),  'Year of Manufacture',     _val(vehicle.year_of_manufacture)),
        ('Cylinder Capacity',           f'{vehicle.cylinder_capacity} cc' if vehicle.cylinder_capacity else '-', 'Fuel Type', _val(vehicle.fuel_type)),
    ]
    id_tbl = doc.add_table(rows=len(id_rows), cols=4)
    id_tbl.style = 'Table Grid'
    _set_col_widths(id_tbl, [1.75, 1.39, 1.75, 1.38])
    for i, (l1, v1, l2, v2) in enumerate(id_rows):
        r = id_tbl.rows[i]
        _cell(r.cells[0], l1, bold=True)
        _cell(r.cells[1], v1)
        _cell(r.cells[2], l2, bold=True)
        _cell(r.cells[3], v2)

    # Block 7 — Parts Row (3 equal cols, 2.09" each)
    parts = doc.add_table(rows=1, cols=3)
    parts.style = 'Table Grid'
    _set_col_widths(parts, [2.09, 2.09, 2.09])
    _cell(parts.rows[0].cells[0], f'Body Parts : {_val(vehicle.body_parts_status)}')
    _cell(parts.rows[0].cells[1], f'Engine Parts : {_val(vehicle.engine_parts_status)}')
    _cell(parts.rows[0].cells[2], f'Accessories : {_val(vehicle.accessories_status)}')

    # Block 8 — Fuel Consumption Row (full-width)
    if vehicle.fuel_consumption:
        fuel_val = f'{vehicle.fuel_consumption} {vehicle.fuel_consumption_unit or ""}'.strip()
    else:
        fuel_val = '-'
    fuel = doc.add_table(rows=1, cols=2)
    fuel.style = 'Table Grid'
    _set_col_widths(fuel, [3.50, 2.77])
    _cell(fuel.rows[0].cells[0], 'Average Of Fuel Consumption :', bold=True)
    _cell(fuel.rows[0].cells[1], fuel_val)

    # Block 9 — Engine Type Row (full-width)
    eng = doc.add_table(rows=1, cols=2)
    eng.style = 'Table Grid'
    _set_col_widths(eng, [3.50, 2.77])
    _cell(eng.rows[0].cells[0], 'Engine :', bold=True)
    _cell(eng.rows[0].cells[1], _val(vehicle.engine_type))

    # Block 10 — Condition Grid (4-col, 6 rows)
    cond_rows = [
        ('Transmission',        _val(vehicle.transmission),        'WD',                 _val(vehicle.wheel_drive)),
        ('Running condition',   _val(vehicle.running_condition),   'Cluth',              _val(vehicle.clutch_status)),
        ('Engine condition',    _val(vehicle.engine_condition),    'Gear box',           _val(vehicle.gear_box_condition)),
        ('Differential',        _val(vehicle.differential_status), 'Gear Selection',     _val(vehicle.gear_selection)),
        ('Body condition',      _val(vehicle.body_condition),      'Chassis condition',  _val(vehicle.chassis_condition)),
        ('Upholstery condition', _val(vehicle.upholstery_condition), 'Underside condition', _val(vehicle.underside_condition)),
    ]
    cond_tbl = doc.add_table(rows=len(cond_rows), cols=4)
    cond_tbl.style = 'Table Grid'
    _set_col_widths(cond_tbl, [1.75, 1.39, 1.75, 1.38])
    for i, (l1, v1, l2, v2) in enumerate(cond_rows):
        r = cond_tbl.rows[i]
        _cell(r.cells[0], l1, bold=True)
        _cell(r.cells[1], v1)
        _cell(r.cells[2], l2, bold=True)
        _cell(r.cells[3], v2)


# ─────────────────────────────────────────────────────────────────────────────
# Page 2
# ─────────────────────────────────────────────────────────────────────────────

def _add_page_2(doc: Document, vehicle: models.Vehicle, report: models.Report) -> None:

    features: dict = safe_get_json_field(vehicle, 'features', {}) or {}
    suspension: dict = safe_get_json_field(vehicle, 'suspension', {}) or {}
    tyres: dict = safe_get_json_field(vehicle, 'tyres', {}) or {}
    electrical: dict = safe_get_json_field(vehicle, 'electrical', {}) or {}
    lights: dict = safe_get_json_field(vehicle, 'lights', {}) or {}
    office_data: dict = safe_get_json_field(vehicle, 'office_data', {}) or {}
    past_valuations: list = safe_get_json_field(vehicle, 'past_valuations', []) or []

    # Normalise Pydantic models → dicts
    def _as_dict(obj: Any) -> dict:
        if isinstance(obj, dict):
            return obj
        if obj is None:
            return {}
        return {k: v for k, v in vars(obj).items() if not k.startswith('_')}

    features = _as_dict(features)
    suspension = _as_dict(suspension)
    electrical = _as_dict(electrical)
    lights = _as_dict(lights)
    office_data = _as_dict(office_data)
    tyres = _as_dict(tyres)

    front_tyre: dict = _as_dict(tyres.get('front')) if tyres.get('front') else {}
    rear_tyre: dict = _as_dict(tyres.get('rear')) if tyres.get('rear') else {}

    # Block P2-1 — Brakes (4-col, 2-row)
    # widths: 1.60 | 1.70 | 1.17 | 1.80
    brk = doc.add_table(rows=2, cols=4)
    brk.style = 'Table Grid'
    _set_col_widths(brk, [1.60, 1.70, 1.17, 1.80])
    _cell(brk.rows[0].cells[0], 'foot brake', bold=True)
    _cell(brk.rows[0].cells[1], _val(vehicle.foot_brake_condition))
    _cell(brk.rows[0].cells[2], 'Disc', bold=True)
    _cell(brk.rows[0].cells[3], _chk(vehicle.disc_brake_available))
    _cell(brk.rows[1].cells[0], 'parking brake', bold=True)
    _cell(brk.rows[1].cells[1], _val(vehicle.parking_brake_condition))
    _cell(brk.rows[1].cells[2], 'ABS', bold=True)
    _cell(brk.rows[1].cells[3], _chk(vehicle.abs_available))

    # Block P2-2 — Features (5-col, 2-row); 6.27 / 5 ≈ 1.254" each
    feat = doc.add_table(rows=2, cols=5)
    feat.style = 'Table Grid'
    _set_col_widths(feat, [1.254, 1.254, 1.254, 1.254, 1.254])
    _cell(feat.rows[0].cells[0], f'Air condition {_chk(features.get("air_condition"))}')
    _cell(feat.rows[0].cells[1], f'Dual Air condition {_chk(features.get("dual_air_condition"))}')
    _cell(feat.rows[0].cells[2], f'Power Mirror {_chk(features.get("power_mirror"))}')
    _cell(feat.rows[0].cells[3], f'Power window {_chk(features.get("power_window"))}')
    _cell(feat.rows[0].cells[4], f'Power Steering {_chk(features.get("power_steering"))}')
    _cell(feat.rows[1].cells[0], f'Air Bag {_chk(features.get("airbag"))}')
    _cell(feat.rows[1].cells[1], f'Number Of Air Bag [{features.get("num_airbags") or ""}]')
    _cell(feat.rows[1].cells[2], f'Seat [{features.get("seats") or ""}]')
    feat.rows[1].cells[3].merge(feat.rows[1].cells[4])
    _cell(feat.rows[1].cells[3], f'Door [{features.get("doors") or ""}]')

    # Block P2-3 — Suspension (3-col, 2-row; col 0 merged vertically)
    # widths: 1.60 | 1.57 | 3.10
    susp = doc.add_table(rows=2, cols=3)
    susp.style = 'Table Grid'
    _set_col_widths(susp, [1.60, 1.57, 3.10])
    susp.cell(0, 0).merge(susp.cell(1, 0))
    _cell(susp.cell(0, 0), 'Suspension', bold=True)
    _cell(susp.rows[0].cells[1], 'Front', bold=True)
    _cell(susp.rows[0].cells[2], _val(suspension.get('front')))
    _cell(susp.rows[1].cells[1], 'Rear', bold=True)
    _cell(susp.rows[1].cells[2], _val(suspension.get('rear')))

    # Block P2-4 — Tyres (5-col, 3-row; col 0 rows 0-1 merged)
    # widths: 1.30 | 0.80 | 1.39 | 1.39 | 1.39
    tyr = doc.add_table(rows=3, cols=5)
    tyr.style = 'Table Grid'
    _set_col_widths(tyr, [1.30, 0.80, 1.39, 1.39, 1.39])
    tyr.cell(0, 0).merge(tyr.cell(1, 0))
    _cell(tyr.cell(0, 0), 'Tyres', bold=True)

    def _tread(t: dict) -> str:
        tp = t.get('tread_percent')
        return f'{tp}%' if tp is not None else _val(t.get('condition'))

    _cell(tyr.rows[0].cells[1], 'Front', bold=True)
    _cell(tyr.rows[0].cells[2], _val(front_tyre.get('brand')))
    _cell(tyr.rows[0].cells[3], _val(front_tyre.get('size')))
    _cell(tyr.rows[0].cells[4], _tread(front_tyre))
    _cell(tyr.rows[1].cells[1], 'Rear', bold=True)
    _cell(tyr.rows[1].cells[2], _val(rear_tyre.get('brand')))
    _cell(tyr.rows[1].cells[3], _val(rear_tyre.get('size')))
    _cell(tyr.rows[1].cells[4], _tread(rear_tyre))
    rear_type = tyres.get('rear_type', '')
    _cell(tyr.rows[2].cells[0], 'Need to Replace', bold=True)
    _cell(tyr.rows[2].cells[1], _yn(tyres.get('need_replacement')))
    _cell(tyr.rows[2].cells[2], 'Rear', bold=True)
    _cell(tyr.rows[2].cells[3], f'Single {_chk(rear_type == "single")}')
    _cell(tyr.rows[2].cells[4], f'Dual {_chk(rear_type == "dual")}')

    # Block P2-5 — Electrical (3-col, 2-row) + Lights (7-col, 1-row) — stacked, no gap
    # Electrical widths: 2.09 × 3
    elec = doc.add_table(rows=2, cols=3)
    elec.style = 'Table Grid'
    _set_col_widths(elec, [2.09, 2.09, 2.09])
    _cell(elec.rows[0].cells[0], f'S. Starter {_chk(electrical.get("starter"))}')
    _cell(elec.rows[0].cells[1], f'Horn {_chk(electrical.get("horn"))}')
    _cell(elec.rows[0].cells[2], f'Wiper {_chk(electrical.get("wiper"))}')
    elec.rows[1].cells[0].merge(elec.rows[1].cells[2])
    batt_cell = elec.rows[1].cells[0]
    batt_cell.text = ''
    batt_para = batt_cell.paragraphs[0]
    br = batt_para.add_run('Battery: ')
    br.bold = True
    br.font.size = Pt(11)
    bv = batt_para.add_run(electrical.get('battery_condition') or '')
    bv.font.size = Pt(11)

    # Lights widths: 0.77 + 0.917×6 = 6.272 ≈ 6.27
    lts = doc.add_table(rows=1, cols=7)
    lts.style = 'Table Grid'
    _set_col_widths(lts, [0.77, 0.917, 0.917, 0.917, 0.917, 0.917, 0.917])
    _cell(lts.rows[0].cells[0], 'Lights', bold=True)
    _cell(lts.rows[0].cells[1], f'Head {_chk(lights.get("head"))}')
    _cell(lts.rows[0].cells[2], f'Dim {_chk(lights.get("dim"))}')
    _cell(lts.rows[0].cells[3], f'Signal {_chk(lights.get("signal"))}')
    _cell(lts.rows[0].cells[4], f'Parking {_chk(lights.get("parking"))}')
    _cell(lts.rows[0].cells[5], f'Reverse {_chk(lights.get("reverse"))}')
    _cell(lts.rows[0].cells[6], f'Meter {_chk(lights.get("meter"))}')

    # Block P2-6 — History (4-col, 2-row)
    # widths: 1.85 | 1.30 | 1.87 | 1.25
    hist = doc.add_table(rows=2, cols=4)
    hist.style = 'Table Grid'
    _set_col_widths(hist, [1.85, 1.30, 1.87, 1.25])
    _cell(hist.rows[0].cells[0], 'Accidents', bold=True)
    _cell(hist.rows[0].cells[1], _yn(vehicle.has_accidents))
    _cell(hist.rows[0].cells[2], 'Repairs', bold=True)
    _cell(hist.rows[0].cells[3], _yn(vehicle.has_repairs))
    _cell(hist.rows[1].cells[0], 'Body parts replaced', bold=True)
    _cell(hist.rows[1].cells[1], _yn(vehicle.body_parts_replaced))
    _cell(hist.rows[1].cells[2], 'Repairs needed within a year', bold=True)
    _cell(hist.rows[1].cells[3], _yn(vehicle.needs_repairs_within_year))

    # Block P2-7 — Summary (bordered box, min 3 cm)
    sum_h = doc.add_paragraph()
    sum_h.paragraph_format.space_before = Pt(8)
    shr = sum_h.add_run('Summary of the Valuation')
    shr.bold = True
    shr.font.size = Pt(12)
    sum_tbl = doc.add_table(rows=1, cols=1)
    sum_tbl.style = 'Table Grid'
    _set_col_widths(sum_tbl, [6.27])
    sum_tbl.cell(0, 0).text = vehicle.valuation_summary or ''
    _set_row_min_height(sum_tbl.rows[0], 3)

    # Block P2-8 — Pricing (2-col, 2-row)
    # widths: 3.50 | 2.77
    price_tbl = doc.add_table(rows=2, cols=2)
    price_tbl.style = 'Table Grid'
    _set_col_widths(price_tbl, [3.50, 2.77])
    _cell(price_tbl.rows[0].cells[0], 'Purchase Price of the vehicle', bold=True)
    _cell(price_tbl.rows[0].cells[1], _fmt_rs(vehicle.purchase_price))
    _cell(price_tbl.rows[1].cells[0], 'Brand New Price of the vehicle', bold=True)
    _cell(price_tbl.rows[1].cells[1], _fmt_rs(vehicle.brand_new_price))

    # Block P2-9 — Previous Assessment (conditional)
    if office_data:
        pa_h = doc.add_paragraph()
        pa_h.paragraph_format.space_before = Pt(10)
        par = pa_h.add_run('Previous Assessment Subject Property')
        par.bold = True
        par.font.size = Pt(12)

        # Metadata table (2-col, 3-row); widths: 2.60 | 3.67
        meta = doc.add_table(rows=3, cols=2)
        meta.style = 'Table Grid'
        _set_col_widths(meta, [2.60, 3.67])
        _cell(meta.rows[0].cells[0], 'Previous Reference (File No)', bold=True)
        _cell(meta.rows[0].cells[1], _val(office_data.get('previous_reference')))
        _cell(meta.rows[1].cells[0], 'Reported Value', bold=True)
        _cell(meta.rows[1].cells[1], _fmt_rs(office_data.get('reported_value')))
        _cell(meta.rows[2].cells[0], 'Reported Date', bold=True)
        _cell(meta.rows[2].cells[1], _val(office_data.get('reported_date')))

        # Past Valuations table
        # 6-col: 0.60 | 1.10 | 1.10 | 0.80 | 1.27 | 1.40
        pv_data = list(past_valuations) if past_valuations else [{}] * 2
        pv_tbl = doc.add_table(rows=len(pv_data) + 2, cols=6)
        pv_tbl.style = 'Table Grid'
        _set_col_widths(pv_tbl, [0.60, 1.10, 1.10, 0.80, 1.27, 1.40])

        # Header row 0
        h0 = pv_tbl.rows[0]
        _cell(h0.cells[0], 'Serial No', bold=True)
        h0.cells[1].merge(h0.cells[2])
        _cell(h0.cells[1], 'Vehicle No (Civil No | Army No)', bold=True)
        h0.cells[3].merge(h0.cells[4])
        _cell(h0.cells[3], 'Past Valuation (Year | Value (Rs))', bold=True)
        _cell(h0.cells[5], 'Market Value (Rs)', bold=True)

        # Header row 1
        h1 = pv_tbl.rows[1]
        _cell(h1.cells[0], '')
        _cell(h1.cells[1], 'Civil No', bold=True)
        _cell(h1.cells[2], 'Army No', bold=True)
        _cell(h1.cells[3], 'Year', bold=True)
        _cell(h1.cells[4], 'Value (Rs)', bold=True)
        _cell(h1.cells[5], '')

        # Data rows
        for i, pv in enumerate(pv_data):
            if isinstance(pv, dict):
                serial = pv.get('serial', i + 1)
                civil = pv.get('civil_no', '')
                military = pv.get('military_no', '')
                year = pv.get('year', '')
                value = pv.get('value')
                mkt = pv.get('market_value')
            else:
                serial = getattr(pv, 'serial', i + 1)
                civil = getattr(pv, 'civil_no', '')
                military = getattr(pv, 'military_no', '')
                year = getattr(pv, 'year', '')
                value = getattr(pv, 'value', None)
                mkt = getattr(pv, 'market_value', None)
            dr = pv_tbl.rows[i + 2]
            _cell(dr.cells[0], str(serial) if serial else '')
            _cell(dr.cells[1], str(civil) if civil else '')
            _cell(dr.cells[2], str(military) if military else '')
            _cell(dr.cells[3], str(year) if year else '')
            _cell(dr.cells[4], _fmt_rs(value) if value else '')
            _cell(dr.cells[5], _fmt_rs(mkt) if mkt else '')

    # Block P2-10 — Footer
    f1 = doc.add_paragraph()
    f1.paragraph_format.space_after = Pt(0)
    f1r = f1.add_run('Valuer / Senior Valuer')
    f1r.font.size = Pt(11)
    f2 = doc.add_paragraph()
    f2.paragraph_format.space_after = Pt(12)
    f2r = f2.add_run('Report submitted for your approval please')
    f2r.font.size = Pt(11)

    # Block P2-11 — Signature Block (borderless 2-col)
    # widths: 3.13 | 3.14
    sig = doc.add_table(rows=2, cols=2)
    _set_col_widths(sig, [3.13, 3.14])
    for col_idx in range(2):
        para = sig.rows[0].cells[col_idx].paragraphs[0]
        para.paragraph_format.space_before = Pt(24)
        run = para.add_run('.' * 35)
        run.font.size = Pt(11)
    _cell(sig.rows[1].cells[0], 'Assistant Valuer / Valuer')
    _cell(sig.rows[1].cells[1], 'Date')


# ─────────────────────────────────────────────────────────────────────────────
# Public entry point
# ─────────────────────────────────────────────────────────────────────────────

def generate_vehicle_report_docx(report: models.Report, user: models.User) -> BytesIO:
    """Generate a formatted DOCX for a vehicle valuation report."""
    try:
        doc = Document()
        template_id = user.preferred_letterhead_template or 'classic'
        template = get_template(template_id)
        template.render_letterhead(doc, user, report)

        vehicle = report.primary_vehicle
        if not vehicle:
            if report.vehicle_associations and len(report.vehicle_associations) > 0:
                vehicle = report.vehicle_associations[0].vehicle
            else:
                raise ValueError('No vehicle found for this report')

        _add_page_1(doc, vehicle, report)
        _add_page_2(doc, vehicle, report)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        logger.info('[DOCX] Vehicle report generated successfully')
        return buffer

    except Exception as e:
        logger.error(f'[DOCX] Error generating vehicle report: {str(e)}')
        import traceback
        logger.error(f'[DOCX] Traceback: {traceback.format_exc()}')
        raise
