"""
Photo grid rendering for property and building photos in DOCX reports.
"""
import base64
import re
import logging
from io import BytesIO

from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

from .styling import PROPERTY_PHOTO_HEIGHT, IMAGE_SPACING_AFTER, FONT_SIZE_CAPTION
from .images import calculate_image_dimensions

logger = logging.getLogger(__name__)


def render_photo_grid(doc, sorted_photos: list, start_idx: int = 0) -> None:
    """
    Render a grid of photos (3 per row) with captions into the document.

    Args:
        doc: python-docx Document
        sorted_photos: Photos already sorted by order; each dict has 'image_data' and 'caption'
        start_idx: Figure number offset (default 0 → figures start at Fig. 1)
    """
    if not sorted_photos:
        return

    num_photos = len(sorted_photos)
    photos_per_row = 3
    photo_idx = 0

    while photo_idx < num_photos:
        remaining = num_photos - photo_idx
        if remaining >= photos_per_row:
            photos_in_row = photos_per_row
        elif remaining == 1 and photo_idx > 0:
            photos_in_row = 1
        else:
            photos_in_row = remaining

        table = doc.add_table(rows=2, cols=photos_in_row)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for row in table.rows:
            for cell in row.cells:
                cell.width = Inches(6.5 / photos_in_row)
                tc = cell._element
                tcPr = tc.get_or_add_tcPr()
                tcBorders = parse_xml(
                    r'<w:tcBorders %s>'
                    r'<w:top w:val="none"/>'
                    r'<w:left w:val="none"/>'
                    r'<w:bottom w:val="none"/>'
                    r'<w:right w:val="none"/>'
                    r'</w:tcBorders>' % nsdecls('w')
                )
                tcPr.append(tcBorders)

        for i in range(photos_in_row):
            if photo_idx >= num_photos:
                break

            photo = sorted_photos[photo_idx]
            try:
                image_data = photo.get('image_data', '')
                caption = photo.get('caption', '')

                if image_data:
                    if image_data.startswith('data:image'):
                        base64_match = re.search(r'base64,(.+)', image_data)
                        if not base64_match:
                            photo_idx += 1
                            continue
                        base64_data = base64_match.group(1)
                    else:
                        base64_data = image_data

                    image_bytes = base64.b64decode(base64_data)
                    image_stream = BytesIO(image_bytes)

                    dimensions = calculate_image_dimensions(
                        image_stream,
                        Inches(2.0),
                        PROPERTY_PHOTO_HEIGHT
                    )

                    image_stream.seek(0)
                    cell = table.rows[0].cells[i]
                    cell_para = cell.paragraphs[0]
                    cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cell_para.add_run().add_picture(image_stream, **dimensions)

                    caption_cell = table.rows[1].cells[i]
                    caption_para = caption_cell.paragraphs[0]
                    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption_para.paragraph_format.space_before = Pt(2)
                    caption_para.paragraph_format.space_after = Pt(2)

                    caption_text = f"Fig. {start_idx + photo_idx + 1}"
                    if caption:
                        caption_text += f": {caption}"

                    caption_run = caption_para.add_run(caption_text)
                    caption_run.font.size = FONT_SIZE_CAPTION
                    caption_run.font.italic = True
                    caption_run.font.color.rgb = RGBColor(60, 60, 60)

                    logger.info(f"Added photo Fig. {start_idx + photo_idx + 1}")

            except Exception as e:
                logger.error(f"Error adding photo {start_idx + photo_idx + 1}: {str(e)}")

            photo_idx += 1

        spacing_para = doc.add_paragraph()
        spacing_para.paragraph_format.space_after = Pt(8)

    final_spacing = doc.add_paragraph()
    final_spacing.paragraph_format.space_after = IMAGE_SPACING_AFTER
