"""
Signature block generator for valuation reports.
"""
from io import BytesIO
from typing import Optional
import requests
import logging

from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .styling import FONT_SIZE_BODY

logger = logging.getLogger(__name__)


def add_signature_block(
    doc,
    user,
    valuer_name: Optional[str],
    valuer_designation: Optional[str],
    certification_date: Optional[str]
):
    """
    Add a standardized signature block to the document.
    """
    doc.add_paragraph("\n")

    sig_line = doc.add_paragraph("_" * 40)
    sig_line.paragraph_format.space_before = Pt(24)
    sig_line.paragraph_format.space_after = Pt(6)

    if hasattr(user, 'signature_image') and user.signature_image:
        try:
            response = requests.get(user.signature_image, timeout=10)
            if response.status_code == 200:
                sig_para = doc.add_paragraph()
                sig_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                sig_para.paragraph_format.space_before = Pt(0)
                sig_para.paragraph_format.space_after = Pt(6)
                image_stream = BytesIO(response.content)
                sig_para.add_run().add_picture(image_stream, width=Inches(2))
        except Exception as e:
            logger.warning(f"Failed to add signature image: {e}")

    if valuer_name:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(valuer_name)
        run.font.bold = True
        run.font.size = FONT_SIZE_BODY

    if valuer_designation:
        p = doc.add_paragraph(valuer_designation)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = FONT_SIZE_BODY

    if certification_date:
        p = doc.add_paragraph(certification_date)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = FONT_SIZE_BODY

    doc.add_paragraph()
