"""
Body-section renderer for single-property DOCX reports.

Renders numbered sections 1.0–9.0 and the invoice, delegating section 4.0
(DESCRIPTION) to description_renderer and section 8.0 (VALUATION) to
valuation_section_renderer.
"""
import logging

import requests
from io import BytesIO

from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT

from .paragraph_builders import add_section_heading, add_subsection_paragraph
from .styling import (
    FONT_SIZE_BODY, BODY_PARA_SPACE_AFTER,
    SUBHEADING_SPACE_BEFORE, SUBHEADING_SPACE_AFTER,
    IMAGE_SPACING_BEFORE, IMAGE_SPACING_AFTER,
    MAP_IMAGE_WIDTH, MAP_IMAGE_MAX_HEIGHT,
)
from .helpers import safe_get_json_field, safe_parse_json_string
from .images import calculate_image_dimensions
from .text_generators import (
    generate_situation_text, generate_access_text, generate_locality_description,
    generate_boundary_summary_text,
    generate_ownership_paragraph, generate_street_lines_paragraph,
    generate_building_limits_paragraph, generate_local_authority_paragraph,
    generate_rent_act_paragraph, generate_land_values_paragraph,
    generate_simplified_certification_text, add_signature_block,
)
from .description_renderer import render_description_section
from .valuation_section_renderer import render_single_property_valuation
from .invoice_generator import generate_invoice_section
from ..utils import format_no_field

logger = logging.getLogger(__name__)


def render_body_sections(doc, report, user) -> None:
    """
    Render all numbered body sections (1.0–9.0) plus invoice for a
    single-property DOCX report.

    Args:
        doc: python-docx Document (letterhead and cover page already rendered)
        report: Report model instance
        user: User model instance (needed for signature block in section 9.0)
    """
    # ===== 1.0 SITUATION SECTION =====
    situation_text = generate_situation_text(report)
    if situation_text:
        add_section_heading(doc, "1.0", "SITUATION")

        situation_para = doc.add_paragraph()
        situation_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        situation_para.paragraph_format.space_before = Pt(0)
        situation_para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
        situation_para.paragraph_format.line_spacing = 0.9
        situation_run = situation_para.add_run(situation_text)
        situation_run.font.size = FONT_SIZE_BODY
        situation_run.font.color.rgb = RGBColor(0, 0, 0)

    # ===== 2.0 ACCESS SECTION =====
    access_text = generate_access_text(report)
    locality_text = generate_locality_description(report)

    if access_text:
        add_section_heading(doc, "2.0", "ACCESS")

        access_para = doc.add_paragraph()
        access_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        access_para.paragraph_format.space_before = Pt(0)
        access_para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
        access_para.paragraph_format.line_spacing = 0.9
        access_run = access_para.add_run(access_text)
        access_run.font.size = FONT_SIZE_BODY
        access_run.font.color.rgb = RGBColor(0, 0, 0)

        # Add coordinates paragraph if available
        if report.property_latitude and report.property_longitude:
            coord_para = doc.add_paragraph()
            coord_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            coord_para.paragraph_format.space_before = Pt(6)
            coord_para.paragraph_format.space_after = Pt(6)
            coord_para.paragraph_format.line_spacing = 0.9

            coord_label = coord_para.add_run("Property Location Coordinates: ")
            coord_label.bold = True
            coord_label.font.size = FONT_SIZE_BODY
            coord_label.font.color.rgb = RGBColor(0, 0, 0)

            lat_value = float(report.property_latitude)
            lng_value = float(report.property_longitude)
            coord_text = coord_para.add_run(f"{lat_value:.6f}, {lng_value:.6f}")
            coord_text.font.size = FONT_SIZE_BODY
            coord_text.font.color.rgb = RGBColor(0, 0, 0)

        # Add map image if available (embedded within ACCESS section)
        if report.location_map_image_data:
            try:
                map_url = report.location_map_image_data
                logger.info(f"[DOCX] Fetching map image from URL (length={len(map_url)})")
                logger.debug(f"[DOCX] Map URL: {map_url[:200]}...")

                response = requests.get(map_url, timeout=30)
                if response.status_code == 200:
                    map_spacing_para = doc.add_paragraph()
                    map_spacing_para.paragraph_format.space_before = IMAGE_SPACING_BEFORE
                    map_spacing_para.paragraph_format.space_after = Pt(0)

                    map_para = doc.add_paragraph()
                    map_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    map_para.paragraph_format.space_before = Pt(0)
                    map_para.paragraph_format.space_after = IMAGE_SPACING_AFTER

                    image_stream = BytesIO(response.content)
                    dimensions = calculate_image_dimensions(
                        image_stream,
                        MAP_IMAGE_WIDTH,
                        MAP_IMAGE_MAX_HEIGHT
                    )
                    map_para.add_run().add_picture(image_stream, **dimensions)

                    logger.info(
                        f"[DOCX] Successfully added map image (size={len(response.content)} bytes)"
                    )
                else:
                    logger.warning(f"[DOCX] Failed to fetch map image: HTTP {response.status_code}")
            except Exception as e:
                logger.error(f"[DOCX] Error adding map image: {str(e)}")
        else:
            logger.info(f"[DOCX] No location_map_image_data available for report {report.id}")

    # ===== 3.0 IDENTIFICATION OF PROPERTY SECTION =====
    has_property_header_data = (
        report.land_traditional_name or
        report.land_extent_formatted or
        report.boundaries or
        report.physical_boundaries_types or
        report.physical_boundaries_description or
        report.boundaries_summary_text
    )

    if has_property_header_data:
        add_section_heading(doc, "3.0", "IDENTIFICATION OF PROPERTY")

        # Traditional Land Name
        if report.land_traditional_name:
            name_para = doc.add_paragraph()
            name_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            name_para.paragraph_format.space_before = Pt(0)
            name_para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
            name_para.paragraph_format.line_spacing = 0.9
            name_label = name_para.add_run("Name of Land: ")
            name_label.bold = True
            name_label.font.size = FONT_SIZE_BODY
            name_label.font.color.rgb = RGBColor(0, 0, 0)
            name_value = name_para.add_run(f'"{report.land_traditional_name}"')
            name_value.font.size = FONT_SIZE_BODY
            name_value.font.color.rgb = RGBColor(0, 0, 0)

        # Survey Plan Information
        if report.lot_number and report.plan_number:
            plan_para = doc.add_paragraph()
            plan_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            plan_para.paragraph_format.space_before = Pt(0)
            plan_para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
            plan_para.paragraph_format.line_spacing = 0.9
            plan_label = plan_para.add_run("Survey Plan: ")
            plan_label.bold = True
            plan_label.font.size = FONT_SIZE_BODY
            plan_label.font.color.rgb = RGBColor(0, 0, 0)

            lot_desc = report.lot_number.strip() if report.lot_number else ''

            prefixes_to_remove = ['plan no', 'plan no:', 'lot plan no', 'lot plan no:']
            lot_desc_lower = lot_desc.lower()
            for prefix in prefixes_to_remove:
                if lot_desc_lower.startswith(prefix):
                    lot_desc = lot_desc[len(prefix):].strip()
                    break

            if not lot_desc.lower().startswith('lot'):
                lot_desc = f"Lot {lot_desc}"

            plan_formatted = format_no_field("Plan", report.plan_number)
            plan_text = f"{lot_desc} in {plan_formatted}"
            if report.plan_date:
                plan_text += f" dated {report.plan_date}"
            if report.licensed_surveyor_name:
                plan_text += f" made by {report.licensed_surveyor_name}, Licensed Surveyor"
            plan_text += "."

            plan_value = plan_para.add_run(plan_text)
            plan_value.font.size = FONT_SIZE_BODY
            plan_value.font.color.rgb = RGBColor(0, 0, 0)

        # Land Extent
        if report.land_extent_formatted:
            extent_para = doc.add_paragraph()
            extent_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            extent_para.paragraph_format.space_before = Pt(0)
            extent_para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
            extent_para.paragraph_format.line_spacing = 0.9
            extent_label = extent_para.add_run("Extent: ")
            extent_label.bold = True
            extent_label.font.size = FONT_SIZE_BODY
            extent_label.font.color.rgb = RGBColor(0, 0, 0)

            extent_text = report.land_extent_formatted
            if report.land_extent_hectares:
                extent_text += f" [{report.land_extent_hectares:.4f} Hectares]"
            if report.land_extent_square_meters:
                extent_text += f" [{report.land_extent_square_meters:.2f} m\u00b2]"

            extent_value = extent_para.add_run(extent_text)
            extent_value.font.size = FONT_SIZE_BODY
            extent_value.font.color.rgb = RGBColor(0, 0, 0)

        # Boundaries — enhanced format with all details
        if report.boundaries:
            boundaries_spacing = doc.add_paragraph()
            boundaries_spacing.paragraph_format.space_before = SUBHEADING_SPACE_BEFORE
            boundaries_spacing.paragraph_format.space_after = SUBHEADING_SPACE_AFTER

            boundaries_heading = doc.add_paragraph()
            boundaries_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            boundaries_heading.paragraph_format.space_before = Pt(0)
            boundaries_heading.paragraph_format.space_after = SUBHEADING_SPACE_AFTER
            boundaries_heading.paragraph_format.line_spacing = 0.9
            boundaries_heading_run = boundaries_heading.add_run("Boundaries:")
            boundaries_heading_run.bold = True
            boundaries_heading_run.font.size = FONT_SIZE_BODY
            boundaries_heading_run.font.color.rgb = RGBColor(0, 0, 0)

            boundary_type_labels = {
                'brick_walls': 'Brick Walls',
                'barbed_wire': 'Barbed Wire Fence',
                'live_fence': 'Live Fence',
                'concrete_posts': 'Concrete Posts',
                'iron_gate': 'Iron Gate',
                'rubble_foundation': 'Rubble Foundation',
                'chain_link': 'Chain Link Fence',
                'wooden_fence': 'Wooden Fence',
                'stone_wall': 'Stone Wall',
                'hedge': 'Hedge',
                'none': 'None'
            }

            directions = ['north', 'northeast', 'east', 'southeast', 'south', 'southwest', 'west', 'northwest']
            direction_labels = ['North', 'North-East', 'East', 'South-East', 'South', 'South-West', 'West', 'North-West']

            boundaries = safe_get_json_field(report, 'boundaries', {})
            if not boundaries:
                boundaries = {}

            for direction, label in zip(directions, direction_labels):
                boundary_data = boundaries.get(direction, {}) if isinstance(boundaries, dict) else {}

                # Skip diagonal boundaries if they have no description
                if direction in ['northeast', 'southeast', 'southwest', 'northwest']:
                    if not boundary_data or not boundary_data.get('description'):
                        continue

                boundary_line = f"{label:<11} : "
                boundary_text = boundary_data.get('description') or 'Not specified'

                boundary_para = doc.add_paragraph()
                boundary_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                boundary_para.paragraph_format.space_before = Pt(0)
                boundary_para.paragraph_format.space_after = Pt(1)
                boundary_para.paragraph_format.line_spacing = 0.9
                boundary_para.paragraph_format.left_indent = Inches(0.5)

                boundary_run = boundary_para.add_run(boundary_line + boundary_text)
                boundary_run.font.size = FONT_SIZE_BODY
                boundary_run.font.color.rgb = RGBColor(0, 0, 0)

                # Additional boundary details (if any)
                additional_details = []

                boundary_types_per_dir = report.boundary_types_per_direction or {}
                physical_type = boundary_types_per_dir.get(direction)
                if physical_type:
                    type_label = boundary_type_labels.get(
                        physical_type, physical_type.replace('_', ' ').title()
                    )
                    additional_details.append(f"Type: {type_label}")

                length = boundary_data.get('length')
                if length:
                    additional_details.append(f"Length: {length}")

                adjoins = boundary_data.get('adjoins')
                if adjoins:
                    additional_details.append(f"Adjoins: {adjoins}")

                notes = boundary_data.get('notes')
                if notes:
                    additional_details.append(f"Notes: {notes}")

                if additional_details:
                    for detail in additional_details:
                        detail_para = doc.add_paragraph()
                        detail_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        detail_para.paragraph_format.space_before = Pt(0)
                        detail_para.paragraph_format.space_after = Pt(1)
                        detail_para.paragraph_format.line_spacing = 0.9
                        detail_para.paragraph_format.left_indent = Inches(0.75)

                        detail_run = detail_para.add_run(f"  \u2022 {detail}")
                        detail_run.font.size = Pt(8)
                        detail_run.font.color.rgb = RGBColor(80, 80, 80)

                    spacing_para = doc.add_paragraph()
                    spacing_para.paragraph_format.space_before = Pt(0)
                    spacing_para.paragraph_format.space_after = Pt(2)

        # Boundary Summary Sentence
        boundary_summary = report.boundaries_summary_text
        if not boundary_summary and (report.boundary_types_per_direction or report.physical_boundaries_types):
            boundary_summary = generate_boundary_summary_text(report)

        if boundary_summary:
            summary_para = doc.add_paragraph()
            summary_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            summary_para.paragraph_format.space_before = Pt(6)
            summary_para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
            summary_para.paragraph_format.line_spacing = 0.9
            summary_run = summary_para.add_run(boundary_summary)
            summary_run.font.size = FONT_SIZE_BODY
            summary_run.font.color.rgb = RGBColor(0, 0, 0)

    # ===== 4.0 DESCRIPTION OF PROPERTY =====
    render_description_section(doc, report)

    # ===== 5.0 LOCALITY SECTION =====
    if locality_text:
        add_section_heading(doc, "5.0", "LOCALITY")

        locality_para = doc.add_paragraph()
        locality_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        locality_para.paragraph_format.space_before = Pt(0)
        locality_para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
        locality_para.paragraph_format.line_spacing = 0.9
        locality_run = locality_para.add_run(locality_text)
        locality_run.font.size = FONT_SIZE_BODY
        locality_run.font.color.rgb = RGBColor(0, 0, 0)

    # ===== 6.0 LEGAL ASPECTS =====
    has_legal_data = (
        report.ownership_type or report.street_lines_status or report.building_limits_status or
        report.local_authority_data or report.rent_act_effectiveness or
        report.title_search_conducted or report.property_encumbered or
        report.street_lines_gazette_ref or report.building_plan_approved or
        report.local_authority_rated
    )

    if has_legal_data:
        add_section_heading(doc, "6.0", "LEGAL ASPECTS")

        if report.ownership_type or hasattr(report, 'deeds') or hasattr(report, 'plan_number'):
            ownership_para = generate_ownership_paragraph(report)
            add_subsection_paragraph(doc, "(a)", "Ownership", ownership_para)

        if report.street_lines_status and report.report_type != 'bare_land':
            street_para = generate_street_lines_paragraph(report)
            add_subsection_paragraph(doc, "(b)", "Street lines", street_para)

        if report.building_limits_status and report.report_type != 'bare_land':
            building_para = generate_building_limits_paragraph(report)
            add_subsection_paragraph(doc, "(c)", "Building limits", building_para)

        if report.local_authority_data or report.pradeshiya_sabha or report.local_authority_rated:
            authority_para = generate_local_authority_paragraph(report)
            add_subsection_paragraph(doc, "(d)", "Local authority data", authority_para)

        if report.rent_act_effectiveness:
            rent_para = generate_rent_act_paragraph(report)
            add_subsection_paragraph(doc, "(e)", "Rent act effectiveness", rent_para)

        doc.add_paragraph()  # Spacing

    # ===== 7.0 LAND VALUES IN THE AREA =====
    if report.comparable_properties or report.land_market_analysis:
        add_section_heading(doc, "7.0", "LAND VALUES IN THE AREA")

        comparables = safe_parse_json_string(report.comparable_properties, [])

        if comparables:
            land_values_text = generate_land_values_paragraph(comparables)

            if land_values_text:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run(land_values_text)
                run.font.size = FONT_SIZE_BODY

                avg_rate = sum(c.get('rate_per_perch', 0) for c in comparables) / len(comparables)
                if avg_rate > 0:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(6)
                    run = p.add_run(f"Average Rate: LKR {avg_rate:,.2f} per perch")
                    run.font.bold = True
                    run.font.size = FONT_SIZE_BODY

        if report.land_market_analysis:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run(report.land_market_analysis)
            run.font.size = FONT_SIZE_BODY

        doc.add_paragraph()

    # ===== 8.0 VALUATION =====
    render_single_property_valuation(doc, report)

    # ===== 9.0 CERTIFICATION =====
    if report.certification_text or report.certification_valuer_name:
        add_section_heading(doc, "9.0", "CERTIFICATION")

        if report.certification_text:
            cert_text = report.certification_text.strip()
        elif report.certification_valuer_name and report.certification_valuer_designation:
            cert_text = generate_simplified_certification_text(
                valuer_name=report.certification_valuer_name,
                valuer_designation=report.certification_valuer_designation,
                lot_number=report.lot_number,
                plan_number=report.plan_number,
                plan_date=report.plan_date,
                licensed_surveyor_name=report.licensed_surveyor_name,
                deeds=safe_get_json_field(report, 'deeds', []),
                property_identification_type=report.property_identification_type
            )
        else:
            cert_text = (
                "I hereby certify that I have personally inspected this property and the "
                "valuation stated herein represents my professional opinion of the market "
                "value as of the date of inspection."
            )

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run(cert_text)
        run.font.size = FONT_SIZE_BODY

        add_signature_block(
            doc=doc,
            user=user,
            valuer_name=report.certification_valuer_name,
            valuer_designation=report.certification_valuer_designation,
            certification_date=report.certification_date
        )

    # ===== INVOICE SECTION =====
    if report.invoice_data:
        logger.info("[DOCX] Generating invoice section")
        generate_invoice_section(doc, report.invoice_data, user, report)
