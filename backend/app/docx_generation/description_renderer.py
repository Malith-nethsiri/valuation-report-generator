"""
Section 4.0 DESCRIPTION OF PROPERTY renderer for single-property DOCX reports.
"""
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .formatting import format_material_list, format_room_count
from .paragraph_builders import add_section_heading, add_inline_field
from .styling import (
    FONT_SIZE_BODY, BODY_PARA_SPACE_AFTER,
    INLINE_FIELD_SPACE_BEFORE, INLINE_FIELD_SPACE_AFTER,
    SUBHEADING_SPACE_BEFORE, SUBHEADING_SPACE_AFTER,
    INDENTED_CONTENT_LEFT_INDENT,
)
from .helpers import safe_get_json_field
from .photo_section import render_photo_grid
from .building_renderer import (
    aggregate_accommodation_across_building,
    render_construction_details,
    render_utilities_and_conveniences,
)
from ..utils.text_helpers import format_list_with_grammar


def render_description_section(doc, report) -> None:
    """
    Render section 4.0 DESCRIPTION OF PROPERTY.

    Covers land description, bare-land features, building details (construction,
    accommodation, floor area, photos, additional structures). Skips silently
    if the report has no relevant description data.
    """
    has_land_data = (
        report.land_description_text or
        report.land_shape or
        report.soil_type or
        report.land_type or
        report.water_table_depth or
        report.flood_risk or
        report.ongoing_construction_notes or
        report.elevation_changes or
        report.drainage_pattern or
        report.vegetation_type or
        report.natural_features
    )

    has_building_data = (
        report.buildings and
        isinstance(report.buildings, list) and
        len(report.buildings) > 0 and
        report.report_type != 'bare_land'
    )

    has_occupier_data = report.occupier_name

    if not (has_land_data or has_building_data or has_occupier_data):
        return

    add_section_heading(doc, "4.0", "DESCRIPTION OF PROPERTY")

    # === LAND DESCRIPTION ===
    if report.land_description_text:
        land_para = doc.add_paragraph()
        land_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        land_para.paragraph_format.space_before = Pt(0)
        land_para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
        land_para.paragraph_format.line_spacing = 0.9
        land_run = land_para.add_run(report.land_description_text)
        land_run.font.size = FONT_SIZE_BODY
        land_run.font.color.rgb = RGBColor(0, 0, 0)
    else:
        # Build land description from individual fields
        land_parts = []

        # Shape and type
        if report.land_shape:
            shape_labels = {
                'rectangular': 'rectangular', 'square': 'square', 'triangular': 'triangular',
                'irregular': 'irregular', 'L_shaped': 'L-shaped', 'trapezoidal': 'trapezoidal'
            }
            land_parts.append(f"The land is {shape_labels.get(report.land_shape, report.land_shape)} in shape")

        if report.land_type:
            type_labels = {
                'flat': 'flat terrain', 'sloped': 'sloped terrain', 'hilly': 'hilly terrain',
                'waterfront': 'waterfront', 'paddy': 'paddy land', 'garden': 'garden land'
            }
            if land_parts:
                land_parts.append(f"with {type_labels.get(report.land_type, report.land_type)}")
            else:
                land_parts.append(f"The land has {type_labels.get(report.land_type, report.land_type)}")

        # Land level
        if report.land_level:
            level_labels = {
                'above_road': 'above road level', 'at_road': 'at road level',
                'below_road': 'below road level', 'varied': 'varied levels'
            }
            land_parts.append(f"and is {level_labels.get(report.land_level, report.land_level)}")
            if report.land_level_difference:
                land_parts.append(f"by approximately {report.land_level_difference} feet")

        # Soil type
        if report.soil_type:
            soil_labels = {
                'sandy': 'sandy', 'clay': 'clay', 'loam': 'loam', 'laterite': 'laterite',
                'rocky': 'rocky', 'gravel': 'gravel', 'alluvial': 'alluvial'
            }
            land_parts.append(f". The soil is {soil_labels.get(report.soil_type, report.soil_type)}")

        # Water table
        if report.water_table_depth:
            land_parts.append(f"with water table at approximately {report.water_table_depth} feet depth")

        # Risks
        risk_parts = []
        if report.flood_risk and report.flood_risk != 'none':
            risk_labels = {'low': 'low', 'moderate': 'moderate', 'high': 'high', 'very_high': 'very high'}
            risk_parts.append(f"flood risk is {risk_labels.get(report.flood_risk, report.flood_risk)}")
        if report.inundation_risk and report.inundation_risk != 'none':
            risk_labels = {'low': 'low', 'moderate': 'moderate', 'high': 'high', 'very_high': 'very high'}
            risk_parts.append(f"inundation risk is {risk_labels.get(report.inundation_risk, report.inundation_risk)}")
        if report.earth_slip_risk and report.earth_slip_risk != 'none':
            risk_labels = {'low': 'low', 'moderate': 'moderate', 'high': 'high', 'very_high': 'very high'}
            risk_parts.append(f"earth slip risk is {risk_labels.get(report.earth_slip_risk, report.earth_slip_risk)}")

        if risk_parts:
            land_parts.append(". The " + ", ".join(risk_parts))

        # Land condition
        if report.land_condition:
            condition_labels = {
                'well_maintained': 'well maintained', 'developed': 'developed',
                'undeveloped': 'undeveloped', 'overgrown': 'overgrown', 'cleared': 'cleared'
            }
            land_parts.append(f". The land is {condition_labels.get(report.land_condition, report.land_condition)}")

        if land_parts:
            land_text = " ".join(land_parts)
            if not land_text.endswith("."):
                land_text += "."

            land_para = doc.add_paragraph()
            land_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            land_para.paragraph_format.space_before = Pt(0)
            land_para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
            land_para.paragraph_format.line_spacing = 0.9
            land_run = land_para.add_run(land_text)
            land_run.font.size = FONT_SIZE_BODY
            land_run.font.color.rgb = RGBColor(0, 0, 0)

    # === DEVELOPMENT FEASIBILITY / ONGOING CONSTRUCTION (Bare Land Reports) ===
    if report.ongoing_construction_notes and report.report_type == 'bare_land':
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
        p.paragraph_format.line_spacing = 0.9

        run = p.add_run("Development Feasibility: ")
        run.font.bold = True
        run.font.size = FONT_SIZE_BODY
        run.font.color.rgb = RGBColor(0, 0, 0)

        run = p.add_run(report.ongoing_construction_notes)
        run.font.size = FONT_SIZE_BODY
        run.font.color.rgb = RGBColor(0, 0, 0)

    # === TOPOGRAPHICAL FEATURES (Bare Land specific) ===
    if report.report_type == 'bare_land':
        topo_parts = []

        if report.elevation_changes:
            topo_parts.append(f"elevation changes are {report.elevation_changes.replace('_', ' ')}")
        if report.drainage_pattern:
            topo_parts.append(f"drainage pattern is {report.drainage_pattern.replace('_', ' ')}")
        if report.vegetation_type:
            topo_parts.append(f"vegetation type is {report.vegetation_type.replace('_', ' ')}")
        if report.natural_features:
            topo_parts.append(f"natural features include {report.natural_features}")

        if topo_parts:
            topo_text = "The land has " + ", ".join(topo_parts) + "."
            topo_para = doc.add_paragraph()
            topo_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            topo_para.paragraph_format.space_before = Pt(4)
            topo_para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
            topo_para.paragraph_format.line_spacing = 0.9

            topo_label = topo_para.add_run("Topographical Features: ")
            topo_label.font.bold = True
            topo_label.font.size = FONT_SIZE_BODY

            topo_value = topo_para.add_run(topo_text)
            topo_value.font.size = FONT_SIZE_BODY

    # === PROPERTY PHOTOS (Bare Land - embedded in Section 4.0) ===
    if report.report_type == 'bare_land' and report.property_photos:
        property_photos = safe_get_json_field(report, 'property_photos', [])
        if property_photos:
            sorted_photos = sorted(property_photos, key=lambda x: x.get('order', 0))
            render_photo_grid(doc, sorted_photos)

    # === BUILDING DETAILS (Direct numbering: 4.1, 4.2, 4.3) ===
    if has_building_data:
        buildings = safe_get_json_field(report, 'buildings', [])
        for idx, building in enumerate(buildings):
            building_number = f"4.{idx + 1}"
            building_name = building.get('building_name', f'Building {idx + 1}')

            # Add building subsection heading
            add_section_heading(doc, building_number, building_name)

            # === CONSTRUCTION DETAILS (STANDALONE PARAGRAPH - NO LABEL) ===
            render_construction_details(doc, building)

            # === PROFESSIONAL STRUCTURED FORMAT (All buildings) ===

            # Use custom description text if provided
            if building.get('building_description_text'):
                opening_para = doc.add_paragraph()
                opening_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                opening_para.paragraph_format.space_before = Pt(0)
                opening_para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
                opening_para.paragraph_format.line_spacing = 0.9
                opening_run = opening_para.add_run(building.get('building_description_text'))
                opening_run.font.size = FONT_SIZE_BODY
                opening_run.font.color.rgb = RGBColor(0, 0, 0)
            else:
                # Build opening description from materials
                roof_types = building.get('roof_types', [])
                wall_types = building.get('wall_types', [])
                floor_types = building.get('floor_types', [])

                if roof_types or wall_types or floor_types:
                    desc_text = "This is constructed with "

                    # Roof
                    if roof_types:
                        roof_labels = {
                            'asbestos': 'asbestos sheets', 'tile': 'tiles',
                            'metal': 'metal sheets', 'concrete': 'concrete flat',
                            'cadjan': 'cadjans', 'zinc': 'zinc sheets'
                        }
                        roof_text = format_material_list(roof_types, roof_labels)
                        desc_text += f"{roof_text} on "

                    # Roof structure
                    stories = building.get('stories', 1)
                    if stories == 1:
                        desc_text += "timber frame roof "
                    else:
                        desc_text += f"{stories} storey structure "

                    # Walls
                    if wall_types:
                        wall_labels = {
                            'brick': 'brick masonry', 'cement_block': 'cement block',
                            'stone': 'stone', 'mud': 'mud', 'timber': 'timber',
                            'cadjan': 'cadjan', 'rcc_columns': 'RCC columns',
                            'rubble_masonry': 'rubble masonry'
                        }
                        wall_text = format_material_list(wall_types, wall_labels)
                        desc_text += f"supported by {wall_text} walls "

                    # Foundation/floor
                    if floor_types:
                        floor_labels = {
                            'cement': 'cement', 'tile': 'tiles', 'terrazzo': 'terrazzo',
                            'timber': 'timber', 'granite': 'granite', 'earth': 'earth',
                            'concrete': 'concrete'
                        }
                        floor_text = format_material_list(floor_types, floor_labels)
                        desc_text += f"and the floor is {floor_text} rendered"

                    desc_text += "."

                    opening_para = doc.add_paragraph()
                    opening_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    opening_para.paragraph_format.space_before = Pt(0)
                    opening_para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
                    opening_para.paragraph_format.line_spacing = 0.9
                    opening_run = opening_para.add_run(desc_text)
                    opening_run.font.size = FONT_SIZE_BODY
                    opening_run.font.color.rgb = RGBColor(0, 0, 0)

            # === ACCOMMODATION ===
            has_new_building_format = building.get('accommodation_summary') is not None
            floors = building.get('floors', [])
            has_old_floor_format = any(floor.get('accommodation_summary') for floor in floors)

            if has_new_building_format or has_old_floor_format:
                aggregated_rooms = aggregate_accommodation_across_building(building)

                room_parts = []

                if aggregated_rooms.get('bedrooms', 0) > 0:
                    count = aggregated_rooms['bedrooms']
                    room_parts.append(format_room_count(count, 'bedroom'))

                if aggregated_rooms.get('bathrooms', 0) > 0:
                    count = aggregated_rooms['bathrooms']
                    room_parts.append(format_room_count(count, 'bathroom'))

                if aggregated_rooms.get('living_rooms', 0) > 0:
                    count = aggregated_rooms['living_rooms']
                    room_parts.append(format_room_count(count, 'living room'))

                if aggregated_rooms.get('dining_rooms', 0) > 0:
                    count = aggregated_rooms['dining_rooms']
                    room_parts.append(format_room_count(count, 'dining room'))

                if aggregated_rooms.get('kitchens', 0) > 0:
                    count = aggregated_rooms['kitchens']
                    room_parts.append(format_room_count(count, 'kitchen'))

                if aggregated_rooms.get('pantries', 0) > 0:
                    count = aggregated_rooms['pantries']
                    room_parts.append(format_room_count(count, 'pantry', 'pantries'))

                if aggregated_rooms.get('verandahs', 0) > 0:
                    count = aggregated_rooms['verandahs']
                    room_parts.append(format_room_count(count, 'verandah'))

                if aggregated_rooms.get('balconies', 0) > 0:
                    count = aggregated_rooms['balconies']
                    room_parts.append(format_room_count(count, 'balcony', 'balconies'))

                if aggregated_rooms.get('garages', 0) > 0:
                    count = aggregated_rooms['garages']
                    room_parts.append(format_room_count(count, 'garage'))

                if aggregated_rooms.get('store_rooms', 0) > 0:
                    count = aggregated_rooms['store_rooms']
                    room_parts.append(format_room_count(count, 'store room'))

                if aggregated_rooms.get('other_rooms', 0) > 0:
                    count = aggregated_rooms['other_rooms']
                    room_parts.append(format_room_count(count, 'other room'))

                if room_parts:
                    accommodation_text = f"The property comprises {format_list_with_grammar(room_parts)}."
                    add_inline_field(doc, "Accommodation", accommodation_text)

            # === FLOOR AREA (CONDITIONAL FORMAT) ===
            floors = building.get('floors', [])
            total_area = building.get('total_floor_area', 0)
            if floors and (total_area or any(f.get('floor_area', 0) > 0 for f in floors)):
                # Filter floors with area > 0
                floors_with_area = [f for f in floors if f.get('floor_area', 0) > 0]

                if len(floors_with_area) == 1:
                    # SINGLE FLOOR: Use inline format
                    floor = floors_with_area[0]
                    floor_area = floor.get('floor_area', 0)
                    area_text = f"{floor_area:,.0f} square feet"
                    add_inline_field(doc, "Floor area", area_text)
                else:
                    # MULTIPLE FLOORS: Use table format
                    floor_area_para = doc.add_paragraph()
                    floor_area_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    floor_area_para.paragraph_format.space_before = INLINE_FIELD_SPACE_BEFORE
                    floor_area_para.paragraph_format.space_after = Pt(2)
                    floor_area_para.paragraph_format.line_spacing = 0.9

                    label_run = floor_area_para.add_run("Floor area:")
                    label_run.bold = True
                    label_run.font.size = FONT_SIZE_BODY
                    label_run.font.color.rgb = RGBColor(0, 0, 0)

                    # Add each floor's area (indented with aligned columns)
                    for floor in floors:
                        floor_name = floor.get('floor_name', 'Floor')
                        floor_area = floor.get('floor_area', 0)

                        if floor_area and floor_area > 0:
                            floor_line_para = doc.add_paragraph()
                            floor_line_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            floor_line_para.paragraph_format.space_before = Pt(0)
                            floor_line_para.paragraph_format.space_after = Pt(1)
                            floor_line_para.paragraph_format.line_spacing = 0.9
                            floor_line_para.paragraph_format.left_indent = Inches(0.5)

                            floor_name_padded = f"{floor_name:<25}"
                            floor_line_run = floor_line_para.add_run(
                                f"{floor_name_padded}{floor_area:>10,.0f} square feet"
                            )
                            floor_line_run.font.size = FONT_SIZE_BODY
                            floor_line_run.font.color.rgb = RGBColor(0, 0, 0)

                    # Add total (indented with aligned columns)
                    if total_area and total_area > 0:
                        total_para = doc.add_paragraph()
                        total_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        total_para.paragraph_format.space_before = Pt(0)
                        total_para.paragraph_format.space_after = INLINE_FIELD_SPACE_AFTER
                        total_para.paragraph_format.line_spacing = 0.9
                        total_para.paragraph_format.left_indent = Inches(0.5)

                        total_padded = f"{'Total':<25}"
                        total_run = total_para.add_run(f"{total_padded}{total_area:>10,.0f} square feet")
                        total_run.bold = True
                        total_run.font.size = FONT_SIZE_BODY
                        total_run.font.color.rgb = RGBColor(0, 0, 0)

            # === AGE AND CONDITION (INLINE FORMAT) ===
            building_age = building.get('building_age')
            condition = building.get('condition', '')
            if building_age or condition:
                condition_labels = {
                    'excellent': 'excellent',
                    'good': 'good',
                    'fair': 'fair',
                    'poor': 'poor',
                    'dilapidated': 'dilapidated'
                }

                parts = []
                if building_age is not None and building_age > 0:
                    year_word = "year" if building_age == 1 else "years"
                    parts.append(f"{building_age} {year_word} old")
                if condition:
                    parts.append(f"condition is {condition_labels.get(condition, condition)}")

                age_text = "; ".join(parts) + "." if parts else "Information not provided."

                add_inline_field(doc, "Age and condition", age_text)

            # === UTILITIES AND CONVENIENCES ===
            render_utilities_and_conveniences(doc, building)

            # === OCCUPATION (BUILDING-LEVEL) ===
            if building.get('occupier_name'):
                relationship = building.get('occupier_relationship')

                # Special case for vacant buildings
                if relationship == 'vacant':
                    occupier_text = "The building is currently vacant."
                else:
                    occupier_text = f"The building is occupied by {building.get('occupier_name')}"

                    if relationship:
                        rel_labels = {
                            'owner': 'the owner',
                            'tenant': 'a tenant',
                            'family_member': 'a family member',
                            'caretaker': 'caretaker'
                        }
                        occupier_text += f" who is {rel_labels.get(relationship, relationship)}."
                    else:
                        occupier_text += "."

                add_inline_field(doc, "Occupation", occupier_text, space_after=Pt(6))

            # Fallback: Support old reports with property-level occupier
            elif report.occupier_name:
                occupier_text = f"The property is occupied by {report.occupier_name}"
                if report.occupier_relationship:
                    rel_labels = {
                        'owner': 'the owner',
                        'tenant': 'a tenant',
                        'family_member': 'a family member',
                        'caretaker': 'caretaker'
                    }
                    occupier_text += (
                        f" who is {rel_labels.get(report.occupier_relationship, report.occupier_relationship)}."
                    )
                else:
                    occupier_text += "."

                add_inline_field(doc, "Occupation", occupier_text, space_after=Pt(6))

            # === BUILDING PHOTOS (3-column grid layout - NO SUBHEADING) ===
            building_photos = building.get('building_photos', [])
            if building_photos:
                sorted_photos = sorted(building_photos, key=lambda x: x.get('order', 0))
                render_photo_grid(doc, sorted_photos)

            # === ADDITIONAL STRUCTURES ===
            additional_structures = building.get('additional_structures_description', '')
            if additional_structures and additional_structures.strip():
                # Add subheading
                add_structures_heading = doc.add_paragraph()
                add_structures_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                add_structures_heading.paragraph_format.space_before = SUBHEADING_SPACE_BEFORE
                add_structures_heading.paragraph_format.space_after = SUBHEADING_SPACE_AFTER
                add_structures_heading.paragraph_format.line_spacing = 0.9
                add_structures_run = add_structures_heading.add_run("Additional Structures")
                add_structures_run.bold = True
                add_structures_run.font.size = FONT_SIZE_BODY
                add_structures_run.font.color.rgb = RGBColor(0, 0, 0)

                # Add description paragraph
                structures_para = doc.add_paragraph()
                structures_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                structures_para.paragraph_format.space_before = Pt(0)
                structures_para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
                structures_para.paragraph_format.line_spacing = 0.9
                structures_para.paragraph_format.left_indent = INDENTED_CONTENT_LEFT_INDENT

                structures_run = structures_para.add_run(additional_structures.strip())
                structures_run.font.size = FONT_SIZE_BODY
                structures_run.font.color.rgb = RGBColor(0, 0, 0)
