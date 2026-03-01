"""
Cover page (opening section) for single-property DOCX reports.
"""
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .styling import FONT_SIZE_BODY, BODY_PARA_SPACE_AFTER, SUBHEADING_SPACE_AFTER
from .text_generators import (
    generate_title_block, generate_applicant_statement,
    generate_organization_side_introduction, generate_deed_description,
    generate_submission_statement,
)


def render_cover_page(doc, report) -> None:
    """
    Render the opening section of a single-property DOCX report.

    Includes: title block, applicant/organization statements, deed description,
    submission statement, inspection date, and special note.
    """
    # Generate and add title block (centered)
    title_lines = generate_title_block(report)
    for i, line in enumerate(title_lines):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)

        # Remove spacing between property description lines (lines 2 and 3)
        if i == 2 or i == 3:  # Property description lines
            para.paragraph_format.line_spacing = 1.0  # Single line spacing, no gap
        else:
            para.paragraph_format.line_spacing = 0.9

        run = para.add_run(line)
        if i == 0:  # "VALUATION REPORT"
            run.bold = True
            run.font.size = Pt(14)
            run.font.underline = True
        elif i == 2 or i == 3:  # Both property description lines
            run.bold = True
            run.font.size = FONT_SIZE_BODY
            run.font.underline = True
        else:
            run.font.size = FONT_SIZE_BODY
            if i == 1:  # "of" line
                run.font.underline = True
        run.font.color.rgb = RGBColor(0, 0, 0)

    # Add spacing before applicant statements
    spacing_para1 = doc.add_paragraph()
    spacing_para1.paragraph_format.space_before = Pt(8)
    spacing_para1.paragraph_format.space_after = Pt(0)

    # Determine which introduction format to use based on request_type
    if report.request_type == 'organization_request':
        # Organization-side format
        org_intro_paragraphs = generate_organization_side_introduction(report)
        for statement in org_intro_paragraphs:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
            para.paragraph_format.line_spacing = 0.9

            # Parse label and value for proper formatting
            if ":-" in statement:
                # This is a label-value pair (Applicant/Address/Contact No)
                label, value = statement.split(":-", 1)

                # Add bold label
                run_label = para.add_run(label.strip())
                run_label.font.size = FONT_SIZE_BODY
                run_label.font.bold = True
                run_label.font.color.rgb = RGBColor(0, 0, 0)

                # Add separator with proper spacing
                run_sep = para.add_run(" : ")
                run_sep.font.size = FONT_SIZE_BODY
                run_sep.font.color.rgb = RGBColor(0, 0, 0)

                # Add value
                run_value = para.add_run(value)
                run_value.font.size = FONT_SIZE_BODY
                run_value.font.color.rgb = RGBColor(0, 0, 0)
            else:
                # This is the introductory paragraph (paragraph 1)
                run = para.add_run(statement)
                run.font.size = FONT_SIZE_BODY
                run.font.color.rgb = RGBColor(0, 0, 0)
    else:
        # Client-side format (default for backward compatibility)
        applicant_statements = generate_applicant_statement(report)
        for statement in applicant_statements:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
            para.paragraph_format.line_spacing = 0.9
            run = para.add_run(statement)
            run.font.size = FONT_SIZE_BODY
            run.font.color.rgb = RGBColor(0, 0, 0)

        # Add deed/certificate description ONLY if not plan-based identification
        should_show_deed_sentence = (
            report.property_identification_type in ['deed', 'certificate_of_sale']
            and report.has_deed_info == "yes"
            and report.deeds
        )

        # Backward compatibility: Old reports (NULL type) with deed data should show sentence
        if not report.property_identification_type and report.has_deed_info == "yes" and report.deeds:
            should_show_deed_sentence = True

        if should_show_deed_sentence:
            deed_text = generate_deed_description(report.deeds)
            if deed_text:
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
                para.paragraph_format.line_spacing = 0.9
                run = para.add_run(deed_text)
                run.font.size = FONT_SIZE_BODY
                run.font.color.rgb = RGBColor(0, 0, 0)

        # Add submission statement
        submission_text = generate_submission_statement(report)
        if submission_text:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
            para.paragraph_format.line_spacing = 0.9
            run = para.add_run(submission_text)
            run.font.size = FONT_SIZE_BODY
            run.font.color.rgb = RGBColor(0, 0, 0)

    # Add inspection date
    if report.inspection_date:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
        para.paragraph_format.line_spacing = 0.9

        # Bold label
        run_label = para.add_run("Date of Inspection:")
        run_label.font.size = FONT_SIZE_BODY
        run_label.font.bold = True
        run_label.font.color.rgb = RGBColor(0, 0, 0)

        # Regular date value
        run_date = para.add_run(f" {report.inspection_date}")
        run_date.font.size = FONT_SIZE_BODY
        run_date.font.color.rgb = RGBColor(0, 0, 0)

    # Add special note if applicable
    if report.has_special_note == "yes" and report.special_note_text:
        # Note label
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = SUBHEADING_SPACE_AFTER
        para.paragraph_format.line_spacing = 0.9
        run = para.add_run("Note:")
        run.bold = True
        run.font.size = FONT_SIZE_BODY
        run.font.color.rgb = RGBColor(0, 0, 0)

        # Note text
        note_para = doc.add_paragraph()
        note_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        note_para.paragraph_format.space_before = Pt(0)
        note_para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
        note_para.paragraph_format.line_spacing = 0.9
        note_run = note_para.add_run(report.special_note_text)
        note_run.font.size = FONT_SIZE_BODY
        note_run.font.color.rgb = RGBColor(0, 0, 0)
