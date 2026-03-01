"""
Valuation section renderer for single-property and multi-property DOCX reports.
"""
import logging

from docx.shared import Pt, Inches
from docx.enum.text import WD_TAB_ALIGNMENT

from .formatting import format_currency, format_currency_aligned, round_for_say
from .paragraph_builders import (
    add_section_heading, add_market_value_line, add_value_rounded_line,
    format_building_valuation_2line, format_addon_compact,
)
from .styling import FONT_SIZE_VALUATION, FONT_SIZE_BODY
from .helpers import safe_get_json_field, safe_parse_json_string, to_float

logger = logging.getLogger(__name__)


def render_valuation_section(doc, prop, report, section_num: int) -> int:
    """
    Add the valuation section (8.0) to the document.

    Renders land valuation, building valuations, addons, market value summary,
    and insurance values. Only renders if the property has valuation data.

    Args:
        doc: python-docx Document
        prop: Property model instance
        report: Report model instance
        section_num: Current section counter

    Returns:
        Updated section_num (incremented by 1 if the section was rendered)
    """
    if not (prop.valuation_total_land_value or prop.valuation_buildings_data):
        return section_num

    add_section_heading(doc, f"{section_num}.0", "VALUATION OF THE PROPERTY")
    section_num += 1

    # Land valuation
    if prop.valuation_total_land_value:
        extent = prop.valuation_land_extent or prop.land_extent_perches or 0
        rate = prop.valuation_rate_per_perch or 0
        land_value = prop.valuation_total_land_value

        p = doc.add_paragraph()
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(6.0), WD_TAB_ALIGNMENT.RIGHT)

        text = f"Land – {extent:,.2f} perches @ {format_currency(rate)} per perch\t= {format_currency_aligned(land_value)}"
        run = p.add_run(text)
        run.font.size = FONT_SIZE_VALUATION
        p.paragraph_format.space_after = Pt(3)

    # Buildings valuation (skip for bare_land)
    total_buildings_value = 0
    buildings_insurance_values = []

    if prop.property_type != 'bare_land' and prop.valuation_buildings_data:
        buildings_data = safe_get_json_field(prop, 'valuation_buildings_data', [])

        for idx, bldg in enumerate(buildings_data, 1):
            building_name = bldg.get('building_name', f'Building {idx}')
            subtotal = bldg.get('subtotal', 0)

            components = bldg.get('components', [])
            total_floor_area = sum(comp.get('floor_area', 0) for comp in components)

            avg_rate = subtotal / total_floor_area if total_floor_area > 0 else 0

            has_depreciation = (
                bldg.get('depreciation_amount') is not None and
                to_float(bldg.get('depreciation_amount', 0)) > 0
            )

            if has_depreciation:
                depreciation_rate = to_float(bldg.get('depreciation_rate_percent', 0))
                depreciated_value = to_float(bldg.get('depreciated_value', subtotal))

                format_building_valuation_2line(
                    doc,
                    building_name,
                    total_floor_area,
                    avg_rate,
                    depreciation_rate,
                    depreciated_value
                )

                building_value = depreciated_value
            else:
                p = doc.add_paragraph()
                tab_stops = p.paragraph_format.tab_stops
                tab_stops.add_tab_stop(Inches(6.0), WD_TAB_ALIGNMENT.RIGHT)

                text = f"{building_name} – {total_floor_area:,.0f} sq.ft @ {format_currency(avg_rate)} per square foot\t= {format_currency_aligned(subtotal)}"
                run = p.add_run(text)
                run.font.size = FONT_SIZE_VALUATION
                p.paragraph_format.space_after = Pt(3)
                building_value = to_float(subtotal)

            total_buildings_value += building_value
            buildings_insurance_values.append({
                'name': building_name,
                'value': to_float(subtotal)
            })

    # Addons/Improvements
    total_addons_value = 0
    if prop.valuation_addons:
        addons = safe_get_json_field(prop, 'valuation_addons', [])
        for addon in addons:
            addon_description = addon.get('description') or addon.get('item_name', 'Add-on')
            addon_value = to_float(addon.get('value', 0))
            format_addon_compact(doc, addon_description, addon_value)
            total_addons_value += addon_value

    # Calculate market values
    land_val = to_float(prop.valuation_total_land_value)
    market_value_calculated = land_val + total_buildings_value + total_addons_value
    market_value_rounded = round_for_say(market_value_calculated)

    has_buildings_or_addons = (total_buildings_value > 0) or (total_addons_value > 0)

    if has_buildings_or_addons:
        add_market_value_line(doc, market_value_calculated, has_blank_before=True)

    add_value_rounded_line(doc, market_value_rounded)

    # Summary of the valuation
    show_forced_sale = report.valuation_type == "Forced Sale Value"

    if show_forced_sale:
        forced_sale_percentage = prop.valuation_forced_sale_percentage or 90
        forced_sale_value = market_value_rounded * (forced_sale_percentage / 100)

    p = doc.add_paragraph()
    run = p.add_run("SUMMARY OF THE VALUATION")
    run.font.bold = True
    run.font.underline = True
    run.font.size = FONT_SIZE_BODY
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(3.5), WD_TAB_ALIGNMENT.LEFT)
    tab_stops.add_tab_stop(Inches(3.7), WD_TAB_ALIGNMENT.LEFT)
    text = f"Open Market Value of the property\t:\t{format_currency(market_value_rounded)}"
    run = p.add_run(text)
    run.font.size = FONT_SIZE_VALUATION
    p.paragraph_format.space_after = Pt(3)

    if show_forced_sale:
        p = doc.add_paragraph()
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(3.5), WD_TAB_ALIGNMENT.LEFT)
        tab_stops.add_tab_stop(Inches(3.7), WD_TAB_ALIGNMENT.LEFT)
        text = f"Forced Sale Value of the property\t:\t{format_currency(forced_sale_value)}"
        run = p.add_run(text)
        run.font.size = FONT_SIZE_VALUATION
        p.paragraph_format.space_after = Pt(3)

    if buildings_insurance_values:
        for building_ins in buildings_insurance_values:
            p = doc.add_paragraph()
            tab_stops = p.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(3.5), WD_TAB_ALIGNMENT.LEFT)
            tab_stops.add_tab_stop(Inches(3.7), WD_TAB_ALIGNMENT.LEFT)
            text = f"Insurance Value of the {building_ins['name']}\t:\t{format_currency(building_ins['value'])}"
            run = p.add_run(text)
            run.font.size = FONT_SIZE_VALUATION
            p.paragraph_format.space_after = Pt(2)

        doc.add_paragraph()  # Final spacing

    return section_num


def render_single_property_valuation(doc, report) -> None:
    """
    Render section 8.0 VALUATION OF THE PROPERTY for single-property reports.

    Renders land/buildings/addons valuation lines, market value summary,
    optional forced-sale value, and insurance values. Skips silently if
    the report has no valuation data.
    """
    if not (report.valuation_total_land_value or report.valuation_buildings_data):
        return

    add_section_heading(doc, "8.0", "VALUATION OF THE PROPERTY")

    # Land valuation line
    if report.valuation_total_land_value:
        extent = report.valuation_land_extent or report.land_extent_perches or 0
        rate = report.valuation_rate_per_perch or 0
        land_value = report.valuation_total_land_value

        p = doc.add_paragraph()
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(6.0), WD_TAB_ALIGNMENT.RIGHT)

        text = (
            f"Land \u2013 {extent:,.2f} perches @ {format_currency(rate)} per perch"
            f"\t= {format_currency_aligned(land_value)}"
        )
        run = p.add_run(text)
        run.font.size = FONT_SIZE_VALUATION
        p.paragraph_format.space_after = Pt(3)

    # Buildings valuation (SKIP for bare_land)
    total_depreciated_buildings_value = 0
    buildings_insurance_values = []

    if report.valuation_buildings_data and report.report_type != 'bare_land':
        buildings_data = safe_parse_json_string(report.valuation_buildings_data, [])

        for idx, bldg in enumerate(buildings_data, 1):
            building_name = bldg.get('building_name', f'Building {idx}')
            subtotal = bldg.get('subtotal', 0)

            components = bldg.get('components', [])
            total_floor_area = sum(comp.get('floor_area', 0) for comp in components)

            avg_rate = subtotal / total_floor_area if total_floor_area > 0 else 0

            has_depreciation = (
                bldg.get('depreciation_amount') is not None and
                to_float(bldg.get('depreciation_amount', 0)) > 0
            )

            if has_depreciation:
                depreciation_rate = to_float(bldg.get('depreciation_rate_percent', 0))
                depreciated_value = to_float(bldg.get('depreciated_value', subtotal))

                format_building_valuation_2line(
                    doc,
                    building_name,
                    total_floor_area,
                    avg_rate,
                    depreciation_rate,
                    depreciated_value
                )

                building_value = depreciated_value
            else:
                p = doc.add_paragraph()
                tab_stops = p.paragraph_format.tab_stops
                tab_stops.add_tab_stop(Inches(6.0), WD_TAB_ALIGNMENT.RIGHT)

                text = (
                    f"{building_name} \u2013 {total_floor_area:,.0f} sq.ft @ "
                    f"{format_currency(avg_rate)} per square foot"
                    f"\t= {format_currency_aligned(subtotal)}"
                )
                run = p.add_run(text)
                run.font.size = FONT_SIZE_VALUATION
                p.paragraph_format.space_after = Pt(3)
                building_value = to_float(subtotal)

            total_depreciated_buildings_value += building_value
            buildings_insurance_values.append({
                'name': building_name,
                'value': to_float(subtotal)  # Replacement cost for insurance
            })

    # Add-ons (compact format)
    total_addons_value = 0
    if report.valuation_addons:
        addons = safe_parse_json_string(report.valuation_addons, [])

        if addons:
            for addon in addons:
                format_addon_compact(
                    doc,
                    addon.get('description', 'Add-on'),
                    to_float(addon.get('value', 0))
                )
                total_addons_value += to_float(addon.get('value', 0))

    # Calculate market values
    land_value = to_float(report.valuation_total_land_value)
    market_value_calculated = (
        land_value + to_float(total_depreciated_buildings_value) + to_float(total_addons_value)
    )
    market_value_rounded = round_for_say(market_value_calculated)

    has_buildings_or_addons = (total_depreciated_buildings_value > 0) or (total_addons_value > 0)

    if has_buildings_or_addons:
        add_market_value_line(doc, market_value_calculated, has_blank_before=True)

    add_value_rounded_line(doc, market_value_rounded)

    show_forced_sale = report.valuation_type == "Forced Sale Value"

    if show_forced_sale:
        forced_sale_percentage = report.valuation_forced_sale_percentage or 90
        forced_sale_value = market_value_rounded * (forced_sale_percentage / 100)

    # === SUMMARY OF THE VALUATION ===
    p = doc.add_paragraph()
    run = p.add_run("SUMMARY OF THE VALUATION")
    run.font.bold = True
    run.font.underline = True
    run.font.size = FONT_SIZE_BODY
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)

    # Open Market Value
    p = doc.add_paragraph()
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(3.5), WD_TAB_ALIGNMENT.LEFT)
    tab_stops.add_tab_stop(Inches(3.7), WD_TAB_ALIGNMENT.LEFT)
    text = f"Open Market Value of the property\t:\t{format_currency(market_value_rounded)}"
    run = p.add_run(text)
    run.font.size = FONT_SIZE_VALUATION
    p.paragraph_format.space_after = Pt(3)

    # Forced Sale Value
    if show_forced_sale:
        p = doc.add_paragraph()
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(3.5), WD_TAB_ALIGNMENT.LEFT)
        tab_stops.add_tab_stop(Inches(3.7), WD_TAB_ALIGNMENT.LEFT)
        text = f"Forced Sale Value of the property\t:\t{format_currency(forced_sale_value)}"
        run = p.add_run(text)
        run.font.size = FONT_SIZE_VALUATION
        p.paragraph_format.space_after = Pt(3)

    # Insurance Value — only show if there are buildings
    if buildings_insurance_values:
        for building_ins in buildings_insurance_values:
            p = doc.add_paragraph()
            tab_stops = p.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(3.5), WD_TAB_ALIGNMENT.LEFT)
            tab_stops.add_tab_stop(Inches(3.7), WD_TAB_ALIGNMENT.LEFT)
            text = (
                f"Insurance Value of the {building_ins['name']}"
                f"\t:\t{format_currency(building_ins['value'])}"
            )
            run = p.add_run(text)
            run.font.size = FONT_SIZE_VALUATION
            p.paragraph_format.space_after = Pt(2)

        doc.add_paragraph()  # Final spacing
