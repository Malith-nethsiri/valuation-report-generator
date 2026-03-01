"""
Property sections rendering for multi-property DOCX reports.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from io import BytesIO
from datetime import datetime
from typing import Any, List, Dict, Optional, Union
import requests
from .. import models
from ..utils import append_label_if_missing, clean_spelling_errors, format_no_field
from ..letterhead_templates import get_template
import logging

from .formatting import (
    format_currency, format_currency_words, format_currency_aligned,
    format_room_count, round_for_say, format_material_list,
)
from .paragraph_builders import (
    add_section_heading, add_market_value_line, add_value_rounded_line,
    add_inline_field, add_subsection_paragraph, format_building_valuation_2line, format_addon_compact,
)
from .styling import (
    add_border_to_paragraph,
    MAP_IMAGE_WIDTH, MAP_IMAGE_MAX_HEIGHT, PROPERTY_PHOTO_WIDTH, PROPERTY_PHOTO_HEIGHT,
    IMAGE_SPACING_BEFORE, IMAGE_SPACING_AFTER, MAJOR_SECTION_SPACE_BEFORE, MAJOR_SECTION_SPACE_AFTER,
    SUBSECTION_SPACE_BEFORE, SUBSECTION_SPACE_AFTER, BODY_PARA_SPACE_BEFORE, BODY_PARA_SPACE_AFTER,
    INLINE_FIELD_SPACE_BEFORE, INLINE_FIELD_SPACE_AFTER, SUBHEADING_SPACE_BEFORE, SUBHEADING_SPACE_AFTER,
    INDENTED_CONTENT_SPACE_BEFORE, INDENTED_CONTENT_SPACE_AFTER, INDENTED_CONTENT_LEFT_INDENT,
    BOUNDARY_LIST_SPACE_AFTER, ACCOMMODATION_ROOM_SPACE_AFTER,
    FONT_SIZE_DOCUMENT_TITLE, FONT_SIZE_SECTION_HEADING, FONT_SIZE_SUBSECTION_HEADING,
    FONT_SIZE_BODY, FONT_SIZE_INLINE_LABEL, FONT_SIZE_VALUATION,
    FONT_SIZE_TABLE_HEADER, FONT_SIZE_TABLE_CELL, FONT_SIZE_INVOICE_TOTAL,
    FONT_SIZE_CAPTION, FONT_SIZE_BANK_HEADER, FONT_SIZE_BANK_DETAILS,
    FONT_SIZE_SIGNATURE, FONT_SIZE_CERTIFICATION,
)
from .images import calculate_image_dimensions, apply_letterbox_to_image
from .helpers import (
    safe_get_json_field, safe_get_array_item, to_float, safe_parse_json_string, safe_get_nested,
)

logger = logging.getLogger(__name__)
Deed_Type = Union[Dict[str, Any], Any]

from ..utils.text_helpers import format_list_with_grammar
from .building_renderer import (
    render_construction_details, render_utilities_and_conveniences,
    aggregate_accommodation_across_building, deduplicate_water_sources,
)
from .photo_section import render_photo_grid
from .valuation_section_renderer import render_valuation_section
from .text_generators import (
    generate_ownership_paragraph, generate_street_lines_paragraph,
    generate_building_limits_paragraph, generate_local_authority_paragraph,
    generate_rent_act_paragraph,
    generate_land_values_paragraph, generate_simplified_certification_text,
    generate_certificate_of_identity_text, add_signature_block, get_pronoun,
    generate_title_block, generate_applicant_statement,
    generate_organization_side_introduction, generate_multi_property_concluding_statement,
    generate_deed_description, generate_submission_statement, generate_situation_text,
    generate_smart_address, generate_access_text, generate_locality_description,
    generate_boundary_summary_text,
)

def generate_property_sections(doc, prop, report, user):
    """
    Generate comprehensive sections for a single property within a multi-property report.

    Generates full professional valuation report structure:
    1. SITUATION
    2. ACCESS (if available)
    3. IDENTIFICATION & BOUNDARIES
    4. DESCRIPTION (land + buildings + photos)
    5. LOCALITY (if available)
    6. LEGAL ASPECTS (if available)
    7. LAND VALUES / COMPARABLES (if available)
    8. VALUATION
    9. CERTIFICATION (per-property)

    Properties flow continuously without page breaks between them.
    """
    section_num = 1

    # ===== 1.0 SITUATION SECTION =====
    situation_text = generate_situation_text(prop)
    if situation_text:
        add_section_heading(doc, f"{section_num}.0", "SITUATION")
        section_num += 1

        situation_para = doc.add_paragraph()
        situation_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        situation_para.paragraph_format.space_before = Pt(0)
        situation_para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
        situation_para.paragraph_format.line_spacing = 0.9
        situation_run = situation_para.add_run(situation_text)
        situation_run.font.size = FONT_SIZE_BODY
        situation_run.font.color.rgb = RGBColor(0, 0, 0)

    # ===== 2.0 ACCESS SECTION (CONDITIONAL) =====
    access_text = generate_access_text(prop)
    if access_text:
        add_section_heading(doc, f"{section_num}.0", "ACCESS")
        section_num += 1

        access_para = doc.add_paragraph()
        access_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        access_para.paragraph_format.space_before = Pt(0)
        access_para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
        access_para.paragraph_format.line_spacing = 0.9
        access_run = access_para.add_run(access_text)
        access_run.font.size = FONT_SIZE_BODY
        access_run.font.color.rgb = RGBColor(0, 0, 0)

        # Add coordinates if available
        if prop.property_latitude and prop.property_longitude:
            coord_para = doc.add_paragraph()
            coord_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            coord_para.paragraph_format.space_before = Pt(6)
            coord_para.paragraph_format.space_after = Pt(6)
            coord_para.paragraph_format.line_spacing = 0.9

            coord_label = coord_para.add_run("Property Location Coordinates: ")
            coord_label.bold = True
            coord_label.font.size = FONT_SIZE_BODY
            coord_label.font.color.rgb = RGBColor(0, 0, 0)

            lat_value = float(prop.property_latitude)
            lng_value = float(prop.property_longitude)
            coord_text = coord_para.add_run(f"{lat_value:.6f}, {lng_value:.6f}")
            coord_text.font.size = FONT_SIZE_BODY
            coord_text.font.color.rgb = RGBColor(0, 0, 0)

        # Add map image if available
        if prop.location_map_image_data:
            try:
                map_url = prop.location_map_image_data
                logger.info(f"[DOCX-MULTI] Fetching map image for property {prop.id} (URL length={len(map_url)})")
                response = requests.get(map_url, timeout=30)
                if response.status_code == 200:
                    map_spacing_para = doc.add_paragraph()
                    map_spacing_para.paragraph_format.space_before = IMAGE_SPACING_BEFORE
                    map_spacing_para.paragraph_format.space_after = Pt(0)

                    map_para = doc.add_paragraph()
                    map_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    map_para.paragraph_format.space_before = Pt(0)
                    map_para.paragraph_format.space_after = IMAGE_SPACING_AFTER

                    image_stream = BytesIO(response.content)
                    dimensions = calculate_image_dimensions(
                        image_stream,
                        MAP_IMAGE_WIDTH,
                        MAP_IMAGE_MAX_HEIGHT
                    )
                    map_para.add_run().add_picture(image_stream, **dimensions)
                    logger.info(f"[DOCX-MULTI] Successfully added map image for property {prop.id}")
                else:
                    logger.warning(f"[DOCX-MULTI] Failed to fetch map image for property {prop.id}: HTTP {response.status_code}")
            except Exception as e:
                logger.warning(f"[DOCX-MULTI] Error adding map image for property {prop.id}: {str(e)}")
        else:
            logger.info(f"[DOCX-MULTI] No location_map_image_data for property {prop.id}")

    # ===== 3.0 IDENTIFICATION OF PROPERTY SECTION =====
    has_property_header_data = (
        prop.land_traditional_name or
        prop.land_extent_formatted or
        prop.boundaries or
        prop.physical_boundaries_types or
        prop.physical_boundaries_description or
        prop.boundaries_summary_text
    )

    if has_property_header_data:
        add_section_heading(doc, f"{section_num}.0", "IDENTIFICATION OF PROPERTY")
        section_num += 1

        # Traditional Land Name
        if prop.land_traditional_name:
            name_para = doc.add_paragraph()
            name_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            name_para.paragraph_format.space_before = Pt(0)
            name_para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
            name_para.paragraph_format.line_spacing = 0.9
            name_label = name_para.add_run("Name of Land: ")
            name_label.bold = True
            name_label.font.size = FONT_SIZE_BODY
            name_label.font.color.rgb = RGBColor(0, 0, 0)
            name_value = name_para.add_run(f'"{prop.land_traditional_name}"')
            name_value.font.size = FONT_SIZE_BODY
            name_value.font.color.rgb = RGBColor(0, 0, 0)

        # Survey Plan Information
        if prop.lot_number and prop.plan_number:
            plan_para = doc.add_paragraph()
            plan_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            plan_para.paragraph_format.space_before = Pt(0)
            plan_para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
            plan_para.paragraph_format.line_spacing = 0.9
            plan_label = plan_para.add_run("Survey Plan: ")
            plan_label.bold = True
            plan_label.font.size = FONT_SIZE_BODY
            plan_label.font.color.rgb = RGBColor(0, 0, 0)

            # Extract just the lot number/identifier (remove "Plan No" prefix if present)
            lot_desc = prop.lot_number.strip() if prop.lot_number else ''

            # Remove common prefixes that shouldn't be in lot description
            prefixes_to_remove = ['plan no', 'plan no:', 'lot plan no', 'lot plan no:']
            lot_desc_lower = lot_desc.lower()
            for prefix in prefixes_to_remove:
                if lot_desc_lower.startswith(prefix):
                    lot_desc = lot_desc[len(prefix):].strip()
                    break

            # Ensure lot description has "Lot" prefix
            if not lot_desc.lower().startswith('lot'):
                lot_desc = f"Lot {lot_desc}"

            # Format as "Lot X in Plan No: Y dated [date] made by [surveyor], Licensed Surveyor"
            plan_formatted = format_no_field("Plan", prop.plan_number)
            plan_text = f"{lot_desc} in {plan_formatted}"
            if prop.plan_date:
                plan_text += f" dated {prop.plan_date}"
            if prop.licensed_surveyor_name:
                plan_text += f" made by {prop.licensed_surveyor_name}, Licensed Surveyor"
            plan_text += "."

            plan_value = plan_para.add_run(plan_text)
            plan_value.font.size = FONT_SIZE_BODY
            plan_value.font.color.rgb = RGBColor(0, 0, 0)

        # Land Extent
        if prop.land_extent_formatted:
            extent_para = doc.add_paragraph()
            extent_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            extent_para.paragraph_format.space_before = Pt(0)
            extent_para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
            extent_para.paragraph_format.line_spacing = 0.9
            extent_label = extent_para.add_run("Extent: ")
            extent_label.bold = True
            extent_label.font.size = FONT_SIZE_BODY
            extent_label.font.color.rgb = RGBColor(0, 0, 0)
            extent_value = extent_para.add_run(prop.land_extent_formatted)
            extent_value.font.size = FONT_SIZE_BODY
            extent_value.font.color.rgb = RGBColor(0, 0, 0)

        # Boundaries - Add detailed boundary descriptions
        boundaries = safe_get_json_field(prop, 'boundaries', None)
        if boundaries:
            boundary_para = doc.add_paragraph()
            boundary_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            boundary_para.paragraph_format.space_before = Pt(6)
            boundary_para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
            boundary_para.paragraph_format.line_spacing = 0.9

            boundary_label = boundary_para.add_run("Boundaries:\n")
            boundary_label.bold = True
            boundary_label.font.size = FONT_SIZE_BODY
            boundary_label.font.color.rgb = RGBColor(0, 0, 0)

            # Add each direction boundary (4 main + 4 optional diagonal)
            for direction in ['North', 'North-East', 'East', 'South-East', 'South', 'South-West', 'West', 'North-West']:
                direction_key = direction.lower().replace('-', '')  # 'North-East' -> 'northeast'
                if isinstance(boundaries, dict) and direction_key in boundaries:
                    boundary_data = boundaries[direction_key]
                    description = boundary_data.get('description', '') if isinstance(boundary_data, dict) else str(boundary_data)

                    # Skip diagonal boundaries if empty
                    if not description and direction_key in ['northeast', 'southeast', 'southwest', 'northwest']:
                        continue

                    if description:
                        boundary_text = f"{direction}: {description}\n"
                        run = boundary_para.add_run(boundary_text)
                        run.font.size = FONT_SIZE_BODY
                        run.font.color.rgb = RGBColor(0, 0, 0)

    # ===== 4.0 DESCRIPTION OF PROPERTY SECTION =====
    add_section_heading(doc, f"{section_num}.0", "DESCRIPTION OF PROPERTY")
    section_num += 1

    # Land Description
    land_desc_text = prop.land_description_text
    if land_desc_text:
        desc_para = doc.add_paragraph()
        desc_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        desc_para.paragraph_format.space_before = Pt(0)
        desc_para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
        desc_para.paragraph_format.line_spacing = 0.9
        desc_run = desc_para.add_run(land_desc_text)
        desc_run.font.size = FONT_SIZE_BODY
        desc_run.font.color.rgb = RGBColor(0, 0, 0)

    # Buildings (if applicable - skip for bare land) - USE SAME FORMAT AS STANDALONE REPORTS
    buildings = safe_get_json_field(prop, 'buildings', [])
    if buildings and prop.property_type != 'bare_land':
        for building_idx, building in enumerate(buildings, 1):
            # Use numbered subsection heading (4.1, 4.2, etc.) - SAME AS STANDALONE
            building_number = f"{section_num - 1}.{building_idx}"
            building_name = building.get('building_name', f'Building {building_idx}')

            # Add building subsection heading using the same helper
            add_section_heading(doc, building_number, building_name)

            # === CONSTRUCTION DETAILS - USE SAME PROFESSIONAL FORMAT AS STANDALONE ===
            render_construction_details(doc, building)

            # === PROFESSIONAL BUILDING DESCRIPTION - SAME AS STANDALONE ===
            # Opening paragraph: Construction materials and general description
            if building.get('building_description_text'):
                opening_para = doc.add_paragraph()
                opening_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                opening_para.paragraph_format.space_before = Pt(0)
                opening_para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
                opening_para.paragraph_format.line_spacing = 0.9
                opening_run = opening_para.add_run(building.get('building_description_text'))
                opening_run.font.size = FONT_SIZE_BODY
                opening_run.font.color.rgb = RGBColor(0, 0, 0)
            else:
                # Build opening description from materials (same logic as standalone)
                roof_types = building.get('roof_types', [])
                wall_types = building.get('wall_types', [])
                floor_types = building.get('floor_types', [])

                if roof_types or wall_types or floor_types:
                    desc_text = "This is constructed with "

                    # Roof
                    if roof_types:
                        roof_labels = {
                            'asbestos': 'asbestos sheets', 'tile': 'tiles',
                            'metal': 'metal sheets', 'concrete': 'concrete flat',
                            'cadjan': 'cadjans', 'zinc': 'zinc sheets'
                        }
                        roof_text = format_material_list(roof_types, roof_labels)
                        desc_text += f"{roof_text} on "

                    # Roof structure
                    stories = building.get('stories', 1)
                    if stories == 1:
                        desc_text += "timber frame roof "
                    else:
                        desc_text += f"{stories} storey structure "

                    # Walls
                    if wall_types:
                        wall_labels = {
                            'brick': 'brick masonry', 'cement_block': 'cement block',
                            'stone': 'stone', 'mud': 'mud', 'timber': 'timber',
                            'cadjan': 'cadjan', 'rcc_columns': 'RCC columns',
                            'rubble_masonry': 'rubble masonry'
                        }
                        wall_text = format_material_list(wall_types, wall_labels)
                        desc_text += f"supported by {wall_text} walls "

                    # Foundation/floor
                    if floor_types:
                        floor_labels = {
                            'cement': 'cement', 'tile': 'tiles', 'terrazzo': 'terrazzo',
                            'timber': 'timber', 'granite': 'granite', 'earth': 'earth',
                            'concrete': 'concrete'
                        }
                        floor_text = format_material_list(floor_types, floor_labels)
                        desc_text += f"and the floor is {floor_text} rendered"

                    desc_text += "."

                    opening_para = doc.add_paragraph()
                    opening_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    opening_para.paragraph_format.space_before = Pt(0)
                    opening_para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
                    opening_para.paragraph_format.line_spacing = 0.9
                    opening_run = opening_para.add_run(desc_text)
                    opening_run.font.size = FONT_SIZE_BODY
                    opening_run.font.color.rgb = RGBColor(0, 0, 0)

            # === ACCOMMODATION - USE SAME AGGREGATED FORMAT AS STANDALONE ===
            # Check if building has accommodation data (new structure at building level or old structure at floor level)
            has_new_building_format = building.get('accommodation_summary') is not None
            floors = building.get('floors', [])
            has_old_floor_format = any(floor.get('accommodation_summary') for floor in floors)

            if has_new_building_format or has_old_floor_format:
                # Get accommodation summary (supports both new and old structures)
                aggregated_rooms = aggregate_accommodation_across_building(building)

                room_parts = []

                if aggregated_rooms.get('bedrooms', 0) > 0:
                    count = aggregated_rooms['bedrooms']
                    room_parts.append(format_room_count(count, 'bedroom'))

                if aggregated_rooms.get('bathrooms', 0) > 0:
                    count = aggregated_rooms['bathrooms']
                    room_parts.append(format_room_count(count, 'bathroom'))

                if aggregated_rooms.get('living_rooms', 0) > 0:
                    count = aggregated_rooms['living_rooms']
                    room_parts.append(format_room_count(count, 'living room'))

                if aggregated_rooms.get('dining_rooms', 0) > 0:
                    count = aggregated_rooms['dining_rooms']
                    room_parts.append(format_room_count(count, 'dining room'))

                if aggregated_rooms.get('kitchens', 0) > 0:
                    count = aggregated_rooms['kitchens']
                    room_parts.append(format_room_count(count, 'kitchen'))

                if aggregated_rooms.get('pantries', 0) > 0:
                    count = aggregated_rooms['pantries']
                    room_parts.append(format_room_count(count, 'pantry', 'pantries'))

                if aggregated_rooms.get('verandahs', 0) > 0:
                    count = aggregated_rooms['verandahs']
                    room_parts.append(format_room_count(count, 'verandah'))

                if aggregated_rooms.get('balconies', 0) > 0:
                    count = aggregated_rooms['balconies']
                    room_parts.append(format_room_count(count, 'balcony', 'balconies'))

                if aggregated_rooms.get('garages', 0) > 0:
                    count = aggregated_rooms['garages']
                    room_parts.append(format_room_count(count, 'garage'))

                if aggregated_rooms.get('store_rooms', 0) > 0:
                    count = aggregated_rooms['store_rooms']
                    room_parts.append(format_room_count(count, 'store room'))

                if aggregated_rooms.get('other_rooms', 0) > 0:
                    count = aggregated_rooms['other_rooms']
                    room_parts.append(format_room_count(count, 'other room'))

                if room_parts:
                    accommodation_text = f"The property comprises {format_list_with_grammar(room_parts)}."
                    add_inline_field(doc, "Accommodation", accommodation_text)

            # === FLOOR AREA BREAKDOWN - CONDITIONAL FORMAT ===
            total_area = building.get('total_floor_area', 0)
            if floors and (total_area or any(f.get('floor_area', 0) > 0 for f in floors)):
                # Filter floors with area > 0
                floors_with_area = [f for f in floors if f.get('floor_area', 0) > 0]

                if len(floors_with_area) == 1:
                    # SINGLE FLOOR: Use inline format
                    floor = floors_with_area[0]
                    floor_area = floor.get('floor_area', 0)
                    area_text = f"{floor_area:,.0f} square feet"
                    add_inline_field(doc, "Floor area", area_text)
                else:
                    # MULTIPLE FLOORS: Use table format
                    floor_area_para = doc.add_paragraph()
                    floor_area_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    floor_area_para.paragraph_format.space_before = INLINE_FIELD_SPACE_BEFORE
                    floor_area_para.paragraph_format.space_after = Pt(2)
                    floor_area_para.paragraph_format.line_spacing = 0.9

                    label_run = floor_area_para.add_run("Floor area:")
                    label_run.bold = True
                    label_run.font.size = FONT_SIZE_BODY
                    label_run.font.color.rgb = RGBColor(0, 0, 0)

                    for floor in floors:
                        floor_name = floor.get('floor_name', 'Floor')
                        floor_area = floor.get('floor_area', 0)

                        if floor_area and floor_area > 0:
                            floor_line_para = doc.add_paragraph()
                            floor_line_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            floor_line_para.paragraph_format.space_before = Pt(0)
                            floor_line_para.paragraph_format.space_after = Pt(1)
                            floor_line_para.paragraph_format.line_spacing = 0.9
                            floor_line_para.paragraph_format.left_indent = Inches(0.5)

                            floor_name_padded = f"{floor_name:<25}"
                            floor_line_run = floor_line_para.add_run(f"{floor_name_padded}{floor_area:>10,.0f} square feet")
                            floor_line_run.font.size = FONT_SIZE_BODY
                            floor_line_run.font.color.rgb = RGBColor(0, 0, 0)

                    if total_area and total_area > 0:
                        total_para = doc.add_paragraph()
                        total_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        total_para.paragraph_format.space_before = Pt(0)
                        total_para.paragraph_format.space_after = INLINE_FIELD_SPACE_AFTER
                        total_para.paragraph_format.line_spacing = 0.9
                        total_para.paragraph_format.left_indent = Inches(0.5)

                        total_padded = f"{'Total':<25}"
                        total_run = total_para.add_run(f"{total_padded}{total_area:>10,.0f} square feet")
                        total_run.bold = True
                        total_run.font.size = FONT_SIZE_BODY
                        total_run.font.color.rgb = RGBColor(0, 0, 0)

            # === AGE AND CONDITION - USE SAME FORMAT AS STANDALONE ===
            building_age = building.get('building_age')
            condition = building.get('condition', '')
            if building_age or condition:
                condition_labels = {
                    'excellent': 'excellent', 'good': 'good', 'fair': 'fair',
                    'poor': 'poor', 'dilapidated': 'dilapidated'
                }

                parts = []
                if building_age is not None and building_age > 0:
                    year_word = "year" if building_age == 1 else "years"
                    parts.append(f"{building_age} {year_word} old")
                if condition:
                    parts.append(f"condition is {condition_labels.get(condition, condition)}")

                age_text = "; ".join(parts) + "." if parts else "Information not provided."
                add_inline_field(doc, "Age and condition", age_text)

            # === UTILITIES AND CONVENIENCES - UNIFIED SECTION ===
            render_utilities_and_conveniences(doc, building)

            # === OCCUPATION (BUILDING-LEVEL) ===
            # Primary: Use building-level occupier data
            if building.get('occupier_name'):
                relationship = building.get('occupier_relationship')

                # Special case for vacant buildings
                if relationship == 'vacant':
                    occupier_text = "The building is currently vacant."
                else:
                    occupier_text = f"The building is occupied by {building.get('occupier_name')}"

                    if relationship:
                        rel_labels = {
                            'owner': 'the owner',
                            'tenant': 'a tenant',
                            'family_member': 'a family member',
                            'caretaker': 'caretaker'
                        }
                        occupier_text += f" who is {rel_labels.get(relationship, relationship)}."
                    else:
                        occupier_text += "."

                add_inline_field(doc, "Occupation", occupier_text, space_after=Pt(6))

            # Fallback: Support old reports with property-level occupier
            elif prop.occupier_name:
                occupier_text = f"The property is occupied by {prop.occupier_name}"
                if prop.occupier_relationship:
                    rel_labels = {
                        'owner': 'the owner',
                        'tenant': 'a tenant',
                        'family_member': 'a family member',
                        'caretaker': 'caretaker'
                    }
                    occupier_text += f" who is {rel_labels.get(prop.occupier_relationship, prop.occupier_relationship)}."
                else:
                    occupier_text += "."

                add_inline_field(doc, "Occupation", occupier_text, space_after=Pt(6))

            # === BUILDING PHOTOS ===
            building_photos = building.get('building_photos', [])
            if building_photos:
                sorted_photos = sorted(building_photos, key=lambda x: x.get('order', 0))
                render_photo_grid(doc, sorted_photos)

    # === DEVELOPMENT FEASIBILITY / ONGOING CONSTRUCTION (Bare Land only) ===
    if prop.property_type == 'bare_land' and prop.ongoing_construction_notes:
        construction_para = doc.add_paragraph()
        construction_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        construction_para.paragraph_format.space_before = Pt(4)
        construction_para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
        construction_para.paragraph_format.line_spacing = 0.9

        label_run = construction_para.add_run("Development Feasibility: ")
        label_run.font.bold = True
        label_run.font.size = FONT_SIZE_BODY
        label_run.font.color.rgb = RGBColor(0, 0, 0)

        value_run = construction_para.add_run(prop.ongoing_construction_notes)
        value_run.font.size = FONT_SIZE_BODY
        value_run.font.color.rgb = RGBColor(0, 0, 0)

    # === TOPOGRAPHICAL FEATURES (Bare Land only) ===
    if prop.property_type == 'bare_land':
        topo_parts = []

        if prop.elevation_changes:
            topo_parts.append(f"elevation changes are {prop.elevation_changes.replace('_', ' ')}")
        if prop.drainage_pattern:
            topo_parts.append(f"drainage pattern is {prop.drainage_pattern.replace('_', ' ')}")
        if prop.vegetation_type:
            topo_parts.append(f"vegetation type is {prop.vegetation_type.replace('_', ' ')}")
        if prop.natural_features:
            topo_parts.append(f"natural features include {prop.natural_features}")

        if topo_parts:
            topo_text = "The land has " + ", ".join(topo_parts) + "."
            topo_para = doc.add_paragraph()
            topo_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            topo_para.paragraph_format.space_before = Pt(4)
            topo_para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
            topo_para.paragraph_format.line_spacing = 0.9

            label_run = topo_para.add_run("Topographical Features: ")
            label_run.font.bold = True
            label_run.font.size = FONT_SIZE_BODY
            label_run.font.color.rgb = RGBColor(0, 0, 0)

            value_run = topo_para.add_run(topo_text)
            value_run.font.size = FONT_SIZE_BODY
            value_run.font.color.rgb = RGBColor(0, 0, 0)

    # Property Photos (embedded in DESCRIPTION section) - USE SAME GRID TABLE FORMAT AS STANDALONE
    # Property Photos (embedded in DESCRIPTION section)
    if prop.property_type == 'bare_land':
        property_photos = safe_get_json_field(prop, 'property_photos', [])
        if property_photos:
            sorted_photos = sorted(property_photos, key=lambda x: x.get('order', 0))
            render_photo_grid(doc, sorted_photos)

    # ===== 5.0 LOCALITY SECTION (CONDITIONAL) =====
    locality_text = generate_locality_description(prop)
    if locality_text or prop.locality_description_text:
        add_section_heading(doc, f"{section_num}.0", "LOCALITY")
        section_num += 1

        final_locality_text = prop.locality_description_text or locality_text

        locality_para = doc.add_paragraph()
        locality_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        locality_para.paragraph_format.space_before = Pt(0)
        locality_para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
        locality_para.paragraph_format.line_spacing = 0.9
        locality_run = locality_para.add_run(final_locality_text)
        locality_run.font.size = FONT_SIZE_BODY
        locality_run.font.color.rgb = RGBColor(0, 0, 0)

    # ===== 6.0 LEGAL ASPECTS SECTION (CONDITIONAL) =====
    has_legal_data = (
        prop.ownership_type or prop.street_lines_status or prop.building_limits_status or
        prop.local_authority_data or prop.rent_act_effectiveness or
        prop.title_search_conducted or prop.property_encumbered or
        prop.street_lines_gazette_ref or prop.building_plan_approved or
        prop.local_authority_rated
    )

    if has_legal_data:
        add_section_heading(doc, f"{section_num}.0", "LEGAL ASPECTS")
        section_num += 1

        # (a) Ownership
        if prop.ownership_type or hasattr(prop, 'deeds') or hasattr(prop, 'plan_number'):
            ownership_para = generate_ownership_paragraph(prop)
            add_subsection_paragraph(doc, "(a)", "Ownership", ownership_para)

        # (b) Street lines (skip for bare_land)
        if prop.street_lines_status and prop.property_type != 'bare_land':
            street_para = generate_street_lines_paragraph(prop)
            add_subsection_paragraph(doc, "(b)", "Street lines", street_para)

        # (c) Building limits (skip for bare_land)
        if prop.building_limits_status and prop.property_type != 'bare_land':
            building_para = generate_building_limits_paragraph(prop)
            add_subsection_paragraph(doc, "(c)", "Building limits", building_para)

        # (d) Local authority data
        if prop.local_authority_data or prop.pradeshiya_sabha or prop.local_authority_rated:
            authority_para = generate_local_authority_paragraph(prop)
            add_subsection_paragraph(doc, "(d)", "Local authority data", authority_para)

        # (e) Rent act effectiveness
        if prop.rent_act_effectiveness:
            rent_para = generate_rent_act_paragraph(prop)
            add_subsection_paragraph(doc, "(e)", "Rent act effectiveness", rent_para)

        doc.add_paragraph()  # Spacing

    # ===== 7.0 LAND VALUES IN THE AREA (CONDITIONAL) =====
    comparables = safe_get_json_field(prop, 'comparable_properties', [])
    # Show section if either comparables OR market analysis exists
    if comparables or prop.land_market_analysis:
        add_section_heading(doc, f"{section_num}.0", "LAND VALUES IN THE AREA")
        section_num += 1

        land_values_text = generate_land_values_paragraph(comparables)
        if land_values_text:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(land_values_text)
            run.font.size = FONT_SIZE_BODY

            # Average rate
            avg_rate = sum(c.get('rate_per_perch', 0) for c in comparables) / len(comparables) if comparables else 0
            if avg_rate > 0:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run(f"Average Rate: LKR {avg_rate:,.2f} per perch")
                run.font.bold = True
                run.font.size = FONT_SIZE_BODY

        if prop.land_market_analysis:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run(prop.land_market_analysis)
            run.font.size = FONT_SIZE_BODY

        doc.add_paragraph()

    # ===== 8.0 VALUATION SECTION =====
    section_num = render_valuation_section(doc, prop, report, section_num)

    # ===== 9.0 CERTIFICATION SECTION (SIMPLIFIED FORMAT - REPORT-LEVEL) =====
    add_section_heading(doc, f"{section_num}.0", "CERTIFICATION")
    section_num += 1

    # Use report-level certification - simplified format
    if report.certification_text:
        # Use custom certification text if provided (override mode)
        cert_text = report.certification_text.strip()
    elif report.certification_valuer_name and report.certification_valuer_designation:
        # Auto-generate simplified certification using report-level data
        # Multi-property uses single unified certification
        cert_text = generate_simplified_certification_text(
            valuer_name=report.certification_valuer_name,
            valuer_designation=report.certification_valuer_designation,
            lot_number=report.lot_number,
            plan_number=report.plan_number,
            plan_date=report.plan_date,
            licensed_surveyor_name=report.licensed_surveyor_name,
            deeds=safe_get_json_field(report, 'deeds', []),
            property_identification_type=report.property_identification_type
        )
    else:
        # Fallback if no valuer info
        cert_text = "I hereby certify that I have personally inspected this property and the valuation stated herein represents my professional opinion of the market value as of the date of inspection."

    # Render certification paragraph
    cert_para = doc.add_paragraph()
    cert_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    cert_para.paragraph_format.line_spacing = 0.9
    cert_para.paragraph_format.space_after = Pt(12)
    cert_run = cert_para.add_run(cert_text)
    cert_run.font.size = FONT_SIZE_BODY
    cert_run.font.color.rgb = RGBColor(0, 0, 0)

    # Signature block
    add_signature_block(
        doc=doc,
        user=user,
        valuer_name=report.certification_valuer_name,
        valuer_designation=report.certification_valuer_designation,
        certification_date=report.certification_date
    )


