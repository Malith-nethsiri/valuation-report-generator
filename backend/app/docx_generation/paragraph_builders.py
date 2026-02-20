from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_UNDERLINE
from docx.shared import Inches, Pt, RGBColor
from docx.document import Document
from .styling import (
    MAJOR_SECTION_SPACE_BEFORE, MAJOR_SECTION_SPACE_AFTER,
    SUBSECTION_SPACE_BEFORE, SUBSECTION_SPACE_AFTER,
    INLINE_FIELD_SPACE_BEFORE, INLINE_FIELD_SPACE_AFTER,
    FONT_SIZE_BODY, FONT_SIZE_SECTION_HEADING,
    FONT_SIZE_SUBSECTION_HEADING, FONT_SIZE_VALUATION
)
from .formatting import format_currency, format_currency_aligned

def add_section_heading(doc: Document, section_number: str, section_title: str):
    """
    Add a hierarchically numbered section heading.

    Args:
        doc: Document object
        section_number: Section number (e.g., "1.0", "4.7.1")
        section_title: Section title text (e.g., "SITUATION", "PHOTOGRAPHS")

    Returns:
        Paragraph object containing the heading
    """
    # Determine if this is a major section (X.0) or subsection (X.Y)
    is_subsection = '.' in section_number and section_number.count('.') >= 1 and not section_number.endswith('.0')

    # Add spacing before section
    spacing_para = doc.add_paragraph()
    if is_subsection:
        spacing_para.paragraph_format.space_before = SUBSECTION_SPACE_BEFORE
    else:
        spacing_para.paragraph_format.space_before = MAJOR_SECTION_SPACE_BEFORE
    spacing_para.paragraph_format.space_after = Pt(0)

    # Create heading paragraph
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading.paragraph_format.space_before = Pt(0)
    if is_subsection:
        heading.paragraph_format.space_after = SUBSECTION_SPACE_AFTER
    else:
        heading.paragraph_format.space_after = MAJOR_SECTION_SPACE_AFTER
    heading.paragraph_format.line_spacing = 0.9

    # Add section number and title
    heading_text = f"{section_number}. {section_title}"
    heading_run = heading.add_run(heading_text)
    heading_run.bold = True
    # Use different font sizes: 13pt for major sections, 12pt for subsections
    if is_subsection:
        heading_run.font.size = FONT_SIZE_SUBSECTION_HEADING
    else:
        heading_run.font.size = FONT_SIZE_SECTION_HEADING
    heading_run.font.color.rgb = RGBColor(0, 0, 0)

    return heading



def format_building_valuation_2line(doc: Document, building_name: str, total_floor_area: float,
                                    avg_rate: float, depreciation_rate: float,
                                    depreciated_value: float) -> None:
    """
    Add building valuation in new 2-line compact format.

    Line 1: Building name – area
    Line 2: @ rate less X% for dep[x.XX] = final value

    Args:
        doc: Document object
        building_name: Name of the building
        total_floor_area: Total floor area in sq.ft
        avg_rate: Average rate per sq.ft (before depreciation)
        depreciation_rate: Depreciation percentage (e.g., 25 for 25%)
        depreciated_value: Final depreciated value
    """
    # Line 1: Building name and area
    p1 = doc.add_paragraph()
    text1 = f"{building_name} – {total_floor_area:,.0f} sq.ft"
    run1 = p1.add_run(text1)
    run1.font.size = FONT_SIZE_VALUATION
    p1.paragraph_format.space_after = Pt(2)

    # Line 2: Rate + depreciation inline with multiplier
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Inches(0.3)

    # Add tab stop for right alignment
    tab_stops = p2.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.0), WD_TAB_ALIGNMENT.RIGHT)

    # Calculate depreciation multiplier
    multiplier = (100 - depreciation_rate) / 100

    text2 = f"@ {format_currency(avg_rate)} per square foot less {depreciation_rate:.0f}% for dep[x{multiplier:.2f}]\t= {format_currency_aligned(depreciated_value)}"
    run2 = p2.add_run(text2)
    run2.font.size = FONT_SIZE_VALUATION
    p2.paragraph_format.space_after = Pt(3)


def add_market_value_line(doc: Document, calculated_value: float, has_blank_before: bool = True) -> None:
    """
    Add "Market Value of the property" line with double-underlined value.

    Args:
        doc: Document object
        calculated_value: The calculated (non-rounded) market value
        has_blank_before: Whether to add blank line before (default True)
    """
    p = doc.add_paragraph()

    # Add spacing before instead of blank line
    if has_blank_before:
        p.paragraph_format.space_before = Pt(6)

    # Add tab stop at 6 inches for right alignment
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.0), WD_TAB_ALIGNMENT.RIGHT)

    # Add label (bold, underlined)
    run_label = p.add_run("Market Value of the property")
    run_label.font.bold = True
    run_label.font.underline = True
    run_label.font.size = FONT_SIZE_BODY

    # Add tab and equals sign
    run_eq = p.add_run("\t= ")
    run_eq.font.bold = True
    run_eq.font.size = FONT_SIZE_BODY

    # Add value (bold, double underlined)
    run_value = p.add_run(format_currency(calculated_value))
    run_value.font.bold = True
    run_value.font.underline = WD_UNDERLINE.DOUBLE
    run_value.font.size = FONT_SIZE_BODY

    p.paragraph_format.space_after = Pt(3)


def add_value_rounded_line(doc: Document, rounded_value: float) -> None:
    """
    Add centered "Value rounded off" line (bold, underlined).

    Args:
        doc: Document object
        rounded_value: The rounded value using round_for_say()
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)

    text = f"Value rounded off to {format_currency(rounded_value)}"
    run = p.add_run(text)
    run.font.bold = True
    run.font.underline = True
    run.font.size = FONT_SIZE_BODY

    p.paragraph_format.space_after = Pt(12)


def format_addon_compact(doc: Document, description: str, value: float) -> None:
    """
    Add add-on in single-line format (description and value on same line).

    Args:
        doc: Document object
        description: Add-on description
        value: Add-on value
    """
    p = doc.add_paragraph()

    # Add tab stop for right alignment
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.0), WD_TAB_ALIGNMENT.RIGHT)

    text = f"{description}\t= {format_currency_aligned(value)}"
    run = p.add_run(text)
    run.font.size = FONT_SIZE_VALUATION
    p.paragraph_format.space_after = Pt(3)


def add_inline_field(doc: Document, label: str, content: str,
                    space_before : float | None =None,
                    space_after : float | None =None) -> None:
    """
    Add a single-line field in format "Label: content" (inline format).

    Args:
        doc: Document object
        label: Bold label text (e.g., "Floor area")
        content: Content text (e.g., "1250 square feet approximately.")
        space_before: Override default spacing before (optional)
        space_after: Override default spacing after (optional)

    Example output:
        "Floor area: 1250 square feet approximately."
        (where "Floor area:" is bold and rest is regular)
    """
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = space_before if space_before is not None else INLINE_FIELD_SPACE_BEFORE
    para.paragraph_format.space_after = space_after if space_after is not None else INLINE_FIELD_SPACE_AFTER
    para.paragraph_format.line_spacing = 0.9

    # Add bold label
    label_run = para.add_run(f"{label}: ")
    label_run.bold = True
    label_run.font.size = FONT_SIZE_BODY
    label_run.font.color.rgb = RGBColor(0, 0, 0)

    # Add regular content
    content_run = para.add_run(content)
    content_run.font.size = FONT_SIZE_BODY
    content_run.font.color.rgb = RGBColor(0, 0, 0)


def add_subsection_paragraph(doc: Document, label: str, heading: str, content: str,
                             space_before : float | None =None,
                             space_after : float | None =None) -> None:
    """
    Add a paragraph with subsection label for narrative content.

    Args:
        doc: Document object
        label: Subsection label (e.g., "(a)", "(b)")
        heading: Section heading (e.g., "Ownership", "Street lines")
        content: Paragraph content text
        space_before: Override default spacing before (optional)
        space_after: Override default spacing after (optional)

    Example output:
        "(a). Ownership:
        Mr. D Indika Harshana Perera claims ownership to the property..."
    """
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = space_before if space_before is not None else INLINE_FIELD_SPACE_BEFORE
    para.paragraph_format.space_after = space_after if space_after is not None else INLINE_FIELD_SPACE_AFTER
    para.paragraph_format.line_spacing = 0.9

    # Add bold label and heading
    label_run = para.add_run(f"{label}. {heading}:\n")
    label_run.bold = True
    label_run.font.size = FONT_SIZE_BODY
    label_run.font.color.rgb = RGBColor(0, 0, 0)

    # Add paragraph content
    content_run = para.add_run(content)
    content_run.font.size = FONT_SIZE_BODY
    content_run.font.color.rgb = RGBColor(0, 0, 0)

