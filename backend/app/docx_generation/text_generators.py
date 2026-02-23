"""
Paragraph text generator functions for Sri Lankan valuation report DOCX generation.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from io import BytesIO
from datetime import datetime
from typing import Any, List, Dict, Optional, Union
import requests
from .. import models
from ..utils import append_label_if_missing, clean_spelling_errors, format_no_field
from ..letterhead_templates import get_template
import logging

from .formatting import (
    format_currency, format_currency_words, format_currency_aligned,
    format_room_count, round_for_say, format_material_list,
)
from .paragraph_builders import (
    add_section_heading, add_market_value_line, add_value_rounded_line,
    add_inline_field, add_subsection_paragraph, format_building_valuation_2line, format_addon_compact,
)
from .styling import (
    add_border_to_paragraph,
    MAP_IMAGE_WIDTH, MAP_IMAGE_MAX_HEIGHT, PROPERTY_PHOTO_WIDTH, PROPERTY_PHOTO_HEIGHT,
    IMAGE_SPACING_BEFORE, IMAGE_SPACING_AFTER, MAJOR_SECTION_SPACE_BEFORE, MAJOR_SECTION_SPACE_AFTER,
    SUBSECTION_SPACE_BEFORE, SUBSECTION_SPACE_AFTER, BODY_PARA_SPACE_BEFORE, BODY_PARA_SPACE_AFTER,
    INLINE_FIELD_SPACE_BEFORE, INLINE_FIELD_SPACE_AFTER, SUBHEADING_SPACE_BEFORE, SUBHEADING_SPACE_AFTER,
    INDENTED_CONTENT_SPACE_BEFORE, INDENTED_CONTENT_SPACE_AFTER, INDENTED_CONTENT_LEFT_INDENT,
    BOUNDARY_LIST_SPACE_AFTER, ACCOMMODATION_ROOM_SPACE_AFTER,
    FONT_SIZE_DOCUMENT_TITLE, FONT_SIZE_SECTION_HEADING, FONT_SIZE_SUBSECTION_HEADING,
    FONT_SIZE_BODY, FONT_SIZE_INLINE_LABEL, FONT_SIZE_VALUATION,
    FONT_SIZE_TABLE_HEADER, FONT_SIZE_TABLE_CELL, FONT_SIZE_INVOICE_TOTAL,
    FONT_SIZE_CAPTION, FONT_SIZE_BANK_HEADER, FONT_SIZE_BANK_DETAILS,
    FONT_SIZE_SIGNATURE, FONT_SIZE_CERTIFICATION,
)
from .images import calculate_image_dimensions, apply_letterbox_to_image
from .helpers import (
    safe_get_json_field, safe_get_array_item, to_float, safe_parse_json_string, safe_get_nested,
)

logger = logging.getLogger(__name__)
Deed_Type = Union[Dict[str, Any], Any]

# ===== LEGAL ASPECTS PARAGRAPH GENERATORS =====

def generate_ownership_paragraph(report: Any) -> str:
    """
    Generate professional ownership paragraph with graceful handling of missing data.

    Template adapts based on available data:
    - Full deed info: Complete ownership statement with deed details
    - Survey plan only: Ownership based on plan identification
    - Minimal data: Fallback to valuation basis statement
    """
    parts: list[str] = []

    # Determine property identification method - USE SAFE HELPERS
    deeds = safe_get_json_field(report, 'deeds', [])
    has_deed = bool(deeds)
    has_plan = bool(safe_get_json_field(report, 'plan_number', None))

    # Owner name (use applicant if available) - USE SAFE HELPER
    owner_name = safe_get_json_field(report, 'applicant_full_name', None)

    # Fallback to placeholder if no owner name
    if not owner_name:
        owner_name = "[Applicant Name Not Provided]"

    # Title/pedigree search statement
    if getattr(report, 'title_search_conducted', None) == 'No' or \
       getattr(report, 'pedigree_search_conducted', None) == 'No':
        parts.append("I did not search regarding the history and pedigree of the property.")

    # Main ownership statement
    if has_deed and owner_name:
        # Deed-based ownership - USE SAFE ARRAY ACCESS
        deed : Deed_Type = safe_get_array_item(deeds, 0, {})
        # Handle both dict and object formats
        deed_type: str = deed.get('deed_type', 'transfer deed') if isinstance(deed, dict) else getattr(deed, 'deed_type', 'transfer deed')
        deed_number = deed.get('deed_number', '') if isinstance(deed, dict) else getattr(deed, 'deed_number', '')
        deed_date = deed.get('deed_date', '') if isinstance(deed, dict) else getattr(deed, 'deed_date', '')
        notary_name = deed.get('notary_name', '') if isinstance(deed, dict) else getattr(deed, 'notary_name', '')
        notary_location = deed.get('notary_location', '') if isinstance(deed, dict) else getattr(deed, 'notary_location', '')

        ownership_stmt = f"{owner_name} claims ownership to the property"

        if deed_number and deed_date:
            ownership_stmt += f" by {deed_type} No:{deed_number} dated {deed_date}"
        else:
            ownership_stmt += f" by {deed_type}"

        if notary_name and notary_location:
            ownership_stmt += f" attested by {notary_name} notary public in {notary_location} district"
        elif notary_name:
            ownership_stmt += f" attested by {notary_name} notary public"

        ownership_stmt += "."
        parts.append(ownership_stmt)

    elif has_plan and owner_name:
        # Plan-based ownership
        plan_number = report.plan_number
        plan_date = getattr(report, 'plan_date', '')
        surveyor = getattr(report, 'licensed_surveyor_name', '')

        ownership_stmt = f"{owner_name} claims ownership to the property by transfer deed for the property identified as per the Survey Plan No: {plan_number}"

        if plan_date:
            ownership_stmt += f" dated {plan_date}"
        if surveyor:
            ownership_stmt += f" prepared by {surveyor}"

        ownership_stmt += "."
        parts.append(ownership_stmt)

    # Encumbrance statement
    property_encumbered = getattr(report, 'property_encumbered', None)
    if property_encumbered == 'Yes':
        encumbrance_type = getattr(report, 'encumbrance_type', 'encumbrance')
        encumbrance_details = getattr(report, 'encumbrance_details', '')

        if encumbrance_type == 'Mortgage' and encumbrance_details:
            parts.append(f"This property is already mortgaged to {encumbrance_details}.")
        elif encumbrance_type == 'Life Interest' and encumbrance_details:
            parts.append(f"This property has a life interest held by {encumbrance_details}.")
        else:
            parts.append(f"This property has {encumbrance_type.lower()} encumbrance.")

    # Valuation basis statement
    ownership_type = getattr(report, 'ownership_type', 'freehold')
    valuation_basis = getattr(report, 'valuation_basis_note', None)

    if not valuation_basis:
        if property_encumbered == 'Yes':
            valuation_basis = "free from all legal encumbrance"
        else:
            valuation_basis = "free from all encumbrances"

    valuation_stmt = f"I valued {ownership_type.lower()} interest of the property {valuation_basis}."
    parts.append(valuation_stmt)

    # Fallback for minimal data
    if not parts or len(parts) == 1:
        return f"I did not search regarding the history and pedigree of the property. Valuation is based on {ownership_type.lower()} title {valuation_basis}."

    return " ".join(parts)


def generate_street_lines_paragraph(report) -> str:
    """
    Generate street lines paragraph with contextual information.
    """
    status = getattr(report, 'street_lines_status', 'not affected')
    status_lower = status.lower()

    parts = []

    # Main status statement
    parts.append(f"Street lines are {status_lower} to the property.")

    # Impact description if provided
    impact_desc = getattr(report, 'street_lines_impact_description', None)
    if impact_desc:
        parts.append(impact_desc)

    # Municipal context
    is_municipal = getattr(report, 'is_municipal_limit', False)
    municipal_context = "within Municipal Council Limits" if is_municipal else "outside Municipal Council Limits"

    # Gazette reference if available
    gazette_ref = getattr(report, 'street_lines_gazette_ref', None)
    gazette_date = getattr(report, 'street_lines_gazette_date', None)

    if gazette_ref and gazette_date:
        parts.append(f"Properties located {municipal_context} are subject to street line regulations as per Gazette No: {gazette_ref} dated {gazette_date}.")
    else:
        # Generic explanatory statement
        if status_lower == 'not affected':
            parts.append("Street lines are affected to the properties which are located along roads within Municipal Council Limits and imposed by a gazette.")
        else:
            parts.append(f"Properties located {municipal_context} are subject to street line regulations.")

    return " ".join(parts)


def generate_building_limits_paragraph(report) -> str:
    """
    Generate building limits paragraph with approval and distance information.
    """
    status = getattr(report, 'building_limits_status', 'not affected')
    status_lower = status.lower()

    parts = []

    # Main status statement
    parts.append(f"Building limits are {status_lower} to the property.")

    # Building plan approval
    plan_approved = getattr(report, 'building_plan_approved', None)
    approval_authority = getattr(report, 'building_approval_authority', None)

    if plan_approved == 'Yes':
        if approval_authority:
            parts.append(f"Building Plan approved by {approval_authority}.")
        else:
            parts.append("Building Plan approved by the local authority.")

    # Distance from road
    distance = getattr(report, 'building_distance_from_road', None)
    if distance:
        parts.append(f"Building limit is {distance} from the access road.")

    # Building within limits
    within_limits = getattr(report, 'building_within_limits', None)
    if within_limits == 'Yes':
        parts.append("Existing building is located inside of the building limits in accordance with the approved building plan.")
    elif within_limits == 'No':
        parts.append("Existing building is located outside of the building limits.")

    # Plan reference
    plan_ref = getattr(report, 'building_plan_reference', None)
    if plan_ref and plan_approved == 'Yes':
        parts.append(f"Building plan reference: {plan_ref}.")

    return " ".join(parts)


def generate_local_authority_paragraph(report) -> str:
    """
    Generate local authority paragraph with rating and administrative information.
    """
    # Check if user provided custom free-text (backward compatibility)
    custom_text = getattr(report, 'local_authority_data', None)
    if custom_text and len(custom_text) > 50:  # User wrote a full paragraph
        return custom_text

    parts = []

    # Administrative location
    pradeshiya_sabha = getattr(report, 'pradeshiya_sabha', None)
    district = getattr(report, 'property_district', None)
    province = getattr(report, 'property_province', None)

    if pradeshiya_sabha or district or province:
        location_parts = []
        if pradeshiya_sabha:
            location_parts.append(pradeshiya_sabha)
        if district:
            location_parts.append(f"{district} District")
        if province:
            location_parts.append(f"{province} Province")

        parts.append(", ".join(location_parts) + ".")

    # Rating status
    is_rated = getattr(report, 'local_authority_rated', None)

    if is_rated == 'Yes':
        parts.append("This area has been rated for the local levy of taxes.")

        # Tax levy details
        tax_levy = getattr(report, 'local_authority_tax_levy', None)
        if tax_levy:
            parts.append(tax_levy)

        # Assessment number
        assessment_num = getattr(report, 'assessment_number', None)
        if assessment_num:
            parts.append(f"Assessment No: {assessment_num}.")
    elif is_rated == 'No':
        parts.append("This area has not been rated for the local levy of taxes.")

    # Fallback to custom text if no structured data
    if not parts and custom_text:
        return custom_text

    return " ".join(parts)


def generate_rent_act_paragraph(report) -> str:
    """
    Generate rent act effectiveness paragraph.
    """
    status = getattr(report, 'rent_act_effectiveness', 'Not affected')

    # Direct mapping from dropdown options
    status_map = {
        'Not affected': "This property is not affected by the rent act.",
        'Subject to Rent Act No. 7 of 1972': "This property is subject to the rent act of No. 7 of 1972.",
        'Subject to Rent Act Amendment No. 26 of 2002': "This property is subject to the rent act of No. 7 of 1972 (Amendment No. 26 of 2002).",
        'Partially affected': "This property is partially affected by the rent control regulations."
    }

    return status_map.get(status, "This property is not affected by the rent act.")


# ===== NARRATIVE GENERATION HELPER FUNCTIONS =====

def _synthesize_location_context(locations: List[str]) -> str:
    """
    Synthesize location context from comparable location descriptions.

    Extracts common patterns and generates appropriate location phrase
    for land values narrative.

    Args:
        locations: List of location descriptions from comparable properties

    Returns:
        Synthesized location context string (e.g., "located in this area",
        "located fronting main highway", "located in Rambukkana area")

    Examples:
        >>> _synthesize_location_context([])
        'in this area'
        >>> _synthesize_location_context(['Rambukkana town', 'Rambukkana center'])
        'located in Rambukkana area'
        >>> _synthesize_location_context(['Fronting Kurunegala Highway'])
        'located fronting Kurunegala Highway'
    """
    if not locations:
        return "in this area"

    # Filter out empty locations
    valid_locations = [loc.strip() for loc in locations if loc and loc.strip()]

    if not valid_locations:
        return "in this area"

    # If single location
    if len(valid_locations) == 1:
        return f"located in {valid_locations[0]}"

    # Find common words (case-insensitive)
    from collections import Counter

    # Tokenize all locations
    all_words = []
    for loc in valid_locations:
        # Extract significant words (ignore common words)
        words = loc.lower().replace(',', '').split()
        all_words.extend(words)

    # Find most common significant words
    word_counts = Counter(all_words)

    # Common stop words to ignore
    stop_words = {'in', 'the', 'a', 'an', 'of', 'to', 'and', 'or', 'near', 'at', 'from', 'this', 'that'}
    significant_words = {word: count for word, count in word_counts.items()
                         if count >= len(valid_locations) * 0.5 and word not in stop_words}

    if significant_words:
        # Use most common word as area reference
        common_area = max(significant_words, key=significant_words.get)
        return f"located in {common_area.title()} area"

    # Check for highway/road frontage mentions
    combined = " ".join(valid_locations).lower()
    if 'highway' in combined:
        # Try to extract highway name if possible
        for loc in valid_locations:
            if 'highway' in loc.lower():
                # Try to extract highway name
                words = loc.split()
                for i, word in enumerate(words):
                    if 'highway' in word.lower() and i > 0:
                        return f"located fronting {' '.join(words[max(0,i-2):i+1])}"
                return "located fronting main highway"

    if 'road' in combined:
        # Check if specific road name is mentioned
        for loc in valid_locations:
            if 'road' in loc.lower() and '-' in loc:
                # Might be like "Kurunegala-Puttalam road"
                return f"located fronting {loc}"
        return "located with road access in this area"

    # Generic fallback
    return "located in this village"


def generate_land_values_paragraph(comparables: List[dict]) -> str:
    """
    Generate professional land values paragraph from comparable properties data.

    Groups comparables by property type and generates narrative description
    following Sri Lankan valuation report conventions.

    Args:
        comparables: List of comparable property dictionaries with keys:
            - property_type: 'Commercial' | 'Residential' | 'Agricultural'
            - location_description: str
            - extent: float (perches)
            - rate_per_perch: float (LKR)

    Returns:
        Formatted paragraph(s) describing land values in natural language.
        Returns empty string if no valid comparables.

    Examples:
        >>> comps = [
        ...     {'property_type': 'Residential', 'extent': 15, 'rate_per_perch': 50000,
        ...      'location_description': 'Rambukkana town'},
        ...     {'property_type': 'Residential', 'extent': 20, 'rate_per_perch': 60000,
        ...      'location_description': 'Rambukkana center'}
        ... ]
        >>> generate_land_values_paragraph(comps)
        'Residential building blocks in extent about 15 – 20 Perches located in Rambukkana area
         are being sold at rates ranging from Rs.50,000/= to Rs.60,000/= Per Perch...'
    """
    # Edge case: Empty or None
    if not comparables:
        return ""

    # Filter out invalid comparables (zero values)
    valid_comparables = [
        c for c in comparables
        if c.get('extent', 0) > 0 and c.get('rate_per_perch', 0) > 0
    ]

    if not valid_comparables:
        return ""  # Return empty string - let market analysis stand alone

    # Handle single comparable specially
    if len(valid_comparables) == 1:
        comp = valid_comparables[0]
        prop_type = comp.get('property_type', 'Residential')

        # Property type descriptors
        type_descriptors = {
            'Residential': 'Residential building blocks',
            'Commercial': 'Commercial building blocks',
            'Agricultural': 'Agricultural land parcels'
        }
        property_descriptor = type_descriptors.get(prop_type, 'Building blocks')

        location = comp.get('location_description', 'this area')
        extent = comp.get('extent', 0)
        rate = comp.get('rate_per_perch', 0)

        return (
            f"{property_descriptor} of about {extent:.0f} Perches "
            f"in {location} is valued at approximately Rs.{rate:,.0f}/= Per Perch."
        )

    # Group by property type
    property_groups = {}
    for comp in valid_comparables:
        prop_type = comp.get('property_type', 'Residential')
        if prop_type not in property_groups:
            property_groups[prop_type] = []
        property_groups[prop_type].append(comp)

    # Property type ordering
    type_order = ['Residential', 'Commercial', 'Agricultural']

    paragraphs = []

    for prop_type in type_order:
        if prop_type not in property_groups:
            continue

        group = property_groups[prop_type]

        # Calculate ranges
        extents = [c['extent'] for c in group]
        rates = [c['rate_per_perch'] for c in group]

        min_extent = min(extents)
        max_extent = max(extents)
        min_rate = min(rates)
        max_rate = max(rates)

        # Format extent range
        if min_extent == max_extent:
            extent_str = f"about {min_extent:.0f} Perches"
        else:
            extent_str = f"about {min_extent:.0f} – {max_extent:.0f} Perches"

        # Format rate range with proper currency formatting
        def format_rate(rate):
            # Format with thousands separator and /= suffix
            return f"Rs.{rate:,.0f}/="

        if min_rate == max_rate:
            rate_str = f"{format_rate(min_rate)} Per Perch"
        else:
            rate_str = f"{format_rate(min_rate)} to {format_rate(max_rate)} Per Perch"

        # Extract location context
        locations = [c.get('location_description', '') for c in group]
        location_context = _synthesize_location_context(locations)

        # Determine property type descriptor
        type_descriptors = {
            'Residential': 'Residential building blocks',
            'Commercial': 'Commercial building blocks',
            'Agricultural': 'Agricultural land parcels'
        }
        property_descriptor = type_descriptors.get(prop_type, 'Building blocks')

        # Calculate rate variance for contextual phrase
        rate_variance = ((max_rate - min_rate) / min_rate * 100) if min_rate > 0 else 0

        if rate_variance < 20:
            context_phrase = "showing stable market conditions"
        elif rate_variance > 50:
            context_phrase = "depending upon the location, topography and nature of the property and convenience etc."
        else:
            context_phrase = "to be influenced by factors affecting the property market"

        # Construct paragraph
        paragraph = (
            f"{property_descriptor} in extent {extent_str} {location_context} "
            f"are being sold at rates ranging from {rate_str}, {context_phrase}."
        )

        paragraphs.append(paragraph)

    return " ".join(paragraphs)


def generate_simplified_certification_text(
    valuer_name: str,
    valuer_designation: str,
    lot_number: Optional[str],
    plan_number: Optional[str],
    plan_date: Optional[str],
    licensed_surveyor_name: Optional[str],
    deeds: Optional[List[dict]],
    property_identification_type: Optional[str]
) -> str:
    """
    Generate simplified single-paragraph certification text.

    New simplified format (replaces old multi-paragraph format):
    "I, [Valuer Name], [Designation], do hereby certify that the property
    inspected by me and valued above is identical to the property depicted as
    [Lot X in] Plan No [number] dated [date] made by [Surveyor], Licensed Surveyor."

    Key behaviors:
    - If lot_number provided and plan exists: "...property depicted as Lot {lot_number} in Plan No..."
    - If no lot_number: "...property depicted as Plan No..."
    - Always use "property depicted as" (not "land depicted as")
    - Only mention plan (not deed) in certification
    - If only deed (no plan): "...property described in {deed_type} No..."
    - If lot_number exists but no plan_number: skip lot, use deed/generic

    Args:
        valuer_name: Name of the valuer
        valuer_designation: Professional designation (e.g., "Chartered Valuer")
        lot_number: Lot number (e.g., "Lot 15", "Lots 1 & 2")
        plan_number: Survey plan number
        plan_date: Survey plan date
        licensed_surveyor_name: Name of surveyor
        deeds: List of deed dictionaries
        property_identification_type: Type of identification

    Returns:
        Formatted single-paragraph certification text
    """
    # Start with valuer identification
    cert_text = f"I, {valuer_name}, {valuer_designation}, do hereby certify that the property inspected by me and valued above is identical to "

    # Determine property identification string
    property_id = None

    # Infer identification type if not provided (backward compatibility)
    if not property_identification_type:
        if plan_number and deeds and len(deeds) > 0:
            property_identification_type = "plan_and_deed"
        elif plan_number:
            property_identification_type = "plan"
        elif deeds and len(deeds) > 0:
            property_identification_type = "deed"

    # Build property identification based on type
    if property_identification_type in ["plan", "plan_and_deed"]:
        # Plan-based identification (ignore deed as per requirements)
        if plan_number:
            # Check if we should include lot number
            # Only include lot if both lot_number and plan_number exist
            if lot_number and lot_number.strip():
                # Format with lot number
                plan_ref = format_no_field("Plan", plan_number)
                property_id = f"the property depicted as Lot {lot_number} in {plan_ref}"
            else:
                # Format without lot number
                plan_ref = format_no_field("Plan", plan_number)
                property_id = f"the property depicted as {plan_ref}"

            # Add plan date
            if plan_date:
                property_id += f" dated {plan_date}"

            # Add surveyor
            if licensed_surveyor_name:
                property_id += f" made by {licensed_surveyor_name}, Licensed Surveyor"

    # If no plan, fall back to deed
    if not property_id and property_identification_type in ["deed", "plan_and_deed", "certificate_of_sale"]:
        # Deed-based identification (fallback)
        has_deed = isinstance(deeds, list) and len(deeds) > 0
        if has_deed:
            deed = deeds[0]
            deed_number = deed.get('deed_number', '') if isinstance(deed, dict) else getattr(deed, 'deed_number', '')
            deed_type = deed.get('deed_type', 'Deed') if isinstance(deed, dict) else getattr(deed, 'deed_type', 'Deed')
            deed_date = deed.get('deed_date', '') if isinstance(deed, dict) else getattr(deed, 'deed_date', '')

            if deed_number:
                property_id = f"the property described in {deed_type} No. {deed_number}"
                if deed_date:
                    property_id += f" dated {deed_date}"

    # Final fallback if no identification data
    if not property_id:
        property_id = "the property described in this report"

    # Complete the certification text
    cert_text += property_id + "."

    return cert_text


# DEPRECATED: Keep for backward compatibility with old code
def generate_certificate_of_identity_text(
    property_identification_type: Optional[str],
    plan_number: Optional[str],
    plan_date: Optional[str],
    deeds: Optional[List[dict]],
    licensed_surveyor_name: Optional[str] = None
) -> Optional[str]:
    """
    DEPRECATED: Use generate_simplified_certification_text instead.

    This function is kept for backward compatibility only.
    Generates old-style Certificate of Identity text.
    """
    if not property_identification_type:
        # Backward compatibility: infer type from available data
        if plan_number and deeds and len(deeds) > 0:
            property_identification_type = "plan_and_deed"
        elif plan_number:
            property_identification_type = "plan"
        elif deeds and len(deeds) > 0:
            property_identification_type = "deed"
        else:
            return None

    # Base text
    base_text = "I certify that the property inspected by me is identical to the property described in "

    # Extract deed information if available
    has_deed = isinstance(deeds, list) and len(deeds) > 0
    deed_text = ""
    if has_deed:
        deed = deeds[0]
        deed_number = deed.get('deed_number', '') if isinstance(deed, dict) else getattr(deed, 'deed_number', '')
        deed_type = deed.get('deed_type', 'Deed') if isinstance(deed, dict) else getattr(deed, 'deed_type', 'Deed')
        deed_date = deed.get('deed_date', '') if isinstance(deed, dict) else getattr(deed, 'deed_date', '')

        if deed_number:
            deed_text = f"{deed_type} No. {deed_number}"
            if deed_date:
                deed_text += f" dated {deed_date}"

    # Generate text based on identification type
    if property_identification_type == "plan":
        if not plan_number:
            return None
        plan_ref = format_no_field("Plan", plan_number)
        text = base_text + plan_ref
        if plan_date:
            text += f" dated {plan_date}"
        if licensed_surveyor_name:
            text += f" made by {licensed_surveyor_name}, Licensed Surveyor"
        text += "."
        return text

    elif property_identification_type == "deed":
        if not deed_text:
            return None
        return base_text + deed_text + "."

    elif property_identification_type == "plan_and_deed":
        if not deed_text and not plan_number:
            return None

        # If only one is available, use whichever exists
        if not deed_text and plan_number:
            plan_ref = format_no_field("Plan", plan_number)
            text = base_text + plan_ref
            if plan_date:
                text += f" dated {plan_date}"
            if licensed_surveyor_name:
                text += f" made by {licensed_surveyor_name}, Licensed Surveyor"
            text += "."
            return text

        if deed_text and not plan_number:
            return base_text + deed_text + "."

        # Both available - combine them
        plan_ref = format_no_field("Plan", plan_number)
        text = base_text + deed_text + " and identified in " + plan_ref
        if plan_date:
            text += f" dated {plan_date}"
        if licensed_surveyor_name:
            text += f" made by {licensed_surveyor_name}, Licensed Surveyor"
        text += "."
        return text

    elif property_identification_type == "certificate_of_sale":
        # Certificate of sale is similar to deed
        if not deed_text:
            return None
        return base_text + deed_text + "."

    # Unknown type or no data
    return None


def add_signature_block(
    doc,
    user,
    valuer_name: Optional[str],
    valuer_designation: Optional[str],
    certification_date: Optional[str]
):
    """
    Add a standardized signature block to the document.

    Format:
    - Spacing before signature
    - Underline ("_" * 40) for signature
    - Optional signature image (if user has one)
    - Valuer name (bold, 9pt)
    - Valuer designation (9pt)
    - Certification date (9pt)

    Args:
        doc: Document object
        user: User model instance
        valuer_name: Name of valuer
        valuer_designation: Professional designation
        certification_date: Date of certification
    """
    # Add spacing before signature
    doc.add_paragraph("\n")

    # Signature line (underline)
    sig_line = doc.add_paragraph("_" * 40)
    sig_line.paragraph_format.space_before = Pt(24)
    sig_line.paragraph_format.space_after = Pt(6)

    # Add signature image if available
    if hasattr(user, 'signature_image') and user.signature_image:
        try:
            response = requests.get(user.signature_image, timeout=10)
            if response.status_code == 200:
                sig_para = doc.add_paragraph()
                sig_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                sig_para.paragraph_format.space_before = Pt(0)
                sig_para.paragraph_format.space_after = Pt(6)
                image_stream = BytesIO(response.content)
                sig_para.add_run().add_picture(image_stream, width=Inches(2))
        except Exception as e:
            logger.warning(f"Failed to add signature image: {e}")

    # Valuer name
    if valuer_name:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(valuer_name)
        run.font.bold = True
        run.font.size = FONT_SIZE_BODY

    # Valuer designation
    if valuer_designation:
        p = doc.add_paragraph(valuer_designation)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = FONT_SIZE_BODY

    # Certification date
    if certification_date:
        p = doc.add_paragraph(certification_date)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = FONT_SIZE_BODY

    doc.add_paragraph()


def get_pronoun(title: Optional[str]) -> Dict[str, str]:
    """
    Determine pronouns based on title.
    Returns: {'subject': 'he'/'she', 'object': 'him'/'her', 'possessive': 'his'/'her'}
    """
    if not title:
        return {'subject': 'he', 'object': 'him', 'possessive': 'his'}

    title_lower = title.lower().strip()

    # Female titles
    if title_lower in ['mrs.', 'mrs', 'miss.', 'miss', 'ms.', 'ms']:
        return {'subject': 'she', 'object': 'her', 'possessive': 'her'}

    # Default to male
    return {'subject': 'he', 'object': 'him', 'possessive': 'his'}

def generate_title_block(report: models.Report) -> str:
    """Generate the title block with support for plan/deed/certificate identification types"""
    lines = []
    lines.append("VALUATION REPORT")
    lines.append("of")

    # Determine identification type (with backward compatibility for old reports)
    id_type = report.property_identification_type
    if not id_type:
        # Infer type from existing data for old reports
        if report.plan_number:
            id_type = "plan"
        elif report.deeds:
            # Old reports with deed data
            id_type = "deed"
        else:
            # Fallback to plan-style if nothing is detected
            id_type = "plan"

    # Generate title based on identification type
    if id_type == "plan" or id_type == "plan_and_deed":
        # Plan-based title (original format)
        # For plan_and_deed: show plan info in header, deed info appears in certification section
        lot_desc = report.lot_number or '[Lot Number]'

        # Remove common prefixes that shouldn't be in lot description
        lot_desc_stripped = lot_desc.strip()
        prefixes_to_remove = ['plan no', 'plan no:', 'lot plan no', 'lot plan no:']
        lot_desc_lower = lot_desc_stripped.lower()
        for prefix in prefixes_to_remove:
            if lot_desc_lower.startswith(prefix):
                lot_desc = lot_desc_stripped[len(prefix):].strip()
                break

        # Ensure it has "Lot" prefix
        if not lot_desc.lower().startswith('lot'):
            lot_desc = f"Lot {lot_desc}"

        plan_num = report.plan_number or '[Plan Number]'
        plan_formatted = format_no_field("Plan", plan_num)

        # Combine date with first line
        plan_date = report.plan_date or '[Date]'
        prop_desc = f"The Property Depicted as {lot_desc} in {plan_formatted} Dated {plan_date}"
        lines.append(prop_desc)

        # Second line starts with "made by"
        surveyor = report.licensed_surveyor_name or '[Surveyor Name]'
        plan_info = f"made by {surveyor} Licensed Surveyor."
        lines.append(plan_info)

    elif id_type == "deed":
        # Deed-based title: show deed information
        deed_type = "Deed"
        deed_number = "[Deed Number]"
        deed_date = "[Date]"

        # Extract deed info from report.deeds JSON array (first deed)
        if report.deeds and isinstance(report.deeds, list) and len(report.deeds) > 0:
            first_deed = report.deeds[0]
            deed_type = first_deed.get('deed_type', 'Deed')
            deed_number = first_deed.get('deed_number', '[Deed Number]')
            deed_date = first_deed.get('deed_date', '[Date]')
        else:
            logger.warning(f"[DOCX] Report {report.id} has deed identification type but no deed data")

        # Format deed number with "No."
        deed_number_formatted = format_no_field("", deed_number, include_label=False)
        prop_desc = f"The Property Depicted as described in {deed_type} No. {deed_number_formatted} dated {deed_date}"
        lines.append(prop_desc)

    elif id_type == "certificate_of_sale":
        # Certificate of sale title: show certificate information (stored as deed internally)
        deed_number = "[Certificate Number]"
        deed_date = "[Date]"

        # Extract certificate info from report.deeds JSON array (certificate stored as deed internally)
        if report.deeds and isinstance(report.deeds, list) and len(report.deeds) > 0:
            first_deed = report.deeds[0]
            deed_number = first_deed.get('deed_number', '[Certificate Number]')
            deed_date = first_deed.get('deed_date', '[Date]')
        else:
            logger.warning(f"[DOCX] Report {report.id} has certificate_of_sale identification type but no deed data")

        # Format certificate number with "No."
        cert_number_formatted = format_no_field("", deed_number, include_label=False)
        prop_desc = f"The Property Depicted as described in Certificate of Sale No. {cert_number_formatted} dated {deed_date}"
        lines.append(prop_desc)

    else:
        # Fallback: Address-based title for old reports without identification type
        address = generate_smart_address(report) or '[Property Address]'
        prop_desc = f"The Property Depicted as {address}"
        lines.append(prop_desc)

    return lines

def generate_applicant_statement(report: models.Report) -> str:
    """Generate the applicant statement paragraph with smart grammar"""

    # Build applicant full description
    applicant_desc = f"{report.applicant_title or ''} {report.applicant_full_name or '[Applicant Name]'}".strip()

    # ID information - handle optional fields gracefully
    if report.applicant_id_number and report.applicant_id_type:
        id_no_formatted = format_no_field("", report.applicant_id_number, include_label=False)
        id_info = f"holder {report.applicant_id_type} {id_no_formatted}"
    else:
        # No ID provided - skip this part of the introduction
        id_info = ""

    # Address parts
    address_parts = []
    if report.applicant_address_line1:
        address_parts.append(report.applicant_address_line1)
    if report.applicant_address_line2:
        address_parts.append(report.applicant_address_line2)
    if report.applicant_district:
        cleaned = clean_spelling_errors(report.applicant_district)
        district_text = append_label_if_missing(cleaned, "District", variants=["district"])
        address_parts.append(district_text)
    if report.applicant_province:
        address_parts.append(report.applicant_province)
    if report.applicant_country:
        address_parts.append(report.applicant_country)

    address_str = ", ".join(address_parts) if address_parts else "[Address]"

    # Get pronouns
    pronouns = get_pronoun(report.applicant_title)

    # Build ownership text (auto-generated based on applicant title)
    ownership_text = f"owned by {pronouns['object']}"

    # Property type
    property_type = report.property_type_valued or "immovable property"

    # Valuation type
    valuation_type = report.valuation_type or "Market Value"

    # Build paragraph - conditionally include ID info
    if id_info:
        paragraph1 = f"This Valuation Report is furnished at the request of {applicant_desc} {id_info} of {address_str}."
    else:
        paragraph1 = f"This Valuation Report is furnished at the request of {applicant_desc} of {address_str}."

    # Second part - wishes to know
    wish_text = f"{applicant_desc} wishes to know the {valuation_type} of {property_type} {ownership_text} in the Democratic Socialist Republic of Sri Lanka."

    # Handle additional owners if any
    if report.has_additional_owner == "yes" and report.additional_owner_names:
        wish_text = wish_text.replace(f"{ownership_text}", f"{ownership_text} & {pronouns['possessive']} family {report.additional_owner_names}")

    return [paragraph1, wish_text]

def generate_organization_side_introduction(report: models.Report) -> List[str]:
    """
    Generate organization-side introduction format.

    Format:
    - Paragraph 1: "At the request of [position] [organization] [address],
                    I am furnishing a Valuation Report of the above property
                    for the [purpose] purpose."
    - Paragraph 2-4: Formatted applicant details
    - Paragraph 5: Inspection date (handled separately in main function)

    Returns:
        List of paragraph strings
    """
    paragraphs = []

    # Paragraph 1: Request statement
    request_parts = []

    if report.submission_recipient_position:
        request_parts.append(report.submission_recipient_position)

    if report.submission_organization:
        request_parts.append(report.submission_organization)

    if report.submission_address:
        request_parts.append(report.submission_address)

    requester_text = ", ".join(request_parts) if request_parts else "[Requesting Organization]"

    # Get purpose text
    purpose = report.valuation_purpose or "[purpose]"

    para1 = f"At the request of {requester_text}, I am furnishing a Valuation Report of the above property for the {purpose} purpose."
    paragraphs.append(para1)

    # Paragraph 2: Applicant info
    applicant_desc = f"{report.applicant_title or ''} {report.applicant_full_name or '[Applicant Name]'}".strip()

    # ID information
    id_text = ""
    if report.applicant_id_number and report.applicant_id_type:
        id_no_formatted = format_no_field("", report.applicant_id_number, include_label=False)
        id_text = f" holder {report.applicant_id_type} {id_no_formatted}"

    para2 = f"Applicant        :-{applicant_desc}{id_text}"
    paragraphs.append(para2)

    # Paragraph 3: Address
    address_parts = []
    if report.applicant_address_line1:
        address_parts.append(report.applicant_address_line1)
    if report.applicant_address_line2:
        address_parts.append(report.applicant_address_line2)
    if report.applicant_district:
        cleaned = clean_spelling_errors(report.applicant_district)
        district_text = append_label_if_missing(cleaned, "District", variants=["district"])
        address_parts.append(district_text)
    if report.applicant_province:
        address_parts.append(report.applicant_province)
    if report.applicant_country:
        address_parts.append(report.applicant_country)

    address_str = ", ".join(address_parts) if address_parts else "[Address]"
    para3 = f"Address          :-{address_str}"
    paragraphs.append(para3)

    # Paragraph 4: Contact Number
    contact = report.applicant_contact_number or "[Contact Number]"
    para4 = f"Contact No       :-{contact}"
    paragraphs.append(para4)

    return paragraphs

def generate_multi_property_concluding_statement(
    report: models.Report,
    user: models.User,
    grand_total: float
) -> List[str]:
    """
    Generate concluding statement for multi-property CLIENT REQUEST reports.

    This function creates the closing statement that appears after the property
    summary table and inspection date, including the total market value in words
    and the valuer's signature.

    Args:
        report: Report model instance
        user: User model instance (for valuer name and credentials)
        grand_total: Total valuation amount from property table

    Returns:
        List of paragraph text strings:
        - [0]: Market value statement with amount in words
        - [1]: Valuer name line (Vlr.[Honorific] [Full Name])
        - [2]: Professional designation line

        Returns empty list if report is organization request.

    Example output:
        [
            "Present Market Value of the Properties claimed by Mr. John Doe & his family...",
            "Vlr.K D A Nimalsiri",
            "Chartered Valuer"
        ]
    """
    # Only generate for client requests
    if report.request_type != 'client_request':
        return []

    # Get pronoun based on applicant title
    pronouns = get_pronoun(report.applicant_title)
    gender_pronoun = pronouns['possessive']  # "his" or "her"

    # Build applicant name
    title = report.applicant_title or "Mr."
    full_name = report.applicant_full_name or "[Applicant Name]"

    # Convert amount to words
    amount_words = format_currency_words(grand_total)

    # Construct statement
    statement = (
        f"Present Market Value of the Properties claimed by {title} {full_name} "
        f"& {gender_pronoun} family in the Democratic Socialist Republic of Sri Lanka "
        f"are in a sum of Lanka Rupees {amount_words} only."
    )

    # Build valuer signature lines
    honorific = user.honorific or ""
    valuer_name = user.full_name or "[Valuer Name]"
    designation = user.professional_designation or "[Professional Designation]"

    # Format valuer name line: "Vlr.Honorific FullName" (remove extra spaces)
    valuer_line = f"Vlr.{honorific} {valuer_name}".replace("  ", " ").strip()

    return [statement, valuer_line, designation]

def generate_deed_description(deeds: Optional[List[Dict]]) -> Optional[str]:
    """Generate deed description text if deeds are provided"""
    if not deeds or len(deeds) == 0:
        return None

    deed_parts = []
    for deed in deeds:
        deed_type = deed.get('deed_type', 'Deed')
        deed_num = deed.get('deed_number', '[Number]')
        deed_date = deed.get('deed_date', '[Date]')

        deed_no_formatted = format_no_field("", deed_num, include_label=False)
        deed_text = f"{deed_type} {deed_no_formatted} dated {deed_date}"

        # Add notary info if available
        if deed.get('notary_name') or deed.get('notary_location'):
            notary_parts = []
            if deed.get('notary_name'):
                notary_parts.append(f"attested by {deed['notary_name']} Notary Public")
            if deed.get('notary_location'):
                cleaned = clean_spelling_errors(deed['notary_location'])
                location_text = append_label_if_missing(cleaned, "District", variants=["district"])
                notary_parts.append(f"in {location_text}")
            if notary_parts:
                deed_text += " " + " ".join(notary_parts)

        deed_parts.append(deed_text)

    if len(deed_parts) == 1:
        return f"requesting a Valuation Report of the immovable properties described in schedules of {deed_parts[0]}."
    else:
        # Multiple deeds - join with &
        deeds_text = " & schedule of ".join(deed_parts)
        return f"requesting a Valuation Report of the immovable properties described in schedules of {deeds_text}."

def generate_submission_statement(report: models.Report) -> Optional[str]:
    """Generate submission destination statement with optional purpose and recipient position"""
    if not report.submission_organization:
        return None

    org_text = ""

    # Add recipient position if available
    if report.submission_recipient_position:
        org_text = f"{report.submission_recipient_position}, "

    org_text += report.submission_organization

    # Add address if available
    if report.submission_address:
        org_text += f", {report.submission_address}"

    statement = f"This Valuation Report is to be submitted to {org_text}"

    # Add purpose if available
    if report.valuation_purpose:
        statement += f" for the purpose of {report.valuation_purpose}"

    statement += "."
    return statement

def generate_situation_text(report: models.Report) -> Optional[str]:
    """
    Generate SITUATION section text from property location data
    Following professional Sri Lankan valuation report format
    """
    print(f"[SITUATION] Report ID: {report.id}")
    print(f"[SITUATION] Village: {report.property_village}")
    print(f"[SITUATION] District: {report.property_district}")
    print(f"[SITUATION] GN Division: {report.grama_niladari_division}")
    print(f"[SITUATION] DS Division: {report.property_divisional_secretariat}")
    print(f"[SITUATION] Korale: {report.korale}")
    print(f"[SITUATION] Pradeshiya Sabha: {report.pradeshiya_sabha}")

    # Need at least some basic location info
    if not (report.property_village or report.property_district):
        print("[SITUATION] No location data found, returning None")
        return None

    # Build the comprehensive situation text parts
    parts = []

    # Start with "The property is situated"
    # Handle Assessment Number if municipal/urban property
    prefix = "The property"
    if report.assessment_number:
        assessment_formatted = format_no_field("Assessment", report.assessment_number, include_label=False)
        prefix += f" to be valued is situated bearing {assessment_formatted}"
        if report.land_traditional_name:
            prefix += f" {report.land_traditional_name}"
    elif report.land_traditional_name:
        prefix += f" is situated at {report.land_traditional_name}"
    else:
        prefix += " is situated"

    # Add village information
    if report.property_village:
        if " is situated" in prefix and prefix.endswith("situated"):
            parts.append(f"in {report.property_village} village")
        else:
            parts.append(f" in {report.property_village} village")

    # Add property number (ward/lot number)
    if report.property_number:
        property_no_formatted = format_no_field("", report.property_number, include_label=False)
        parts.append(f"in {property_no_formatted}")

    # Add location direction if available (e.g., "north-east")
    direction_suffix = ""
    if report.location_direction:
        direction_suffix = f"-{report.location_direction}"

    # Add Grama Niladari Division (with intelligent label handling)
    if report.grama_niladari_division:
        cleaned = clean_spelling_errors(report.grama_niladari_division)
        gn_text = append_label_if_missing(
            cleaned,
            "Grama Niladari division",
            variants=["grama niladari division", "Grama Niladari Division", "GN division", "GN Division"]
        )
        # Inject direction suffix before "Grama Niladari division" if it exists
        if direction_suffix and "Grama Niladari division" in gn_text:
            gn_text = gn_text.replace(" Grama Niladari division", f"{direction_suffix} Grama Niladari division")
        elif direction_suffix:
            # If variant was used, just prepend the direction
            gn_text = f"{gn_text}{direction_suffix}"
        parts.append(gn_text)

    # Add Divisional Secretariat (with intelligent label handling)
    if report.property_divisional_secretariat:
        cleaned = clean_spelling_errors(report.property_divisional_secretariat)
        ds_text = append_label_if_missing(
            cleaned,
            "Divisional Secretariat",
            variants=[
                "divisional secretariat division",
                "Divisional Secretariat Division",
                "divisional secretariat",
                "DS division",
                "DS Division"
            ]
        )
        # Add "division of" wrapper, but check if "division" is already in the text
        if "division" in ds_text.lower():
            parts.append(f"within the {ds_text} of")
        else:
            parts.append(f"within the {ds_text} division of")

    # Add Hathpaththuwa (with intelligent label handling)
    if report.hathpaththuwa:
        cleaned = clean_spelling_errors(report.hathpaththuwa)
        hathpaththuwa_text = append_label_if_missing(cleaned, "Hathpaththuwa", variants=["hathpaththuwa", "hatpattu"])
        parts.append(f"in {hathpaththuwa_text}")

    # Add Korale (with intelligent label handling)
    if report.korale:
        cleaned = clean_spelling_errors(report.korale)
        korale_text = append_label_if_missing(cleaned, "Korale", variants=["korale"])
        parts.append(f"in {korale_text}")

    # Add Pradeshiya Sabha or Municipal Council (with intelligent label handling)
    if report.is_municipal_limit and report.property_district:
        district_cleaned = clean_spelling_errors(report.property_district)
        district_text = append_label_if_missing(district_cleaned, "District", variants=["district"])
        parts.append(f"within the Municipal Council limit of {district_text}")
    elif report.pradeshiya_sabha:
        cleaned = clean_spelling_errors(report.pradeshiya_sabha)
        ps_text = append_label_if_missing(cleaned, "Pradeshiya Sabha", variants=["pradeshiya sabha", "PS"])
        parts.append(f"of {ps_text}")

    # Add District (only if not in municipal limit, with intelligent label handling)
    if report.property_district and not report.is_municipal_limit:
        cleaned = clean_spelling_errors(report.property_district)
        district_text = append_label_if_missing(cleaned, "District", variants=["district"])
        parts.append(f"in {district_text}")

    # Add Province (with intelligent label handling)
    if report.property_province:
        cleaned = clean_spelling_errors(report.property_province)
        province_text = append_label_if_missing(cleaned, "Province", variants=["province"])
        parts.append(province_text)

    # Add "of Sri Lanka" at the end
    if parts:
        parts.append("of Sri Lanka")

    # Assemble the complete sentence
    if parts:
        situation_text = prefix + " " + " ".join(parts) + "."
    else:
        # Fallback to smart address if no detailed location data
        smart_address = generate_smart_address(report)
        if smart_address:
            situation_text = f"The property is situated at {smart_address}."
        else:
            print("[SITUATION] No situation text could be generated")
            return None

    print(f"[SITUATION] Generated text: {situation_text}")
    return situation_text

def generate_smart_address(report: models.Report) -> Optional[str]:
    """
    Generate formatted property address from address components.

    Format: [Property Number], [Village] Village, [GN Division] Grama Niladari Division,
            [DS Division] Divisional Secretariat Division, [District] District, [Province]

    Only includes components that are present.
    """
    components = []

    # Property number (e.g., "No: 1202")
    if report.property_number:
        components.append(report.property_number)

    # Village name
    if report.property_village:
        components.append(f"{report.property_village} Village")

    # Grama Niladari Division
    if report.grama_niladari_division:
        components.append(f"{report.grama_niladari_division} Grama Niladari Division")

    # Divisional Secretariat (if available)
    if hasattr(report, 'property_divisional_secretariat') and report.property_divisional_secretariat:
        components.append(f"{report.property_divisional_secretariat} Divisional Secretariat Division")

    # Hathpaththuwa (if available)
    if report.hathpaththuwa:
        components.append(report.hathpaththuwa)

    # Korale (if available)
    if report.korale:
        components.append(report.korale)

    # Pradeshiya Sabha (if available)
    if report.pradeshiya_sabha:
        components.append(report.pradeshiya_sabha)

    # District (usually required)
    if report.property_district:
        components.append(f"{report.property_district} District")

    # Province
    if report.property_province:
        components.append(report.property_province)

    if not components:
        return None

    # Join with commas
    return ", ".join(components)

def generate_access_text(report: models.Report) -> Optional[str]:
    """Generate ACCESS section text from access directions data"""
    print(f"[ACCESS] Report ID: {report.id}")
    print(f"[ACCESS] Directions Text: {report.access_directions_text}")
    print(f"[ACCESS] Starting Point: {report.access_starting_point_name}")
    print(f"[ACCESS] Distance: {report.access_distance_km}")
    print(f"[ACCESS] Map Image URL: {report.location_map_image_data}")

    if not report.access_directions_text:
        print("[ACCESS] No access directions text found, returning None")
        return None

    print(f"[ACCESS] Returning access text: {report.access_directions_text}")
    return report.access_directions_text


def generate_locality_description(report: models.Report) -> Optional[str]:
    """Generate LOCALITY section text from locality data"""
    print(f"[LOCALITY] Report ID: {report.id}")
    print(f"[LOCALITY] Description Text: {report.locality_description_text}")
    print(f"[LOCALITY] Major Town: {report.major_town_name} ({report.distance_to_major_town_km} km)")
    print(f"[LOCALITY] Nearby Facilities: {len(report.nearby_facilities) if report.nearby_facilities else 0}")

    if not report.locality_description_text:
        print("[LOCALITY] No locality description text found, returning None")
        return None

    print(f"[LOCALITY] Returning locality text: {report.locality_description_text}")
    return report.locality_description_text

def generate_boundary_summary_text(report: models.Report) -> Optional[str]:
    """
    Generate professional boundary summary sentence from boundary types.
    Example output: "The boundaries referred above are in accordance with those shown in plan
    and are well defined and demarcated by brick masonry parapet walls constructed on rubble
    masonry foundation on the north, east, south and the west. There is an Auto roller gate
    entrance to the property."
    """
    # Map boundary type IDs to readable labels
    boundary_type_labels = {
        'brick_walls': 'brick masonry parapet walls',
        'barbed_wire': 'barbed wire fence',
        'live_fence': 'live fence',
        'concrete_posts': 'concrete posts',
        'iron_gate': 'iron gate',
        'rubble_foundation': 'rubble masonry foundation',
        'chain_link': 'chain link fence',
        'wooden_fence': 'wooden fence',
        'stone_wall': 'stone wall',
        'hedge': 'hedge'
    }

    # Entrance type labels
    entrance_labels = {
        'auto_roller_gate': 'an Auto roller gate',
        'iron_gate': 'an Iron gate',
        'wooden_gate': 'a Wooden gate',
        'sliding_gate': 'a Sliding gate',
        'swing_gate': 'a Swing gate',
        'no_gate': None
    }

    parts = []

    # Opening statement
    parts.append("The boundaries referred above are in accordance with those shown in plan and are well defined")

    # Determine boundary types per direction or use general physical boundaries
    boundary_types_per_dir = report.boundary_types_per_direction or {}
    physical_types = report.physical_boundaries_types or []

    if boundary_types_per_dir:
        # Group directions by boundary type
        type_to_directions = {}
        directions_order = ['north', 'northeast', 'east', 'southeast', 'south', 'southwest', 'west', 'northwest']

        for direction in directions_order:
            b_type = boundary_types_per_dir.get(direction)
            if b_type:
                if b_type not in type_to_directions:
                    type_to_directions[b_type] = []
                type_to_directions[b_type].append(direction)

        # Build description for each boundary type group
        boundary_descriptions = []
        for b_type, directions in type_to_directions.items():
            type_label = boundary_type_labels.get(b_type, b_type)
            dir_labels = [d for d in directions]

            if len(dir_labels) == 8:
                dir_text = "on all eight sides"
            elif len(dir_labels) == 4:
                # Check if it's the main 4 directions
                main_dirs = {'north', 'south', 'east', 'west'}
                if set(dir_labels) == main_dirs:
                    dir_text = "on all four sides"
                else:
                    dir_text = "on the " + ", ".join(dir_labels[:-1]) + f" and {dir_labels[-1]}"
            elif len(dir_labels) == 1:
                dir_text = f"on the {dir_labels[0]}"
            else:
                # Format as "north, northeast, east and south"
                dir_text = "on the " + ", ".join(dir_labels[:-1]) + f" and {dir_labels[-1]}"

            boundary_descriptions.append(f"demarcated by {type_label} {dir_text}")

        if boundary_descriptions:
            parts.append(" and " + ", ".join(boundary_descriptions))

    elif physical_types:
        # Use general physical boundary types
        type_labels = [boundary_type_labels.get(t, t) for t in physical_types[:3]]

        if 'rubble_foundation' in physical_types:
            # Special case: "brick walls constructed on rubble foundation"
            main_types = [t for t in type_labels if 'rubble' not in t]
            if main_types:
                parts.append(f" and demarcated by {main_types[0]} constructed on rubble masonry foundation on the north, east, south and the west")
            else:
                parts.append(f" and demarcated by rubble masonry foundation on all sides")
        elif type_labels:
            parts.append(f" and demarcated by {', '.join(type_labels)}")

    parts.append(".")

    # Add entrance/gate information
    entrance_type = report.entrance_type
    if entrance_type and entrance_type != 'no_gate':
        entrance_label = entrance_labels.get(entrance_type)
        if entrance_label:
            parts.append(f" There is {entrance_label} entrance to the property.")

    result = "".join(parts)
    return result if result != "The boundaries referred above are in accordance with those shown in plan and are well defined." else None


def format_list_with_grammar(items: List[str]) -> str:
    """
    Format a list of items with proper Oxford comma grammar.

    Examples:
        ['item1'] -> 'item1'
        ['item1', 'item2'] -> 'item1 and item2'
        ['item1', 'item2', 'item3'] -> 'item1, item2, and item3'
    """
    if not items:
        return ""
    if len(items) == 1:
        return items[0]
    elif len(items) == 2:
        return f"{items[0]} and {items[1]}"
    else:
        return ", ".join(items[:-1]) + f", and {items[-1]}"


