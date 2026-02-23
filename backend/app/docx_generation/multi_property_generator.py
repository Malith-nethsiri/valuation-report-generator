"""
Multi-property report DOCX generation.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from io import BytesIO
from datetime import datetime
from typing import Any, List, Dict, Optional, Union
import requests
from .. import models
from ..utils import append_label_if_missing, clean_spelling_errors, format_no_field
from ..letterhead_templates import get_template
import logging

from .formatting import (
    format_currency, format_currency_words, format_currency_aligned,
    format_room_count, round_for_say, format_material_list,
)
from .paragraph_builders import (
    add_section_heading, add_market_value_line, add_value_rounded_line,
    add_inline_field, add_subsection_paragraph, format_building_valuation_2line, format_addon_compact,
)
from .styling import (
    add_border_to_paragraph,
    MAP_IMAGE_WIDTH, MAP_IMAGE_MAX_HEIGHT, PROPERTY_PHOTO_WIDTH, PROPERTY_PHOTO_HEIGHT,
    IMAGE_SPACING_BEFORE, IMAGE_SPACING_AFTER, MAJOR_SECTION_SPACE_BEFORE, MAJOR_SECTION_SPACE_AFTER,
    SUBSECTION_SPACE_BEFORE, SUBSECTION_SPACE_AFTER, BODY_PARA_SPACE_BEFORE, BODY_PARA_SPACE_AFTER,
    INLINE_FIELD_SPACE_BEFORE, INLINE_FIELD_SPACE_AFTER, SUBHEADING_SPACE_BEFORE, SUBHEADING_SPACE_AFTER,
    INDENTED_CONTENT_SPACE_BEFORE, INDENTED_CONTENT_SPACE_AFTER, INDENTED_CONTENT_LEFT_INDENT,
    BOUNDARY_LIST_SPACE_AFTER, ACCOMMODATION_ROOM_SPACE_AFTER,
    FONT_SIZE_DOCUMENT_TITLE, FONT_SIZE_SECTION_HEADING, FONT_SIZE_SUBSECTION_HEADING,
    FONT_SIZE_BODY, FONT_SIZE_INLINE_LABEL, FONT_SIZE_VALUATION,
    FONT_SIZE_TABLE_HEADER, FONT_SIZE_TABLE_CELL, FONT_SIZE_INVOICE_TOTAL,
    FONT_SIZE_CAPTION, FONT_SIZE_BANK_HEADER, FONT_SIZE_BANK_DETAILS,
    FONT_SIZE_SIGNATURE, FONT_SIZE_CERTIFICATION,
)
from .images import calculate_image_dimensions, apply_letterbox_to_image
from .helpers import (
    safe_get_json_field, safe_get_array_item, to_float, safe_parse_json_string, safe_get_nested,
)

logger = logging.getLogger(__name__)
Deed_Type = Union[Dict[str, Any], Any]

from .property_section_generator import generate_property_sections
from .invoice_generator import generate_invoice_section
from .text_generators import (
    generate_ownership_paragraph, generate_street_lines_paragraph,
    generate_building_limits_paragraph, generate_local_authority_paragraph,
    generate_rent_act_paragraph, _synthesize_location_context,
    generate_land_values_paragraph, generate_simplified_certification_text,
    generate_certificate_of_identity_text, add_signature_block, get_pronoun,
    generate_title_block, generate_applicant_statement,
    generate_organization_side_introduction, generate_multi_property_concluding_statement,
    generate_deed_description, generate_submission_statement, generate_situation_text,
    generate_smart_address, generate_access_text, generate_locality_description,
    generate_boundary_summary_text, format_list_with_grammar,
)
from .building_renderer import (
    aggregate_accommodation_across_building, deduplicate_water_sources,
    render_construction_details, render_utilities_and_conveniences,
)

def generate_multi_property_report_docx(report: models.Report, user: models.User) -> BytesIO:
    """
    Generate a multi-property valuation report DOCX file.

    Structure:
    1. Summary Page - Property listing with totals (COMPLETED properties only)
    2. Individual Property Sections - Full details for each COMPLETED property
       - Each property includes its own certification section
    3. Invoice Section - Professional fees

    Features:
    - Filters properties by status='completed' (draft properties excluded)
    - Per-property certification (allows different valuers per property)
    - Properties appear in user-defined order (from property_order field)

    Args:
        report: Report model instance with is_multi_property=True
        user: User model instance

    Returns:
        BytesIO object containing the DOCX document

    Raises:
        ValueError: If report has no properties or all properties are drafts
    """
    try:
        logger.info(f"[MULTI-PROPERTY DOCX] Starting generation for report ID {report.id}")
        logger.info(f"[MULTI-PROPERTY DOCX] Property count: {report.property_count}")

        # Get all properties ordered by property_order
        all_properties = report.properties  # Uses the @property helper method
        if not all_properties:
            raise ValueError("Multi-property report has no properties")

        # Filter to only completed properties (exclude drafts)
        properties = [prop for prop in all_properties if prop.status == 'completed']

        logger.info(f"[MULTI-PROPERTY DOCX] Total properties: {len(all_properties)}, Completed: {len(properties)}, Draft: {len(all_properties) - len(properties)}")

        # Validate that at least one completed property exists
        if not properties:
            raise ValueError("Cannot generate report: All properties are in draft status. Please complete at least one property before generating the report.")

        # Create document
        doc = Document()

        # ===== LETTERHEAD =====
        template_id = user.preferred_letterhead_template or 'classic'
        template = get_template(template_id)
        template.render_letterhead(doc, user, report)

        # ===== SUMMARY PAGE =====
        logger.info("[MULTI-PROPERTY DOCX] Generating summary page")

        # Title: "VALUATION REPORT" - Centered and bold (matching residential/bare land format)
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.paragraph_format.space_before = Pt(0)
        title_para.paragraph_format.space_after = Pt(12)
        title_run = title_para.add_run("VALUATION REPORT")
        title_run.bold = True
        title_run.font.size = Pt(14)  # Updated from 11 to 14
        title_run.font.underline = True  # Add underline
        title_run.font.color.rgb = RGBColor(0, 0, 0)

        # Add spacing before applicant statements (matching residential/bare land)
        spacing_para1 = doc.add_paragraph()
        spacing_para1.paragraph_format.space_before = Pt(8)
        spacing_para1.paragraph_format.space_after = Pt(0)

        # Determine which introduction format to use based on request_type
        if report.request_type == 'organization_request':
            # Organization-side format
            org_intro_paragraphs = generate_organization_side_introduction(report)
            for statement in org_intro_paragraphs:
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
                para.paragraph_format.line_spacing = 0.9

                # Parse label and value for proper formatting
                if ":-" in statement:
                    # This is a label-value pair (Applicant/Address/Contact No)
                    label, value = statement.split(":-", 1)

                    # Add bold label
                    run_label = para.add_run(label.strip())
                    run_label.font.size = FONT_SIZE_BODY
                    run_label.font.bold = True
                    run_label.font.color.rgb = RGBColor(0, 0, 0)

                    # Add separator with proper spacing
                    run_sep = para.add_run(" : ")
                    run_sep.font.size = FONT_SIZE_BODY
                    run_sep.font.color.rgb = RGBColor(0, 0, 0)

                    # Add value
                    run_value = para.add_run(value)
                    run_value.font.size = FONT_SIZE_BODY
                    run_value.font.color.rgb = RGBColor(0, 0, 0)
                else:
                    # This is the introductory paragraph (paragraph 1)
                    run = para.add_run(statement)
                    run.font.size = FONT_SIZE_BODY
                    run.font.color.rgb = RGBColor(0, 0, 0)
        else:
            # Client-side format (default for backward compatibility)
            applicant_statements = generate_applicant_statement(report)
            for statement in applicant_statements:
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
                para.paragraph_format.line_spacing = 0.9
                run = para.add_run(statement)
                run.font.size = FONT_SIZE_BODY
                run.font.color.rgb = RGBColor(0, 0, 0)

            # Add deed/certificate description if applicable (matching residential/bare land logic)
            should_show_deed_sentence = (
                report.property_identification_type in ['deed', 'certificate_of_sale']
                and report.has_deed_info == "yes"
                and report.deeds
            )

            # Backward compatibility: Old reports (NULL type) with deed data should show sentence
            if not report.property_identification_type and report.has_deed_info == "yes" and report.deeds:
                should_show_deed_sentence = True

            if should_show_deed_sentence:
                deed_text = generate_deed_description(report.deeds)
                if deed_text:
                    para = doc.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
                    para.paragraph_format.line_spacing = 0.9
                    run = para.add_run(deed_text)
                    run.font.size = FONT_SIZE_BODY
                    run.font.color.rgb = RGBColor(0, 0, 0)

            # Add submission statement (matching residential/bare land format)
            submission_text = generate_submission_statement(report)
            if submission_text:
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
                para.paragraph_format.line_spacing = 0.9
                run = para.add_run(submission_text)
                run.font.size = FONT_SIZE_BODY
                run.font.color.rgb = RGBColor(0, 0, 0)

        # Add multi-property specific statement about the summary table
        summary_intro_para = doc.add_paragraph()
        summary_intro_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        summary_intro_para.paragraph_format.space_before = Pt(6)
        summary_intro_para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
        summary_intro_para.paragraph_format.line_spacing = 0.9
        summary_intro_run = summary_intro_para.add_run(
            f"This report comprises the valuation of {len(properties)} properties. "
            "The summary of properties and their respective valuations are presented in the following table:"
        )
        summary_intro_run.font.size = FONT_SIZE_BODY
        summary_intro_run.font.color.rgb = RGBColor(0, 0, 0)

        # Property list table
        logger.info("[MULTI-PROPERTY DOCX] Creating property list table")

        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Set custom column widths: #(8%), Location(25%), Description(35%), Extent(12%), Value(20%)
        column_widths = [Inches(0.52), Inches(1.625), Inches(2.275), Inches(0.78), Inches(1.3)]
        for idx_col, width in enumerate(column_widths):
            table.columns[idx_col].width = width

        # Header row
        header_cells = table.rows[0].cells
        headers = ['#', 'Location', 'Description', 'Extent', 'Value (LKR)']
        for i, header_text in enumerate(headers):
            cell = header_cells[i]
            cell.text = header_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = FONT_SIZE_TABLE_HEADER
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Property rows
        grand_total = 0
        for idx, prop in enumerate(properties, start=1):
            row_cells = table.add_row().cells

            # Property number
            row_cells[0].text = str(idx)
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Location
            location = prop.property_village or "N/A"
            if prop.property_district and prop.property_district != location:
                location += f", {prop.property_district}"
            row_cells[1].text = location

            # Description (Lot/Plan)
            desc = ""
            if prop.lot_number:
                lot = prop.lot_number
                desc = f"Lot {lot}"
            if prop.plan_number:
                if desc:
                    desc += f", Plan {prop.plan_number}"
                else:
                    desc = f"Plan {prop.plan_number}"
            if not desc:
                desc = "N/A"
            row_cells[2].text = desc

            # Extent
            extent = prop.land_extent_formatted or "N/A"
            row_cells[3].text = extent

            # Value
            value = to_float(prop.valuation_market_value)
            grand_total += value
            row_cells[4].text = format_currency(value)
            row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Format all cells
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = FONT_SIZE_TABLE_CELL

        # Total row
        total_row_cells = table.add_row().cells
        total_row_cells[0].merge(total_row_cells[3])
        total_row_cells[0].text = "GRAND TOTAL:"
        total_row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for paragraph in total_row_cells[0].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = FONT_SIZE_BODY

        total_row_cells[4].text = format_currency(grand_total)
        total_row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for paragraph in total_row_cells[4].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = FONT_SIZE_BODY

        # Add spacing after table
        doc.add_paragraph()

        # Inspection date (matching residential/bare land format)
        if report.inspection_date:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
            para.paragraph_format.line_spacing = 0.9

            # Bold label
            run_label = para.add_run("Date of Inspection")
            run_label.font.size = FONT_SIZE_BODY
            run_label.font.bold = True
            run_label.font.color.rgb = RGBColor(0, 0, 0)

            # Separator with proper spacing
            run_sep = para.add_run(" : ")
            run_sep.font.size = FONT_SIZE_BODY
            run_sep.font.color.rgb = RGBColor(0, 0, 0)

            # Regular date value
            run_date = para.add_run(report.inspection_date)
            run_date.font.size = FONT_SIZE_BODY
            run_date.font.color.rgb = RGBColor(0, 0, 0)

        # Add concluding statement for CLIENT REQUESTS only (after Date of Inspection)
        concluding_parts = generate_multi_property_concluding_statement(report, user, grand_total)
        if concluding_parts:
            # Statement paragraph (reduced spacing after Date of Inspection)
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
            para.paragraph_format.line_spacing = 0.9
            run = para.add_run(concluding_parts[0])
            run.font.size = FONT_SIZE_BODY
            run.font.color.rgb = RGBColor(0, 0, 0)

            # Valuer name (bold) - increased space for signature
            name_para = doc.add_paragraph()
            name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            name_para.paragraph_format.space_before = Pt(30)
            name_para.paragraph_format.space_after = Pt(2)
            name_run = name_para.add_run(concluding_parts[1])
            name_run.font.size = FONT_SIZE_BODY
            name_run.font.bold = True
            name_run.font.color.rgb = RGBColor(0, 0, 0)

            # Professional designation
            desig_para = doc.add_paragraph()
            desig_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            desig_para.paragraph_format.space_before = Pt(0)
            desig_para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
            desig_run = desig_para.add_run(concluding_parts[2])
            desig_run.font.size = FONT_SIZE_BODY
            desig_run.font.color.rgb = RGBColor(0, 0, 0)

        # Add special note if applicable (matching residential/bare land format)
        if report.has_special_note == "yes" and report.special_note_text:
            # Note label
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = SUBHEADING_SPACE_AFTER
            para.paragraph_format.line_spacing = 0.9
            run = para.add_run("Note:")
            run.bold = True
            run.font.size = FONT_SIZE_BODY
            run.font.color.rgb = RGBColor(0, 0, 0)

            # Note text
            note_para = doc.add_paragraph()
            note_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            note_para.paragraph_format.space_before = Pt(0)
            note_para.paragraph_format.space_after = BODY_PARA_SPACE_AFTER
            note_para.paragraph_format.line_spacing = 0.9
            note_run = note_para.add_run(report.special_note_text)
            note_run.font.size = FONT_SIZE_BODY
            note_run.font.color.rgb = RGBColor(0, 0, 0)

        # REMOVED: Forced page break - let natural pagination handle page breaks
        # This allows better space utilization, especially for organization requests
        # where the table ends earlier and there's often significant space remaining on page 1
        # doc.add_page_break()

        # ===== INDIVIDUAL PROPERTY SECTIONS =====
        logger.info("[MULTI-PROPERTY DOCX] Generating individual property sections")

        for prop_idx, prop in enumerate(properties, start=1):
            logger.info(f"[MULTI-PROPERTY DOCX] Generating section for property {prop_idx}/{len(properties)}")

            # Professional property description (NO "PROPERTY 1" header - removed per user request)
            # Format: "The Property Depicted as Lot [X] in Plan No: [Y]"
            if prop.lot_number and prop.plan_number:
                # Extract just the lot number/identifier (remove "Plan No" prefix if present)
                lot_desc = prop.lot_number.strip() if prop.lot_number else ''

                # Remove common prefixes that shouldn't be in lot description
                prefixes_to_remove = ['plan no', 'plan no:', 'lot plan no', 'lot plan no:']
                lot_desc_lower = lot_desc.lower()
                for prefix in prefixes_to_remove:
                    if lot_desc_lower.startswith(prefix):
                        lot_desc = lot_desc[len(prefix):].strip()
                        break

                # Ensure lot description has "Lot" prefix (if it doesn't already)
                if not lot_desc.lower().startswith('lot'):
                    lot_desc = f"Lot {lot_desc}"

                plan_formatted = format_no_field("Plan", prop.plan_number)

                # Combine date with first line
                plan_date = prop.plan_date or ''
                if plan_date:
                    prop_desc = f"The Property Depicted as {lot_desc} in {plan_formatted} Dated {plan_date}"
                else:
                    prop_desc = f"The Property Depicted as {lot_desc} in {plan_formatted}"

                prop_desc_para = doc.add_paragraph()
                prop_desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                prop_desc_para.paragraph_format.space_before = Pt(0)
                prop_desc_para.paragraph_format.space_after = Pt(0)  # No spacing
                prop_desc_para.paragraph_format.line_spacing = 1.0  # Single line spacing
                prop_desc_run = prop_desc_para.add_run(prop_desc)
                prop_desc_run.bold = True
                prop_desc_run.font.size = Pt(12)  # Updated from 10 to 12
                prop_desc_run.font.underline = True  # Add underline
                prop_desc_run.font.color.rgb = RGBColor(0, 0, 0)

                # Second line starts with "made by" - no spacing between lines
                if prop.licensed_surveyor_name:
                    plan_info = f"made by {prop.licensed_surveyor_name} Licensed Surveyor."
                    plan_info_para = doc.add_paragraph()
                    plan_info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    plan_info_para.paragraph_format.space_before = Pt(0)  # No spacing
                    plan_info_para.paragraph_format.space_after = Pt(12)
                    plan_info_para.paragraph_format.line_spacing = 1.0  # Single line spacing
                    plan_info_run = plan_info_para.add_run(plan_info)
                    plan_info_run.bold = True
                    plan_info_run.font.size = Pt(12)  # Updated from 10 to 12
                    plan_info_run.font.underline = True  # Add underline
                    plan_info_run.font.color.rgb = RGBColor(0, 0, 0)
                else:
                    # Just add spacing if no plan info
                    spacing_para = doc.add_paragraph()
                    spacing_para.paragraph_format.space_after = Pt(12)
            else:
                # Fallback to location-based description if no plan info
                location = prop.property_village or "N/A"
                if prop.property_district:
                    location += f", {prop.property_district}"
                prop_loc_para = doc.add_paragraph()
                prop_loc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                prop_loc_para.paragraph_format.space_after = Pt(12)
                prop_loc_run = prop_loc_para.add_run(location)
                prop_loc_run.font.size = FONT_SIZE_BODY
                prop_loc_run.font.color.rgb = RGBColor(0, 0, 0)

            # Generate property sections (comprehensive full sections)
            generate_property_sections(doc, prop, report, user)

            # NO page breaks between properties - continuous flow
            # User will manually adjust page breaks as needed

        # ===== INVOICE SECTION =====
        if report.invoice_data:
            logger.info("[MULTI-PROPERTY DOCX] Generating invoice section")
            # Page break handled inside generate_invoice_section()
            generate_invoice_section(doc, report.invoice_data, user, report)

        # Note: Per-property certification now included in individual property sections
        # This allows different valuers to certify different properties in the same report

        # ===== SAVE DOCUMENT =====
        logger.info("[MULTI-PROPERTY DOCX] Saving document")
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        if file_stream.getvalue() == b'':
            raise ValueError("Generated multi-property document is empty")

        logger.info(f"[MULTI-PROPERTY DOCX] Document generated successfully, size: {len(file_stream.getvalue())} bytes")
        return file_stream

    except Exception as e:
        logger.error(f"[MULTI-PROPERTY DOCX] Error: {e}", exc_info=True)
        raise


